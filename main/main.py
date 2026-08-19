import matplotlib
matplotlib.use("TkAgg")

import sys
import os
import json
import re
import threading
import calendar
import webbrowser
import requests
import tempfile
import subprocess
from datetime import datetime

import customtkinter as ctk
from tkinter import messagebox, filedialog

import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from auth_backend import (
    login, registro, llamar_gemini, llamar_groq, token_guardado, borrar_token,
    actualizar_hogar_principal, borrar_cuenta_usuario, obtener_chats_cloud, guardar_chat_cloud, borrar_chat_cloud,
    enviar_mensaje_soporte, obtener_mensajes_soporte, admin_listar_usuarios,
    admin_aplicar_ban, admin_desbanear, admin_ver_mensajes_raw,
    admin_obtener_tickets_soporte, admin_responder_ticket_soporte,
    obtener_cuentas_guardadas, cambiar_a_cuenta, eliminar_cuenta_switcher,
    agregar_cuenta_secundaria, actualizar_rol_usuario,
    verificar_estado_baneo, admin_borrar_ticket_soporte,
    tutoria_listar_profesores, tutoria_enviar_solicitud, tutoria_obtener_estado_alumno,
    tutoria_profesor_obtener_solicitudes_y_alumnos, tutoria_profesor_responder_solicitud,
    tutoria_enviar_mensaje, tutoria_obtener_mensajes_chat, tutoria_desvincular
)
from config_manager import (
    obtener_ajustes_tts, guardar_ajustes_tts, obtener_idioma, guardar_idioma
)
from i18n import t, fijar_idioma, IDIOMAS_DISPONIBLES, obtener_idioma_activo
from tts_engine import tts_engine, VOICES_DISPONIBLES, VELOCIDADES_DISPONIBLES

# ─────────────────────────────────────────────
#  INICIALIZACIÓN DE IDIOMA Y CONFIGURACIÓN GLOBAL
# ─────────────────────────────────────────────
fijar_idioma(obtener_idioma())

VERSION_APP = "1.6"

def es_version_superior(remota: str, local: str) -> bool:
    """Compara si la versión remota (ej: 'v1.2') es superior a la local (ej: '1.1')."""
    try:
        def parse_nums(v):
            return [int(x) for x in re.findall(r'\d+', str(v))]
        v_remota = parse_nums(remota)
        v_local = parse_nums(local)
        return v_remota > v_local
    except Exception:
        return remota.lstrip('vV').strip() != local.lstrip('vV').strip()

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

# Constantes de diseño armonizadas con la web
COLOR_BG_DARK       = "#050811"  # Fondo base de la ventana
COLOR_BG_SIDEBAR    = "#070c18"  # Fondo sidebar lateral
COLOR_BG_CARD       = "#0a1124"  # Paneles y tarjetas base
COLOR_BG_CARD_LIGHT = "#0f1a35"  # Entradas de texto y visores
COLOR_BG_SURFACE    = "#152449"  # Superficies activas
COLOR_BORDER        = "#1e3a6a"  # Bordes azulados sutiles
COLOR_BORDER_GLOW   = "#3b82f6"  # Borde con resplandor activo

COLOR_ACCENT_PRIMARY = "#2563eb" # Azul Eléctrico principal
COLOR_ACCENT_HOVER   = "#3b82f6" # Azul hover brillante
COLOR_ACCENT_CYAN    = "#06b6d4" # Cian brillante
COLOR_ACCENT_CYAN_HOVER = "#0891b2"
COLOR_ACCENT_SKY     = "#38bdf8" # Celeste
COLOR_ACCENT_PURPLE  = "#6366f1" # Indigo (Docentes)
COLOR_ACCENT_PURPLE_HOVER = "#4f46e5"

COLOR_TEXT_MAIN      = "#f8fafc" # Blanco nítido
COLOR_TEXT_MUTED     = "#94a3b8" # Gris azulado secundario
COLOR_TEXT_DIM       = "#64748b" # Gris tenue
COLOR_SUCCESS        = "#10b981" # Verde esmeralda
COLOR_SUCCESS_HOVER  = "#059669"
COLOR_WARNING        = "#f59e0b" # Ámbar
COLOR_DANGER         = "#ef4444" # Rojo coral
COLOR_DANGER_HOVER   = "#dc2626"

def construir_prompt(instrucciones, historial=None):
    partes = [instrucciones.strip()]
    if historial:
        for msg in historial:
            rol = "IA" if msg.get("role") == "assistant" else msg.get("role", "Usuario").capitalize()
            partes.append(f"{rol}: {msg.get('content', '')}")
    return "\n\n".join(partes)


# ══════════════════════════════════════════════════════════════════════════════
#  MODAL DE PROTECCIÓN DE HOGAR PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
class VentanaConfirmacionHogar(ctk.CTkToplevel):
    """Modal estilo Netflix para confirmar ubicación/red principal de estudio."""
    def __init__(self, parent, hogar_info: dict, on_finalizar):
        super().__init__(parent)
        self.title("🏠 Hogar Principal de Estudio – KernossIA")
        self.geometry("580x480")
        self.minsize(520, 420)
        self.configure(fg_color=COLOR_BG_DARK)
        self.transient(parent)
        self.grab_set()
        self.hogar_info = hogar_info
        self.on_finalizar = on_finalizar
        self._build_ui()

    def _build_ui(self):
        # Cabecera
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", padx=25, pady=(22, 10))

        ctk.CTkLabel(header, text="🏠 ¿Estás en tu Hogar Principal?",
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(anchor="w")
        ctk.CTkLabel(header, text="Control de ubicación y protección de cuenta compartida",
                     font=("Segoe UI", 12), text_color=COLOR_TEXT_MUTED).pack(anchor="w", pady=(2, 0))

        # Tarjeta informativa
        card = ctk.CTkFrame(self, fg_color=COLOR_BG_CARD, corner_radius=14,
                            border_width=1, border_color=COLOR_BORDER)
        card.pack(fill="both", expand=True, padx=25, pady=(5, 15))

        nombre_hogar = self.hogar_info.get("hogar_nombre", "Hogar Principal")
        ip_actual = self.hogar_info.get("ip_actual", "Nueva Red")

        ctk.CTkLabel(card, text="📍 Ubicación o Red No Habitual Detectada",
                     font=("Segoe UI", 14, "bold"), text_color=COLOR_WARNING).pack(anchor="w", padx=18, pady=(16, 6))

        desc = (
            f"Tu cuenta de KernossIA tiene registrado como Hogar Principal: '{nombre_hogar}'.\n\n"
            f"Hemos detectado que estás iniciando sesión desde una red o ubicación diferente ({ip_actual}). "
            "Para evitar el uso compartido indebido y proteger tu cuenta:\n\n"
            "• Si estás en la biblioteca, cafetería o de viaje, puedes continuar normalmente.\n"
            "• Si te has mudado o esta es tu nueva red fija, puedes actualizar tu Hogar Principal."
        )
        ctk.CTkLabel(card, text=desc, font=("Segoe UI", 12), text_color=COLOR_TEXT_MAIN,
                     wraplength=480, justify="left").pack(anchor="w", padx=18, pady=(0, 15))

        # Botones de Acción
        btns = ctk.CTkFrame(self, fg_color="transparent")
        btns.pack(fill="x", padx=25, pady=(0, 20))

        ctk.CTkButton(
            btns, text="✈️ Estudiar Fuera de Casa (De viaje)",
            height=40, font=("Segoe UI", 12, "bold"),
            fg_color=COLOR_BG_SURFACE, hover_color=COLOR_BORDER,
            command=self._continuar_temporal
        ).pack(fill="x", pady=(0, 8))

        ctk.CTkButton(
            btns, text="🏡 Establecer Esta Red como mi Hogar Principal",
            height=40, font=("Segoe UI", 12, "bold"),
            fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
            command=self._actualizar_hogar
        ).pack(fill="x")

    def _continuar_temporal(self):
        self.destroy()
        if self.on_finalizar:
            self.on_finalizar()

    def _actualizar_hogar(self):
        exito, msg = actualizar_hogar_principal("Hogar Principal de Estudio")
        if exito:
            messagebox.showinfo("Hogar Actualizado", "Se ha registrado esta red como tu Hogar Principal de Estudio.")
        self.destroy()
        if self.on_finalizar:
            self.on_finalizar()


# ══════════════════════════════════════════════════════════════════════════════
#  PANTALLA DE LOGIN / REGISTRO
# ══════════════════════════════════════════════════════════════════════════════
class PantallaLogin(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("KernossIA – Acceso")
        self.geometry("520x680")
        self.resizable(False, False)
        self.configure(fg_color=COLOR_BG_DARK)
        self.usuario_autenticado = None
        self._build_ui()

    def _al_cambiar_idioma(self, nombre_idioma):
        for code, name in IDIOMAS_DISPONIBLES.items():
            if name == nombre_idioma:
                fijar_idioma(code)
                guardar_idioma(code)
                break
        # Reconstruir la UI en el nuevo idioma seleccionado
        for w in self.winfo_children():
            w.destroy()
        self._build_ui()

    def _build_ui(self):
        # Barra superior con Selector de Idioma visible antes de iniciar sesión o registrarse
        bar_top_lang = ctk.CTkFrame(self, fg_color="transparent")
        bar_top_lang.pack(fill="x", padx=35, pady=(15, 0))

        ctk.CTkLabel(
            bar_top_lang, text=t("lbl_selecciona_idioma"),
            font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED
        ).pack(side="left")

        idiomas_nombres = list(IDIOMAS_DISPONIBLES.values())
        idioma_actual = IDIOMAS_DISPONIBLES.get(obtener_idioma_activo(), "🇪🇸 Español")

        self.combo_idioma_login = ctk.CTkComboBox(
            bar_top_lang, values=idiomas_nombres,
            font=("Segoe UI", 11), width=150, height=28,
            fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER,
            command=self._al_cambiar_idioma
        )
        self.combo_idioma_login.set(idioma_actual)
        self.combo_idioma_login.pack(side="right")

        # Header con logo
        header_frame = ctk.CTkFrame(self, fg_color="transparent")
        header_frame.pack(pady=(20, 15))

        ctk.CTkLabel(header_frame, text=t("app_nombre"),
                     font=("Segoe UI", 36, "bold"), text_color=COLOR_ACCENT_SKY).pack()
        ctk.CTkLabel(header_frame, text=t("app_subtitulo"),
                     font=("Segoe UI", 12), text_color=COLOR_TEXT_MUTED).pack(pady=(2, 0))

        # Tarjeta contenedora de Login
        self.frame = ctk.CTkFrame(self, corner_radius=18, fg_color=COLOR_BG_CARD,
                                  border_width=1, border_color=COLOR_BORDER)
        self.frame.pack(padx=45, fill="x")

        self.tab = ctk.CTkTabview(self.frame, height=340, fg_color=COLOR_BG_CARD_LIGHT,
                                  segmented_button_selected_color=COLOR_ACCENT_PRIMARY,
                                  segmented_button_selected_hover_color=COLOR_ACCENT_HOVER,
                                  segmented_button_unselected_color=COLOR_BG_CARD,
                                  segmented_button_unselected_hover_color=COLOR_BG_SURFACE)
        self.tab.pack(fill="x", padx=16, pady=16)
        self.tab.add(t("tab_login"))
        self.tab.add(t("tab_registro"))

        # ── LOGIN ──
        login_tab = self.tab.tab(t("tab_login"))
        ctk.CTkLabel(login_tab, text=t("lbl_email"), anchor="w", font=("Segoe UI", 12, "bold"),
                     text_color=COLOR_TEXT_MAIN).pack(fill="x", padx=10, pady=(10, 2))
        self.entry_login_email = ctk.CTkEntry(login_tab, placeholder_text=t("placeholder_email"), height=38,
                                              fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_login_email.pack(fill="x", padx=10)

        ctk.CTkLabel(login_tab, text=t("lbl_pass"), anchor="w", font=("Segoe UI", 12, "bold"),
                     text_color=COLOR_TEXT_MAIN).pack(fill="x", padx=10, pady=(10, 2))
        self.entry_login_pass = ctk.CTkEntry(login_tab, placeholder_text=t("placeholder_pass"), show="•", height=38,
                                             fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_login_pass.pack(fill="x", padx=10)
        self.entry_login_pass.bind("<Return>", lambda e: self._login())

        ctk.CTkButton(login_tab, text=t("btn_login"), height=42,
                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                      font=("Segoe UI", 13, "bold"),
                      command=self._login).pack(fill="x", padx=10, pady=(20, 5))
        self.lbl_login_error = ctk.CTkLabel(login_tab, text="", text_color=COLOR_DANGER, font=("Segoe UI", 11))
        self.lbl_login_error.pack()

        # ── REGISTRO ──
        reg = self.tab.tab(t("tab_registro"))
        ctk.CTkLabel(reg, text=t("lbl_nombre"), anchor="w", font=("Segoe UI", 12, "bold"),
                     text_color=COLOR_TEXT_MAIN).pack(fill="x", padx=10, pady=(6, 2))
        self.entry_reg_nombre = ctk.CTkEntry(reg, placeholder_text=t("placeholder_nombre"), height=36,
                                             fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_reg_nombre.pack(fill="x", padx=10)

        ctk.CTkLabel(reg, text=t("lbl_email"), anchor="w", font=("Segoe UI", 12, "bold"),
                     text_color=COLOR_TEXT_MAIN).pack(fill="x", padx=10, pady=(6, 2))
        self.entry_reg_email = ctk.CTkEntry(reg, placeholder_text=t("placeholder_email"), height=36,
                                            fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_reg_email.pack(fill="x", padx=10)

        ctk.CTkLabel(reg, text=t("lbl_pass"), anchor="w", font=("Segoe UI", 12, "bold"),
                     text_color=COLOR_TEXT_MAIN).pack(fill="x", padx=10, pady=(6, 2))
        self.entry_reg_pass = ctk.CTkEntry(reg, placeholder_text=t("placeholder_pass_reg"), show="•", height=36,
                                           fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_reg_pass.pack(fill="x", padx=10)

        ctk.CTkLabel(reg, text=t("lbl_rol"), anchor="w", font=("Segoe UI", 12, "bold"),
                     text_color=COLOR_TEXT_MAIN).pack(fill="x", padx=10, pady=(6, 2))
        self.combo_rol = ctk.CTkOptionMenu(reg, values=[t("lbl_rol_alumno"), t("lbl_rol_profesor")], height=36,
                                           fg_color=COLOR_ACCENT_PRIMARY,
                                           button_color=COLOR_ACCENT_HOVER,
                                           dropdown_fg_color=COLOR_BG_CARD)
        self.combo_rol.pack(fill="x", padx=10)

        ctk.CTkButton(reg, text=t("btn_registro"), height=40,
                      fg_color=COLOR_ACCENT_CYAN, hover_color=COLOR_ACCENT_CYAN_HOVER,
                      font=("Segoe UI", 13, "bold"),
                      command=self._registrar).pack(fill="x", padx=10, pady=(12, 4))
        self.lbl_reg_error = ctk.CTkLabel(reg, text="", text_color=COLOR_DANGER, font=("Segoe UI", 11))
        self.lbl_reg_error.pack()

        ctk.CTkLabel(self, text=t("login_privacidad"),
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_DIM).pack(pady=(16, 0))

    def _login(self):
        email    = self.entry_login_email.get().strip().lower()
        password = self.entry_login_pass.get()
        if not email or not password:
            self.lbl_login_error.configure(text=t("login_error_campos"))
            return
        self.lbl_login_error.configure(text=t("login_conectando"), text_color=COLOR_ACCENT_SKY)
        self.update()
        exito, error, sesion, hogar_info = login(email, password)
        if not exito:
            self.lbl_login_error.configure(text=error, text_color=COLOR_DANGER)
            return

        self.usuario_autenticado = sesion

        # Comprobar si se conecta desde fuera de su Hogar Principal (estilo Netflix)
        if hogar_info.get("hogar_estado") == "fuera_de_hogar":
            def _cerrar_y_entrar():
                self.destroy()
            VentanaConfirmacionHogar(self, hogar_info, on_finalizar=_cerrar_y_entrar)
        else:
            self.destroy()

    def _registrar(self):
        nombre   = self.entry_reg_nombre.get().strip()
        email    = self.entry_reg_email.get().strip().lower()
        password = self.entry_reg_pass.get()
        rol_txt  = self.combo_rol.get()
        # Mapear rol si está en otro idioma
        rol = "Profesor" if rol_txt in [t("lbl_rol_profesor"), "Profesor", "Teacher", "Lehrer / Dozent"] else "Alumno"

        if not nombre or not email or not password:
            self.lbl_reg_error.configure(text=t("login_error_campos"))
            return
        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            self.lbl_reg_error.configure(text=t("login_error_correo"))
            return
        if len(password) < 6:
            self.lbl_reg_error.configure(text=t("login_error_pass_len"))
            return
        self.lbl_reg_error.configure(text=t("reg_creando"), text_color=COLOR_ACCENT_SKY)
        self.update()
        exito, error, sesion = registro(nombre, email, password, rol)
        if not exito:
            self.lbl_reg_error.configure(text=error, text_color=COLOR_DANGER)
            return
        self.usuario_autenticado = sesion
        self.destroy()


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: CALCULADOR DE MEDIAS
# ══════════════════════════════════════════════════════════════════════════════
class ModuloCalculador(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.notas_finales  = []
        self.nombres_notas  = []
        self.porcentajes    = []
        self.canvas_grafico = None
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=1)

        frame_titulo = ctk.CTkFrame(self, fg_color="transparent")
        frame_titulo.pack(fill="x", padx=20, pady=(20, 15))
        ctk.CTkLabel(frame_titulo, text="📊 Calculadora de Medias Ponderadas",
                     font=("Segoe UI", 28, "bold"),
                     text_color=COLOR_ACCENT_SKY).pack()
        ctk.CTkLabel(frame_titulo, text="Organiza, pondera y visualiza el progreso de tus calificaciones",
                     font=("Segoe UI", 13),
                     text_color=COLOR_TEXT_MUTED).pack(pady=(3, 0))

        self.frame_entrada = ctk.CTkFrame(self, corner_radius=14, border_width=1,
                                          fg_color=COLOR_BG_CARD,
                                          border_color=COLOR_BORDER)
        self.frame_entrada.pack(padx=25, pady=15, fill="both", expand=False)
        self.frame_entrada.grid_columnconfigure((0, 1), weight=1)

        ctk.CTkLabel(self.frame_entrada, text="Materia o Asignatura",
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=0, padx=15, pady=(15, 5), sticky="w")
        self.entrada_nombre = ctk.CTkEntry(self.frame_entrada, placeholder_text="Ej: Matemáticas, Física...",
                                           height=40, font=("Segoe UI", 12),
                                           fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entrada_nombre.grid(row=1, column=0, padx=15, pady=(0, 12), sticky="ew")

        ctk.CTkLabel(self.frame_entrada, text="Nota Directa (0-10)",
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=1, padx=15, pady=(15, 5), sticky="w")
        self.entrada_nota_directa = ctk.CTkEntry(self.frame_entrada, placeholder_text="Ej: 9.5",
                                                  height=40, font=("Segoe UI", 12),
                                                  fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entrada_nota_directa.grid(row=1, column=1, padx=15, pady=(0, 12), sticky="ew")

        ctk.CTkLabel(self.frame_entrada,
                     text="% del Total (Ej: 60% Pruebas / 30% Proyectos / 10% Actitud)",
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=2, column=0, padx=15, pady=(0, 5), sticky="w")
        self.entrada_porcentaje = ctk.CTkEntry(self.frame_entrada,
                                               placeholder_text="Ej: 40",
                                               height=40, font=("Segoe UI", 12),
                                               fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entrada_porcentaje.grid(row=3, column=0, columnspan=2, padx=15, pady=(0, 15), sticky="ew")

        frame_botones = ctk.CTkFrame(self, fg_color="transparent")
        frame_botones.pack(fill="x", padx=25, pady=8)
        frame_botones.grid_columnconfigure((0, 1), weight=1)

        ctk.CTkButton(frame_botones, text="➕ Guardar Nota Directa",
                      command=self.agregar_nota_principal,
                      fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                      height=40, font=("Segoe UI", 12, "bold")).grid(row=0, column=0, padx=(0, 6), sticky="ew")
        ctk.CTkButton(frame_botones, text="📂 Agregar Bloques / Subnotas",
                      command=self.gestionar_subnotas,
                      fg_color=COLOR_BG_SURFACE, border_width=1,
                      border_color=COLOR_ACCENT_CYAN,
                      hover_color=COLOR_ACCENT_PRIMARY,
                      height=40, font=("Segoe UI", 12, "bold")).grid(row=0, column=1, padx=(6, 0), sticky="ew")

        frame_principal = ctk.CTkFrame(self, fg_color="transparent")
        frame_principal.pack(padx=25, pady=12, fill="both", expand=True)
        frame_principal.grid_columnconfigure((0, 1), weight=1)
        frame_principal.grid_rowconfigure(1, weight=1)

        ctk.CTkLabel(frame_principal, text="Registro Detallado",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=0, pady=(0, 8), padx=(0, 10), sticky="w")
        self.salida_texto = ctk.CTkTextbox(frame_principal, width=500, height=320,
                                           font=("Consolas", 12), corner_radius=12,
                                           fg_color=COLOR_BG_CARD, border_width=1,
                                           border_color=COLOR_BORDER, state="disabled")
        self.salida_texto.grid(row=1, column=0, padx=(0, 10), sticky="nsew")

        ctk.CTkLabel(frame_principal, text="Gráfica de Rendimiento",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=1, pady=(0, 8), padx=(10, 0), sticky="w")
        self.frame_grafico = ctk.CTkFrame(frame_principal, corner_radius=12, border_width=1,
                                          fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.frame_grafico.grid(row=1, column=1, padx=(10, 0), sticky="nsew")

        frame_acciones = ctk.CTkFrame(self, fg_color="transparent")
        frame_acciones.pack(padx=25, pady=15, fill="x")
        frame_acciones.grid_columnconfigure((0, 1, 2), weight=1)

        ctk.CTkButton(frame_acciones, text="📊 Calcular Promedio Final",
                      command=self.calcular_total_final,
                      height=44, font=("Segoe UI", 13, "bold"),
                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER).grid(row=0, column=0, padx=(0, 6), sticky="ew")
        ctk.CTkButton(frame_acciones, text="🧹 Limpiar Todo",
                      command=self.limpiar_datos,
                      fg_color=COLOR_BG_SURFACE, hover_color=COLOR_DANGER_HOVER,
                      border_width=1, border_color=COLOR_BORDER,
                      height=44, font=("Segoe UI", 13, "bold")).grid(row=0, column=1, padx=3, sticky="ew")
        ctk.CTkButton(frame_acciones, text="📄 Exportar a Word (.docx)",
                      command=self.exportar_a_word,
                      fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                      height=44, font=("Segoe UI", 13, "bold")).grid(row=0, column=2, padx=(6, 0), sticky="ew")

    def calcular_media(self, lista):
        if not lista:
            return 0
        return sum(lista) / len(lista)

    def actualizar_grafico(self):
        if not self.notas_finales:
            return
        if self.canvas_grafico:
            self.canvas_grafico.get_tk_widget().destroy()
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
        fig.patch.set_facecolor('#0a1124')
        colores = ['#10b981' if n >= 7 else '#f59e0b' if n >= 5 else '#ef4444' for n in self.notas_finales]
        ax1.bar(range(len(self.nombres_notas)), self.notas_finales, color=colores, alpha=0.85, edgecolor='#3b82f6')
        ax1.set_xticks(range(len(self.nombres_notas)))
        ax1.set_xticklabels(self.nombres_notas, rotation=35, ha='right', fontsize=9, color='#cbd5e1')
        ax1.set_ylabel('Calificación', color='#94a3b8', fontsize=10)
        ax1.set_title('Notas por Asignatura', color='#f8fafc', fontsize=12, fontweight='bold')
        ax1.set_ylim(0, 10)
        ax1.set_facecolor('#0f1a35')
        ax1.tick_params(colors='#94a3b8')
        ax1.grid(axis='y', alpha=0.2, color='#3b82f6')

        promedio = self.calcular_media(self.notas_finales)
        por_encima = sum(1 for n in self.notas_finales if n >= promedio)
        por_debajo = sum(1 for n in self.notas_finales if n < promedio)
        ax2.pie([max(por_encima, 0.001), max(por_debajo, 0.001)],
                labels=[f'Arriba media\n({por_encima})', f'Debajo media\n({por_debajo})'],
                colors=['#10b981', '#ef4444'], autopct='%1.1f%%', startangle=90,
                textprops={'color': '#f8fafc', 'fontsize': 10})
        ax2.set_facecolor('#0a1124')
        ax2.set_title(f'Distribución (Promedio: {promedio:.2f})', color='#38bdf8', fontsize=12, fontweight='bold')
        plt.tight_layout()
        self.canvas_grafico = FigureCanvasTkAgg(fig, master=self.frame_grafico)
        self.canvas_grafico.draw()
        self.canvas_grafico.get_tk_widget().pack(fill="both", expand=True)

    def gestionar_subnotas(self):
        nombre_principal = self.entrada_nombre.get().strip()
        if not nombre_principal:
            messagebox.showwarning("Atención", "Escribe el nombre de la asignatura primero.")
            return
        d = ctk.CTkInputDialog(text=f"¿Cuántos bloques tiene {nombre_principal}?\n(Ej: 3 → Exámenes, Prácticas, Trabajo)",
                               title="Bloques de calificación")
        res = d.get_input()
        self.focus_force(); self.lift()
        if not res or not res.isdigit():
            return
        num_bloques = int(res)
        self.salida_texto.configure(state="normal")
        self.salida_texto.insert("end", f"\n{'─'*55}\n  📚 {nombre_principal.upper()}\n{'─'*55}\n")
        nota_final_asignatura = 0.0
        suma_pesos_bloques = 0.0
        for i in range(num_bloques):
            d = ctk.CTkInputDialog(text=f"Nombre del bloque {i+1}:", title="Nombre del bloque")
            nombre_bloque = d.get_input(); self.focus_force(); self.lift()
            if nombre_bloque is None: break
            d = ctk.CTkInputDialog(text=f"¿Cuánto vale '{nombre_bloque}' en % sobre el total?", title="% del bloque")
            peso_bloque_str = d.get_input(); self.focus_force(); self.lift()
            if peso_bloque_str is None: break
            try:
                peso_bloque = float(peso_bloque_str)
            except ValueError:
                messagebox.showerror("Error", "Número inválido."); continue
            d = ctk.CTkInputDialog(text=f"¿Cuántas notas hay dentro de '{nombre_bloque}'?", title="Notas del bloque")
            res_sub = d.get_input(); self.focus_force(); self.lift()
            if not res_sub or not res_sub.isdigit(): continue
            num_sub = int(res_sub)
            self.salida_texto.insert("end", f"\n  📂 {nombre_bloque} ({peso_bloque:.0f}% del total)\n")
            nota_bloque_ponderada = 0.0; suma_pesos_sub = 0.0
            for j in range(num_sub):
                d = ctk.CTkInputDialog(text=f"Nombre de la nota {j+1}:", title="Nota")
                nombre_sub = d.get_input(); self.focus_force(); self.lift()
                if nombre_sub is None: break
                d = ctk.CTkInputDialog(text=f"¿Cuánto vale '{nombre_sub}' en % dentro de '{nombre_bloque}'?", title=f"% dentro de {nombre_bloque}")
                peso_sub_str = d.get_input(); self.focus_force(); self.lift()
                if peso_sub_str is None: break
                try:
                    peso_sub = float(peso_sub_str)
                except ValueError:
                    messagebox.showerror("Error", "Número inválido."); continue
                d = ctk.CTkInputDialog(text=f"Calificación de '{nombre_sub}':", title="Calificación")
                valor_sub_str = d.get_input(); self.focus_force(); self.lift()
                if valor_sub_str is None: break
                try:
                    valor_sub = float(valor_sub_str)
                except ValueError:
                    messagebox.showerror("Error", "Número inválido."); continue
                nota_bloque_ponderada += valor_sub * (peso_sub / 100)
                suma_pesos_sub += peso_sub
                self.salida_texto.insert("end", f"      • {nombre_sub:20} {peso_sub:.0f}% → {valor_sub:.2f}\n")
            if suma_pesos_sub > 0 and suma_pesos_sub != 100:
                nota_bloque_ponderada = nota_bloque_ponderada / (suma_pesos_sub / 100)
            self.salida_texto.insert("end", f"    ✓ Nota bloque '{nombre_bloque}': {nota_bloque_ponderada:.2f}\n")
            nota_final_asignatura += nota_bloque_ponderada * (peso_bloque / 100)
            suma_pesos_bloques += peso_bloque
        if suma_pesos_bloques > 0 and suma_pesos_bloques != 100:
            nota_final_asignatura = nota_final_asignatura / (suma_pesos_bloques / 100)
        self.nombres_notas.append(nombre_principal)
        self.notas_finales.append(nota_final_asignatura)
        self.porcentajes.append(100.0)
        self.salida_texto.insert("end", f"\n  {'═'*50}\n  ✅ NOTA FINAL {nombre_principal.upper()}: {nota_final_asignatura:.2f}\n  {'═'*50}\n\n")
        self.salida_texto.configure(state="disabled")
        self.entrada_nombre.delete(0, "end")
        self.actualizar_grafico()

    def agregar_nota_principal(self):
        nombre = self.entrada_nombre.get().strip()
        nota_str = self.entrada_nota_directa.get().strip()
        porcentaje_str = self.entrada_porcentaje.get().strip()
        if nombre and nota_str:
            try:
                nota = float(nota_str)
                if porcentaje_str:
                    porcentaje_sobre_total = float(porcentaje_str)
                else:
                    total = len(self.notas_finales) + 1
                    porcentaje_sobre_total = (1 / total) * 100
                self.nombres_notas.append(nombre)
                self.notas_finales.append(nota)
                self.porcentajes.append(porcentaje_sobre_total)
                self.salida_texto.configure(state="normal")
                self.salida_texto.insert("end", f"  ✓ {nombre:25} → {nota:.2f} ({porcentaje_sobre_total:.1f}%)\n")
                self.salida_texto.configure(state="disabled")
                self.entrada_nombre.delete(0, "end")
                self.entrada_nota_directa.delete(0, "end")
                self.entrada_porcentaje.delete(0, "end")
                self.actualizar_grafico()
            except ValueError:
                messagebox.showerror("Error", "La calificación y el porcentaje deben ser números válidos.")
        else:
            messagebox.showwarning("Campos incompletos", "Por favor completa materia y calificación.")

    def calcular_total_final(self):
        if not self.notas_finales:
            messagebox.showinfo("Sin datos", "No hay calificaciones guardadas para calcular.")
            return
        suma_ponderada = sum(n * (p / 100) for n, p in zip(self.notas_finales, self.porcentajes))
        total_porcentaje = sum(self.porcentajes)
        media_total = suma_ponderada / (total_porcentaje / 100) if total_porcentaje != 100 else suma_ponderada
        self.salida_texto.configure(state="normal")
        self.salida_texto.insert("end", f"\n{'═'*55}\n  📊 RESULTADO FINAL\n")
        self.salida_texto.insert("end", f"  Asignaturas: {len(self.nombres_notas)}\n")
        self.salida_texto.insert("end", f"  Peso total asignado: {total_porcentaje:.1f}%\n")
        self.salida_texto.insert("end", f"  Promedio Ponderado General: {media_total:.2f}\n{'═'*55}\n\n")
        self.salida_texto.configure(state="disabled")
        self.salida_texto.see("end")

    def limpiar_datos(self):
        if messagebox.askyesno("Confirmar", "¿Deseas limpiar todas las notas?"):
            self.notas_finales.clear(); self.nombres_notas.clear(); self.porcentajes.clear()
            self.salida_texto.configure(state="normal")
            self.salida_texto.delete("1.0", "end")
            self.salida_texto.configure(state="disabled")
            self.entrada_nombre.delete(0, "end")
            self.entrada_nota_directa.delete(0, "end")
            self.entrada_porcentaje.delete(0, "end")
            if self.canvas_grafico:
                self.canvas_grafico.get_tk_widget().destroy()
                self.canvas_grafico = None

    def exportar_a_word(self):
        if not self.notas_finales:
            messagebox.showwarning("Sin datos", "No hay calificaciones para exportar.")
            return
        doc = Document()
        titulo = doc.add_heading('Reporte de Calificaciones – KernossIA', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha = doc.add_paragraph(f'Generado el: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        tabla = doc.add_table(rows=len(self.nombres_notas) + 1, cols=3)
        tabla.style = 'Light Grid Accent 1'
        enc = tabla.rows[0].cells
        enc[0].text = '#'; enc[1].text = 'Asignatura'; enc[2].text = 'Calificación'
        for i, (nombre, nota) in enumerate(zip(self.nombres_notas, self.notas_finales), start=1):
            fila = tabla.rows[i].cells
            fila[0].text = str(i); fila[1].text = nombre; fila[2].text = f"{nota:.2f}"
        doc.add_paragraph()
        promedio = self.calcular_media(self.notas_finales)
        p = doc.add_paragraph(f'Promedio General: {promedio:.2f}')
        p.runs[0].font.bold = True; p.runs[0].font.size = Pt(14)
        ruta = os.path.join(os.path.expanduser("~/Documents"),
                            f"Reporte_Calificaciones_{datetime.now().strftime('%d%m%Y_%H%M%S')}.docx")
        doc.save(ruta)
        messagebox.showinfo("Éxito", f"Documento exportado a:\n{ruta}")


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: APUNTADOR DE NOTAS
# ══════════════════════════════════════════════════════════════════════════════
class ModuloApuntador(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.archivo_actual = None
        self.notas_guardadas = {}
        self._cargar_notas()
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Panel lateral
        self.frame_lateral = ctk.CTkFrame(self, width=270, corner_radius=0,
                                          fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        self.frame_lateral.grid(row=0, column=0, sticky="nsew")
        ctk.CTkLabel(self.frame_lateral, text="📝 Mis Notas",
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=20, padx=10)

        self.btn_nueva_nota = ctk.CTkButton(self.frame_lateral, text="➕ Nueva Nota",
                                            fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                            height=38, font=("Segoe UI", 12, "bold"),
                                            command=self.nueva_nota)
        self.btn_nueva_nota.pack(fill="x", padx=15, pady=(0, 10))

        self.lista_notas_frame = ctk.CTkScrollableFrame(self.frame_lateral, fg_color="transparent")
        self.lista_notas_frame.pack(fill="both", expand=True, padx=10, pady=5)
        self.actualizar_listbox()

        # Panel principal de edición
        self.frame_editor = ctk.CTkFrame(self, fg_color="transparent")
        self.frame_editor.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.frame_editor.grid_rowconfigure(1, weight=1)
        self.frame_editor.grid_columnconfigure(0, weight=1)

        frame_top = ctk.CTkFrame(self.frame_editor, fg_color="transparent")
        frame_top.grid(row=0, column=0, sticky="ew", pady=(0, 12))
        self.label_nota_abierta = ctk.CTkLabel(frame_top, text="Selecciona o crea una nota",
                                               font=("Segoe UI", 18, "bold"), text_color=COLOR_TEXT_MAIN)
        self.label_nota_abierta.pack(side="left")

        btn_bar = ctk.CTkFrame(frame_top, fg_color="transparent")
        btn_bar.pack(side="right")

        self.btn_tts_nota = ctk.CTkButton(btn_bar, text="🔊 Leer", width=85, height=36,
                                          fg_color=COLOR_BG_SURFACE, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                          hover_color=COLOR_ACCENT_HOVER,
                                          command=self._toggle_tts)
        self.btn_tts_nota.pack(side="left", padx=4)

        ctk.CTkButton(btn_bar, text="💾 Guardar", width=95, height=36,
                      fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                      command=self.guardar_nota).pack(side="left", padx=4)
        ctk.CTkButton(btn_bar, text="📄 Word", width=95, height=36,
                      fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                      command=self.exportar_nota_word).pack(side="left", padx=4)
        ctk.CTkButton(btn_bar, text="🗑️ Borrar", width=95, height=36,
                      fg_color=COLOR_DANGER, hover_color=COLOR_DANGER_HOVER,
                      command=self.eliminar_nota).pack(side="left", padx=4)

        self.editor_texto = ctk.CTkTextbox(self.frame_editor, font=("Segoe UI", 14),
                                           fg_color=COLOR_BG_CARD_LIGHT, border_width=1,
                                           border_color=COLOR_BORDER, wrap="word")
        self.editor_texto.grid(row=1, column=0, sticky="nsew")

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_nota.configure(text="🔊 Leer", fg_color=COLOR_BG_SURFACE)
        else:
            texto = self.editor_texto.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin texto", "No hay texto escrito en la nota para leer.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_nota.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_nota.configure(text="🔊 Leer", fg_color=COLOR_BG_SURFACE)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def nueva_nota(self):
        d = ctk.CTkInputDialog(text="Nombre de la nueva nota:", title="Nueva Nota")
        nombre = d.get_input()
        if nombre and nombre.strip():
            nombre = nombre.strip()
            self.archivo_actual = nombre
            self.notas_guardadas[nombre] = ""
            self._guardar_notas()
            self.actualizar_listbox()
            self.abrir_nota(nombre)

    def abrir_nota(self, nombre):
        self.archivo_actual = nombre
        self.label_nota_abierta.configure(text=nombre)
        self.editor_texto.delete("1.0", "end")
        self.editor_texto.insert("1.0", self.notas_guardadas.get(nombre, ""))

    def actualizar_listbox(self):
        for w in self.lista_notas_frame.winfo_children():
            w.destroy()
        for nombre in self.notas_guardadas:
            ctk.CTkButton(self.lista_notas_frame, text=f"• {nombre}", anchor="w",
                          fg_color="transparent", text_color=COLOR_TEXT_MAIN,
                          hover_color=COLOR_BG_SURFACE, height=32,
                          command=lambda n=nombre: self.abrir_nota(n)).pack(fill="x", pady=2)

    def guardar_nota(self):
        if self.archivo_actual:
            self.notas_guardadas[self.archivo_actual] = self.editor_texto.get("1.0", "end-1c")
            self._guardar_notas()
            messagebox.showinfo("Guardado", "Nota guardada correctamente.")

    def eliminar_nota(self):
        if self.archivo_actual and messagebox.askyesno("Confirmar", f"¿Eliminar '{self.archivo_actual}'?"):
            del self.notas_guardadas[self.archivo_actual]
            self._guardar_notas()
            self.archivo_actual = None
            self.editor_texto.delete("1.0", "end")
            self.label_nota_abierta.configure(text="Seleccione una nota")
            self.actualizar_listbox()

    def exportar_nota_word(self):
        if not self.archivo_actual:
            return
        doc = Document()
        doc.add_heading(self.archivo_actual, 0)
        doc.add_paragraph(self.editor_texto.get("1.0", "end-1c"))
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"{self.archivo_actual}.docx")
        if path:
            doc.save(path)
            messagebox.showinfo("Éxito", "Exportado correctamente.")

    def _guardar_notas(self):
        ruta = os.path.expanduser("~/.apuntador_notas.json")
        with open(ruta, 'w', encoding='utf-8') as f:
            json.dump(self.notas_guardadas, f, ensure_ascii=False, indent=2)

    def _cargar_notas(self):
        ruta = os.path.expanduser("~/.apuntador_notas.json")
        if os.path.exists(ruta):
            with open(ruta, 'r', encoding='utf-8') as f:
                self.notas_guardadas = json.load(f)


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: RESUMIDOR DE TEXTOS AI
# ══════════════════════════════════════════════════════════════════════════════
class ModuloResumidor(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.modelo = "groq"
        self.instrucciones = (
            "Eres un experto en el tema proporcionado. Tu conocimiento se basa estrictamente en hechos reales. "
            "REGLA DE SEGURIDAD ABSOLUTA: Solo puedes responder a temas que pertenezcan al ámbito educativo, "
            "académico, histórico o laboral. Si el usuario te pide algo fuera de estos ámbitos, DEBES responder "
            "ÚNICAMENTE con la frase: 'ERROR: La petición no pertenece al ámbito educativo o laboral.' "
            "Si la petición es válida, redacta un texto muy extenso, preciso y con párrafos bien estructurados "
            "explicando el contexto, las causas y las consecuencias. No inventes datos bajo ninguna circunstancia."
        )
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        header = ctk.CTkFrame(self, fg_color="transparent")
        header.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        ctk.CTkLabel(header, text="🔍 Resumidor de Textos con IA",
                     font=("Segoe UI", 28, "bold"), text_color=COLOR_ACCENT_SKY).pack(side="left")
        self.entry_nombre = ctk.CTkEntry(header, placeholder_text="Tu nombre...", width=200,
                                         fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_nombre.pack(side="right", padx=10)

        main = ctk.CTkFrame(self, fg_color="transparent")
        main.grid(row=1, column=0, padx=20, sticky="nsew")
        main.grid_columnconfigure((0, 1), weight=1)
        main.grid_rowconfigure(0, weight=1)

        input_f = ctk.CTkFrame(main, fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER, corner_radius=14)
        input_f.grid(row=0, column=0, padx=(0, 10), sticky="nsew")
        ctk.CTkLabel(input_f, text="Pega tus apuntes o tema aquí:", font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=10)
        self.txt_input = ctk.CTkTextbox(input_f, font=("Segoe UI", 13), fg_color=COLOR_BG_CARD_LIGHT, wrap="word")
        self.txt_input.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        output_f = ctk.CTkFrame(main, fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER, corner_radius=14)
        output_f.grid(row=0, column=1, padx=(10, 0), sticky="nsew")
        ctk.CTkLabel(output_f, text="Resumen Riguroso Generado por IA:", font=("Segoe UI", 13, "bold"), text_color=COLOR_ACCENT_CYAN).pack(pady=10)
        self.txt_output = ctk.CTkTextbox(output_f, font=("Segoe UI", 13), fg_color=COLOR_BG_CARD_LIGHT, wrap="word")
        self.txt_output.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        footer = ctk.CTkFrame(self, fg_color="transparent")
        footer.grid(row=2, column=0, padx=20, pady=20, sticky="ew")
        self.progress_bar = ctk.CTkProgressBar(footer, progress_color=COLOR_ACCENT_CYAN)
        self.progress_bar.pack(fill="x", pady=(0, 15))
        self.progress_bar.set(0)

        self.btn_procesar = ctk.CTkButton(footer, text="⚡ Generar Resumen Riguroso",
                                          height=45, font=("Segoe UI", 14, "bold"),
                                          fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                          command=self.iniciar_proceso)
        self.btn_procesar.pack(side="left", fill="x", expand=True, padx=(0, 10))

        self.btn_tts_resumen = ctk.CTkButton(footer, text="🔊 Escuchar", height=45, width=140,
                                            font=("Segoe UI", 12, "bold"),
                                            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                            hover_color=COLOR_ACCENT_HOVER,
                                            command=self._toggle_tts)
        self.btn_tts_resumen.pack(side="left", padx=(0, 10))

        ctk.CTkButton(footer, text="📄 Guardar en Word", height=45, width=180,
                      fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                      font=("Segoe UI", 13, "bold"),
                      command=self.exportar_word).pack(side="right")

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_resumen.configure(text="🔊 Escuchar", fg_color=COLOR_BG_CARD)
        else:
            texto = self.txt_output.get("1.0", "end-1c").strip()
            if not texto or "ERROR:" in texto:
                messagebox.showinfo("Sin resumen", "Primero genera un resumen para escucharlo en voz alta.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_resumen.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_resumen.configure(text="🔊 Escuchar", fg_color=COLOR_BG_CARD)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def iniciar_proceso(self):
        texto = self.txt_input.get("1.0", "end-1c").strip()
        if not texto:
            messagebox.showwarning("Atención", "Por favor, introduce el texto que deseas resumir.")
            return
        self.txt_output.delete("1.0", "end")
        self.btn_procesar.configure(state="disabled")
        self.progress_bar.configure(mode="indeterminate")
        self.progress_bar.start()
        threading.Thread(target=self._ejecutar_ia, args=(texto,), daemon=True).start()

    def _ejecutar_ia(self, texto):
        try:
            resultado = llamar_groq(
                f"{self.instrucciones}\n\nDesarrolla o resume de manera extensa y rigurosa: {texto}"
            )
            self.after(0, self._escribir_output, resultado)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error Cloud", f"Fallo al conectar con IA: {e}"))
        finally:
            self.after(0, self._finalizar)

    def _escribir_output(self, char):
        self.txt_output.insert("end", char)
        self.txt_output.see("end")

    def _finalizar(self):
        self.progress_bar.stop()
        self.progress_bar.set(1)
        self.btn_procesar.configure(state="normal")

    def exportar_word(self):
        contenido = self.txt_output.get("1.0", "end-1c").strip()
        if not contenido or "ERROR:" in contenido:
            messagebox.showwarning("No se puede guardar", "No hay un resumen válido para exportar.")
            return
        nombre = self.entry_nombre.get().strip() or "Usuario"
        ruta = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"Resumen_{nombre}.docx")
        if ruta:
            try:
                doc = Document()
                doc.add_heading(f'Resumen Académico de {nombre}', 0)
                doc.add_paragraph(contenido)
                doc.save(ruta)
                messagebox.showinfo("Éxito", f"Archivo guardado en:\n{ruta}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar: {e}")


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: GENERADOR DE EXÁMENES AI
# ══════════════════════════════════════════════════════════════════════════════
class ModuloExamen(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.examen_en_memoria = ""
        self.historial_conversacion = []
        self.instrucciones_base = (
            "Eres un evaluador académico profesional. El examen debe tener:\n"
            "1. Un título relevante.\n"
            "2. Preguntas de opción múltiple (A-E) o completar huecos.\n"
            "3. La mitad de preguntas de desarrollo.\n"
            "REGLA CRÍTICA: No des las respuestas hasta que el usuario responda. "
            "No pongas las respuestas correctas en el examen."
        )
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        sidebar = ctk.CTkFrame(self, width=280, corner_radius=0,
                               fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        sidebar.grid(row=0, column=0, sticky="nsew")
        ctk.CTkLabel(sidebar, text="🎯 Configurar Examen", font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=20)
        self.entry_nombre = ctk.CTkEntry(sidebar, placeholder_text="Tu nombre...",
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_nombre.pack(fill="x", padx=20, pady=8)

        ctk.CTkLabel(sidebar, text="🟢 Servidor IA Listo", text_color=COLOR_SUCCESS,
                     font=("Segoe UI", 12, "bold")).pack(fill="x", padx=20, pady=5)
        ctk.CTkLabel(sidebar, text="Tema del examen:", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(15, 0), anchor="w", padx=20)
        self.txt_tema = ctk.CTkTextbox(sidebar, height=130, fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.txt_tema.pack(fill="x", padx=20, pady=8)

        self.btn_generar = ctk.CTkButton(sidebar, text="✨ Generar Examen",
                                         fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                         height=40, font=("Segoe UI", 12, "bold"),
                                         command=self.iniciar_generacion)
        self.btn_generar.pack(fill="x", padx=20, pady=8)

        self.btn_tts_examen = ctk.CTkButton(sidebar, text="🔊 Escuchar Examen", fg_color=COLOR_BG_CARD_LIGHT,
                                            border_width=1, border_color=COLOR_ACCENT_CYAN,
                                            hover_color=COLOR_ACCENT_HOVER,
                                            height=36, font=("Segoe UI", 11, "bold"),
                                            command=self._toggle_tts)
        self.btn_tts_examen.pack(fill="x", padx=20, pady=4)

        ctk.CTkButton(sidebar, text="📄 Exportar Word", fg_color=COLOR_BG_SURFACE, border_width=1,
                      border_color=COLOR_BORDER, hover_color=COLOR_ACCENT_HOVER,
                      height=36, command=self.exportar_word).pack(fill="x", padx=20, pady=4)

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_examen.configure(text="🔊 Escuchar Examen", fg_color=COLOR_BG_CARD_LIGHT)
        else:
            texto = self.output_text.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin examen", "Primero genera un examen para escucharlo.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_examen.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_examen.configure(text="🔊 Escuchar Examen", fg_color=COLOR_BG_CARD_LIGHT)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

        self.status_label = ctk.CTkLabel(sidebar, text="🟢 Listo para crear",
                                         text_color=COLOR_SUCCESS, font=("Segoe UI", 11))
        self.status_label.pack(side="bottom", pady=15)

        main_f = ctk.CTkFrame(self, fg_color="transparent")
        main_f.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        main_f.grid_rowconfigure(0, weight=1)
        main_f.grid_columnconfigure(0, weight=1)

        self.output_text = ctk.CTkTextbox(main_f, font=("Consolas", 13),
                                          fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.output_text.grid(row=0, column=0, sticky="nsew", pady=(0, 10))

        input_f = ctk.CTkFrame(main_f, fg_color="transparent")
        input_f.grid(row=1, column=0, sticky="ew")
        input_f.grid_columnconfigure(0, weight=1)
        self.entry_respuesta = ctk.CTkEntry(input_f, placeholder_text="Escribe tus respuestas aquí...", height=42,
                                            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_respuesta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_respuesta.bind("<Return>", lambda e: self.enviar_respuesta())
        ctk.CTkButton(input_f, text="Enviar Respuestas", width=140, height=42,
                      fg_color=COLOR_ACCENT_CYAN, hover_color=COLOR_ACCENT_CYAN_HOVER,
                      font=("Segoe UI", 12, "bold"),
                      command=self.enviar_respuesta).grid(row=0, column=1)

    def iniciar_generacion(self):
        tema = self.txt_tema.get("1.0", "end-1c").strip()
        if not tema:
            messagebox.showwarning("Atención", "Escribe un tema para el examen.")
            return
        self.output_text.delete("1.0", "end")
        self.output_text.insert("end", f"Generando examen sobre: {tema}...\n\n")
        self.btn_generar.configure(state="disabled")
        self.status_label.configure(text="🟡 Elaborando Examen...", text_color="orange")
        threading.Thread(target=self._proceso_groq, args=(tema,), daemon=True).start()

    def _proceso_groq(self, tema):
        try:
            prompt = (
                f"{self.instrucciones_base}\n\nHazme un examen sobre: {tema}"
            )
            full = llamar_groq(prompt)
            self.examen_en_memoria = full
            self.historial_conversacion = [
                {'role': 'system', 'content': self.instrucciones_base},
                {'role': 'assistant', 'content': full}
            ]
            self.after(0, self._update_output, full)
            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status_label.configure(text="🟢 Examen Listo", text_color=COLOR_SUCCESS))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error Cloud", str(e)))
            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status_label.configure(text="🔴 Error de red", text_color=COLOR_DANGER))

    def enviar_respuesta(self):
        msg = self.entry_respuesta.get().strip()
        if not msg:
            return
        self.output_text.insert("end", f"\n\n👤 TÚ: {msg}\n\n🤖 CORRECCIÓN IA: ")
        self.entry_respuesta.delete(0, "end")
        self.status_label.configure(text="🟡 Corrigiendo...", text_color="orange")
        self.historial_conversacion.append({'role': 'user', 'content': msg})
        threading.Thread(target=self._proceso_respuesta, daemon=True).start()

    def _proceso_respuesta(self):
        try:
            full = llamar_groq(construir_prompt(self.instrucciones_base, self.historial_conversacion))
            self.historial_conversacion.append({'role': 'assistant', 'content': full})
            self.after(0, self._update_output, full)
            self.after(0, lambda: self.status_label.configure(text="🟢 Corrección Completada", text_color=COLOR_SUCCESS))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", str(e)))

    def _update_output(self, texto):
        self.output_text.insert("end", texto)
        self.output_text.see("end")

    def exportar_word(self):
        texto = self.output_text.get("1.0", "end-1c")
        if not texto.strip():
            messagebox.showwarning("Vacío", "No hay contenido para exportar.")
            return
        nombre = self.entry_nombre.get().strip() or "Usuario"
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"Examen_{nombre}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Examen de {nombre} – KernossIA", 0)
            doc.add_paragraph(texto)
            doc.save(path)
            messagebox.showinfo("Éxito", f"Archivo guardado en:\n{path}")


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: AYUDANTE DE PROBLEMAS AI (CHAT DE RESOLUCIÓN)
# ══════════════════════════════════════════════════════════════════════════════
class ModuloAyudador(ctk.CTkFrame):
    def __init__(self, master, sesion=None):
        super().__init__(master, fg_color="transparent")
        self.sesion = sesion or {}
        self.modelo_actual  = "groq"
        self.historial_conversacion = []
        self.instrucciones_groq = (
            "Eres un asistente BÁSICO y RÁPIDO (Groq). "
            "Tu objetivo es ayudar y resolver de la mejor manera posible lo que te pida el usuario. "
            "Sé directo pero si hace falta explicar cualquier cosa hazlo."
        )
        self.instrucciones_gemini = (
            "Eres un asistente AVANZADO y PROFUNDO (Gemini). "
            "Tu función es resolver de la forma más inteligente cualquier cosa que te pidan. "
            "Siempre explica el desarrollo y el porqué de las cosas."
        )
        suffix = self.sesion.get("email", "default").replace("@", "_").replace(".", "_")
        self.ruta_historial = os.path.expanduser(f"~/.historial_solver_{suffix}.json")
        self.chat_actual_id = None
        self.todo_el_historial = self._cargar_historial()
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        sidebar = ctk.CTkFrame(self, width=280, corner_radius=0,
                               fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        sidebar.grid(row=0, column=0, sticky="nsew")
        ctk.CTkLabel(sidebar, text="🤖 Solver IA", font=("Segoe UI", 22, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(20, 5))
        self.entry_nombre = ctk.CTkEntry(sidebar, placeholder_text="¿Cuál es tu nombre?",
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_nombre.pack(fill="x", padx=15, pady=8)
        if self.sesion.get("nombre"):
            self.entry_nombre.insert(0, self.sesion["nombre"])

        ctk.CTkLabel(sidebar, text="⚙️ Instrucciones del Sistema",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(10, 2))
        self.txt_instrucciones = ctk.CTkTextbox(sidebar, height=110, wrap="word", font=("Segoe UI", 11),
                                                fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.txt_instrucciones.pack(fill="x", padx=15, pady=5)
        self.txt_instrucciones.insert("1.0", self.instrucciones_groq)

        ctk.CTkLabel(sidebar, text="Motor de IA", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(10, 2))
        frame_modelo = ctk.CTkFrame(sidebar, fg_color=COLOR_BG_CARD_LIGHT, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
        frame_modelo.pack(fill="x", padx=15, pady=5)
        self.btn_basico = ctk.CTkButton(frame_modelo, text="⚡ Básico (Groq)", height=36,
                                         fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                         command=lambda: self._cambiar_modelo("groq"))
        self.btn_basico.pack(fill="x", padx=8, pady=(8, 4))
        self.btn_avanzado = ctk.CTkButton(frame_modelo, text="🧠 Avanzado (Gemini)", height=36,
                                           fg_color="transparent", border_width=1, border_color=COLOR_ACCENT_PURPLE,
                                           hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                           command=lambda: self._cambiar_modelo("gemini"))
        self.btn_avanzado.pack(fill="x", padx=8, pady=(4, 8))
        self.lbl_limite = ctk.CTkLabel(sidebar, text="ℹ️ Groq: ~1.000 msgs/día",
                                        font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED)
        self.lbl_limite.pack(pady=(0, 5))
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=10, pady=8)

        ctk.CTkButton(sidebar, text="📄 Exportar a Word", fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_PRIMARY,
                      border_width=1, border_color=COLOR_BORDER,
                      command=self._exportar_word).pack(fill="x", padx=15, pady=4)
        ctk.CTkButton(sidebar, text="🧹 Nuevo Chat", fg_color="transparent", border_width=1,
                      border_color=COLOR_BORDER, hover_color=COLOR_BG_SURFACE,
                      command=self._limpiar_chat).pack(fill="x", padx=15, pady=4)

        ctk.CTkLabel(sidebar, text="🕒 Historial de Consultas", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(12, 3))
        self.frame_historial = ctk.CTkScrollableFrame(sidebar, fg_color="transparent")
        self.frame_historial.pack(fill="both", expand=True, padx=10, pady=5)
        self.status_label = ctk.CTkLabel(sidebar, text="🟢 Groq Listo",
                                          text_color=COLOR_SUCCESS, font=("Segoe UI", 11))
        self.status_label.pack(side="bottom", pady=(0, 15))
        self._actualizar_historial_ui()

        chat_frame = ctk.CTkFrame(self, fg_color="transparent")
        chat_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        chat_frame.grid_rowconfigure(1, weight=1)
        chat_frame.grid_columnconfigure(0, weight=1)

        self.lbl_banner = ctk.CTkLabel(chat_frame,
                                        text="Modelo activo: Groq  •  Límite aprox: 1.000 mensajes / día",
                                        font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED)
        self.lbl_banner.grid(row=0, column=0, sticky="w", pady=(0, 6))

        self.txt_chat = ctk.CTkTextbox(chat_frame, font=("Segoe UI", 14), state="disabled", wrap="word",
                                       fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_chat.grid(row=1, column=0, sticky="nsew", pady=(0, 12))

        input_frame = ctk.CTkFrame(chat_frame, fg_color="transparent")
        input_frame.grid(row=2, column=0, sticky="ew")
        input_frame.grid_columnconfigure(0, weight=1)

        self.entry_pregunta = ctk.CTkEntry(input_frame, placeholder_text="Describe tu duda o problema aquí...", height=44,
                                           fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_pregunta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_pregunta.bind("<Return>", lambda e: self._enviar())

        ctk.CTkButton(input_frame, text="⚡ Analizar con IA", width=140, height=44,
                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                      font=("Segoe UI", 12, "bold"),
                      command=self._enviar).grid(row=0, column=1, padx=(0, 6))

        self.btn_tts_ayudador = ctk.CTkButton(input_frame, text="🔊 Escuchar", width=110, height=44,
                                              font=("Segoe UI", 12, "bold"),
                                              fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                              hover_color=COLOR_ACCENT_HOVER,
                                              command=self._toggle_tts)
        self.btn_tts_ayudador.grid(row=0, column=2)

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_ayudador.configure(text="🔊 Escuchar", fg_color=COLOR_BG_CARD)
        else:
            texto = getattr(self, "ultima_respuesta_ayudador", "")
            if not texto:
                texto = self.txt_chat.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin respuesta", "No hay respuesta para escuchar en voz alta.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_ayudador.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_ayudador.configure(text="🔊 Escuchar", fg_color=COLOR_BG_CARD)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def _cambiar_modelo(self, modelo):
        texto_actual = self.txt_instrucciones.get("1.0", "end-1c").strip()
        if self.modelo_actual == "groq":
            self.instrucciones_groq = texto_actual
        else:
            self.instrucciones_gemini = texto_actual
        self.modelo_actual = modelo
        self.txt_instrucciones.delete("1.0", "end")
        if modelo == "groq":
            self.btn_basico.configure(fg_color=COLOR_ACCENT_PRIMARY)
            self.btn_avanzado.configure(fg_color="transparent")
            self.lbl_limite.configure(text="ℹ️ Groq: ~1.000 msgs/día")
            self.lbl_banner.configure(text="Modelo activo: Groq (Básico)  •  Límite aprox: 1.000 mensajes / día")
            self.status_label.configure(text="🟢 Groq Listo", text_color=COLOR_SUCCESS)
            self.txt_instrucciones.insert("1.0", self.instrucciones_groq)
        else:
            self.btn_avanzado.configure(fg_color=COLOR_ACCENT_PURPLE)
            self.btn_basico.configure(fg_color="transparent")
            self.lbl_limite.configure(text="ℹ️ Gemini: 5 msgs/minuto")
            self.lbl_banner.configure(text="Modelo activo: Gemini  •  Límite: 5 mensajes / minuto")
            self.status_label.configure(text="🟣 Gemini Listo", text_color="#a5b4fc")
            self.txt_instrucciones.insert("1.0", self.instrucciones_gemini)

    def _agregar_texto(self, emisor, texto):
        self.txt_chat.configure(state="normal")
        if emisor == "IA":
            self.txt_chat.insert("end", f"\n 🤖 IA ({self.modelo_actual.upper()}):\n{texto}\n")
        else:
            self.txt_chat.insert("end", f"\n\n 👤 {emisor}:\n{texto}\n")
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")

    def _enviar(self):
        pregunta = self.entry_pregunta.get().strip()
        nombre   = self.entry_nombre.get().strip() or "Tú"
        if not pregunta:
            return
        self._agregar_texto(nombre, pregunta)
        self.entry_pregunta.delete(0, "end")
        self.status_label.configure(text="🟡 Pensando...", text_color="orange")
        self.historial_conversacion.append({"role": "user", "content": pregunta})
        threading.Thread(target=self._proceso_ia, daemon=True).start()

    def _proceso_ia(self):
        try:
            self.after(0, lambda: self._stream_update(f"\n 🤖 IA ({self.modelo_actual.upper()}):\n"))
            instrucciones = self.txt_instrucciones.get("1.0", "end-1c").strip()
            prompt = construir_prompt(instrucciones, self.historial_conversacion)
            if self.modelo_actual == "groq":
                respuesta_completa = llamar_groq(prompt)
            else:
                respuesta_completa = llamar_gemini(prompt)
            self.ultima_respuesta_ayudador = respuesta_completa
            self.historial_conversacion.append({"role": "assistant", "content": respuesta_completa})
            self.after(0, self._stream_update, respuesta_completa)
            self._guardar_historial()
            color = COLOR_SUCCESS if self.modelo_actual == "groq" else "#a5b4fc"
            texto = "🟢 Groq Listo" if self.modelo_actual == "groq" else "🟣 Gemini Listo"
            self.after(0, lambda: self.status_label.configure(text=texto, text_color=color))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            self.after(0, lambda: self.status_label.configure(text="🔴 Error", text_color=COLOR_DANGER))
        finally:
            self.after(0, lambda: self.txt_chat.configure(state="disabled"))

    def _stream_update(self, contenido):
        self.txt_chat.configure(state="normal")
        self.txt_chat.insert("end", contenido)
        self.txt_chat.see("end")
        self.txt_chat.configure(state="disabled")

    def _limpiar_chat(self):
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        self.txt_chat.configure(state="disabled")
        self.historial_conversacion = []
        self.chat_actual_id = None

    def _exportar_word(self):
        if not self.historial_conversacion:
            messagebox.showwarning("Vacío", "No hay nada que exportar.")
            return
        nombre = self.entry_nombre.get().strip() or "Usuario"
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             initialfile=f"Solucion_{nombre}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Informe de Solución – {nombre} (KernossIA)", 0)
            for msg in self.historial_conversacion:
                rol = "Tú" if msg["role"] == "user" else "IA"
                p = doc.add_paragraph()
                p.add_run(f"{rol}: ").bold = True
                p.add_run(msg["content"])
            doc.save(path)
            messagebox.showinfo("Éxito", f"Guardado en:\n{path}")

    def _cargar_historial(self):
        if os.path.exists(self.ruta_historial):
            try:
                with open(self.ruta_historial, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def _guardar_historial(self):
        if not self.historial_conversacion:
            return
        if self.chat_actual_id is None:
            primer = next((m["content"] for m in self.historial_conversacion if m["role"] == "user"), "Chat")
            resumen = primer[:25] + "..." if len(primer) > 25 else primer
            self.chat_actual_id = f"[{datetime.now().strftime('%H:%M')}] {resumen}"
        self.todo_el_historial[self.chat_actual_id] = self.historial_conversacion
        try:
            with open(self.ruta_historial, "w", encoding="utf-8") as f:
                json.dump(self.todo_el_historial, f, ensure_ascii=False, indent=2)
            self.after(0, self._actualizar_historial_ui)
        except Exception as e:
            print(f"Error guardando historial: {e}")

    def _actualizar_historial_ui(self):
        for w in self.frame_historial.winfo_children():
            w.destroy()
        for chat_id in reversed(list(self.todo_el_historial.keys())):
            fila = ctk.CTkFrame(self.frame_historial, fg_color="transparent")
            fila.pack(fill="x", pady=3)
            ctk.CTkButton(fila, text=chat_id, fg_color="transparent",
                          text_color=COLOR_TEXT_MAIN, anchor="w", hover_color=COLOR_BG_SURFACE, height=32,
                          command=lambda cid=chat_id: self._cargar_chat(cid)).pack(side="left", fill="x", expand=True, padx=(0, 4))
            ctk.CTkButton(fila, text="❌", width=32, height=32,
                          fg_color=COLOR_DANGER, hover_color=COLOR_DANGER_HOVER,
                          command=lambda cid=chat_id: self._borrar_chat(cid)).pack(side="right")

    def _cargar_chat(self, chat_id):
        self.chat_actual_id = chat_id
        self.historial_conversacion = self.todo_el_historial[chat_id]
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        nombre = self.entry_nombre.get().strip() or "Tú"
        for msg in self.historial_conversacion:
            if msg["role"] == "user":
                self.txt_chat.insert("end", f"\n\n 👤 {nombre}:\n{msg['content']}\n")
            else:
                self.txt_chat.insert("end", f"\n 🤖 IA:\n{msg['content']}\n")
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")

    def _borrar_chat(self, chat_id):
        if messagebox.askyesno("Confirmar", f"¿Borrar este chat?\n'{chat_id}'"):
            del self.todo_el_historial[chat_id]
            try:
                with open(self.ruta_historial, "w", encoding="utf-8") as f:
                    json.dump(self.todo_el_historial, f, ensure_ascii=False, indent=2)
            except Exception:
                pass
            if self.chat_actual_id == chat_id:
                self._limpiar_chat()
            self._actualizar_historial_ui()


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: AGENDA Y CALENDARIO
# ══════════════════════════════════════════════════════════════════════════════
class ModuloCalendario(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.ruta_datos = os.path.expanduser("~/.agenda_estudios.json")
        self.eventos = self._cargar_eventos()
        self.hoy = datetime.now()
        self.año_actual = self.hoy.year
        self.mes_actual = self.hoy.month
        self.dia_seleccionado = f"{self.año_actual}-{self.mes_actual:02d}-{self.hoy.day:02d}"
        self.nombres_meses = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
                               "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]
        self.botones_dias = []
        self._build_ui()
        self._actualizar_calendario()
        self._cargar_evento_en_editor()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=2)
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Panel izquierdo: calendario
        frame_izq = ctk.CTkFrame(self, fg_color="transparent")
        frame_izq.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")

        self.frame_nav = ctk.CTkFrame(frame_izq, fg_color=COLOR_BG_CARD, height=60, corner_radius=12,
                                      border_width=1, border_color=COLOR_BORDER)
        self.frame_nav.pack(fill="x", pady=(0, 15))
        self.frame_nav.pack_propagate(False)
        ctk.CTkButton(self.frame_nav, text="◀", width=40, font=("Segoe UI", 16),
                      fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_HOVER,
                      command=self._mes_anterior).pack(side="left", padx=15, pady=10)
        self.lbl_mes_año = ctk.CTkLabel(self.frame_nav, text="", font=("Segoe UI", 20, "bold"),
                                         text_color=COLOR_ACCENT_SKY)
        self.lbl_mes_año.pack(side="left", expand=True)
        ctk.CTkButton(self.frame_nav, text="▶", width=40, font=("Segoe UI", 16),
                      fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_HOVER,
                      command=self._mes_siguiente).pack(side="right", padx=15, pady=10)

        self.frame_dias = ctk.CTkFrame(frame_izq, fg_color=COLOR_BG_CARD, corner_radius=15,
                                        border_width=1, border_color=COLOR_BORDER)
        self.frame_dias.pack(fill="both", expand=True)
        for i in range(7):
            self.frame_dias.grid_columnconfigure(i, weight=1, uniform="dias")
        for i in range(7):
            self.frame_dias.grid_rowconfigure(i, weight=1, uniform="semanas")
        for col, dia in enumerate(["Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]):
            ctk.CTkLabel(self.frame_dias, text=dia, font=("Segoe UI", 13, "bold"),
                         text_color=COLOR_ACCENT_CYAN).grid(row=0, column=col, pady=8, sticky="nsew")

        # Panel derecho: editor
        frame_der = ctk.CTkFrame(self, fg_color=COLOR_BG_CARD, corner_radius=15,
                                 border_width=1, border_color=COLOR_BORDER)
        frame_der.grid(row=0, column=1, sticky="nsew", padx=(0, 20), pady=20)
        ctk.CTkLabel(frame_der, text="📅 Tareas del día", font=("Segoe UI", 18, "bold"),
                     text_color=COLOR_ACCENT_SKY).pack(pady=(25, 5), padx=20, anchor="w")
        self.lbl_fecha_actual = ctk.CTkLabel(frame_der, text="", font=("Segoe UI", 13), text_color=COLOR_TEXT_MUTED)
        self.lbl_fecha_actual.pack(pady=(0, 15), padx=20, anchor="w")
        self.txt_tareas = ctk.CTkTextbox(frame_der, font=("Segoe UI", 13), border_width=1,
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.txt_tareas.pack(fill="both", expand=True, padx=20, pady=10)
        ctk.CTkButton(frame_der, text="💾 Guardar Tareas", font=("Segoe UI", 13, "bold"),
                      fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                      command=self._guardar_evento).pack(fill="x", padx=20, pady=(10, 5))
        ctk.CTkButton(frame_der, text="🗑️ Borrar Tareas", font=("Segoe UI", 13),
                      fg_color=COLOR_DANGER, hover_color=COLOR_DANGER_HOVER,
                      command=self._borrar_evento).pack(fill="x", padx=20, pady=(5, 25))

    def _actualizar_calendario(self):
        for btn in self.botones_dias:
            btn.destroy()
        self.botones_dias.clear()
        self.lbl_mes_año.configure(text=f"{self.nombres_meses[self.mes_actual]} {self.año_actual}")
        primer_dia_semana, dias_en_mes = calendar.monthrange(self.año_actual, self.mes_actual)
        fila = 1; columna = primer_dia_semana
        for dia in range(1, dias_en_mes + 1):
            fecha_clave = f"{self.año_actual}-{self.mes_actual:02d}-{dia:02d}"
            esta_sel = (fecha_clave == self.dia_seleccionado)
            es_hoy = (self.hoy.year == self.año_actual and self.hoy.month == self.mes_actual and self.hoy.day == dia)
            tiene_tareas = fecha_clave in self.eventos and self.eventos[fecha_clave].strip()
            if esta_sel:
                fg = COLOR_ACCENT_PRIMARY; tc = "white"
            elif es_hoy:
                fg = COLOR_ACCENT_CYAN; tc = "white"
            elif tiene_tareas:
                fg = "#064e3b"; tc = "#6ee7b7"
            else:
                fg = COLOR_BG_CARD_LIGHT; tc = COLOR_TEXT_MAIN
            btn = ctk.CTkButton(self.frame_dias, text=str(dia), font=("Segoe UI", 12, "bold" if (es_hoy or tiene_tareas) else "normal"),
                                fg_color=fg, hover_color=COLOR_ACCENT_HOVER, text_color=tc,
                                corner_radius=8, command=lambda d=dia: self._seleccionar_dia(d))
            btn.grid(row=fila, column=columna, padx=4, pady=4, sticky="nsew")
            self.botones_dias.append(btn)
            columna += 1
            if columna > 6:
                columna = 0; fila += 1

    def _seleccionar_dia(self, dia):
        self.dia_seleccionado = f"{self.año_actual}-{self.mes_actual:02d}-{dia:02d}"
        self._actualizar_calendario()
        self._cargar_evento_en_editor()

    def _mes_anterior(self):
        self.mes_actual -= 1
        if self.mes_actual < 1:
            self.mes_actual = 12; self.año_actual -= 1
        self._actualizar_calendario()

    def _mes_siguiente(self):
        self.mes_actual += 1
        if self.mes_actual > 12:
            self.mes_actual = 1; self.año_actual += 1
        self._actualizar_calendario()

    def _cargar_eventos(self):
        if os.path.exists(self.ruta_datos):
            try:
                with open(self.ruta_datos, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def _cargar_evento_en_editor(self):
        partes = self.dia_seleccionado.split("-")
        fecha_objeto = datetime(int(partes[0]), int(partes[1]), int(partes[2]))
        self.lbl_fecha_actual.configure(text=fecha_objeto.strftime("%d de %B de %Y").title())
        self.txt_tareas.delete("1.0", "end")
        if self.dia_seleccionado in self.eventos:
            self.txt_tareas.insert("1.0", self.eventos[self.dia_seleccionado])

    def _guardar_evento(self):
        contenido = self.txt_tareas.get("1.0", "end-1c").strip()
        if contenido:
            self.eventos[self.dia_seleccionado] = contenido
        elif self.dia_seleccionado in self.eventos:
            del self.eventos[self.dia_seleccionado]
        try:
            with open(self.ruta_datos, 'w', encoding='utf-8') as f:
                json.dump(self.eventos, f, ensure_ascii=False, indent=2)
            self._actualizar_calendario()
            messagebox.showinfo("Guardado", "Agenda actualizada correctamente.")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar:\n{e}")

    def _borrar_evento(self):
        if self.dia_seleccionado in self.eventos:
            if messagebox.askyesno("Confirmar", "¿Seguro que quieres borrar todas las tareas de este día?"):
                del self.eventos[self.dia_seleccionado]
                self.txt_tareas.delete("1.0", "end")
                with open(self.ruta_datos, 'w', encoding='utf-8') as f:
                    json.dump(self.eventos, f, ensure_ascii=False, indent=2)
                self._actualizar_calendario()


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: CREADOR DE EJERCICIOS (PROFESOR)
# ══════════════════════════════════════════════════════════════════════════════
INSTRUCCIONES_EJERCICIO = """Eres un profesor experto creando material educativo.
Crea ejercicios bien estructurados con:
1. Título claro y objetivo de aprendizaje.
2. Enunciado detallado.
3. Apartados o preguntas numerados.
4. Nivel de dificultad indicado.
5. Tiempo estimado.
NO incluyas las soluciones a menos que se te pida explícitamente.
El formato debe ser limpio y listo para imprimir o enviar a alumnos."""

INSTRUCCIONES_SOLUCIONES = """Eres un profesor. Genera el solucionario completo y detallado
del ejercicio que te proporcionan. Explica cada paso y por qué es correcto."""


class ModuloCreadorEjercicios(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.modelo_actual = "groq"
        self.ejercicio_actual = ""
        self.solucionario_actual = ""
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        sidebar = ctk.CTkFrame(self, width=290, corner_radius=0,
                               fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        sidebar.grid(row=0, column=0, sticky="nsew")
        sidebar.grid_propagate(False)
        ctk.CTkLabel(sidebar, text="✏️ Creador de\nEjercicios",
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_PURPLE).pack(pady=(20, 5))
        ctk.CTkLabel(sidebar, text="Herramienta Docente Exclusiva",
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(pady=(0, 15))

        ctk.CTkLabel(sidebar, text="Motor de IA:", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        fm = ctk.CTkFrame(sidebar, fg_color="transparent")
        fm.pack(fill="x", padx=15, pady=5)
        self.btn_groq = ctk.CTkButton(fm, text="⚡ Básico\n(Groq)", height=45,
                                       fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                       font=("Segoe UI", 10, "bold"),
                                       command=lambda: self._set_modelo("groq"))
        self.btn_groq.pack(side="left", expand=True, fill="x", padx=(0, 3))
        self.btn_gemini = ctk.CTkButton(fm, text="🧠 Avanzado\n(Gemini)", height=45,
                                         fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                         font=("Segoe UI", 10, "bold"),
                                         command=lambda: self._set_modelo("gemini"))
        self.btn_gemini.pack(side="left", expand=True, fill="x", padx=(3, 0))
        self.lbl_limite = ctk.CTkLabel(sidebar, text="Groq: ~1.000 msgs/día",
                                        font=("Segoe UI", 9), text_color=COLOR_TEXT_MUTED)
        self.lbl_limite.pack(pady=(2, 10))
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=5)

        ctk.CTkLabel(sidebar, text="Tema / Materia", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15, pady=(5, 2))
        self.entry_tema = ctk.CTkEntry(sidebar, placeholder_text="Ej: Ecuaciones de 2º grado",
                                       fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_tema.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Nivel educativo", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.combo_nivel = ctk.CTkComboBox(sidebar, values=[
            "Primaria", "1º ESO", "2º ESO", "3º ESO", "4º ESO",
            "1º Bachillerato", "2º Bachillerato", "Universidad"],
            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.combo_nivel.set("2º ESO")
        self.combo_nivel.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Tipo de ejercicio", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.combo_tipo = ctk.CTkComboBox(sidebar, values=[
            "Problemas paso a paso", "Preguntas teóricas", "Opción múltiple (A-D)",
            "Completar huecos", "Verdadero / Falso", "Ejercicio práctico", "Mixto (teoría + práctica)"],
            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.combo_tipo.set("Mixto (teoría + práctica)")
        self.combo_tipo.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Nº de preguntas", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.spin_preguntas = ctk.CTkSlider(sidebar, from_=3, to=20, number_of_steps=17, progress_color=COLOR_ACCENT_PURPLE)
        self.spin_preguntas.set(8)
        self.spin_preguntas.pack(fill="x", padx=15, pady=(0, 2))
        self.lbl_n_preguntas = ctk.CTkLabel(sidebar, text="8 preguntas", font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED)
        self.lbl_n_preguntas.pack()
        self.spin_preguntas.configure(command=lambda v: self.lbl_n_preguntas.configure(text=f"{int(v)} preguntas"))

        self.check_soluciones = ctk.CTkCheckBox(sidebar, text="Incluir solucionario al final",
                                                fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER)
        self.check_soluciones.pack(padx=15, pady=8, anchor="w")
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=5)

        self.btn_generar = ctk.CTkButton(sidebar, text="✨ Generar Ejercicio", height=42,
                                          font=("Segoe UI", 13, "bold"),
                                          fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                          command=self.generar_ejercicio)
        self.btn_generar.pack(fill="x", padx=15, pady=5)

        self.btn_regenerar = ctk.CTkButton(sidebar, text="🔄 Regenerar", height=36,
                                            fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_HOVER,
                                            command=self.generar_ejercicio, state="disabled")
        self.btn_regenerar.pack(fill="x", padx=15, pady=2)

        self.btn_editar = ctk.CTkButton(sidebar, text="✏️ Editar manualmente", height=36,
                                         fg_color="transparent", border_width=1, border_color=COLOR_BORDER,
                                         command=self.activar_edicion, state="disabled")
        self.btn_editar.pack(fill="x", padx=15, pady=2)

        self.btn_solucionario = ctk.CTkButton(sidebar, text="💡 Generar Solucionario", height=36,
                                               fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                                               command=self.generar_solucionario, state="disabled")
        self.btn_solucionario.pack(fill="x", padx=15, pady=2)

        self.btn_tts_ejercicio = ctk.CTkButton(sidebar, text="🔊 Escuchar Ejercicio", height=36,
                                              fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                              hover_color=COLOR_ACCENT_HOVER,
                                              command=self._toggle_tts, state="disabled")
        self.btn_tts_ejercicio.pack(fill="x", padx=15, pady=2)

        self.btn_exportar = ctk.CTkButton(sidebar, text="📄 Exportar con solución", height=36,
                                           fg_color=COLOR_BG_CARD_LIGHT, hover_color=COLOR_BG_SURFACE,
                                           command=lambda: self.exportar_word(True), state="disabled")
        self.btn_exportar.pack(fill="x", padx=15, pady=2)

        self.status = ctk.CTkLabel(sidebar, text="Listo", text_color=COLOR_TEXT_MUTED, font=("Segoe UI", 10))
        self.status.pack(side="bottom", pady=8)

        panel = ctk.CTkFrame(self, fg_color="transparent")
        panel.grid(row=0, column=1, sticky="nsew", padx=15, pady=15)
        panel.grid_rowconfigure(1, weight=1)
        panel.grid_columnconfigure(0, weight=1)

        cab = ctk.CTkFrame(panel, fg_color="transparent")
        cab.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        ctk.CTkLabel(cab, text="Ejercicio Pedagógico Generado", font=("Segoe UI", 18, "bold"), text_color=COLOR_TEXT_MAIN).pack(side="left")
        self.lbl_modo_edicion = ctk.CTkLabel(cab, text="", font=("Segoe UI", 11), text_color=COLOR_WARNING)
        self.lbl_modo_edicion.pack(side="right")

        self.txt_ejercicio = ctk.CTkTextbox(panel, font=("Consolas", 13), fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_ejercicio.grid(row=1, column=0, sticky="nsew")
        self.txt_ejercicio.insert("end", "Configura los parámetros en el panel izquierdo y pulsa 'Generar Ejercicio'.")
        self.txt_ejercicio.configure(state="disabled")

        self.entry_cambio = ctk.CTkEntry(panel, placeholder_text="Ajustes con IA: Ej: 'Hazlo más difícil' / 'Añade 2 preguntas teóricas'", height=42,
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_cambio.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        self.entry_cambio.bind("<Return>", lambda e: self.aplicar_cambio_ia())

        btn_row = ctk.CTkFrame(panel, fg_color="transparent")
        btn_row.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        ctk.CTkButton(btn_row, text="⚡ Aplicar cambio con IA", fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                      height=38, font=("Segoe UI", 12, "bold"),
                      command=self.aplicar_cambio_ia).pack(side="left", padx=(0, 8))

    def _set_modelo(self, m):
        self.modelo_actual = m
        if m == "groq":
            self.btn_groq.configure(fg_color=COLOR_ACCENT_PRIMARY); self.btn_gemini.configure(fg_color=COLOR_BG_SURFACE)
            self.lbl_limite.configure(text="Groq: ~1.000 msgs/día")
        else:
            self.btn_groq.configure(fg_color=COLOR_BG_SURFACE); self.btn_gemini.configure(fg_color=COLOR_ACCENT_PURPLE)
            self.lbl_limite.configure(text="Gemini: 15 msgs/minuto")

    def _llamar_ia(self, prompt):
        full_prompt = f"{INSTRUCCIONES_EJERCICIO}\n\n{prompt}"
        if self.modelo_actual == "groq":
            return llamar_groq(full_prompt)
        return llamar_gemini(full_prompt)

    def _construir_prompt(self):
        tema  = self.entry_tema.get().strip() or "tema libre"
        nivel = self.combo_nivel.get()
        tipo  = self.combo_tipo.get()
        n     = int(self.spin_preguntas.get())
        inc   = self.check_soluciones.get()
        return (f"Crea un ejercicio de {tipo} sobre '{tema}' para alumnos de {nivel}. "
                f"Debe tener exactamente {n} preguntas/apartados. "
                f"{'Incluye el solucionario completo al final.' if inc else 'NO incluyas las soluciones.'} "
                f"Formato limpio y listo para entregar a los alumnos.")

    def generar_ejercicio(self):
        if not self.entry_tema.get().strip():
            messagebox.showwarning("Tema vacío", "Introduce el tema del ejercicio.")
            return
        self.btn_generar.configure(state="disabled")
        self.status.configure(text="Generando...", text_color="orange")
        threading.Thread(target=self._thread_generar, daemon=True).start()

    def _thread_generar(self):
        try:
            resultado = self._llamar_ia(self._construir_prompt())
            self.ejercicio_actual = resultado
            self.after(0, self._mostrar_ejercicio, resultado)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    def _mostrar_ejercicio(self, texto):
        self.txt_ejercicio.configure(state="normal")
        self.txt_ejercicio.delete("1.0", "end")
        self.txt_ejercicio.insert("end", texto)
        self.txt_ejercicio.configure(state="disabled")
        for btn in [self.btn_regenerar, self.btn_editar, self.btn_solucionario,
                    self.btn_tts_ejercicio, self.btn_exportar]:
            btn.configure(state="normal")

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_ejercicio.configure(text="🔊 Escuchar Ejercicio", fg_color=COLOR_BG_CARD_LIGHT)
        else:
            texto = self.txt_ejercicio.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin ejercicio", "Primero genera un ejercicio para escucharlo.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_ejercicio.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_ejercicio.configure(text="🔊 Escuchar Ejercicio", fg_color=COLOR_BG_CARD_LIGHT)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def activar_edicion(self):
        self.txt_ejercicio.configure(state="normal")
        self.lbl_modo_edicion.configure(text="✏️ Modo edición manual activo")
        self.btn_editar.configure(text="Bloquear edición", command=self.desactivar_edicion)

    def desactivar_edicion(self):
        self.ejercicio_actual = self.txt_ejercicio.get("1.0", "end-1c")
        self.txt_ejercicio.configure(state="disabled")
        self.lbl_modo_edicion.configure(text="")
        self.btn_editar.configure(text="Editar manualmente", command=self.activar_edicion)

    def aplicar_cambio_ia(self):
        cambio = self.entry_cambio.get().strip()
        if not cambio:
            messagebox.showwarning("Atención", "Escribe qué quieres cambiar.")
            return
        if not self.ejercicio_actual:
            messagebox.showwarning("Atención", "Primero genera un ejercicio.")
            return
        self.status.configure(text="Aplicando cambio...", text_color="orange")
        prompt = (f"Este es el ejercicio actual:\n\n{self.ejercicio_actual}\n\n"
                  f"Aplica el siguiente cambio y devuelve el ejercicio completo actualizado: {cambio}")
        self.entry_cambio.delete(0, "end")
        threading.Thread(target=lambda: self._thread_cambio(prompt), daemon=True).start()

    def _thread_cambio(self, prompt):
        try:
            resultado = self._llamar_ia(prompt)
            self.ejercicio_actual = resultado
            self.after(0, self._mostrar_ejercicio, resultado)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    def generar_solucionario(self):
        if not self.ejercicio_actual:
            messagebox.showwarning("Atención", "Primero genera un ejercicio.")
            return
        self.status.configure(text="Generando solucionario...", text_color="orange")
        prompt = f"{INSTRUCCIONES_SOLUCIONES}\n\nGenera el solucionario completo de este ejercicio:\n\n{self.ejercicio_actual}"
        def _thread():
            try:
                sol = llamar_groq(prompt) if self.modelo_actual == "groq" else llamar_gemini(prompt)
                self.solucionario_actual = sol
                def _mostrar():
                    ven = ctk.CTkToplevel(self)
                    ven.title("Solucionario – KernossIA"); ven.geometry("800x600")
                    ven.configure(fg_color=COLOR_BG_DARK)
                    ctk.CTkLabel(ven, text="💡 Solucionario Completo",
                                 font=("Segoe UI", 18, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=15)
                    txt = ctk.CTkTextbox(ven, font=("Consolas", 12), fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
                    txt.pack(fill="both", expand=True, padx=15, pady=(0, 10))
                    txt.insert("end", sol); txt.configure(state="disabled")
                    ctk.CTkButton(ven, text="📄 Exportar Solucionario a Word",
                                  fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                                  command=lambda: self._exportar_solucionario(sol)).pack(fill="x", padx=15, pady=(0, 15))
                self.after(0, _mostrar)
            except Exception as e:
                self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            finally:
                self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))
        threading.Thread(target=_thread, daemon=True).start()

    def exportar_word(self, con_solucion=True):
        texto = self.txt_ejercicio.get("1.0", "end-1c")
        if not texto.strip():
            messagebox.showwarning("Vacío", "No hay ejercicio que exportar.")
            return
        tema   = self.entry_tema.get().strip() or "Ejercicio"
        sufijo = "ConSolucion" if con_solucion else "SinSolucion"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Ejercicio_{tema}_{sufijo}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return
        doc = Document()
        doc.add_heading(f"Ejercicio — {tema} (KernossIA)", 0)
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()} | Tipo: {self.combo_tipo.get()}")
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(""); doc.add_paragraph(texto)
        if con_solucion and self.solucionario_actual:
            doc.add_page_break()
            doc.add_heading("SOLUCIONARIO", level=1)
            doc.add_paragraph(self.solucionario_actual)
        doc.save(path)
        messagebox.showinfo("Exportado", f"Guardado en:\n{path}")

    def _exportar_solucionario(self, sol):
        tema = self.entry_tema.get().strip() or "Ejercicio"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Solucionario_{tema}_{datetime.now().strftime('%Y%m%d')}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Solucionario — {tema} (KernossIA)", 0)
            doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
            doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph(""); doc.add_paragraph(sol)
            doc.save(path)
            messagebox.showinfo("Exportado", f"Solucionario guardado en:\n{path}")


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: CORRECTOR DE EXÁMENES (PROFESOR)
# ══════════════════════════════════════════════════════════════════════════════
INSTRUCCIONES_CORRECTOR = """Eres un profesor corrector experto y riguroso.
Tu tarea es evaluar las respuestas de un alumno comparándolas con el ejercicio/criterios dados.
Para cada pregunta debes:
1. Indicar si está CORRECTA, PARCIALMENTE CORRECTA o INCORRECTA.
2. Puntuación obtenida sobre la puntuación máxima de esa pregunta.
3. Comentario breve explicando el acierto o el error.
4. Sugerencia de mejora si aplica.

Al final incluye:
- NOTA TOTAL (sobre 10).
- Resumen general de puntos fuertes y débiles del alumno.
- Recomendaciones de estudio personalizadas.

Sé justo, constructivo y motivador en el tono."""


class ModuloCorrectorExamenes(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.modelo_actual = "groq"
        self.correccion_actual = ""
        self.alumnos = {}
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=1)
        self.grid_rowconfigure(0, weight=1)

        sidebar = ctk.CTkFrame(self, width=270, corner_radius=0,
                               fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        sidebar.grid(row=0, column=0, sticky="nsew")
        sidebar.grid_propagate(False)
        ctk.CTkLabel(sidebar, text="📋 Corrector de\nExámenes",
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(20, 5))
        ctk.CTkLabel(sidebar, text="Modo Docente",
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(pady=(0, 15))

        ctk.CTkLabel(sidebar, text="Motor de IA:", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        fm = ctk.CTkFrame(sidebar, fg_color="transparent")
        fm.pack(fill="x", padx=15, pady=5)
        self.btn_groq = ctk.CTkButton(fm, text="⚡ Groq", height=42,
                                       fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                       font=("Segoe UI", 10, "bold"),
                                       command=lambda: self._set_modelo("groq"))
        self.btn_groq.pack(side="left", expand=True, fill="x", padx=(0, 3))
        self.btn_gemini = ctk.CTkButton(fm, text="🧠 Gemini", height=42,
                                         fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                         font=("Segoe UI", 10, "bold"),
                                         command=lambda: self._set_modelo("gemini"))
        self.btn_gemini.pack(side="left", expand=True, fill="x", padx=(3, 0))
        self.lbl_limite = ctk.CTkLabel(sidebar, text="Groq: ~1.000 msgs/día",
                                        font=("Segoe UI", 9), text_color=COLOR_TEXT_MUTED)
        self.lbl_limite.pack(pady=(2, 8))
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=5)

        ctk.CTkLabel(sidebar, text="Materia / Examen", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15, pady=(5, 2))
        self.entry_materia = ctk.CTkEntry(sidebar, placeholder_text="Ej: Historia — Tema 4",
                                          fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_materia.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Nivel educativo", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.combo_nivel = ctk.CTkComboBox(sidebar, values=[
            "Primaria", "1º ESO", "2º ESO", "3º ESO", "4º ESO",
            "1º Bachillerato", "2º Bachillerato", "Universidad"],
            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.combo_nivel.set("2º ESO")
        self.combo_nivel.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Nombre del alumno", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.entry_alumno = ctk.CTkEntry(sidebar, placeholder_text="Nombre y apellidos",
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_alumno.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Criterios de puntuación", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.entry_criterios = ctk.CTkEntry(sidebar, placeholder_text="Ej: P1=2pts, P2=3pts, P3=5pts",
                                            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_criterios.pack(fill="x", padx=15, pady=(0, 8))
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=5)

        self.btn_corregir = ctk.CTkButton(sidebar, text="✅ Corregir Examen",
                                           height=44, font=("Segoe UI", 13, "bold"),
                                           fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                                           command=self.corregir_examen)
        self.btn_corregir.pack(fill="x", padx=15, pady=5)

        self.btn_guardar_alumno = ctk.CTkButton(sidebar, text="💾 Guardar alumno",
                                                 height=36, fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_HOVER,
                                                 command=self.guardar_alumno, state="disabled")
        self.btn_guardar_alumno.pack(fill="x", padx=15, pady=2)

        self.btn_tts_corrector = ctk.CTkButton(sidebar, text="🔊 Escuchar Corrección", height=36,
                                              fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                              hover_color=COLOR_ACCENT_HOVER,
                                              command=self._toggle_tts, state="disabled")
        self.btn_tts_corrector.pack(fill="x", padx=15, pady=2)

        self.btn_exportar_uno = ctk.CTkButton(sidebar, text="📄 Exportar (Word)",
                                               height=36, fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                               command=self.exportar_correccion, state="disabled")
        self.btn_exportar_uno.pack(fill="x", padx=15, pady=2)

        self.btn_exportar_clase = ctk.CTkButton(sidebar, text="📊 Informe de clase",
                                                 height=36, fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                                 command=self.exportar_informe_clase, state="disabled")
        self.btn_exportar_clase.pack(fill="x", padx=15, pady=2)

        ctk.CTkButton(sidebar, text="🧹 Nuevo examen", height=32,
                      fg_color="transparent", border_width=1, border_color=COLOR_BORDER,
                      command=self.limpiar_todo).pack(fill="x", padx=15, pady=2)

        ctk.CTkLabel(sidebar, text="Alumnos corregidos",
                     font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED).pack(pady=(8, 2))
        self.frame_alumnos = ctk.CTkScrollableFrame(sidebar, fg_color="transparent", height=100)
        self.frame_alumnos.pack(fill="x", padx=10, pady=2)
        self.status = ctk.CTkLabel(sidebar, text="Listo", text_color=COLOR_TEXT_MUTED, font=("Segoe UI", 10))
        self.status.pack(side="bottom", pady=6)

        # Panel central — Enunciado
        panel_izq = ctk.CTkFrame(self, fg_color="transparent")
        panel_izq.grid(row=0, column=1, sticky="nsew", padx=(15, 7), pady=15)
        panel_izq.grid_rowconfigure(1, weight=1)
        panel_izq.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(panel_izq, text="📝 Enunciado y Criterios del examen",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.txt_enunciado = ctk.CTkTextbox(panel_izq, font=("Consolas", 12), fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_enunciado.grid(row=1, column=0, sticky="nsew")
        self.txt_enunciado.insert("end",
            "Pega aquí el enunciado del examen y/o los criterios de corrección...\n\n"
            "Ejemplo:\n"
            "Pregunta 1 (2 pts): ¿Qué es la fotosíntesis? Explícala.\n"
            "Pregunta 2 (3 pts): Nombra 3 diferencias entre células animales y vegetales.")

        # Panel derecho — Respuestas y corrección
        panel_der = ctk.CTkFrame(self, fg_color="transparent")
        panel_der.grid(row=0, column=2, sticky="nsew", padx=(7, 15), pady=15)
        panel_der.grid_rowconfigure(1, weight=2)
        panel_der.grid_rowconfigure(3, weight=3)
        panel_der.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(panel_der, text="✍️ Respuestas del alumno",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.txt_respuestas = ctk.CTkTextbox(panel_der, font=("Consolas", 12), fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_respuestas.grid(row=1, column=0, sticky="nsew", pady=(0, 10))
        self.txt_respuestas.insert("end", "Pega o escribe aquí las respuestas del alumno...")

        sep_frame = ctk.CTkFrame(panel_der, fg_color="transparent")
        sep_frame.grid(row=2, column=0, sticky="ew", pady=4)
        ctk.CTkLabel(sep_frame, text="📋 Corrección Inteligente",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_ACCENT_CYAN).pack(side="left")
        self.lbl_nota = ctk.CTkLabel(sep_frame, text="", font=("Segoe UI", 16, "bold"), text_color=COLOR_SUCCESS)
        self.lbl_nota.pack(side="right")

        self.txt_correccion = ctk.CTkTextbox(panel_der, font=("Consolas", 12), state="disabled",
                                            fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_correccion.grid(row=3, column=0, sticky="nsew")

    def _set_modelo(self, m):
        self.modelo_actual = m
        if m == "groq":
            self.btn_groq.configure(fg_color=COLOR_ACCENT_PRIMARY); self.btn_gemini.configure(fg_color=COLOR_BG_SURFACE)
            self.lbl_limite.configure(text="Groq: ~1.000 msgs/día")
        else:
            self.btn_groq.configure(fg_color=COLOR_BG_SURFACE); self.btn_gemini.configure(fg_color=COLOR_ACCENT_PURPLE)
            self.lbl_limite.configure(text="Gemini: 15 msgs/minuto")

    def _llamar_ia(self, prompt):
        full_prompt = f"{INSTRUCCIONES_CORRECTOR}\n\n{prompt}"
        if self.modelo_actual == "groq":
            return llamar_groq(full_prompt)
        return llamar_gemini(full_prompt)

    def corregir_examen(self):
        enunciado = self.txt_enunciado.get("1.0", "end-1c").strip()
        respuestas = self.txt_respuestas.get("1.0", "end-1c").strip()
        alumno = self.entry_alumno.get().strip() or "Alumno sin nombre"
        materia = self.entry_materia.get().strip() or "Examen"
        criterios = self.entry_criterios.get().strip()
        if not enunciado or "Pega aquí" in enunciado:
            messagebox.showwarning("Falta enunciado", "Introduce el enunciado del examen.")
            return
        if not respuestas or "Pega o escribe" in respuestas:
            messagebox.showwarning("Faltan respuestas", "Introduce las respuestas del alumno.")
            return
        self.btn_corregir.configure(state="disabled")
        self.status.configure(text="Corrigiendo...", text_color="orange")
        self.lbl_nota.configure(text="")
        prompt = (f"MATERIA: {materia}\nNIVEL: {self.combo_nivel.get()}\nALUMNO: {alumno}\n"
                  + (f"CRITERIOS DE PUNTUACIÓN: {criterios}\n" if criterios else "")
                  + f"\nENUNCIADO DEL EXAMEN:\n{enunciado}\n\nRESPUESTAS DEL ALUMNO:\n{respuestas}")
        threading.Thread(target=self._thread_corregir, args=(prompt, alumno, respuestas), daemon=True).start()

    def _thread_corregir(self, prompt, alumno, respuestas):
        try:
            resultado = self._llamar_ia(prompt)
            self.correccion_actual = resultado
            nota_str = ""
            for linea in resultado.splitlines():
                if "NOTA TOTAL" in linea.upper() or "NOTA FINAL" in linea.upper():
                    partes = linea.split(":")
                    if len(partes) > 1:
                        nota_str = partes[-1].strip().split()[0]
                    break
            def _mostrar():
                self.txt_correccion.configure(state="normal")
                self.txt_correccion.delete("1.0", "end")
                self.txt_correccion.insert("end", resultado)
                self.txt_correccion.configure(state="disabled")
                if nota_str:
                    self.lbl_nota.configure(text=f"Nota: {nota_str}")
                for btn in [self.btn_guardar_alumno, self.btn_exportar_uno, self.btn_tts_corrector]:
                    btn.configure(state="normal")
            self.after(0, _mostrar)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.btn_corregir.configure(state="normal"))
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_corrector.configure(text="🔊 Escuchar Corrección", fg_color=COLOR_BG_CARD_LIGHT)
        else:
            texto = self.txt_correccion.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin corrección", "Primero realiza una corrección para escucharla en voz alta.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_corrector.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_corrector.configure(text="🔊 Escuchar Corrección", fg_color=COLOR_BG_CARD_LIGHT)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def guardar_alumno(self):
        alumno = self.entry_alumno.get().strip() or "Alumno sin nombre"
        if not self.correccion_actual:
            messagebox.showwarning("Sin corrección", "Primero corrige un examen.")
            return
        self.alumnos[alumno] = {
            "respuestas": self.txt_respuestas.get("1.0", "end-1c"),
            "correccion": self.correccion_actual,
            "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M")}
        self._actualizar_lista_alumnos()
        self.btn_exportar_clase.configure(state="normal")
        messagebox.showinfo("Guardado", f"Resultado de '{alumno}' guardado.\nTotal alumnos: {len(self.alumnos)}")

    def _actualizar_lista_alumnos(self):
        for w in self.frame_alumnos.winfo_children():
            w.destroy()
        for nombre in self.alumnos:
            fila = ctk.CTkFrame(self.frame_alumnos, fg_color="transparent")
            fila.pack(fill="x", pady=2)
            ctk.CTkButton(fila, text=nombre, fg_color="transparent",
                          text_color=COLOR_TEXT_MAIN, anchor="w", hover_color=COLOR_BG_SURFACE, height=28,
                          command=lambda n=nombre: self._cargar_alumno(n)).pack(side="left", fill="x", expand=True)
            ctk.CTkButton(fila, text="❌", width=28, height=28,
                          fg_color=COLOR_DANGER, hover_color=COLOR_DANGER_HOVER,
                          command=lambda n=nombre: self._borrar_alumno(n)).pack(side="right")

    def _cargar_alumno(self, nombre):
        datos = self.alumnos[nombre]
        self.entry_alumno.delete(0, "end"); self.entry_alumno.insert(0, nombre)
        self.txt_respuestas.delete("1.0", "end"); self.txt_respuestas.insert("end", datos["respuestas"])
        self.txt_correccion.configure(state="normal")
        self.txt_correccion.delete("1.0", "end"); self.txt_correccion.insert("end", datos["correccion"])
        self.txt_correccion.configure(state="disabled")
        self.correccion_actual = datos["correccion"]
        for btn in [self.btn_guardar_alumno, self.btn_exportar_uno]:
            btn.configure(state="normal")

    def _borrar_alumno(self, nombre):
        if messagebox.askyesno("Confirmar", f"¿Borrar resultado de '{nombre}'?"):
            del self.alumnos[nombre]
            self._actualizar_lista_alumnos()
            if not self.alumnos:
                self.btn_exportar_clase.configure(state="disabled")

    def limpiar_todo(self):
        self.txt_respuestas.delete("1.0", "end")
        self.txt_correccion.configure(state="normal")
        self.txt_correccion.delete("1.0", "end")
        self.txt_correccion.configure(state="disabled")
        self.entry_alumno.delete(0, "end")
        self.correccion_actual = ""
        self.lbl_nota.configure(text="")
        for btn in [self.btn_guardar_alumno, self.btn_exportar_uno]:
            btn.configure(state="disabled")

    def exportar_correccion(self):
        if not self.correccion_actual:
            messagebox.showwarning("Vacío", "No hay corrección que exportar.")
            return
        alumno  = self.entry_alumno.get().strip() or "Alumno"
        materia = self.entry_materia.get().strip() or "Examen"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Correccion_{alumno}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return
        doc = Document()
        doc.add_heading(f"Corrección — {materia} (KernossIA)", 0)
        doc.add_paragraph(f"Alumno: {alumno}")
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
        doc.add_paragraph(f"Fecha corrección: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")
        doc.add_heading("Respuestas del alumno", level=1)
        doc.add_paragraph(self.txt_respuestas.get("1.0", "end-1c"))
        doc.add_page_break()
        doc.add_heading("Corrección y calificación", level=1)
        doc.add_paragraph(self.correccion_actual)
        doc.save(path)
        messagebox.showinfo("Exportado", f"Corrección guardada en:\n{path}")

    def exportar_informe_clase(self):
        if not self.alumnos:
            messagebox.showwarning("Sin datos", "No hay alumnos guardados todavía.")
            return
        materia = self.entry_materia.get().strip() or "Examen"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"InformeClase_{materia}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return
        doc = Document()
        doc.add_heading(f"Informe de Clase — {materia} (KernossIA)", 0)
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
        doc.add_paragraph(f"Total alumnos corregidos: {len(self.alumnos)}")
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")
        for nombre, datos in self.alumnos.items():
            doc.add_heading(nombre, level=1)
            doc.add_paragraph(f"Corregido: {datos['timestamp']}")
            doc.add_paragraph("")
            doc.add_heading("Corrección", level=2)
            doc.add_paragraph(datos["correccion"])
            doc.add_page_break()
        doc.save(path)
        messagebox.showinfo("Exportado", f"Informe de clase guardado en:\n{path}")


# ══════════════════════════════════════════════════════════════════════════════
#  MÓDULO: GENERADOR Y EDITOR DE MAPAS MENTALES CON IA
# ══════════════════════════════════════════════════════════════════════════════
INSTRUCCIONES_MAPA_MENTAL = """
Eres un pedagogo experto en síntesis visual y mapas conceptuales educativos.
Tu objetivo es transformar el tema académico en una estructura de Mapa Mental jerárquica, clara y visualmente estructurada.

Debes responder ÚNICAMENTE con un objeto JSON válido (sin texto antes ni después) estructurado de la siguiente forma:
{
  "tema_central": "Título conciso del tema central",
  "descripcion_general": "Resumen conceptual sintético de 2-3 frases.",
  "ramas": [
    {
      "titulo": "Rama 1 (ej: Origen, Definición, Fases, etc.)",
      "descripcion": "Explicación breve de la rama.",
      "sub_conceptos": [
        {
          "nombre": "Subconcepto 1.1",
          "detalle": "Dato o definición clave."
        },
        {
          "nombre": "Subconcepto 1.2",
          "detalle": "Dato o definición clave."
        }
      ]
    },
    {
      "titulo": "Rama 2",
      "descripcion": "Explicación breve.",
      "sub_conceptos": [
        {
          "nombre": "Subconcepto 2.1",
          "detalle": "Dato o definición clave."
        }
      ]
    }
  ]
}

Genera entre 4 y 6 ramas principales coherentes con el nivel educativo indicado. Asegúrate de que las definiciones sean precisas y claras.
"""

class ModuloMapaMental(ctk.CTkFrame):
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.modelo_actual = "groq"
        self.datos_mapa = None
        self.fig = None
        self.ax = None
        self.canvas_grafico = None
        self._build_ui()

    def _build_ui(self):
        # 2 Columnas principales: Izquierda (Entradas + Editor) y Derecha (Mapa Visual Interactivo)
        self.grid_columnconfigure(0, weight=4, minsize=420)
        self.grid_columnconfigure(1, weight=6, minsize=550)
        self.grid_rowconfigure(0, weight=1)

        # ── PANEL IZQUIERDO: FORMULARIO Y EDITOR ──
        panel_izq = ctk.CTkFrame(self, corner_radius=14, fg_color=COLOR_BG_CARD,
                                 border_width=1, border_color=COLOR_BORDER)
        panel_izq.grid(row=0, column=0, sticky="nsew", padx=(20, 10), pady=20)
        panel_izq.grid_rowconfigure(7, weight=1)
        panel_izq.grid_columnconfigure(0, weight=1)

        # Encabezado
        frame_header = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_header.grid(row=0, column=0, sticky="ew", padx=18, pady=(16, 10))

        ctk.CTkLabel(frame_header, text="🧠 Generador de Mapas Mentales",
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(anchor="w")
        ctk.CTkLabel(frame_header, text="Genera esquemas conceptuales con IA, edítalos y expórtalos",
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(anchor="w", pady=(2, 0))

        # Campo: Tema Principal
        ctk.CTkLabel(panel_izq, text="Tema o Materia a Sintetizar:",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=1, column=0, sticky="w", padx=18, pady=(4, 2))
        self.entry_tema = ctk.CTkEntry(panel_izq, placeholder_text="Ej: La Célula y Fotosíntesis, Guerra Fría, Vectores...",
                                       height=38, font=("Segoe UI", 12),
                                       fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_tema.grid(row=2, column=0, sticky="ew", padx=18, pady=(0, 10))

        # Fila: Curso / Nivel + Enfoque
        frame_opts = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_opts.grid(row=3, column=0, sticky="ew", padx=18, pady=(0, 10))
        frame_opts.grid_columnconfigure((0, 1), weight=1)

        ctk.CTkLabel(frame_opts, text="Nivel / Curso:", font=("Segoe UI", 11, "bold"),
                     text_color=COLOR_TEXT_MUTED).grid(row=0, column=0, sticky="w", padx=(0, 6), pady=(0, 2))
        self.combo_nivel = ctk.CTkComboBox(frame_opts,
                                           values=["Secundaria / ESO", "Bachillerato", "Universidad / FP", "Primaria", "General"],
                                           font=("Segoe UI", 11), height=34,
                                           fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.combo_nivel.set("Secundaria / ESO")
        self.combo_nivel.grid(row=1, column=0, sticky="ew", padx=(0, 6))

        ctk.CTkLabel(frame_opts, text="Enfoque Específico (Opcional):", font=("Segoe UI", 11, "bold"),
                     text_color=COLOR_TEXT_MUTED).grid(row=0, column=1, sticky="w", padx=(6, 0), pady=(0, 2))
        self.entry_enfoque = ctk.CTkEntry(frame_opts, placeholder_text="Ej: Énfasis en fórmulas...",
                                          font=("Segoe UI", 11), height=34,
                                          fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_enfoque.grid(row=1, column=1, sticky="ew", padx=(6, 0))

        # Fila: Selector de Modelo IA + Botón Generar
        frame_ia_bar = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_ia_bar.grid(row=4, column=0, sticky="ew", padx=18, pady=(0, 10))
        frame_ia_bar.grid_columnconfigure(0, weight=1)

        frame_model_switch = ctk.CTkFrame(frame_ia_bar, fg_color=COLOR_BG_CARD_LIGHT,
                                          border_width=1, border_color=COLOR_BORDER, corner_radius=8)
        frame_model_switch.pack(side="left")

        self.btn_groq = ctk.CTkButton(frame_model_switch, text="⚡ Groq", height=28, width=75,
                                      font=("Segoe UI", 10, "bold"),
                                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                      command=lambda: self._set_modelo("groq"))
        self.btn_groq.pack(side="left", padx=2, pady=2)

        self.btn_gemini = ctk.CTkButton(frame_model_switch, text="🧠 Gemini", height=28, width=75,
                                        font=("Segoe UI", 10, "bold"),
                                        fg_color="transparent", hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                        command=lambda: self._set_modelo("gemini"))
        self.btn_gemini.pack(side="left", padx=2, pady=2)

        self.btn_generar = ctk.CTkButton(frame_ia_bar, text="✨ Generar con IA", height=34,
                                         font=("Segoe UI", 12, "bold"),
                                         fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                         command=self.generar_mapa_mental)
        self.btn_generar.pack(side="right", fill="x", expand=True, padx=(10, 0))

        # Estado
        self.lbl_status = ctk.CTkLabel(panel_izq, text="Listo para generar", font=("Segoe UI", 11),
                                       text_color=COLOR_TEXT_DIM)
        self.lbl_status.grid(row=5, column=0, sticky="w", padx=18, pady=(0, 6))

        # Editor de Estructura / JSON en tiempo real
        frame_edit_header = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_edit_header.grid(row=6, column=0, sticky="ew", padx=18, pady=(4, 4))
        ctk.CTkLabel(frame_edit_header, text="📝 Editor de Estructura:",
                     font=("Segoe UI", 11, "bold"), text_color=COLOR_ACCENT_CYAN).pack(side="left")

        self.btn_redibujar = ctk.CTkButton(frame_edit_header, text="🔄 Redibujar Mapa", height=24, width=120,
                                           font=("Segoe UI", 10, "bold"),
                                           fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_HOVER,
                                           border_width=1, border_color=COLOR_BORDER,
                                           command=self.redibujar_desde_editor)
        self.btn_redibujar.pack(side="right")

        self.txt_estructura = ctk.CTkTextbox(panel_izq, font=("Consolas", 11), wrap="word",
                                             fg_color=COLOR_BG_CARD_LIGHT, border_width=1,
                                             border_color=COLOR_BORDER, corner_radius=10)
        self.txt_estructura.grid(row=7, column=0, sticky="nsew", padx=18, pady=(0, 10))

        # Botones de Acción (Exportar Word, Exportar Imagen)
        frame_acciones = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_acciones.grid(row=8, column=0, sticky="ew", padx=18, pady=(0, 16))
        frame_acciones.grid_columnconfigure((0, 1), weight=1)

        self.btn_exportar_word = ctk.CTkButton(frame_acciones, text="📄 Exportar a Word (.docx)", height=36,
                                               font=("Segoe UI", 11, "bold"),
                                               fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                               state="disabled",
                                               command=self.exportar_word)
        self.btn_exportar_word.grid(row=0, column=0, sticky="ew", padx=(0, 4))

        self.btn_exportar_img = ctk.CTkButton(frame_acciones, text="🖼️ Guardar Imagen (.png)", height=36,
                                              font=("Segoe UI", 11, "bold"),
                                              fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                                              state="disabled",
                                              command=self.exportar_imagen)
        self.btn_exportar_img.grid(row=0, column=1, sticky="ew", padx=(4, 0))


        # ── PANEL DERECHO: VISUALIZADOR GRÁFICO INTERACTIVO ──
        self.panel_der = ctk.CTkFrame(self, corner_radius=14, fg_color=COLOR_BG_CARD,
                                      border_width=1, border_color=COLOR_BORDER)
        self.panel_der.grid(row=0, column=1, sticky="nsew", padx=(10, 20), pady=20)
        self.panel_der.grid_rowconfigure(1, weight=1)
        self.panel_der.grid_columnconfigure(0, weight=1)

        # Header del visualizador
        header_der = ctk.CTkFrame(self.panel_der, fg_color="transparent")
        header_der.grid(row=0, column=0, sticky="ew", padx=18, pady=(16, 6))

        ctk.CTkLabel(header_der, text="🎨 Vista Gráfica del Mapa Mental",
                     font=("Segoe UI", 16, "bold"), text_color=COLOR_TEXT_MAIN).pack(side="left")

        self.lbl_info_ramas = ctk.CTkLabel(header_der, text="Sin mapa generado",
                                           font=("Segoe UI", 11), text_color=COLOR_TEXT_DIM)
        self.lbl_info_ramas.pack(side="right")

        # Contenedor del Canvas de Matplotlib
        self.frame_canvas = ctk.CTkFrame(self.panel_der, fg_color="#070c18", corner_radius=10,
                                         border_width=1, border_color=COLOR_BORDER)
        self.frame_canvas.grid(row=1, column=0, sticky="nsew", padx=18, pady=(0, 18))
        self.frame_canvas.grid_rowconfigure(0, weight=1)
        self.frame_canvas.grid_columnconfigure(0, weight=1)

        # Inicializar figura de matplotlib
        self._inicializar_canvas_vacio()

    def _set_modelo(self, modelo):
        self.modelo_actual = modelo
        if modelo == "groq":
            self.btn_groq.configure(fg_color=COLOR_ACCENT_PRIMARY)
            self.btn_gemini.configure(fg_color="transparent")
        else:
            self.btn_groq.configure(fg_color="transparent")
            self.btn_gemini.configure(fg_color=COLOR_ACCENT_PURPLE)

    def _inicializar_canvas_vacio(self):
        plt.style.use("dark_background")
        self.fig, self.ax = plt.subplots(figsize=(8, 7), facecolor="#070c18")
        self.ax.set_facecolor("#070c18")
        self.ax.text(0, 0, "🧠 Escribe un tema y haz clic en\n'✨ Generar con IA' para crear tu mapa",
                     color="#64748b", fontsize=12, ha="center", va="center",
                     bbox=dict(boxstyle="round,pad=0.8", facecolor="#0a1124", edgecolor="#1e3a6a", lw=1.5))
        self.ax.set_xlim(-6, 6)
        self.ax.set_ylim(-6, 6)
        self.ax.axis("off")

        self.canvas_grafico = FigureCanvasTkAgg(self.fig, master=self.frame_canvas)
        self.canvas_grafico.get_tk_widget().pack(fill="both", expand=True)
        self.canvas_grafico.draw()

    def generar_mapa_mental(self):
        tema = self.entry_tema.get().strip()
        if not tema:
            messagebox.showwarning("Tema requerido", "Introduce el tema o concepto para crear el mapa mental.")
            return

        nivel = self.combo_nivel.get()
        enfoque = self.entry_enfoque.get().strip()

        self.btn_generar.configure(state="disabled")
        self.lbl_status.configure(text="✨ Creando estructura con IA...", text_color=COLOR_ACCENT_SKY)

        prompt = f"TEMA PRINCIPAL: {tema}\nNIVEL EDUCATIVO: {nivel}\n"
        if enfoque:
            prompt += f"ENFOQUE Y PUNTOS CLAVE: {enfoque}\n"

        threading.Thread(target=self._thread_generar_mapa, args=(prompt,), daemon=True).start()

    def _thread_generar_mapa(self, prompt):
        try:
            full_prompt = f"{INSTRUCCIONES_MAPA_MENTAL}\n\n{prompt}"
            if self.modelo_actual == "groq":
                respuesta = llamar_groq(full_prompt)
            else:
                respuesta = llamar_gemini(full_prompt)

            datos = self._extraer_json(respuesta)
            if not datos:
                datos = {
                    "tema_central": self.entry_tema.get().strip() or "Tema Principal",
                    "descripcion_general": "Esquema conceptual estructurado.",
                    "ramas": [
                        {"titulo": "Conceptos Clave", "descripcion": "Puntos esenciales.", "sub_conceptos": [{"nombre": "Definición", "detalle": "Concepto central."}]},
                        {"titulo": "Características", "descripcion": "Propiedades fundamentales.", "sub_conceptos": [{"nombre": "Propiedad 1", "detalle": "Detalle explicativo."}]},
                        {"titulo": "Aplicaciones", "descripcion": "Uso práctico.", "sub_conceptos": [{"nombre": "Ejemplo práctico", "detalle": "Demostración."}]}
                    ]
                }

            self.after(0, lambda: self._mostrar_mapa_generado(datos))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error al Generar", f"No se pudo generar el mapa mental: {e}"))
        finally:
            self.after(0, lambda: [
                self.btn_generar.configure(state="normal"),
                self.lbl_status.configure(text="Listo", text_color=COLOR_SUCCESS)
            ])

    def _extraer_json(self, texto):
        try:
            match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", texto, re.DOTALL)
            if match:
                return json.loads(match.group(1))
            start = texto.find("{")
            end = texto.rfind("}")
            if start != -1 and end != -1:
                return json.loads(texto[start:end+1])
        except Exception:
            pass
        return None

    def _mostrar_mapa_generado(self, datos):
        self.datos_mapa = datos

        self.txt_estructura.delete("1.0", "end")
        self.txt_estructura.insert("end", json.dumps(datos, ensure_ascii=False, indent=2))

        self._dibujar_mapa_visual(datos)

        self.btn_exportar_word.configure(state="normal")
        self.btn_exportar_img.configure(state="normal")
        ramas_count = len(datos.get("ramas", []))
        self.lbl_info_ramas.configure(text=f"✨ {ramas_count} Ramas Principales Generadas", text_color=COLOR_ACCENT_CYAN)

    def redibujar_desde_editor(self):
        contenido = self.txt_estructura.get("1.0", "end-1c").strip()
        if not contenido:
            return
        try:
            datos = json.loads(contenido)
            self.datos_mapa = datos
            self._dibujar_mapa_visual(datos)
            self.lbl_status.configure(text="Mapa visual actualizado", text_color=COLOR_SUCCESS)
        except Exception as e:
            messagebox.showerror("Error JSON", f"El formato JSON no es válido:\n{e}")

    def _dibujar_mapa_visual(self, datos):
        import numpy as np

        self.ax.clear()
        self.ax.set_facecolor("#070c18")

        tema_central = datos.get("tema_central", "Tema Central")
        ramas = datos.get("ramas", [])
        num_ramas = len(ramas)

        if num_ramas == 0:
            self._inicializar_canvas_vacio()
            return

        paleta_colores = ["#06b6d4", "#38bdf8", "#6366f1", "#10b981", "#f59e0b", "#ec4899", "#8b5cf6"]

        # 1. Dibujar Nodo Central
        tema_fmt = "\n".join([tema_central[i:i+16] for i in range(0, len(tema_central), 16)])
        self.ax.text(0, 0, f"🌟\n{tema_fmt}",
                     color="#ffffff", fontsize=11, fontweight="bold", ha="center", va="center",
                     bbox=dict(boxstyle="round,pad=0.7", facecolor="#2563eb", edgecolor="#38bdf8", lw=2.5, alpha=0.95))

        # 2. Dibujar Ramas y Subconceptos
        r_rama = 3.2
        r_sub = 5.2

        for i, rama in enumerate(ramas):
            color = paleta_colores[i % len(paleta_colores)]
            angulo = (2 * np.pi * i / num_ramas) + (np.pi / (num_ramas * 2))

            x_rama = r_rama * np.cos(angulo)
            y_rama = r_rama * np.sin(angulo)

            # Conector Centro -> Rama (Curva suave)
            self.ax.annotate("", xy=(x_rama, y_rama), xytext=(0, 0),
                             arrowprops=dict(arrowstyle="-", color=color, lw=2.2, alpha=0.85,
                                             connectionstyle="arc3,rad=0.1"))

            # Nodo Rama
            titulo_rama = rama.get("titulo", f"Rama {i+1}")
            titulo_fmt = "\n".join([titulo_rama[k:k+14] for k in range(0, len(titulo_rama), 14)])

            self.ax.text(x_rama, y_rama, titulo_fmt,
                         color="#f8fafc", fontsize=9.5, fontweight="bold", ha="center", va="center",
                         bbox=dict(boxstyle="round,pad=0.5", facecolor="#0f1a35", edgecolor=color, lw=2.0, alpha=0.95))

            # Sub-conceptos
            subs = rama.get("sub_conceptos", [])
            num_subs = len(subs)

            for j, sub in enumerate(subs):
                offset_ang = (j - (num_subs - 1) / 2) * (0.35 if num_subs > 1 else 0)
                sub_ang = angulo + offset_ang

                x_sub = r_sub * np.cos(sub_ang)
                y_sub = r_sub * np.sin(sub_ang)

                # Conector Rama -> Subconcepto
                self.ax.annotate("", xy=(x_sub, y_sub), xytext=(x_rama, y_rama),
                                 arrowprops=dict(arrowstyle="-", color=color, lw=1.2, ls="--", alpha=0.6,
                                                 connectionstyle="arc3,rad=-0.08"))

                # Nodo Subconcepto
                nombre_sub = sub.get("nombre", f"Punto {j+1}")
                nombre_fmt = "\n".join([nombre_sub[k:k+16] for k in range(0, len(nombre_sub), 16)])

                self.ax.text(x_sub, y_sub, nombre_fmt,
                             color="#cbd5e1", fontsize=8, ha="center", va="center",
                             bbox=dict(boxstyle="round,pad=0.35", facecolor="#070c18", edgecolor=color, lw=1.0, alpha=0.9))

        self.ax.set_xlim(-6.8, 6.8)
        self.ax.set_ylim(-6.8, 6.8)
        self.ax.set_aspect("equal")
        self.ax.axis("off")
        self.canvas_grafico.draw()

    def exportar_imagen(self):
        if not self.datos_mapa or not self.fig:
            messagebox.showwarning("Sin mapa", "Primero genera o redibuja un mapa mental.")
            return

        tema = self.datos_mapa.get("tema_central", "Mapa_Mental").replace(" ", "_")
        path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("Imagen PNG", "*.png"), ("Todos los archivos", "*.*")],
            initialfile=f"MapaMental_{tema}_{datetime.now().strftime('%Y%m%d')}.png"
        )
        if not path:
            return

        try:
            self.fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="#070c18", edgecolor="none")
            messagebox.showinfo("Imagen Guardada", f"Mapa mental exportado con éxito en alta resolución (300 DPI):\n{path}")
        except Exception as e:
            messagebox.showerror("Error al Guardar", f"No se pudo guardar la imagen:\n{e}")

    def exportar_word(self):
        if not self.datos_mapa:
            messagebox.showwarning("Sin mapa", "Primero genera o redibuja un mapa mental.")
            return

        tema = self.datos_mapa.get("tema_central", "Mapa Mental")
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documento Word", "*.docx"), ("Todos los archivos", "*.*")],
            initialfile=f"MapaMental_{tema.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        )
        if not path:
            return

        try:
            doc = Document()
            # Título principal
            titulo_p = doc.add_heading(f"Mapa Mental: {tema}", 0)
            doc.add_paragraph(f"Nivel Académico: {self.combo_nivel.get()}  |  Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            
            desc_gen = self.datos_mapa.get("descripcion_general", "")
            if desc_gen:
                p_res = doc.add_paragraph()
                p_res.add_run("Resumen Conceptual: ").bold = True
                p_res.add_run(desc_gen)

            # Guardar imagen temporal e incrustar
            temp_img = os.path.expanduser("~/.temp_mapa_mental_export.png")
            self.fig.savefig(temp_img, dpi=300, bbox_inches="tight", facecolor="#070c18", edgecolor="none")

            doc.add_paragraph("")
            doc.add_heading("Estructura Gráfica del Mapa Mental", level=1)
            doc.add_picture(temp_img, width=Inches(6.2))
            doc.add_paragraph("")

            if os.path.exists(temp_img):
                os.remove(temp_img)

            # Desglose de ramas y conceptos
            doc.add_heading("Desglose Detallado de Ramas y Conceptos", level=1)

            for rama in self.datos_mapa.get("ramas", []):
                doc.add_heading(f"📌 {rama.get('titulo', 'Rama')}", level=2)
                if rama.get("descripcion"):
                    doc.add_paragraph(rama.get("descripcion"))

                subs = rama.get("sub_conceptos", [])
                if subs:
                    for sub in subs:
                        p_sub = doc.add_paragraph(style="List Bullet")
                        r_bold = p_sub.add_run(f"{sub.get('nombre', 'Concepto')}: ")
                        r_bold.bold = True
                        p_sub.add_run(sub.get("detalle", ""))

            doc.save(path)
            messagebox.showinfo("Exportado a Word", f"Documento Word creado con éxito con gráfico HD e información detallada:\n{path}")
        except Exception as e:
            messagebox.showerror("Error al Exportar", f"No se pudo generar el documento Word:\n{e}")


# ══════════════════════════════════════════════════════════════════════════════
#  VENTANA MODAL: AJUSTES DE LA APLICACIÓN, IDIOMA & MODELO DE VOZ IA (TTS)
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
#  MODAL: AGREGAR CUENTA AL GESTOR DE MULTICUENTAS
# ══════════════════════════════════════════════════════════════════════════════
class ModalAgregarCuenta(ctk.CTkToplevel):
    """Modal para iniciar sesión con una cuenta adicional y añadirla al switcher."""
    def __init__(self, parent, on_cuenta_agregada):
        super().__init__(parent)
        self.title("➕ Agregar Cuenta a KernossAI")
        self.geometry("450x380")
        self.minsize(400, 340)
        self.configure(fg_color=COLOR_BG_DARK)
        self.transient(parent)
        self.grab_set()
        self.on_cuenta_agregada = on_cuenta_agregada
        self._build_ui()

    def _build_ui(self):
        ctk.CTkLabel(self, text="👥 Iniciar Sesión con Otra Cuenta", font=("Segoe UI", 16, "bold"), text_color=COLOR_ACCENT_SKY).pack(padx=20, pady=(20, 4))
        ctk.CTkLabel(self, text="Agrega una cuenta adicional para alternar entre ellas sin cerrar sesión.", font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED).pack(padx=20, pady=(0, 15))

        f_card = ctk.CTkFrame(self, fg_color=COLOR_BG_CARD, corner_radius=12, border_width=1, border_color=COLOR_BORDER)
        f_card.pack(fill="both", expand=True, padx=20, pady=(0, 15))

        ctk.CTkLabel(f_card, text="Correo Electrónico:", font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=16, pady=(12, 2))
        self.entry_email = ctk.CTkEntry(f_card, height=36, font=("Segoe UI", 11), placeholder_text="ej: usuario@correo.com")
        self.entry_email.pack(fill="x", padx=16, pady=(0, 8))

        ctk.CTkLabel(f_card, text="Contraseña:", font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=16, pady=(0, 2))
        self.entry_pass = ctk.CTkEntry(f_card, height=36, font=("Segoe UI", 11), show="•", placeholder_text="••••••••")
        self.entry_pass.pack(fill="x", padx=16, pady=(0, 10))
        self.entry_pass.bind("<Return>", lambda e: self._guardar())

        self.lbl_estado = ctk.CTkLabel(f_card, text="", font=("Segoe UI", 10))
        self.lbl_estado.pack(anchor="w", padx=16, pady=(0, 6))

        f_btns = ctk.CTkFrame(self, fg_color="transparent")
        f_btns.pack(fill="x", padx=20, pady=(0, 18))

        self.btn_guardar = ctk.CTkButton(f_btns, text="➕ Agregar Cuenta", height=38, font=("Segoe UI", 11, "bold"),
                                          fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER, command=self._guardar)
        self.btn_guardar.pack(side="left", fill="x", expand=True, padx=(0, 6))

        ctk.CTkButton(f_btns, text="Cancelar", height=38, width=85, font=("Segoe UI", 11),
                      fg_color=COLOR_BG_SURFACE, hover_color=COLOR_BORDER, command=self.destroy).pack(side="right")

    def _guardar(self):
        email = self.entry_email.get().strip().lower()
        password = self.entry_pass.get()
        if not email or not password:
            self.lbl_estado.configure(text="Completa todos los campos.", text_color=COLOR_DANGER)
            return

        self.lbl_estado.configure(text="⏳ Verificando credenciales...", text_color=COLOR_ACCENT_SKY)
        self.btn_guardar.configure(state="disabled")

        def _thread():
            ok, msg, sesion = agregar_cuenta_secundaria(email, password)
            if ok:
                def _exito():
                    if self.on_cuenta_agregada:
                        self.on_cuenta_agregada()
                    self.destroy()
                    messagebox.showinfo("Cuenta Agregada", f"La cuenta '{email}' se ha guardado en el gestor de multicuentas.")
                self.after(0, _exito)
            else:
                def _error():
                    self.lbl_estado.configure(text=msg, text_color=COLOR_DANGER)
                    self.btn_guardar.configure(state="normal")
                self.after(0, _error)

        threading.Thread(target=_thread, daemon=True).start()


# ══════════════════════════════════════════════════════════════════════════════
#  MODAL DE AJUSTES & PREFERENCIAS (MULTICUENTAS, ROL, TTS, IDIOMA)
# ══════════════════════════════════════════════════════════════════════════════
class VentanaAjustes(ctk.CTkToplevel):
    def __init__(self, master):
        super().__init__(master)
        self.title(t("ajustes_titulo"))
        self.geometry("620x740")
        self.minsize(540, 600)
        self.configure(fg_color=COLOR_BG_DARK)
        self.transient(master)
        self.grab_set()
        
        self.sesion_actual = getattr(master, "sesion", {})
        self.email_actual = self.sesion_actual.get("email", "").lower()
        self.rol_original = self.sesion_actual.get("rol", "Alumno")
        
        self._build_ui()

    def _build_ui(self):
        # Header
        frame_header = ctk.CTkFrame(self, fg_color="transparent")
        frame_header.pack(fill="x", padx=25, pady=(18, 8))

        ctk.CTkLabel(frame_header, text=t("ajustes_titulo"),
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(anchor="w")
        ctk.CTkLabel(frame_header, text=t("ajustes_subtitulo"),
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(anchor="w", pady=(2, 0))

        # Tarjeta scrollable de configuración principal
        self.scroll_tarjeta = ctk.CTkScrollableFrame(self, fg_color=COLOR_BG_CARD, corner_radius=14,
                                                     border_width=1, border_color=COLOR_BORDER)
        self.scroll_tarjeta.pack(fill="both", expand=True, padx=25, pady=(5, 12))

        # ── SECCIÓN 1: IDIOMA DE LA INTERFAZ ──
        ctk.CTkLabel(self.scroll_tarjeta, text=f"🌐 {t('ajustes_sec_idioma')}:",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=18, pady=(12, 4))
        
        self.map_idiomas_a_id = {v: k for k, v in IDIOMAS_DISPONIBLES.items()}
        nombres_idiomas = list(IDIOMAS_DISPONIBLES.values())
        idioma_actual = obtener_idioma()
        nombre_idioma_actual = IDIOMAS_DISPONIBLES.get(idioma_actual, nombres_idiomas[0])

        self.combo_idioma = ctk.CTkComboBox(self.scroll_tarjeta, values=nombres_idiomas, height=36,
                                             font=("Segoe UI", 12), fg_color=COLOR_BG_CARD_LIGHT,
                                             border_color=COLOR_BORDER)
        self.combo_idioma.set(nombre_idioma_actual)
        self.combo_idioma.pack(fill="x", padx=18, pady=(0, 10))

        # Separador
        ctk.CTkFrame(self.scroll_tarjeta, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=18, pady=(2, 10))

        # ── SECCIÓN 2: ROL ACADÉMICO (ALUMNO / PROFESOR) ──
        ctk.CTkLabel(self.scroll_tarjeta, text="🎓 Modo de Cuenta / Rol Académico:",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=18, pady=(0, 2))
        ctk.CTkLabel(self.scroll_tarjeta, text="Elige entre 'Alumno' (estudio y deberes) o 'Profesor' (herramientas docentes).",
                     font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=18, pady=(0, 4))

        roles_disp = ["Alumno", "Profesor"]
        rol_activo = self.rol_original if self.rol_original in roles_disp else "Alumno"

        self.combo_rol = ctk.CTkComboBox(self.scroll_tarjeta, values=roles_disp, height=36,
                                         font=("Segoe UI", 12), fg_color=COLOR_BG_CARD_LIGHT,
                                         border_color=COLOR_BORDER)
        self.combo_rol.set(rol_activo)
        self.combo_rol.pack(fill="x", padx=18, pady=(0, 10))

        # Separador
        ctk.CTkFrame(self.scroll_tarjeta, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=18, pady=(2, 10))

        # ── SECCIÓN 3: GESTIÓN DE MULTICUENTAS (INICIAR CON MÁS DE 1 CUENTA) ──
        ctk.CTkLabel(self.scroll_tarjeta, text="👥 Gestión de Multicuentas (Alternar Cuentas):",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=18, pady=(0, 2))
        ctk.CTkLabel(self.scroll_tarjeta, text="Inicia sesión con varias cuentas y cambia de una a otra al instante.",
                     font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=18, pady=(0, 6))

        self.frame_multicuentas_lista = ctk.CTkFrame(self.scroll_tarjeta, fg_color=COLOR_BG_SURFACE, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
        self.frame_multicuentas_lista.pack(fill="x", padx=18, pady=(0, 8))

        self._cargar_multicuentas()

        self.btn_add_cuenta = ctk.CTkButton(
            self.scroll_tarjeta, text="➕ Iniciar Sesión con Otra Cuenta",
            font=("Segoe UI", 11, "bold"), height=32,
            fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
            command=self._abrir_modal_agregar_cuenta
        )
        self.btn_add_cuenta.pack(fill="x", padx=18, pady=(0, 10))

        # Separador
        ctk.CTkFrame(self.scroll_tarjeta, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=18, pady=(2, 10))

        # ── SECCIÓN 4: LECTOR DE VOZ IA (TTS) ──
        ctk.CTkLabel(self.scroll_tarjeta, text=f"🔊 {t('ajustes_sec_voz')} - {t('ajustes_lbl_voz')}",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=18, pady=(0, 4))

        voz_actual, vel_actual = obtener_ajustes_tts()

        self.map_voces_a_id = {v: k for k, v in VOICES_DISPONIBLES.items()}
        nombres_voces = list(VOICES_DISPONIBLES.values())
        nombre_voz_actual = VOICES_DISPONIBLES.get(voz_actual, nombres_voces[0])

        self.combo_voz = ctk.CTkComboBox(self.scroll_tarjeta, values=nombres_voces, height=36,
                                         font=("Segoe UI", 12), fg_color=COLOR_BG_CARD_LIGHT,
                                         border_color=COLOR_BORDER)
        self.combo_voz.set(nombre_voz_actual)
        self.combo_voz.pack(fill="x", padx=18, pady=(0, 10))

        # Velocidad de habla
        ctk.CTkLabel(self.scroll_tarjeta, text=f"⚡ {t('ajustes_lbl_velocidad')}",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=18, pady=(0, 4))

        self.map_vel_a_id = {v: k for k, v in VELOCIDADES_DISPONIBLES.items()}
        nombres_vel = list(VELOCIDADES_DISPONIBLES.values())
        nombre_vel_actual = VELOCIDADES_DISPONIBLES.get(vel_actual, nombres_vel[1])

        self.combo_vel = ctk.CTkComboBox(self.scroll_tarjeta, values=nombres_vel, height=36,
                                         font=("Segoe UI", 12), fg_color=COLOR_BG_CARD_LIGHT,
                                         border_color=COLOR_BORDER)
        self.combo_vel.set(nombre_vel_actual)
        self.combo_vel.pack(fill="x", padx=18, pady=(0, 10))

        # Botón Probar Voz
        self.btn_probar = ctk.CTkButton(self.scroll_tarjeta, text=t("ajustes_btn_probar"),
                                        font=("Segoe UI", 11, "bold"), height=32,
                                        fg_color=COLOR_BG_SURFACE, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                        hover_color=COLOR_ACCENT_PRIMARY,
                                        command=self._probar_voz)
        self.btn_probar.pack(fill="x", padx=18, pady=(0, 10))

        # Separador
        ctk.CTkFrame(self.scroll_tarjeta, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=18, pady=(2, 10))

        # ── SECCIÓN 5: ZONA DE CUENTA & PRIVACIDAD (BORRAR CUENTA) ──
        ctk.CTkLabel(self.scroll_tarjeta, text="⚠️ Zona de Privacidad — Cuenta",
                     font=("Segoe UI", 11, "bold"), text_color=COLOR_DANGER).pack(anchor="w", padx=18, pady=(0, 1))

        ctk.CTkLabel(self.scroll_tarjeta, text="Elimina de forma permanente tu usuario, historial y acceso del servidor.",
                     font=("Segoe UI", 9), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=18, pady=(0, 6))

        self.btn_eliminar_cuenta = ctk.CTkButton(
            self.scroll_tarjeta, text="🗑️ Borrar Mi Cuenta Definitivamente",
            font=("Segoe UI", 11, "bold"), height=34,
            fg_color="#7f1d1d", hover_color="#991b1b",
            border_width=1, border_color=COLOR_DANGER,
            command=self._solicitar_eliminar_cuenta
        )
        self.btn_eliminar_cuenta.pack(fill="x", padx=18, pady=(0, 12))

        # Botones Inferiores: Guardar y Cancelar
        frame_btns = ctk.CTkFrame(self, fg_color="transparent")
        frame_btns.pack(fill="x", padx=25, pady=(0, 18))

        ctk.CTkButton(frame_btns, text=t("ajustes_btn_guardar"), height=40,
                      font=("Segoe UI", 12, "bold"),
                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                      command=self._guardar).pack(side="left", fill="x", expand=True, padx=(0, 6))

        ctk.CTkButton(frame_btns, text=t("ajustes_btn_cancelar"), height=40, width=90,
                      font=("Segoe UI", 12),
                      fg_color=COLOR_BG_SURFACE, hover_color=COLOR_BORDER,
                      command=self._cerrar).pack(side="right")

    def _cargar_multicuentas(self):
        for w in self.frame_multicuentas_lista.winfo_children():
            w.destroy()

        cuentas = obtener_cuentas_guardadas()
        if not cuentas:
            if self.email_actual:
                cuentas = [{
                    "email": self.email_actual,
                    "nombre": self.sesion_actual.get("nombre", "Usuario"),
                    "rol": self.sesion_actual.get("rol", "Alumno")
                }]

        for c in cuentas:
            em = c.get("email", "").lower()
            nom = c.get("nombre", em)
            rol = c.get("rol", "Alumno")
            es_activa = (em == self.email_actual)

            row = ctk.CTkFrame(self.frame_multicuentas_lista, fg_color=COLOR_BG_CARD if es_activa else "transparent", corner_radius=6)
            row.pack(fill="x", padx=6, pady=3)

            badge = "🟢 (Activa)" if es_activa else "👤"
            ctk.CTkLabel(row, text=f"{badge} {nom} ({em}) — {rol}", font=("Segoe UI", 10, "bold" if es_activa else "normal"),
                         text_color=COLOR_ACCENT_SKY if es_activa else COLOR_TEXT_MAIN).pack(side="left", padx=8, pady=6)

            if not es_activa:
                ctk.CTkButton(row, text="🗑️", width=30, height=26, font=("Segoe UI", 10),
                              fg_color="transparent", hover_color=COLOR_DANGER_HOVER, text_color=COLOR_TEXT_MUTED,
                              command=lambda e=em: self._quitar_cuenta_switcher(e)).pack(side="right", padx=(2, 6), pady=4)

                ctk.CTkButton(row, text="🔄 Cambiar", width=80, height=26, font=("Segoe UI", 10, "bold"),
                              fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                              command=lambda e=em: self._switch_to_account(e)).pack(side="right", padx=2, pady=4)

    def _abrir_modal_agregar_cuenta(self):
        ModalAgregarCuenta(self, on_cuenta_agregada=self._cargar_multicuentas)

    def _switch_to_account(self, email):
        ok, msg, nueva_sesion = cambiar_a_cuenta(email)
        if ok:
            messagebox.showinfo("Sesión Cambiada", f"Has cambiado a la cuenta de: {nueva_sesion.get('nombre', email)}")
            self._cerrar()
            if hasattr(self.master, "_al_cambiar_cuenta"):
                self.master._al_cambiar_cuenta(nueva_sesion)
        else:
            messagebox.showerror("Error", msg)

    def _quitar_cuenta_switcher(self, email):
        eliminar_cuenta_switcher(email)
        self._cargar_multicuentas()

    def _probar_voz(self):
        nombre_voz = self.combo_voz.get()
        voz_id = self.map_voces_a_id.get(nombre_voz, "es-ES-AlvaroNeural")
        nombre_vel = self.combo_vel.get()
        vel_id = self.map_vel_a_id.get(nombre_vel, "+0%")

        guardar_ajustes_tts(voz_id, vel_id)

        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_probar.configure(text=t("ajustes_btn_probar"), fg_color=COLOR_BG_SURFACE)
            return

        def _callback(reproduciendo):
            if reproduciendo:
                self.btn_probar.configure(text="⏹️ Detener Prueba", fg_color=COLOR_DANGER)
            else:
                self.btn_probar.configure(text=t("ajustes_btn_probar"), fg_color=COLOR_BG_SURFACE)

        frases_prueba = {
            "es": "Hola, esta es una prueba de la voz del asistente en KernossIA. ¿Qué te parece?",
            "en": "Hello, this is a test of the assistant voice in KernossAI. How does it sound?",
            "de": "Hallo, dies ist ein Test der KI-Stimme in KernossAI. Wie gefällt sie dir?",
            "fr": "Bonjour, ceci est un test de la voix de l'assistant dans KernossAI. Qu'en pensez-vous ?"
        }
        idioma = obtener_idioma()
        frase = frases_prueba.get(idioma, frases_prueba["es"])

        tts_engine.hablar(
            frase,
            callback_estado=lambda r: self.after(0, lambda: _callback(r))
        )

    def _guardar(self):
        # 1. Guardar Idioma
        nombre_idioma = self.combo_idioma.get()
        idioma_id = self.map_idiomas_a_id.get(nombre_idioma, "es")
        idioma_anterior = obtener_idioma()
        guardar_idioma(idioma_id)
        fijar_idioma(idioma_id)

        # 2. Guardar Rol (Alumno / Profesor)
        nuevo_rol = self.combo_rol.get()
        if nuevo_rol != self.rol_original:
            ok_rol, msg_rol = actualizar_rol_usuario(nuevo_rol)
            if ok_rol:
                if hasattr(self.master, "sesion"):
                    self.master.sesion["rol"] = nuevo_rol
                if hasattr(self.master, "_al_cambiar_cuenta"):
                    self.master._al_cambiar_cuenta(self.master.sesion)

        # 3. Guardar TTS
        nombre_voz = self.combo_voz.get()
        voz_id = self.map_voces_a_id.get(nombre_voz, "es-ES-AlvaroNeural")
        nombre_vel = self.combo_vel.get()
        vel_id = self.map_vel_a_id.get(nombre_vel, "+0%")
        guardar_ajustes_tts(voz_id, vel_id)

        if idioma_id != idioma_anterior:
            messagebox.showinfo(t("ajustes_titulo"), t("ajustes_aviso_reinicio"))
        else:
            messagebox.showinfo(t("ajustes_titulo"), t("ajustes_guardado_ok"))
        self._cerrar()

    def _solicitar_eliminar_cuenta(self):
        confirm = messagebox.askyesno(
            "⚠️ Eliminar Cuenta",
            "¿Estás seguro de que deseas ELIMINAR tu cuenta de KernossAI?\n\n"
            "• Tu usuario y contraseña se borrarán de Supabase.\n"
            "• Se eliminarán todos tus chats y registros.\n"
            "• Esta acción es IRREVERSIBLE.\n\n"
            "¿Deseas continuar?"
        )
        if not confirm:
            return

        def _thread():
            ok, msg = borrar_cuenta_usuario()
            if ok:
                def _exito():
                    messagebox.showinfo("Cuenta Eliminada", "Tu cuenta ha sido eliminada con éxito del sistema.")
                    self._cerrar()
                    if hasattr(self.master, "_cerrar_sesion"):
                        self.master._cerrar_sesion()
                    else:
                        self.master.destroy()
                self.after(0, _exito)
            else:
                self.after(0, lambda: messagebox.showerror("Error", msg))

        threading.Thread(target=_thread, daemon=True).start()

    def _cerrar(self):
        tts_engine.detener()
        self.destroy()


# ══════════════════════════════════════════════════════════════════════════════
#  MODAL DE NOVEDADES CON RESUMEN INTELIGENTE IA Y LECTOR DE VOZ
# ══════════════════════════════════════════════════════════════════════════════
class VentanaNovedadesIA(ctk.CTkToplevel):
    """Ventana modal interactiva con el resumen de novedades y cambios generado automáticamente por IA."""
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.title(f"✨ Novedades y Cambios de Versión (v{VERSION_APP})")
        self.geometry("720x660")
        self.minsize(600, 520)
        self.configure(fg_color=COLOR_BG_DARK)
        self.transient(parent)
        self.grab_set()

        self._notas_cambios_detectadas = ""
        self._build_ui()

        # Iniciar análisis automático con IA al abrir la ventana
        self.after(300, self._generar_resumen_ia)

    def _build_ui(self):
        # Cabecera
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", padx=25, pady=(20, 10))

        ctk.CTkLabel(header, text=t("modal_novedades_titulo", version=VERSION_APP),
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(anchor="w")
        ctk.CTkLabel(header, text=t("modal_novedades_subtitulo"),
                     font=("Segoe UI", 12), text_color=COLOR_TEXT_MUTED).pack(anchor="w", pady=(2, 0))

        # Panel scrollable con tarjeta de resumen IA y lista de cambios
        self.scroll = ctk.CTkScrollableFrame(self, fg_color=COLOR_BG_CARD,
                                             border_width=1, border_color=COLOR_BORDER,
                                             corner_radius=14)
        self.scroll.pack(fill="both", expand=True, padx=25, pady=(0, 14))

        # Cuadro de Resumen Explicativo Inteligente (IA)
        frame_ia = ctk.CTkFrame(self.scroll, fg_color=COLOR_BG_CARD_LIGHT, corner_radius=10,
                                border_width=1, border_color=COLOR_ACCENT_CYAN)
        frame_ia.pack(fill="x", padx=8, pady=8)

        bar_ia = ctk.CTkFrame(frame_ia, fg_color="transparent")
        bar_ia.pack(fill="x", padx=12, pady=(10, 4))
        
        lbl_ia_title = "🤖 Explicación Inteligente de la IA" if obtener_idioma_activo() == "es" else (
            "🤖 AI Smart Summary" if obtener_idioma_activo() == "en" else (
                "🤖 KI-Zusammenfassung" if obtener_idioma_activo() == "de" else "🤖 Résumé Intelligent IA"
            )
        )
        ctk.CTkLabel(bar_ia, text=lbl_ia_title,
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_ACCENT_CYAN).pack(side="left")

        self.btn_resumir_ia = ctk.CTkButton(
            bar_ia, text=t("btn_consultar_ia"), height=28, width=130,
            font=("Segoe UI", 11, "bold"), fg_color=COLOR_ACCENT_PRIMARY,
            hover_color=COLOR_ACCENT_HOVER, command=self._generar_resumen_ia
        )
        self.btn_resumir_ia.pack(side="right", padx=(6, 0))

        self.btn_tts_novedades = ctk.CTkButton(
            bar_ia, text=t("btn_escuchar"), height=28, width=95,
            font=("Segoe UI", 11, "bold"), fg_color=COLOR_BG_SURFACE,
            border_width=1, border_color=COLOR_BORDER, hover_color=COLOR_ACCENT_HOVER,
            command=self._toggle_tts
        )
        self.btn_tts_novedades.pack(side="right")

        self.txt_resumen_ia = ctk.CTkTextbox(frame_ia, font=("Segoe UI", 12), height=130, wrap="word",
                                             fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        self.txt_resumen_ia.pack(fill="x", padx=12, pady=(4, 12))
        
        texto_inicial = "⏳ Conectando con la IA..." if obtener_idioma_activo() == "es" else (
            "⏳ Connecting to AI..." if obtener_idioma_activo() == "en" else (
                "⏳ Verbindung zur KI..." if obtener_idioma_activo() == "de" else "⏳ Connexion à l'IA..."
            )
        )
        self.txt_resumen_ia.insert("1.0", texto_inicial)
        self.txt_resumen_ia.configure(state="disabled")

        # Lista de versiones y cambios detallados
        lbl_hist_title = "📋 Historial Detallado de Versiones" if obtener_idioma_activo() == "es" else (
            "📋 Version History & Changelog" if obtener_idioma_activo() == "en" else (
                "📋 Versionsverlauf" if obtener_idioma_activo() == "de" else "📋 Historique des Versions"
            )
        )
        ctk.CTkLabel(self.scroll, text=lbl_hist_title,
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=8, pady=(10, 4))

        tag_actual = "Versión Actual" if obtener_idioma_activo() == "es" else (
            "Current Version" if obtener_idioma_activo() == "en" else (
                "Aktuelle Version" if obtener_idioma_activo() == "de" else "Version Actuelle"
            )
        )
        self.versiones_catalogo = [
            (f"v{VERSION_APP} ({tag_actual})", [
                (t("nov_item_soporte_tit"), t("nov_item_soporte_desc")),
                (t("nov_item_multi_tit"), t("nov_item_multi_desc")),
                (t("nov_item_rol_tit"), t("nov_item_rol_desc")),
                (t("nov_item_mod_tit"), t("nov_item_mod_desc")),
                (t("nov_item_cloud_tit"), t("nov_item_cloud_desc")),
                (t("nov_item_hwid_tit"), t("nov_item_hwid_desc"))
            ], COLOR_ACCENT_PRIMARY),
            ("v1.5", [
                ("🌐 Internacionalización & Multi-idioma", "Interfaz traducida al Español (🇪🇸), Inglés (🇬🇧), Alemán (🇩🇪) y Francés (🇫🇷) con selector en Ajustes."),
                ("☁️ Sincronización Cloud Multi-dispositivo", "Tus conversaciones y consultas con la IA se sincronizan entre todos tus ordenadores de forma instantánea."),
                ("🏡 Protección de Hogar Principal de Estudio", "Detección inteligente de red local para proteger tu cuenta y modo de estudio fuera de casa."),
                ("🔊 Lector en Voz Alta con IA (TTS Humano)", "Síntesis de voz natural en resúmenes, tutor, exámenes, notas y ejercicios.")
            ], COLOR_BG_CARD_LIGHT),
            ("v1.4", [
                ("🔊 Motor de Voz Neuronal TTS", "Integración de voces humanas HD para estudio auditivo sin cansar la vista."),
                ("⚙️ Panel de Ajustes y Configuración", "Configuración centralizada de modelos de voz, cadencia y preferencias."),
                ("✨ Botón Permanente de Novedades", "Acceso directo en la barra superior para consultar cambios.")
            ], COLOR_BG_CARD_LIGHT),
            ("v1.3", [
                ("🧠 Generador y Editor de Mapas Mentales", "Crea esquemas conceptuales con IA a partir de cualquier tema y nivel, edítalos y expórtalos a Word (.docx) e imagen (.png)."),
                ("🎨 Estética Cósmica Azul", "Diseño unificado con la nueva web oficial y mayor contraste visual.")
            ], COLOR_BG_CARD_LIGHT),
            ("v1.2", [
                ("💬 Chat Directo de IA con Historial", "Acceso rápido a los modelos de IA desde el inicio con conversaciones guardadas en tu perfil."),
                ("⚡ Barra de Acceso Rápido", "Navegación instantánea entre tus herramientas de estudio en un solo clic.")
            ], COLOR_BG_CARD_LIGHT),
            ("v1.1 / v1.0", [
                ("📚 Suite Académica Base", "Calculador de medias ponderadas, apuntador de notas, resumidor inteligente, generador de exámenes y panel docente.")
            ], COLOR_BG_CARD_LIGHT)
        ]

        for ver_titulo, items, color_bg in self.versiones_catalogo:
            card = ctk.CTkFrame(self.scroll, fg_color=color_bg, corner_radius=10,
                                border_width=1, border_color=COLOR_BORDER)
            card.pack(fill="x", padx=8, pady=5)
            
            ctk.CTkLabel(card, text=ver_titulo, font=("Segoe UI", 13, "bold"),
                         text_color=COLOR_ACCENT_CYAN if tag_actual in ver_titulo else COLOR_TEXT_MUTED).pack(anchor="w", padx=14, pady=(8, 4))

            for titulo_cambio, desc in items:
                ctk.CTkLabel(card, text=f"• {titulo_cambio}: {desc}", font=("Segoe UI", 11),
                             text_color=COLOR_TEXT_MAIN, wraplength=570, justify="left").pack(anchor="w", padx=14, pady=(1, 5))

        # Botón Cerrar
        btn_txt = "Entendido" if obtener_idioma_activo() == "es" else (
            "Got it" if obtener_idioma_activo() == "en" else (
                "Verstanden" if obtener_idioma_activo() == "de" else "Compris"
            )
        )
        btn_cerrar = ctk.CTkButton(self, text=btn_txt, height=38,
                                   font=("Segoe UI", 12, "bold"),
                                   fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                   command=self._cerrar)
        btn_cerrar.pack(fill="x", padx=25, pady=(0, 18))

    def _obtener_notas_cambios(self) -> str:
        """Obtiene las mejoras concretas de la versión actual combinando el catálogo y notas online."""
        items_v_actual = self.versiones_catalogo[0][1]
        lineas_catalogo = [f"• {tit}: {desc}" for tit, desc in items_v_actual]
        notas_base = f"Novedades oficiales de KernossIA v{VERSION_APP}:\n" + "\n".join(lineas_catalogo)

        # Si hay texto adicional en el release de GitHub, anexarlo si no es solo un commit hash
        try:
            res = requests.get("https://api.github.com/repos/LDPCICNM2012/KernossAI/releases/latest", timeout=3)
            if res.status_code == 200:
                data = res.json()
                body = data.get("body", "").strip()
                if body and len(body) > 30 and "What's Changed" not in body:
                    notas_base += f"\n\nNotas adicionales del desarrollador:\n{body}"
        except Exception:
            pass

        return notas_base

    def _generar_resumen_ia(self):
        self.btn_resumir_ia.configure(state="disabled", text="Analizando...")
        self.txt_resumen_ia.configure(state="normal")
        self.txt_resumen_ia.delete("1.0", "end")
        self.txt_resumen_ia.insert("1.0", "🤖 Analizando las mejoras reales de esta versión...")
        self.txt_resumen_ia.configure(state="disabled")
        threading.Thread(target=self._thread_ia_resumen, daemon=True).start()

    def _thread_ia_resumen(self):
        try:
            notas = self._obtener_notas_cambios()
            idioma = obtener_idioma_activo()
            nombres_idioma = {"es": "Español", "en": "English", "de": "Deutsch", "fr": "Français"}
            idioma_str = nombres_idioma.get(idioma, "Español")

            prompt = (
                f"Eres el asistente oficial de KernossAI. El usuario acaba de abrir la app en la versión v{VERSION_APP}.\n\n"
                f"Estas son las novedades EXACTAS y REALES implementadas en esta versión:\n"
                f"'''\n{notas}\n'''\n\n"
                "INSTRUCCIONES OBLIGATORIAS:\n"
                "- NO des respuestas genéricas ni digas 'los detalles están bajo el capó' ni hables de cosas no mencionadas.\n"
                "- Explica de forma concisa y directa (máximo 120-150 palabras) las mejoras listadas arriba (Multi-idioma ES/EN/DE/FR, Sincronización en la nube, Protección de Hogar, y Lector de voz TTS).\n"
                "- Da 1 consejo práctico y breve para estudiar o preparar clases con estas funciones.\n"
                f"- Responde exclusivamente en el idioma: {idioma_str}."
            )
            try:
                resumen = llamar_gemini(prompt)
            except Exception:
                resumen = llamar_groq(prompt)

            def _actualizar():
                self.txt_resumen_ia.configure(state="normal")
                self.txt_resumen_ia.delete("1.0", "end")
                self.txt_resumen_ia.insert("1.0", resumen)
                self.txt_resumen_ia.configure(state="disabled")
                self.btn_resumir_ia.configure(state="normal", text="🔄 Reanalizar")

            self.after(0, _actualizar)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            self.after(0, lambda: self.btn_resumir_ia.configure(state="normal", text="⚡ Analizar con IA"))

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_novedades.configure(text="🔊 Escuchar", fg_color=COLOR_BG_SURFACE)
        else:
            texto = self.txt_resumen_ia.get("1.0", "end-1c").strip()
            if not texto or texto.startswith("🤖") or texto.startswith("⏳"):
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_novedades.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_novedades.configure(text="🔊 Escuchar", fg_color=COLOR_BG_SURFACE)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def _cerrar(self):
        tts_engine.detener()
        self.destroy()


# ══════════════════════════════════════════════════════════════════════════════
#  VENTANA MODAL DE SOPORTE TÉCNICO CON CIFRADO E2EE
# ══════════════════════════════════════════════════════════════════════════════
class VentanaSoporteE2EE(ctk.CTkToplevel):
    """Canal oficial de atención al usuario y soporte técnico con cifrado de extremo a extremo (E2EE)."""
    def __init__(self, parent, sesion):
        super().__init__(parent)
        self.parent = parent
        self.sesion = sesion
        self.email = sesion.get("email", "")
        self.nombre = sesion.get("nombre", "Usuario")
        self.title("🛡️ Soporte Oficial KernossAI — Chat Cifrado E2EE")
        self.geometry("780x680")
        self.minsize(640, 520)
        self.configure(fg_color=COLOR_BG_DARK)
        self.transient(parent)

        self._mensajes = []
        self._polling_activo = True
        self._build_ui()
        self._iniciar_polling()

    def _build_ui(self):
        # Header con candado de seguridad
        header = ctk.CTkFrame(self, fg_color=COLOR_BG_SURFACE, height=75, corner_radius=0)
        header.pack(fill="x")
        
        frame_h_info = ctk.CTkFrame(header, fg_color="transparent")
        frame_h_info.pack(side="left", padx=20, pady=12)
        
        ctk.CTkLabel(frame_h_info, text="🛡️ Soporte Oficial KernossAI (kernossai@support.com)",
                     font=("Segoe UI", 14, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")
        
        badge_sec = ctk.CTkLabel(frame_h_info, text="🔒 Cifrado de Extremo a Extremo (E2EE) Activo • Cero Conocimiento en Servidor",
                                 font=("Segoe UI", 10, "bold"), text_color=COLOR_SUCCESS)
        badge_sec.pack(anchor="w")

        btn_cerrar = ctk.CTkButton(header, text="✕", width=36, height=36, fg_color=COLOR_BG_CARD,
                                   hover_color=COLOR_DANGER, font=("Segoe UI", 14, "bold"), command=self._cerrar)
        btn_cerrar.pack(side="right", padx=16)

        # Scroll de mensajes
        self.scroll_chat = ctk.CTkScrollableFrame(self, fg_color=COLOR_BG_DARK)
        self.scroll_chat.pack(fill="both", expand=True, padx=20, pady=(12, 10))

        # Barra de entrada de texto
        barra_input = ctk.CTkFrame(self, fg_color=COLOR_BG_SURFACE, corner_radius=0)
        barra_input.pack(fill="x", side="bottom")

        # Aviso en pequeño de función beta
        aviso_beta = ctk.CTkFrame(barra_input, fg_color="transparent")
        aviso_beta.pack(fill="x", padx=18, pady=(8, 2))

        ctk.CTkLabel(
            aviso_beta,
            text="🧪 Función beta • Cualquier error, comuníquelo en GitHub",
            font=("Segoe UI", 9, "bold"),
            text_color=COLOR_TEXT_MUTED
        ).pack(side="left")

        # Fila de selección de Motivo de la consulta
        fila_motivo = ctk.CTkFrame(barra_input, fg_color="transparent")
        fila_motivo.pack(fill="x", padx=18, pady=(2, 2))

        ctk.CTkLabel(fila_motivo, text="📌 Motivo:", font=("Segoe UI", 10, "bold"), text_color=COLOR_TEXT_MAIN).pack(side="left", padx=(0, 8))

        self.combo_motivo = ctk.CTkComboBox(
            fila_motivo,
            values=[
                "Duda Académica / Tareas",
                "Problema con la IA",
                "Error o Bug del Programa",
                "Problema con mi Cuenta",
                "Sugerencia / Idea",
                "Consulta General"
            ],
            height=28, width=240, font=("Segoe UI", 10), fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER
        )
        self.combo_motivo.set("Duda Académica / Tareas")
        self.combo_motivo.pack(side="left")

        # Fila del input de texto y botón enviar
        fila_entrada = ctk.CTkFrame(barra_input, fg_color="transparent")
        fila_entrada.pack(fill="x", padx=18, pady=(4, 12))

        self.entry_msg = ctk.CTkEntry(
            fila_entrada, font=("Segoe UI", 12), height=42,
            placeholder_text="Escribe tu mensaje a soporte (se cifrará antes de enviarse)...",
            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER
        )
        self.entry_msg.pack(side="left", fill="x", expand=True, padx=(0, 10))
        self.entry_msg.bind("<Return>", lambda e: self._enviar())

        self.btn_enviar = ctk.CTkButton(
            fila_entrada, text="Enviar 📤", width=110, height=42,
            font=("Segoe UI", 12, "bold"), fg_color=COLOR_ACCENT_PRIMARY,
            hover_color=COLOR_ACCENT_HOVER, command=self._enviar
        )
        self.btn_enviar.pack(side="right")

    def _iniciar_polling(self):
        def _loop():
            while self._polling_activo:
                try:
                    ok, msgs = obtener_mensajes_soporte()
                    if ok:
                        if len(msgs) != len(self._mensajes):
                            self._mensajes = msgs
                            self.after(0, self._renderizar_mensajes)
                except Exception:
                    pass
                threading.Event().wait(3.0)

        threading.Thread(target=_loop, daemon=True).start()

    def _renderizar_mensajes(self):
        for w in self.scroll_chat.winfo_children():
            w.destroy()

        if not self._mensajes:
            card_empty = ctk.CTkFrame(self.scroll_chat, fg_color=COLOR_BG_CARD, corner_radius=12, border_width=1, border_color=COLOR_BORDER)
            card_empty.pack(fill="x", padx=10, pady=20)
            ctk.CTkLabel(card_empty, text="👋 ¡Bienvenido al Canal Directo de Soporte Oficial!", font=("Segoe UI", 14, "bold"), text_color=COLOR_TEXT_MAIN).pack(padx=20, pady=(15, 4))
            ctk.CTkLabel(card_empty, text="¿Tienes alguna duda académica, sugerencia o problema técnico?\nEscribe tu mensaje a continuación. Tu consulta viajará 100% cifrada hacia nuestro equipo.",
                         font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED, justify="center").pack(padx=20, pady=(0, 15))
            return

        for m in self._mensajes:
            es_mio = m.get("es_mio", False)
            emisor = "Tú" if es_mio else m.get("emisor_nombre", "Soporte KernossAI")
            hora = m.get("timestamp", "")[11:16] if m.get("timestamp") else ""
            texto = m.get("texto", "")

            frame_burbuja = ctk.CTkFrame(
                self.scroll_chat,
                fg_color=COLOR_ACCENT_PRIMARY if es_mio else COLOR_BG_SURFACE,
                corner_radius=12,
                border_width=0 if es_mio else 1,
                border_color=COLOR_BORDER
            )
            side_align = "e" if es_mio else "w"
            frame_burbuja.pack(anchor=side_align, padx=10, pady=4)

            hdr_txt = f"{emisor} • {hora}" if not es_mio else f"Tú • {hora}"
            ctk.CTkLabel(frame_burbuja, text=hdr_txt, font=("Segoe UI", 9, "bold"),
                         text_color="#93c5fd" if es_mio else COLOR_TEXT_MUTED).pack(anchor="w", padx=12, pady=(6, 1))

            ctk.CTkLabel(frame_burbuja, text=texto, font=("Segoe UI", 11),
                         text_color="#ffffff" if es_mio else COLOR_TEXT_MAIN,
                         wraplength=480, justify="left").pack(anchor="w", padx=12, pady=(0, 8))

        self.scroll_chat._parent_canvas.yview_moveto(1.0)

    def _enviar(self):
        txt = self.entry_msg.get().strip()
        if not txt:
            return
        
        motivo = self.combo_motivo.get().strip()
        msg_con_motivo = f"📌 [Motivo: {motivo}]\n{txt}" if motivo else txt

        self.entry_msg.delete(0, "end")
        self.btn_enviar.configure(state="disabled")

        def _thread():
            ok, res = enviar_mensaje_soporte(msg_con_motivo)
            if ok:
                ok2, msgs = obtener_mensajes_soporte()
                if ok2:
                    self._mensajes = msgs
                    self.after(0, self._renderizar_mensajes)
            else:
                self.after(0, lambda: messagebox.showerror("Error Soporte", res))
            self.after(0, lambda: self.btn_enviar.configure(state="normal"))

        threading.Thread(target=_thread, daemon=True).start()

    def _cerrar(self):
        self._polling_activo = False
        self.destroy()


# ══════════════════════════════════════════════════════════════════════════════
#  VENTANA MODAL DE ADMINISTRACIÓN, HARDWARE-BAN & VISUALIZADOR DE RENDER
# ══════════════════════════════════════════════════════════════════════════════
class VentanaAdminModeracion(ctk.CTkToplevel):
    """Panel de Control de Moderación, Hardware-Ban, IP-Ban y visualizador de Base de Datos Cifrada en Render."""
    def __init__(self, parent, sesion):
        super().__init__(parent)
        self.parent = parent
        self.sesion = sesion
        self.title("👑 Panel Maestro de Administración & Moderación (Render)")
        self.geometry("980x740")
        self.minsize(800, 600)
        self.configure(fg_color=COLOR_BG_DARK)
        self.transient(parent)

        self._build_ui()
        self._cargar_usuarios()

    def _build_ui(self):
        header = ctk.CTkFrame(self, fg_color=COLOR_BG_SURFACE, height=65, corner_radius=0)
        header.pack(fill="x")
        
        ctk.CTkLabel(header, text="👑 Panel Maestro de Administración & Moderación",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).pack(side="left", padx=20, pady=16)

        btn_refrescar = ctk.CTkButton(header, text="🔄 Recargar Servidor", width=140, height=32,
                                      font=("Segoe UI", 11, "bold"), fg_color=COLOR_ACCENT_PRIMARY,
                                      hover_color=COLOR_ACCENT_HOVER, command=self._recargar_todo)
        btn_refrescar.pack(side="right", padx=(10, 20))

        self.tabview = ctk.CTkTabview(self, fg_color=COLOR_BG_DARK)
        self.tabview.pack(fill="both", expand=True, padx=20, pady=10)

        self.tab_soporte_inbox = self.tabview.add("📬 Bandeja de Soporte")
        self.tab_usuarios = self.tabview.add("👥 Usuarios & Bans")
        self.tab_db_raw = self.tabview.add("🔒 Ver Base de Datos Cifrada (Render)")
        self.tab_ban_manual = self.tabview.add("🚫 Aplicar Ban Manual (IP / HWID)")

        self._build_tab_soporte_inbox()
        self._build_tab_usuarios()
        self._build_tab_db_raw()
        self._build_tab_ban_manual()

    def _recargar_todo(self):
        self._cargar_usuarios()
        self._cargar_tickets_soporte()

    def _build_tab_soporte_inbox(self):
        # Frame contenedor con dos columnas: Lista de Tickets a la izquierda y Conversación a la derecha
        self.frame_inbox_split = ctk.CTkFrame(self.tab_soporte_inbox, fg_color="transparent")
        self.frame_inbox_split.pack(fill="both", expand=True, padx=4, pady=4)
        self.frame_inbox_split.grid_columnconfigure(0, weight=4)
        self.frame_inbox_split.grid_columnconfigure(1, weight=6)
        self.frame_inbox_split.grid_rowconfigure(0, weight=1)

        # Columna Izquierda: Lista de Tickets
        col_izq = ctk.CTkFrame(self.frame_inbox_split, fg_color=COLOR_BG_SURFACE, corner_radius=12, border_width=1, border_color=COLOR_BORDER)
        col_izq.grid(row=0, column=0, sticky="nsew", padx=(0, 6), pady=4)
        
        ctk.CTkLabel(col_izq, text="📬 Mensajes Recibidos de Alumnos", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=14, pady=(12, 6))

        self.scroll_tickets = ctk.CTkScrollableFrame(col_izq, fg_color="transparent")
        self.scroll_tickets.pack(fill="both", expand=True, padx=6, pady=(0, 8))

        # Columna Derecha: Vista del Chat y Respuesta Oficial
        self.col_der = ctk.CTkFrame(self.frame_inbox_split, fg_color=COLOR_BG_SURFACE, corner_radius=12, border_width=1, border_color=COLOR_BORDER)
        self.col_der.grid(row=0, column=1, sticky="nsew", padx=(6, 0), pady=4)
        self.col_der.grid_rowconfigure(1, weight=1)
        self.col_der.grid_columnconfigure(0, weight=1)

        # Header del chat seleccionado
        self.hdr_chat_soporte = ctk.CTkFrame(self.col_der, fg_color=COLOR_BG_CARD, height=55, corner_radius=8)
        self.hdr_chat_soporte.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 6))
        
        self.lbl_chat_usr_info = ctk.CTkLabel(self.hdr_chat_soporte, text="Selecciona un ticket a la izquierda para responder",
                                              font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MUTED)
        self.lbl_chat_usr_info.pack(side="left", padx=14, pady=12)

        self.btn_borrar_chat = ctk.CTkButton(
            self.hdr_chat_soporte, text="🗑️ Vaciar Chat", width=105, height=30,
            font=("Segoe UI", 10, "bold"), fg_color="#7f1d1d", hover_color="#991b1b",
            border_width=1, border_color=COLOR_DANGER, command=self._borrar_conversacion_ticket
        )

        # Scroll de mensajes de la conversación
        self.scroll_mensajes_soporte = ctk.CTkScrollableFrame(self.col_der, fg_color=COLOR_BG_DARK)
        self.scroll_mensajes_soporte.grid(row=1, column=0, sticky="nsew", padx=10, pady=4)

        # Barra de respuesta oficial
        barra_resp = ctk.CTkFrame(self.col_der, fg_color="transparent")
        barra_resp.grid(row=2, column=0, sticky="ew", padx=10, pady=(6, 12))

        self.entry_resp_soporte = ctk.CTkEntry(
            barra_resp, font=("Segoe UI", 11), height=38,
            placeholder_text="Escribe la respuesta oficial al alumno...",
            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER
        )
        self.entry_resp_soporte.pack(side="left", fill="x", expand=True, padx=(0, 8))
        self.entry_resp_soporte.bind("<Return>", lambda e: self._enviar_respuesta_soporte())

        self.btn_enviar_resp = ctk.CTkButton(
            barra_resp, text="Responder 📤", width=115, height=38,
            font=("Segoe UI", 11, "bold"), fg_color=COLOR_SUCCESS,
            hover_color="#15803d", command=self._enviar_respuesta_soporte,
            state="disabled"
        )
        self.btn_enviar_resp.pack(side="right")

        self._ticket_activo_email = None
        self._cargar_tickets_soporte()

    def _cargar_tickets_soporte(self):
        for w in self.scroll_tickets.winfo_children():
            w.destroy()

        def _thread():
            ok, tickets = admin_obtener_tickets_soporte()
            if not ok:
                return

            def _render():
                if not tickets:
                    ctk.CTkLabel(self.scroll_tickets, text="No hay tickets de soporte aún.",
                                 font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(pady=30)
                    return

                for t in tickets:
                    u_email = t.get("usuario_email", "")
                    u_nombre = t.get("usuario_nombre", u_email)
                    u_rol = t.get("usuario_rol", "Alumno")
                    u_fecha = t.get("ultimo_timestamp", "")[:16].replace("T", " ") if t.get("ultimo_timestamp") else ""
                    ultimo_txt = t.get("ultimo_texto", "")

                    card_t = ctk.CTkFrame(self.scroll_tickets, fg_color=COLOR_BG_CARD, corner_radius=8, border_width=1, border_color=COLOR_BORDER)
                    card_t.pack(fill="x", padx=2, pady=3)

                    f_info = ctk.CTkFrame(card_t, fg_color="transparent")
                    f_info.pack(fill="x", padx=10, pady=8)

                    ctk.CTkLabel(f_info, text=f"👤 {u_nombre} ({u_rol})", font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")
                    ctk.CTkLabel(f_info, text=f"✉️ {u_email} • 🕒 {u_fecha}", font=("Segoe UI", 9), text_color=COLOR_ACCENT_CYAN).pack(anchor="w", pady=(1, 3))
                    ctk.CTkLabel(f_info, text=f"💬 \"{ultimo_txt[:60]}...\"", font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED, wraplength=260, justify="left").pack(anchor="w")

                    btn_abrir = ctk.CTkButton(
                        card_t, text="💬 Abrir Ticket", height=26, font=("Segoe UI", 10, "bold"),
                        fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                        command=lambda u=u_email, n=u_nombre, r=u_rol, msgs=t.get("mensajes", []): self._seleccionar_ticket(u, n, r, msgs)
                    )
                    btn_abrir.pack(fill="x", padx=10, pady=(0, 8))

            self.after(0, _render)

        threading.Thread(target=_thread, daemon=True).start()

    def _seleccionar_ticket(self, u_email: str, u_nombre: str, u_rol: str, mensajes: list):
        self._ticket_activo_email = u_email
        self.lbl_chat_usr_info.configure(text=f"🛡️ Conversación con: {u_nombre} ({u_email}) — {u_rol}", text_color=COLOR_TEXT_MAIN)
        self.btn_borrar_chat.pack(side="right", padx=10, pady=10)
        self.btn_enviar_resp.configure(state="normal")
        self.entry_resp_soporte.focus_set()

        for w in self.scroll_mensajes_soporte.winfo_children():
            w.destroy()

        for m in mensajes:
            es_soporte = m.get("es_soporte", False)
            emisor = "🛡️ Soporte Oficial" if es_soporte else f"👤 {m.get('emisor_nombre', u_nombre)}"
            hora = m.get("timestamp", "")[:19].replace("T", " ") if m.get("timestamp") else ""
            texto = m.get("texto", "")

            frame_b = ctk.CTkFrame(
                self.scroll_mensajes_soporte,
                fg_color="#065f46" if es_soporte else COLOR_BG_CARD,
                corner_radius=10,
                border_width=1,
                border_color="#10b981" if es_soporte else COLOR_BORDER
            )
            frame_b.pack(anchor="e" if es_soporte else "w", padx=8, pady=4)

            hdr = f"{emisor} • 🕒 {hora}"
            ctk.CTkLabel(frame_b, text=hdr, font=("Segoe UI", 9, "bold"),
                         text_color="#6ee7b7" if es_soporte else COLOR_TEXT_MUTED).pack(anchor="w", padx=10, pady=(6, 1))

            ctk.CTkLabel(frame_b, text=texto, font=("Segoe UI", 11),
                         text_color="#ffffff" if es_soporte else COLOR_TEXT_MAIN,
                         wraplength=420, justify="left").pack(anchor="w", padx=10, pady=(0, 8))

        self.scroll_mensajes_soporte._parent_canvas.yview_moveto(1.0)

    def _borrar_conversacion_ticket(self):
        if not self._ticket_activo_email:
            return
        
        usr = self._ticket_activo_email
        confirm = messagebox.askyesno(
            "🗑️ Vaciar Conversación",
            f"¿Estás seguro de que deseas ELIMINAR todo el historial de soporte con '{usr}'?\n\n"
            "• Se borrarán todos los mensajes cifrados de este ticket en Supabase.\n"
            "• Esta acción es permanente (estilo WhatsApp).\n\n"
            "¿Deseas continuar?"
        )
        if not confirm:
            return
        
        def _thread():
            ok, msg = admin_borrar_ticket_soporte(usr)
            if ok:
                def _exito():
                    self._ticket_activo_email = None
                    self.lbl_chat_usr_info.configure(text="Selecciona un ticket a la izquierda para responder", text_color=COLOR_TEXT_MUTED)
                    self.btn_enviar_resp.configure(state="disabled")
                    self.btn_borrar_chat.pack_forget()
                    for w in self.scroll_mensajes_soporte.winfo_children():
                        w.destroy()
                    self._cargar_tickets_soporte()
                    messagebox.showinfo("Conversación Eliminada", msg)
                self.after(0, _exito)
            else:
                self.after(0, lambda: messagebox.showerror("Error al borrar", msg))
        
        threading.Thread(target=_thread, daemon=True).start()

    def _enviar_respuesta_soporte(self):
        if not self._ticket_activo_email:
            return
        texto = self.entry_resp_soporte.get().strip()
        if not texto:
            return
        
        self.entry_resp_soporte.delete(0, "end")
        self.btn_enviar_resp.configure(state="disabled")

        def _thread():
            ok, res = admin_responder_ticket_soporte(self._ticket_activo_email, texto)
            if ok:
                # Recargar conversación
                ok2, tickets = admin_obtener_tickets_soporte()
                if ok2:
                    for t in tickets:
                        if t.get("usuario_email", "").lower() == self._ticket_activo_email.lower():
                            self.after(0, lambda: self._seleccionar_ticket(
                                self._ticket_activo_email, t.get("usuario_nombre", ""), t.get("usuario_rol", ""), t.get("mensajes", [])
                            ))
                            break
                    self.after(0, self._cargar_tickets_soporte)
            else:
                self.after(0, lambda: messagebox.showerror("Error Soporte", res))
            self.after(0, lambda: self.btn_enviar_resp.configure(state="normal"))

        threading.Thread(target=_thread, daemon=True).start()

    def _build_tab_usuarios(self):
        # Barra superior de búsqueda rápida
        frame_busqueda = ctk.CTkFrame(self.tab_usuarios, fg_color="transparent")
        frame_busqueda.pack(fill="x", padx=6, pady=(6, 8))

        self.entry_buscar_usr = ctk.CTkEntry(
            frame_busqueda, height=36, font=("Segoe UI", 11),
            placeholder_text="🔍 Buscar usuario por nombre, email, rol, IP o HWID...",
            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER
        )
        self.entry_buscar_usr.pack(fill="x", expand=True)
        self.entry_buscar_usr.bind("<KeyRelease>", lambda e: self._filtrar_usuarios())

        self.scroll_usr = ctk.CTkScrollableFrame(self.tab_usuarios, fg_color=COLOR_BG_DARK)
        self.scroll_usr.pack(fill="both", expand=True, padx=4, pady=4)
        self._lista_usuarios_cache = []

    def _build_tab_db_raw(self):
        frame_top = ctk.CTkFrame(self.tab_db_raw, fg_color="transparent")
        frame_top.pack(fill="x", padx=10, pady=(10, 6))

        ctk.CTkLabel(frame_top, text="🔍 Contenido en bruto de 'mensajes' almacenado en el servidor:",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(side="left")

        btn_ver_raw = ctk.CTkButton(frame_top, text="📡 Consultar Raw", width=160, height=30,
                                    font=("Segoe UI", 11, "bold"), fg_color=COLOR_ACCENT_CYAN,
                                    text_color="#000", hover_color="#38bdf8", command=self._cargar_raw_db)
        btn_ver_raw.pack(side="right")

        self.txt_raw = ctk.CTkTextbox(self.tab_db_raw, font=("Consolas", 11), wrap="word",
                                      fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        self.txt_raw.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        self.txt_raw.insert("1.0", "Haz clic en 'Consultar Raw' para comprobar en vivo que los mensajes están cifrados en el servidor.")

    def _build_tab_ban_manual(self):
        frame = ctk.CTkFrame(self.tab_ban_manual, fg_color=COLOR_BG_SURFACE, corner_radius=12, border_width=1, border_color=COLOR_BORDER)
        frame.pack(fill="x", padx=30, pady=25)

        ctk.CTkLabel(frame, text="🚫 Bloqueo Manual de Usuarios, IP o Dispositivo Físico (HWID)",
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=20, pady=(15, 10))

        ctk.CTkLabel(frame, text="Tipo de Sanción:", font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=20)
        self.combo_tipo_ban = ctk.CTkComboBox(frame, values=["usuario", "ip", "hwid"], font=("Segoe UI", 11))
        self.combo_tipo_ban.pack(fill="x", padx=20, pady=(2, 10))

        ctk.CTkLabel(frame, text="Objetivo (Email / Dirección IP / HWID Hexadecimal):", font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=20)
        self.entry_ban_obj = ctk.CTkEntry(frame, font=("Segoe UI", 11), placeholder_text="ej: usuario@correo.com o 192.168.1.1 o 657ed7dbd298...")
        self.entry_ban_obj.pack(fill="x", padx=20, pady=(2, 10))

        ctk.CTkLabel(frame, text="Motivo del Baneo:", font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=20)
        self.entry_ban_motivo = ctk.CTkEntry(frame, font=("Segoe UI", 11), placeholder_text="ej: Uso indebido del sistema o multicuentas")
        self.entry_ban_motivo.pack(fill="x", padx=20, pady=(2, 15))

        btn_ejecutar_ban = ctk.CTkButton(frame, text="⛔ Ejecutar Baneo Inmediato", height=38,
                                         font=("Segoe UI", 12, "bold"), fg_color=COLOR_DANGER,
                                         hover_color="#b91c1c", command=self._aplicar_ban_manual)
        btn_ejecutar_ban.pack(fill="x", padx=20, pady=(0, 20))

    def _cargar_usuarios(self):
        def _thread():
            ok, usuarios = admin_listar_usuarios()
            if not ok:
                self.after(0, lambda: messagebox.showerror("Error", "No se pudo conectar con el servidor."))
                return

            def _render():
                self._lista_usuarios_cache = usuarios or []
                self._filtrar_usuarios()

            self.after(0, _render)

        threading.Thread(target=_thread, daemon=True).start()

    def _filtrar_usuarios(self):
        for w in self.scroll_usr.winfo_children():
            w.destroy()

        filtro = self.entry_buscar_usr.get().strip().lower()
        usuarios = [
            u for u in self._lista_usuarios_cache
            if not filtro or
            filtro in u.get("email", "").lower() or
            filtro in u.get("nombre", "").lower() or
            filtro in u.get("rol", "").lower() or
            filtro in u.get("ip", "").lower() or
            filtro in u.get("hwid", "").lower()
        ]

        if not usuarios:
            msg = "🔍 No se encontraron usuarios que coincidan con la búsqueda." if filtro else "No hay usuarios registrados aún."
            ctk.CTkLabel(self.scroll_usr, text=msg, font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(pady=30)
            return

        for u in usuarios:
            email = u.get("email", "")
            nombre = u.get("nombre", "")
            rol = u.get("rol", "Alumno")
            ip = u.get("ip", "N/D")
            hwid = u.get("hwid", "N/D")
            baneado = u.get("baneado", False)
            ip_baneada = u.get("ip_baneada", False)
            hwid_baneado = u.get("hwid_baneado", False)

            card = ctk.CTkFrame(self.scroll_usr, fg_color=COLOR_BG_CARD, corner_radius=10, border_width=1,
                                border_color=COLOR_DANGER if (baneado or ip_baneada or hwid_baneado) else COLOR_BORDER)
            card.pack(fill="x", padx=6, pady=4)

            f_info = ctk.CTkFrame(card, fg_color="transparent")
            f_info.pack(side="left", padx=12, pady=10, fill="x", expand=True)

            status_str = "⛔ BANEADO" if baneado else "🟢 ACTIVO"
            ctk.CTkLabel(f_info, text=f"{nombre} ({email}) — {rol} [{status_str}]",
                         font=("Segoe UI", 12, "bold"),
                         text_color=COLOR_DANGER if baneado else COLOR_TEXT_MAIN).pack(anchor="w")

            det = f"🌐 IP: {ip} {'(IP-BANEADA)' if ip_baneada else ''}  |  💻 HWID: {hwid[:16]}... {'(HWID-BANEADO)' if hwid_baneado else ''}"
            ctk.CTkLabel(f_info, text=det, font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED).pack(anchor="w", pady=(2, 0))

            f_actions = ctk.CTkFrame(card, fg_color="transparent")
            f_actions.pack(side="right", padx=10, pady=10)

            if not baneado:
                ctk.CTkButton(f_actions, text="🚫 Ban Cuenta", width=95, height=28,
                              font=("Segoe UI", 10, "bold"), fg_color=COLOR_DANGER, hover_color="#b91c1c",
                              command=lambda em=email: self._ban_quick(em, "usuario")).pack(side="left", padx=2)
            else:
                ctk.CTkButton(f_actions, text="✅ Desban Cuenta", width=110, height=28,
                              font=("Segoe UI", 10, "bold"), fg_color=COLOR_SUCCESS, hover_color="#15803d",
                              command=lambda em=email: self._desban_quick(em, "usuario")).pack(side="left", padx=2)

            if not ip_baneada and ip != "N/D":
                ctk.CTkButton(f_actions, text="🌐 IP-Ban", width=75, height=28,
                              font=("Segoe UI", 10, "bold"), fg_color="#7c3aed", hover_color="#6d28d9",
                              command=lambda ip_val=ip: self._ban_quick(ip_val, "ip")).pack(side="left", padx=2)
            elif ip_baneada:
                ctk.CTkButton(f_actions, text="✅ Desban IP", width=85, height=28,
                              font=("Segoe UI", 10, "bold"), fg_color=COLOR_SUCCESS, hover_color="#15803d",
                              command=lambda ip_val=ip: self._desban_quick(ip_val, "ip")).pack(side="left", padx=2)

            if not hwid_baneado and hwid != "N/D" and hwid != "HWID_NO_REPORTADO":
                ctk.CTkButton(f_actions, text="💻 HW-Ban", width=80, height=28,
                              font=("Segoe UI", 10, "bold"), fg_color="#b45309", hover_color="#92400e",
                              command=lambda hw=hwid: self._ban_quick(hw, "hwid")).pack(side="left", padx=2)
            elif hwid_baneado:
                ctk.CTkButton(f_actions, text="✅ Desban HW", width=90, height=28,
                              font=("Segoe UI", 10, "bold"), fg_color=COLOR_SUCCESS, hover_color="#15803d",
                              command=lambda hw=hwid: self._desban_quick(hw, "hwid")).pack(side="left", padx=2)

    def _ban_quick(self, objetivo, tipo):
        dialogo = ctk.CTkInputDialog(
            text=f"Introduce el motivo de la sanción para {tipo} '{objetivo}':\n(Se mostrará al usuario en su pantalla roja de bloqueo)",
            title=f"⛔ Motivo del Baneo de {tipo.upper()}"
        )
        motivo = dialogo.get_input()
        if motivo is None:
            return  # Cancelado por el admin

        motivo_limpio = motivo.strip() or "Infracción de las normas del sistema"
        ok, res = admin_aplicar_ban(objetivo, tipo, motivo_limpio)
        if ok:
            messagebox.showinfo("Baneo Aplicado", f"{res}\n\n📋 Motivo: {motivo_limpio}")
            self._cargar_usuarios()
        else:
            messagebox.showerror("Error", res)

    def _desban_quick(self, objetivo, tipo):
        ok, res = admin_desbanear(objetivo, tipo)
        if ok:
            messagebox.showinfo("Desbaneo Aplicado", res)
            self._cargar_usuarios()
        else:
            messagebox.showerror("Error", res)

    def _aplicar_ban_manual(self):
        tipo = self.combo_tipo_ban.get()
        obj = self.entry_ban_obj.get().strip()
        motivo = self.entry_ban_motivo.get().strip() or "Infracción de normas"
        if not obj:
            messagebox.showwarning("Atención", "Escribe el objetivo a banear.")
            return

        ok, res = admin_aplicar_ban(obj, tipo, motivo)
        if ok:
            messagebox.showinfo("Baneo Exitoso", res)
            self.entry_ban_obj.delete(0, "end")
            self._cargar_usuarios()
        else:
            messagebox.showerror("Error", res)

    def _cargar_raw_db(self):
        self.txt_raw.delete("1.0", "end")
        self.txt_raw.insert("1.0", "⏳ Consultando base de datos en Render...")

        def _thread():
            ok, res = admin_ver_mensajes_raw()
            if ok:
                txt = json.dumps(res, ensure_ascii=False, indent=2)
                self.after(0, lambda: (self.txt_raw.delete("1.0", "end"), self.txt_raw.insert("1.0", txt)))
            else:
                self.after(0, lambda: (self.txt_raw.delete("1.0", "end"), self.txt_raw.insert("1.0", "Error al consultar Render.")))

        threading.Thread(target=_thread, daemon=True).start()


# ══════════════════════════════════════════════════════════════════════════════
#  VENTANA MODAL DE VINCULACIÓN Y TUTORÍA (ALUMNO <-> PROFESOR)
# ══════════════════════════════════════════════════════════════════════════════
class VentanaTutoriaAlumnoProfesor(ctk.CTkToplevel):
    """
    Sistema integral de tutoría académica que permite:
    1. Alumno: Explorar la lista de profesores registrados, solicitar vinculación y chatear con su tutor.
    2. Profesor: Gestionar solicitudes pendientes de alumnos y resolver dudas mediante chat individual.
    """
    def __init__(self, parent, sesion):
        super().__init__(parent)
        self.parent = parent
        self.sesion = sesion
        self.rol = sesion.get("rol", "Alumno")
        self.email = sesion.get("email", "").strip().lower()
        self.nombre = sesion.get("nombre", self.email)
        
        titulo = t("tutoria_modal_titulo_alumno") if self.rol == "Alumno" else t("tutoria_modal_titulo_profesor")
        self.title(titulo)
        self.geometry("920x720")
        self.minsize(780, 560)
        self.configure(fg_color=COLOR_BG_DARK)
        self.transient(parent)

        self._polling_activo = True
        self._profesores = []
        self._estado_vinculacion = None  # Para el alumno
        self._solicitudes_pendientes = []  # Para el profesor
        self._alumnos_vinculados = []      # Para el profesor
        self._alumno_activo_chat = None    # Para el profesor (chat activo)
        self._mensajes_chat = []

        self._build_ui()
        self._iniciar_carga()

    def _build_ui(self):
        # Header común
        header = ctk.CTkFrame(self, fg_color=COLOR_BG_SURFACE, height=75, corner_radius=0)
        header.pack(fill="x")

        frame_h_info = ctk.CTkFrame(header, fg_color="transparent")
        frame_h_info.pack(side="left", padx=20, pady=12)

        tit_header = "👨‍🏫 Mi Profesor Tutor & Consultas Académicas" if self.rol == "Alumno" else "🎓 Panel Docente de Tutoría & Alumnos"
        ctk.CTkLabel(frame_h_info, text=tit_header, font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")

        ctk.CTkLabel(frame_h_info, text="🔒 Canal de comunicación directo y cifrado E2EE • Consultas y resolución de dudas",
                     font=("Segoe UI", 10, "bold"), text_color=COLOR_ACCENT_CYAN).pack(anchor="w")

        btn_cerrar = ctk.CTkButton(header, text="✕", width=36, height=36, fg_color=COLOR_BG_CARD,
                                   hover_color=COLOR_DANGER, font=("Segoe UI", 14, "bold"), command=self._cerrar)
        btn_cerrar.pack(side="right", padx=16)

        # Contenedor central según el ROL
        self.contenedor_principal = ctk.CTkFrame(self, fg_color="transparent")
        self.contenedor_principal.pack(fill="both", expand=True, padx=15, pady=12)

        if self.rol == "Alumno":
            self._build_ui_alumno()
        else:
            self._build_ui_profesor()

    # ─────────────────────────────────────────────────────────
    #  VISTA DEL ALUMNO
    # ─────────────────────────────────────────────────────────
    def _build_ui_alumno(self):
        self.frame_alumno_dinamico = ctk.CTkFrame(self.contenedor_principal, fg_color="transparent")
        self.frame_alumno_dinamico.pack(fill="both", expand=True)

    def _renderizar_alumno_estado(self):
        for w in self.frame_alumno_dinamico.winfo_children():
            w.destroy()

        vinc = self._estado_vinculacion
        estado = vinc.get("estado") if vinc else None

        # CASO 1: Sin vinculación ni solicitud activa -> Mostrar listado de profesores para elegir
        if not vinc or estado == "rechazada":
            self._render_alumno_selector_profesores()
        # CASO 2: Solicitud enviada pendiente de aceptación
        elif estado == "pendiente":
            self._render_alumno_solicitud_pendiente(vinc)
        # CASO 3: Vinculación aceptada -> Mostrar Chat Directo con el Tutor
        elif estado == "aceptada":
            self._render_alumno_chat_tutor(vinc)

    def _render_alumno_selector_profesores(self):
        card_info = ctk.CTkFrame(self.frame_alumno_dinamico, fg_color=COLOR_BG_CARD, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
        card_info.pack(fill="x", pady=(0, 10))
        ctk.CTkLabel(card_info, text="👨‍🏫 Vinculación con un Profesor Tutor", font=("Segoe UI", 13, "bold"), text_color=COLOR_ACCENT_SKY).pack(anchor="w", padx=14, pady=(10, 2))
        ctk.CTkLabel(card_info, text="Selecciona a continuación al profesor con el que deseas vincularte para enviarle preguntas, tareas y dudas académicas directamente.",
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=14, pady=(0, 10))

        frame_busq = ctk.CTkFrame(self.frame_alumno_dinamico, fg_color="transparent")
        frame_busq.pack(fill="x", pady=(0, 8))
        self.entry_busq_profe = ctk.CTkEntry(
            frame_busq, height=36, font=("Segoe UI", 11),
            placeholder_text="🔍 Buscar profesor por nombre o correo electrónico...",
            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER
        )
        self.entry_busq_profe.pack(side="left", fill="x", expand=True, padx=(0, 8))
        self.entry_busq_profe.bind("<KeyRelease>", lambda e: self._filtrar_profesores_ui())

        ctk.CTkButton(
            frame_busq, text="🔄 Actualizar", width=95, height=36,
            font=("Segoe UI", 11, "bold"), fg_color=COLOR_BG_SURFACE,
            hover_color=COLOR_ACCENT_HOVER, command=self._cargar_profesores_alumno
        ).pack(side="right")

        self.scroll_profes = ctk.CTkScrollableFrame(self.frame_alumno_dinamico, fg_color=COLOR_BG_SURFACE, corner_radius=10)
        self.scroll_profes.pack(fill="both", expand=True)
        self._filtrar_profesores_ui()

    def _filtrar_profesores_ui(self):
        for w in self.scroll_profes.winfo_children():
            w.destroy()

        filtro = self.entry_busq_profe.get().strip().lower() if hasattr(self, "entry_busq_profe") else ""
        profes_filtrados = [
            p for p in self._profesores
            if filtro in p.get("nombre", "").lower() or filtro in p.get("email", "").lower()
        ]

        if not profes_filtrados:
            ctk.CTkLabel(self.scroll_profes, text=t("tutoria_sin_profes"),
                         font=("Segoe UI", 12), text_color=COLOR_TEXT_MUTED).pack(pady=40)
            return

        for p in profes_filtrados:
            p_email = p.get("email", "").strip()
            p_nombre = p.get("nombre", p_email)

            card = ctk.CTkFrame(self.scroll_profes, fg_color=COLOR_BG_CARD, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
            card.pack(fill="x", padx=6, pady=4)

            f_left = ctk.CTkFrame(card, fg_color="transparent")
            f_left.pack(side="left", fill="x", expand=True, padx=14, pady=10)

            ctk.CTkLabel(f_left, text=f"👨‍🏫 {p_nombre}", font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")
            ctk.CTkLabel(f_left, text=f"✉️ {p_email}   •   🎓 Rol: Docente Verificado", font=("Segoe UI", 10), text_color=COLOR_ACCENT_CYAN).pack(anchor="w", pady=(2, 0))

            btn_solicitar = ctk.CTkButton(
                card, text="📨 Solicitar Vinculación", height=32, width=170,
                font=("Segoe UI", 11, "bold"), fg_color=COLOR_ACCENT_PRIMARY,
                hover_color=COLOR_ACCENT_HOVER,
                command=lambda em=p_email, nm=p_nombre: self._modal_enviar_solicitud(em, nm)
            )
            btn_solicitar.pack(side="right", padx=14, pady=10)

    def _modal_enviar_solicitud(self, profesor_email: str, profesor_nombre: str):
        dialog = ctk.CTkToplevel(self)
        dialog.title("Enviar Solicitud de Tutoría")
        dialog.geometry("480x280")
        dialog.resizable(False, False)
        dialog.configure(fg_color=COLOR_BG_DARK)
        dialog.transient(self)
        dialog.grab_set()

        ctk.CTkLabel(dialog, text=f"📨 Vincularme con {profesor_nombre}", font=("Segoe UI", 14, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(16, 4))
        ctk.CTkLabel(dialog, text=f"Docente: {profesor_email}", font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(pady=(0, 10))

        ctk.CTkLabel(dialog, text="Mensaje o presentación para el profesor (opcional):", font=("Segoe UI", 10, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=25, pady=(2, 2))
        entry_msg = ctk.CTkEntry(dialog, height=36, font=("Segoe UI", 11), fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        entry_msg.pack(fill="x", padx=25, pady=(0, 15))
        entry_msg.insert(0, f"Hola profesor {profesor_nombre}, me gustaría vincularme contigo para resolver dudas académicas.")

        def _enviar():
            msg_texto = entry_msg.get().strip()
            dialog.destroy()
            def _thread():
                ok, res = tutoria_enviar_solicitud(profesor_email, profesor_nombre, msg_texto)
                if ok:
                    self.after(0, lambda: messagebox.showinfo("Solicitud Enviada", res))
                    self._iniciar_carga()
                else:
                    self.after(0, lambda: messagebox.showerror("Error", res))
            threading.Thread(target=_thread, daemon=True).start()

        frame_btns = ctk.CTkFrame(dialog, fg_color="transparent")
        frame_btns.pack(fill="x", padx=25, pady=(0, 15))
        ctk.CTkButton(frame_btns, text="Enviar Solicitud 🚀", font=("Segoe UI", 12, "bold"), fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER, height=36, command=_enviar).pack(side="left", fill="x", expand=True, padx=(0, 8))
        ctk.CTkButton(frame_btns, text="Cancelar", font=("Segoe UI", 11), fg_color=COLOR_BG_SURFACE, hover_color=COLOR_BORDER, height=36, width=80, command=dialog.destroy).pack(side="right")

    def _render_alumno_solicitud_pendiente(self, vinc: dict):
        p_nombre = vinc.get("profesor_nombre", "Profesor")
        p_email = vinc.get("profesor_email", "")
        fecha = vinc.get("timestamp", "")[:16].replace("T", " ")

        card = ctk.CTkFrame(self.frame_alumno_dinamico, fg_color=COLOR_BG_CARD, corner_radius=12, border_width=1, border_color="#eab308")
        card.pack(fill="x", pady=20, padx=20)

        ctk.CTkLabel(card, text="⏳ Solicitud de Vinculación en Espera", font=("Segoe UI", 16, "bold"), text_color="#fde047").pack(pady=(18, 6))
        ctk.CTkLabel(card, text=f"Has enviado una solicitud de tutoría al profesor {p_nombre} ({p_email}).", font=("Segoe UI", 12), text_color=COLOR_TEXT_MAIN).pack(pady=(0, 4))
        ctk.CTkLabel(card, text=f"🕒 Enviada el: {fecha}\nEl docente debe aceptar tu solicitud desde su panel para habilitar el chat.", font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED, justify="center").pack(pady=(0, 14))

        frame_btns = ctk.CTkFrame(card, fg_color="transparent")
        frame_btns.pack(pady=(0, 18))

        ctk.CTkButton(
            frame_btns, text="🔄 Comprobar Estado", font=("Segoe UI", 11, "bold"),
            fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER, height=34,
            command=self._iniciar_carga
        ).pack(side="left", padx=6)

        ctk.CTkButton(
            frame_btns, text="❌ Cancelar Solicitud", font=("Segoe UI", 11),
            fg_color="#7f1d1d", hover_color="#991b1b", border_width=1, border_color=COLOR_DANGER, height=34,
            command=lambda: self._desvincular_confirm(p_email)
        ).pack(side="left", padx=6)

    def _render_alumno_chat_tutor(self, vinc: dict):
        p_nombre = vinc.get("profesor_nombre", "Profesor")
        p_email = vinc.get("profesor_email", "")

        hdr_chat = ctk.CTkFrame(self.frame_alumno_dinamico, fg_color=COLOR_BG_SURFACE, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
        hdr_chat.pack(fill="x", pady=(0, 8))

        f_tit = ctk.CTkFrame(hdr_chat, fg_color="transparent")
        f_tit.pack(side="left", padx=14, pady=10)
        ctk.CTkLabel(f_tit, text=f"👨‍🏫 Profesor Tutor: {p_nombre}", font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")
        ctk.CTkLabel(f_tit, text=f"✉️ {p_email}   •   🟢 Vinculación Activa", font=("Segoe UI", 10), text_color=COLOR_SUCCESS).pack(anchor="w")

        ctk.CTkButton(
            hdr_chat, text="⚠️ Desvincular", font=("Segoe UI", 10, "bold"),
            fg_color="#7f1d1d", hover_color="#991b1b", height=28, width=100,
            command=lambda: self._desvincular_confirm(p_email)
        ).pack(side="right", padx=14)

        self.scroll_chat_alumno = ctk.CTkScrollableFrame(self.frame_alumno_dinamico, fg_color=COLOR_BG_DARK)
        self.scroll_chat_alumno.pack(fill="both", expand=True, pady=(0, 8))

        f_input = ctk.CTkFrame(self.frame_alumno_dinamico, fg_color=COLOR_BG_SURFACE, corner_radius=10)
        f_input.pack(fill="x")

        self.entry_msg_alumno = ctk.CTkEntry(
            f_input, height=40, font=("Segoe UI", 11),
            placeholder_text=f"Escribe tu duda o pregunta académica para el profesor {p_nombre}...",
            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER
        )
        self.entry_msg_alumno.pack(side="left", fill="x", expand=True, padx=10, pady=10)
        self.entry_msg_alumno.bind("<Return>", lambda e: self._enviar_msg_alumno(p_email))

        self.btn_enviar_alumno = ctk.CTkButton(
            f_input, text="Enviar 📤", width=110, height=40,
            font=("Segoe UI", 11, "bold"), fg_color=COLOR_ACCENT_PRIMARY,
            hover_color=COLOR_ACCENT_HOVER, command=lambda: self._enviar_msg_alumno(p_email)
        )
        self.btn_enviar_alumno.pack(side="right", padx=(0, 10), pady=10)

        self._cargar_mensajes_alumno(p_email)

    def _enviar_msg_alumno(self, p_email: str):
        texto = self.entry_msg_alumno.get().strip()
        if not texto:
            return
        self.entry_msg_alumno.delete(0, "end")
        self.btn_enviar_alumno.configure(state="disabled")

        def _thread():
            ok, _ = tutoria_enviar_mensaje(p_email, texto)
            if ok:
                self._cargar_mensajes_alumno(p_email)
            self.after(0, lambda: self.btn_enviar_alumno.configure(state="normal"))

        threading.Thread(target=_thread, daemon=True).start()

    def _cargar_mensajes_alumno(self, p_email: str):
        def _thread():
            ok, msgs = tutoria_obtener_mensajes_chat(p_email)
            if ok:
                self._mensajes_chat = msgs
                self.after(0, self._render_mensajes_alumno_ui)

        threading.Thread(target=_thread, daemon=True).start()

    def _render_mensajes_alumno_ui(self):
        if not hasattr(self, "scroll_chat_alumno"):
            return
        for w in self.scroll_chat_alumno.winfo_children():
            w.destroy()

        if not self._mensajes_chat:
            card_empty = ctk.CTkFrame(self.scroll_chat_alumno, fg_color=COLOR_BG_CARD, corner_radius=10)
            card_empty.pack(fill="x", padx=10, pady=20)
            ctk.CTkLabel(card_empty, text="👋 ¡Bienvenido al Canal de Tutoría!", font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(12, 2))
            ctk.CTkLabel(card_empty, text="Escribe a continuación tus dudas sobre ejercicios, temarios o exámenes para que tu profesor pueda ayudarte.", font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(pady=(0, 12))
            return

        for m in self._mensajes_chat:
            es_mio = m.get("es_mio", False)
            emisor = "Tú (Alumno)" if es_mio else f"👨‍🏫 {m.get('emisor_nombre', 'Profesor')}"
            hora = m.get("timestamp", "")[11:16] if m.get("timestamp") else ""
            texto = m.get("texto", "")

            frame_b = ctk.CTkFrame(
                self.scroll_chat_alumno,
                fg_color=COLOR_ACCENT_PRIMARY if es_mio else COLOR_BG_CARD,
                corner_radius=10,
                border_width=1,
                border_color=COLOR_ACCENT_HOVER if es_mio else COLOR_BORDER
            )
            frame_b.pack(anchor="e" if es_mio else "w", padx=8, pady=4)

            hdr = f"{emisor} • 🕒 {hora}"
            ctk.CTkLabel(frame_b, text=hdr, font=("Segoe UI", 9, "bold"),
                         text_color="#bae6fd" if es_mio else COLOR_ACCENT_CYAN).pack(anchor="w", padx=10, pady=(6, 1))

            ctk.CTkLabel(frame_b, text=texto, font=("Segoe UI", 11),
                         text_color="#ffffff", wraplength=460, justify="left").pack(anchor="w", padx=10, pady=(0, 8))

        self.scroll_chat_alumno._parent_canvas.yview_moveto(1.0)

    # ─────────────────────────────────────────────────────────
    #  VISTA DEL PROFESOR
    # ─────────────────────────────────────────────────────────
    def _build_ui_profesor(self):
        self.tabview = ctk.CTkTabview(self.contenedor_principal, fg_color=COLOR_BG_CARD,
                                      segmented_button_selected_color=COLOR_ACCENT_PURPLE,
                                      segmented_button_selected_hover_color="#7e22ce")
        self.tabview.pack(fill="both", expand=True)

        self.tab_solicitudes = self.tabview.add("🔔 Solicitudes Pendientes")
        self.tab_alumnos = self.tabview.add("👥 Mis Alumnos & Consultas")

        self._build_tab_profesor_solicitudes()
        self._build_tab_profesor_alumnos()

    def _build_tab_profesor_solicitudes(self):
        header_s = ctk.CTkFrame(self.tab_solicitudes, fg_color="transparent")
        header_s.pack(fill="x", padx=6, pady=(4, 6))

        self.lbl_num_solicitudes = ctk.CTkLabel(
            header_s, text="🔔 Solicitudes de Alumnos Pendientes de Aprobación",
            font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN
        )
        self.lbl_num_solicitudes.pack(side="left")

        ctk.CTkButton(
            header_s, text="🔄 Actualizar", width=95, height=30,
            font=("Segoe UI", 10, "bold"), fg_color=COLOR_BG_SURFACE,
            hover_color=COLOR_ACCENT_HOVER, command=self._iniciar_carga
        ).pack(side="right")

        self.scroll_solicitudes = ctk.CTkScrollableFrame(self.tab_solicitudes, fg_color=COLOR_BG_SURFACE, corner_radius=10)
        self.scroll_solicitudes.pack(fill="both", expand=True, padx=4, pady=4)

    def _render_profesor_solicitudes(self):
        for w in self.scroll_solicitudes.winfo_children():
            w.destroy()

        n = len(self._solicitudes_pendientes)
        self.lbl_num_solicitudes.configure(text=f"🔔 Solicitudes de Alumnos Pendientes de Aprobación ({n})")
        if "🔔 Solicitudes Pendientes" in self.tabview._segmented_button._buttons_dict:
            self.tabview._segmented_button._buttons_dict["🔔 Solicitudes Pendientes"].configure(text=f"🔔 Solicitudes ({n})")

        if not self._solicitudes_pendientes:
            ctk.CTkLabel(self.scroll_solicitudes, text="✨ No tienes solicitudes pendientes de alumnos en este momento.",
                         font=("Segoe UI", 12), text_color=COLOR_TEXT_MUTED).pack(pady=40)
            return

        for s in self._solicitudes_pendientes:
            a_email = s.get("alumno_email", "")
            a_nombre = s.get("alumno_nombre", a_email)
            msg = s.get("mensaje", "")
            fecha = s.get("timestamp", "")[:16].replace("T", " ")

            card = ctk.CTkFrame(self.scroll_solicitudes, fg_color=COLOR_BG_CARD, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
            card.pack(fill="x", padx=4, pady=4)

            f_info = ctk.CTkFrame(card, fg_color="transparent")
            f_info.pack(side="left", fill="x", expand=True, padx=14, pady=10)

            ctk.CTkLabel(f_info, text=f"🎓 {a_nombre}", font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")
            ctk.CTkLabel(f_info, text=f"✉️ {a_email}   •   🕒 {fecha}", font=("Segoe UI", 10), text_color=COLOR_ACCENT_CYAN).pack(anchor="w", pady=(1, 3))
            if msg:
                ctk.CTkLabel(f_info, text=f"💬 \"{msg}\"", font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED, wraplength=420, justify="left").pack(anchor="w")

            f_actions = ctk.CTkFrame(card, fg_color="transparent")
            f_actions.pack(side="right", padx=12, pady=10)

            ctk.CTkButton(
                f_actions, text="✅ Aceptar Alumno", width=125, height=32,
                font=("Segoe UI", 11, "bold"), fg_color=COLOR_SUCCESS, hover_color="#15803d",
                command=lambda em=a_email: self._responder_solicitud_profesor(em, True)
            ).pack(side="left", padx=4)

            ctk.CTkButton(
                f_actions, text="❌ Rechazar", width=95, height=32,
                font=("Segoe UI", 11), fg_color="#7f1d1d", hover_color="#991b1b", border_width=1, border_color=COLOR_DANGER,
                command=lambda em=a_email: self._responder_solicitud_profesor(em, False)
            ).pack(side="left", padx=4)

    def _responder_solicitud_profesor(self, alumno_email: str, aceptar: bool):
        def _thread():
            ok, res = tutoria_profesor_responder_solicitud(alumno_email, aceptar)
            if ok:
                self.after(0, lambda: messagebox.showinfo("Tutoría", res))
                self._iniciar_carga()
            else:
                self.after(0, lambda: messagebox.showerror("Error", res))

        threading.Thread(target=_thread, daemon=True).start()

    def _build_tab_profesor_alumnos(self):
        split = ctk.CTkFrame(self.tab_alumnos, fg_color="transparent")
        split.pack(fill="both", expand=True)
        split.grid_columnconfigure(0, weight=4)
        split.grid_columnconfigure(1, weight=6)
        split.grid_rowconfigure(0, weight=1)

        # Columna Izquierda: Alumnos
        col_izq = ctk.CTkFrame(split, fg_color=COLOR_BG_SURFACE, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
        col_izq.grid(row=0, column=0, sticky="nsew", padx=(0, 6), pady=4)

        ctk.CTkLabel(col_izq, text="👥 Alumnos Vinculados", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=14, pady=(10, 6))

        self.scroll_alumnos_profesor = ctk.CTkScrollableFrame(col_izq, fg_color="transparent")
        self.scroll_alumnos_profesor.pack(fill="both", expand=True, padx=4, pady=(0, 8))

        # Columna Derecha: Chat
        self.col_der_profesor = ctk.CTkFrame(split, fg_color=COLOR_BG_SURFACE, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
        self.col_der_profesor.grid(row=0, column=1, sticky="nsew", padx=(6, 0), pady=4)
        self.col_der_profesor.grid_rowconfigure(1, weight=1)
        self.col_der_profesor.grid_columnconfigure(0, weight=1)

        # Header del chat
        self.hdr_chat_profe = ctk.CTkFrame(self.col_der_profesor, fg_color=COLOR_BG_CARD, height=50, corner_radius=8)
        self.hdr_chat_profe.grid(row=0, column=0, sticky="ew", padx=8, pady=(8, 4))

        self.lbl_chat_alumno_info = ctk.CTkLabel(
            self.hdr_chat_profe, text="Selecciona un alumno a la izquierda para ver y responder dudas",
            font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED
        )
        self.lbl_chat_alumno_info.pack(side="left", padx=12, pady=10)

        self.btn_desvincular_alumno_profe = ctk.CTkButton(
            self.hdr_chat_profe, text="⚠️ Desvincular", width=100, height=28,
            font=("Segoe UI", 10, "bold"), fg_color="#7f1d1d", hover_color="#991b1b",
            command=self._desvincular_alumno_desde_profe
        )

        # Scroll mensajes
        self.scroll_mensajes_profe = ctk.CTkScrollableFrame(self.col_der_profesor, fg_color=COLOR_BG_DARK)
        self.scroll_mensajes_profe.grid(row=1, column=0, sticky="nsew", padx=8, pady=4)

        # Barra respuesta
        barra_resp = ctk.CTkFrame(self.col_der_profesor, fg_color="transparent")
        barra_resp.grid(row=2, column=0, sticky="ew", padx=8, pady=(4, 10))

        self.entry_resp_profe = ctk.CTkEntry(
            barra_resp, font=("Segoe UI", 11), height=38,
            placeholder_text="Escribe tu respuesta académica al alumno...",
            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER
        )
        self.entry_resp_profe.pack(side="left", fill="x", expand=True, padx=(0, 8))
        self.entry_resp_profe.bind("<Return>", lambda e: self._enviar_msg_profesor())

        self.btn_enviar_resp_profe = ctk.CTkButton(
            barra_resp, text="Responder 📤", width=115, height=38,
            font=("Segoe UI", 11, "bold"), fg_color=COLOR_ACCENT_PURPLE,
            hover_color="#7e22ce", command=self._enviar_msg_profesor,
            state="disabled"
        )
        self.btn_enviar_resp_profe.pack(side="right")

    def _render_profesor_alumnos(self):
        for w in self.scroll_alumnos_profesor.winfo_children():
            w.destroy()

        if not self._alumnos_vinculados:
            ctk.CTkLabel(self.scroll_alumnos_profesor, text="No tienes alumnos vinculados aún.\nCuando aceptes solicitudes aparecerán aquí.",
                         font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED, justify="center").pack(pady=30)
            return

        for a in self._alumnos_vinculados:
            a_email = a.get("alumno_email", "")
            a_nombre = a.get("alumno_nombre", a_email)

            card = ctk.CTkFrame(self.scroll_alumnos_profesor, fg_color=COLOR_BG_CARD, corner_radius=8, border_width=1, border_color=COLOR_BORDER)
            card.pack(fill="x", padx=2, pady=3)

            f_txt = ctk.CTkFrame(card, fg_color="transparent")
            f_txt.pack(fill="x", padx=8, pady=6)

            ctk.CTkLabel(f_txt, text=f"🎓 {a_nombre}", font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")
            ctk.CTkLabel(f_txt, text=a_email, font=("Segoe UI", 9), text_color=COLOR_ACCENT_CYAN).pack(anchor="w")

            ctk.CTkButton(
                card, text="💬 Abrir Tutoría", height=26, font=("Segoe UI", 10, "bold"),
                fg_color=COLOR_ACCENT_PURPLE, hover_color="#7e22ce",
                command=lambda em=a_email, nm=a_nombre: self._seleccionar_alumno_chat_profe(em, nm)
            ).pack(fill="x", padx=8, pady=(0, 6))

    def _seleccionar_alumno_chat_profe(self, a_email: str, a_nombre: str):
        self._alumno_activo_chat = {"email": a_email, "nombre": a_nombre}
        self.lbl_chat_alumno_info.configure(text=f"🎓 Alumno: {a_nombre} ({a_email})", text_color=COLOR_TEXT_MAIN)
        self.btn_desvincular_alumno_profe.pack(side="right", padx=10, pady=8)
        self.btn_enviar_resp_profe.configure(state="normal")
        self.entry_resp_profe.focus_set()
        self._cargar_mensajes_profesor_ui()

    def _cargar_mensajes_profesor_ui(self):
        if not self._alumno_activo_chat:
            return
        a_email = self._alumno_activo_chat["email"]

        def _thread():
            ok, msgs = tutoria_obtener_mensajes_chat(a_email)
            if ok:
                def _render():
                    for w in self.scroll_mensajes_profe.winfo_children():
                        w.destroy()

                    for m in msgs:
                        es_mio = m.get("es_mio", False)
                        emisor = "Tú (Profesor)" if es_mio else f"🎓 {m.get('emisor_nombre', 'Alumno')}"
                        hora = m.get("timestamp", "")[11:16] if m.get("timestamp") else ""
                        texto = m.get("texto", "")

                        frame_b = ctk.CTkFrame(
                            self.scroll_mensajes_profe,
                            fg_color=COLOR_ACCENT_PURPLE if es_mio else COLOR_BG_CARD,
                            corner_radius=10,
                            border_width=1,
                            border_color="#a855f7" if es_mio else COLOR_BORDER
                        )
                        frame_b.pack(anchor="e" if es_mio else "w", padx=8, pady=4)

                        hdr = f"{emisor} • 🕒 {hora}"
                        ctk.CTkLabel(frame_b, text=hdr, font=("Segoe UI", 9, "bold"),
                                     text_color="#e9d5ff" if es_mio else COLOR_ACCENT_CYAN).pack(anchor="w", padx=10, pady=(6, 1))

                        ctk.CTkLabel(frame_b, text=texto, font=("Segoe UI", 11),
                                     text_color="#ffffff", wraplength=380, justify="left").pack(anchor="w", padx=10, pady=(0, 8))

                    self.scroll_mensajes_profe._parent_canvas.yview_moveto(1.0)
                self.after(0, _render)

        threading.Thread(target=_thread, daemon=True).start()

    def _enviar_msg_profesor(self):
        if not self._alumno_activo_chat:
            return
        a_email = self._alumno_activo_chat["email"]
        texto = self.entry_resp_profe.get().strip()
        if not texto:
            return

        self.entry_resp_profe.delete(0, "end")
        self.btn_enviar_resp_profe.configure(state="disabled")

        def _thread():
            ok, _ = tutoria_enviar_mensaje(a_email, texto)
            if ok:
                self._cargar_mensajes_profesor_ui()
            self.after(0, lambda: self.btn_enviar_resp_profe.configure(state="normal"))

        threading.Thread(target=_thread, daemon=True).start()

    def _desvincular_alumno_desde_profe(self):
        if not self._alumno_activo_chat:
            return
        a_email = self._alumno_activo_chat["email"]
        self._desvincular_confirm(a_email)

    def _desvincular_confirm(self, otro_email: str):
        confirm = messagebox.askyesno(
            "⚠️ Finalizar Vinculación",
            f"¿Estás seguro de que deseas finalizar la vinculación con '{otro_email}'?\n\n"
            "• Se cerrará el canal de tutoría actual.\n"
            "• El alumno podrá solicitar unirse a otro profesor.\n\n"
            "¿Deseas continuar?"
        )
        if not confirm:
            return

        def _thread():
            ok, msg = tutoria_desvincular(otro_email)
            if ok:
                self.after(0, lambda: messagebox.showinfo("Tutoría", msg))
                self.after(0, self._iniciar_carga)
            else:
                self.after(0, lambda: messagebox.showerror("Error", msg))

        threading.Thread(target=_thread, daemon=True).start()

    # ─────────────────────────────────────────────────────────
    #  CARGA DE DATOS Y POLLING
    # ─────────────────────────────────────────────────────────
    def _iniciar_carga(self):
        if self.rol == "Alumno":
            self._cargar_datos_alumno()
        else:
            self._cargar_datos_profesor()

    def _cargar_datos_alumno(self):
        def _thread():
            ok_v, vinc = tutoria_obtener_estado_alumno()
            self._estado_vinculacion = vinc if ok_v else None

            if not self._estado_vinculacion or self._estado_vinculacion.get("estado") == "rechazada":
                ok_p, profes = tutoria_listar_profesores()
                self._profesores = profes if ok_p else []

            self.after(0, self._renderizar_alumno_estado)

        threading.Thread(target=_thread, daemon=True).start()

    def _cargar_profesores_alumno(self):
        def _thread():
            ok_p, profes = tutoria_listar_profesores()
            if ok_p:
                self._profesores = profes
                self.after(0, self._filtrar_profesores_ui)

        threading.Thread(target=_thread, daemon=True).start()

    def _cargar_datos_profesor(self):
        def _thread():
            ok, pend, acept = tutoria_profesor_obtener_solicitudes_y_alumnos()
            if ok:
                self._solicitudes_pendientes = pend
                self._alumnos_vinculados = acept
                self.after(0, lambda: (
                    self._render_profesor_solicitudes(),
                    self._render_profesor_alumnos()
                ))

        threading.Thread(target=_thread, daemon=True).start()

    def _cerrar(self):
        self._polling_activo = False
        self.destroy()


# ══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD PRINCIPAL (HOME CON CHAT DIRECTO IA, HISTORIAL Y CAMBIO DE MÓDULOS)
# ══════════════════════════════════════════════════════════════════════════════
class DashboardEstudios(ctk.CTk):
    def __init__(self, sesion):
        super().__init__()
        self.sesion  = sesion
        self.rol     = sesion.get("rol", "Alumno")
        self.nombre  = sesion.get("nombre", "Usuario")
        self.email   = sesion.get("email", "")

        self.title(f"KernossIA – {self.rol}: {self.nombre}")
        self.geometry("1420x840")
        self.minsize(1050, 680)
        self.configure(fg_color=COLOR_BG_DARK)

        # Módulos instanciados pero inicialmente ocultos
        self._modulos: dict[str, ctk.CTkFrame] = {}
        self._modulo_activo = None

        # Estado del Chat Directo de Bienvenida (Home)
        self.modelo_chat_home = "groq"
        self.historial_chat_home = []
        suffix = self.email.replace("@", "_").replace(".", "_") if self.email else "default"
        self.ruta_historial_home = os.path.expanduser(f"~/.historial_home_{suffix}.json")
        self.todo_el_historial_home = self._cargar_historial_home()
        self.chat_home_id_actual = None

        self._botones_modulos = {}
        self._datos_actualizacion = None
        self.frame_banner_update = None
        self._build_ui()
        self._mostrar_home_chat()

        # Sincronizar chats de la nube en segundo plano (Multi-dispositivo)
        threading.Thread(target=self._sincronizar_chats_cloud, daemon=True).start()

        # Comprobar si hay nueva versión de KernossIA en segundo plano
        threading.Thread(target=self._comprobar_actualizaciones, daemon=True).start()

        # Iniciar vigilante de baneos en tiempo real (Cierre de sesión y bloqueo inmediato)
        self._vigilante_activo = True
        self._iniciar_vigilante_baneos()

        # Comprobar si corresponde mostrar el Changelog de actualización (Solo 1 vez tras actualizar)
        self.after(600, self._comprobar_changelog_post_actualizacion)

    def _build_ui(self):
        # ── SIDEBAR LATERAL ──
        self.sidebar = ctk.CTkFrame(self, width=285, corner_radius=0,
                                   fg_color=COLOR_BG_SIDEBAR,
                                   border_width=1, border_color=COLOR_BORDER)
        self.sidebar.pack(side="left", fill="y")
        self.sidebar.pack_propagate(False)

        # Cabecera de la Marca
        frame_brand = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        frame_brand.pack(fill="x", padx=18, pady=(24, 10))

        ctk.CTkLabel(frame_brand, text=t("app_nombre"),
                     font=("Segoe UI", 24, "bold"), text_color=COLOR_ACCENT_SKY).pack(anchor="w")
        ctk.CTkLabel(frame_brand, text=t("app_tagline"),
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_DIM).pack(anchor="w")

        # Perfil de usuario con badge
        frame_user = ctk.CTkFrame(self.sidebar, fg_color=COLOR_BG_CARD, corner_radius=12,
                                  border_width=1, border_color=COLOR_BORDER)
        frame_user.pack(fill="x", padx=15, pady=(10, 10))
        icono = "🎓" if self.rol == "Alumno" else "👨‍🏫"
        self.lbl_perfil_nombre = ctk.CTkLabel(frame_user, text=f"{icono} {self.nombre}",
                                              font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN, anchor="w")
        self.lbl_perfil_nombre.pack(fill="x", padx=12, pady=(10, 2))

        self.lbl_perfil_email = ctk.CTkLabel(frame_user, text=self.email,
                                             font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED, anchor="w")
        self.lbl_perfil_email.pack(fill="x", padx=12)

        rol_texto = t("lbl_rol_alumno") if self.rol == "Alumno" else t("lbl_rol_profesor")
        badge_color = COLOR_ACCENT_PRIMARY if self.rol == "Alumno" else COLOR_ACCENT_PURPLE
        self.lbl_perfil_rol = ctk.CTkLabel(frame_user, text=f"  {rol_texto}  ",
                                           font=("Segoe UI", 10, "bold"), fg_color=badge_color,
                                           corner_radius=8, text_color="white")
        self.lbl_perfil_rol.pack(anchor="w", padx=12, pady=(6, 10))

        # ── BOTÓN HOME (INICIO / CHAT DIRECTO) ──
        self.btn_home = ctk.CTkButton(
            self.sidebar, text=t("btn_home"),
            font=("Segoe UI", 13, "bold"), height=42, anchor="w",
            fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
            command=self._mostrar_home_chat
        )
        self.btn_home.pack(fill="x", padx=15, pady=(4, 6))

        # ── HISTORIAL DE CHAT DESPLEGABLE EN SIDEBAR ──
        frame_historial_header = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        frame_historial_header.pack(fill="x", padx=16, pady=(8, 2))
        ctk.CTkLabel(frame_historial_header, text=t("hdr_historial"),
                     font=("Segoe UI", 10, "bold"), text_color=COLOR_ACCENT_CYAN).pack(side="left")

        ctk.CTkButton(frame_historial_header, text=t("btn_nuevo_chat"), width=55, height=22,
                      font=("Segoe UI", 10, "bold"), fg_color=COLOR_BG_SURFACE,
                      hover_color=COLOR_ACCENT_HOVER,
                      command=self._nuevo_chat_home).pack(side="right")

        self.scroll_historial_home = ctk.CTkScrollableFrame(self.sidebar, height=95, fg_color=COLOR_BG_CARD,
                                                           border_width=1, border_color=COLOR_BORDER,
                                                           corner_radius=8)
        self.scroll_historial_home.pack(fill="x", padx=15, pady=(2, 8))
        self._actualizar_historial_home_ui()

        ctk.CTkFrame(self.sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=4)

        # ── LISTADO DE MÓDULOS ──
        ctk.CTkLabel(self.sidebar, text=t("hdr_modulos_estudio"),
                     font=("Segoe UI", 10, "bold"), text_color=COLOR_TEXT_DIM).pack(anchor="w", padx=20, pady=(6, 2))

        self._btn(t("mod_mapas"),        "mapa_mental")
        self._btn(t("mod_calculador"),  "calculador")
        self._btn(t("mod_apuntador"),     "apuntador")
        self._btn(t("mod_resumidor"), "resumidor")
        self._btn(t("mod_examenes"),  "examen")
        self._btn(t("mod_ayudante"),  "ayudador")
        self._btn(t("mod_agenda"),     "calendario")

        if self.rol == "Profesor":
            ctk.CTkFrame(self.sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=6)
            ctk.CTkLabel(self.sidebar, text=t("hdr_herramientas_docente"),
                         font=("Segoe UI", 10, "bold"), text_color=COLOR_ACCENT_PURPLE).pack(anchor="w", padx=20, pady=(2, 2))
            self._btn(t("mod_ejercicios"),  "creador",  color=COLOR_BG_SURFACE)
            self._btn(t("mod_corrector"),  "corrector", color=COLOR_BG_SURFACE)

        # Botón Ajustes / Configuración de Voz IA
        ctk.CTkButton(self.sidebar, text=t("btn_ajustes"), height=36,
                      fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER,
                      text_color=COLOR_TEXT_MAIN, hover_color=COLOR_ACCENT_HOVER,
                      command=self._abrir_ajustes).pack(fill="x", padx=15, pady=(4, 6), side="bottom")

        # Botón Cerrar Sesión fijo abajo
        ctk.CTkButton(self.sidebar, text=t("btn_cerrar_sesion"), height=36,
                      fg_color="transparent", border_width=1, border_color=COLOR_BORDER,
                      text_color=COLOR_TEXT_MUTED, hover_color=COLOR_DANGER_HOVER,
                      command=self._cerrar_sesion).pack(fill="x", padx=15, pady=(4, 16), side="bottom")

        # ── CONTENEDOR PRINCIPAL DERECHO ──
        self.contenedor = ctk.CTkFrame(self, corner_radius=0, fg_color=COLOR_BG_DARK)
        self.contenedor.pack(side="right", fill="both", expand=True)
        self.contenedor.grid_rowconfigure(1, weight=1)
        self.contenedor.grid_columnconfigure(0, weight=1)

        # Barra Superior Permanente (Siempre visible con Novedades arriba a la derecha)
        self.header_top = ctk.CTkFrame(self.contenedor, fg_color=COLOR_BG_DARK, height=44)
        self.header_top.grid(row=0, column=0, sticky="ew", padx=25, pady=(12, 4))
        self.header_top.grid_columnconfigure(0, weight=1)

        self.lbl_seccion_actual = ctk.CTkLabel(
            self.header_top, text=t("btn_home"),
            font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MUTED
        )
        self.lbl_seccion_actual.grid(row=0, column=0, sticky="w")

        frame_top_derecha = ctk.CTkFrame(self.header_top, fg_color="transparent")
        frame_top_derecha.grid(row=0, column=1, sticky="e")

        # Botón Maestro de Administración (Solo para kernossai@support.com o admin)
        es_admin = (self.email.lower() in ["kernossai@support.com", "admin@kernosai.com", "soporte@kernosai.com"]) or self.sesion.get("is_premium", False)
        if es_admin:
            self.btn_admin_top = ctk.CTkButton(
                frame_top_derecha, text="👑 Moderación & Bans",
                font=("Segoe UI", 11, "bold"), height=32,
                fg_color="#3b0764", border_width=1, border_color="#c084fc",
                text_color="#f3e8ff", hover_color="#6b21a8",
                command=self._abrir_admin_moderacion
            )
            self.btn_admin_top.pack(side="right", padx=(6, 0))

        # Botón Tutoría Alumno <-> Profesor (Al lado de Soporte Oficial)
        texto_tutoria = t("btn_tutoria_alumno") if self.rol == "Alumno" else t("btn_tutoria_profesor")
        self.btn_tutoria_top = ctk.CTkButton(
            frame_top_derecha, text=texto_tutoria,
            font=("Segoe UI", 11, "bold"), height=32,
            fg_color="#1e1b4b" if self.rol == "Alumno" else "#312e81",
            border_width=1, border_color="#818cf8",
            text_color="#e0e7ff", hover_color="#4338ca",
            command=self._abrir_tutoria_profesor
        )
        self.btn_tutoria_top.pack(side="right", padx=(6, 0))

        # Botón Soporte Oficial con Cifrado E2EE
        self.btn_soporte_top = ctk.CTkButton(
            frame_top_derecha, text="🛡️ Soporte Oficial",
            font=("Segoe UI", 11, "bold"), height=32,
            fg_color="#064e3b", border_width=1, border_color="#34d399",
            text_color="#a7f3d0", hover_color="#059669",
            command=self._abrir_soporte_e2ee
        )
        self.btn_soporte_top.pack(side="right", padx=(6, 0))

        self.btn_novedades_top = ctk.CTkButton(
            frame_top_derecha, text=t("btn_novedades"),
            font=("Segoe UI", 11, "bold"), height=32,
            fg_color="#0c234a", border_width=1, border_color="#38bdf8",
            text_color="#38bdf8", hover_color="#0284c7",
            command=self._abrir_modal_novedades_ia
        )
        self.btn_novedades_top.pack(side="right", padx=(6, 0))

        # Frame donde se montan las pantallas y módulos
        self.frame_contenido = ctk.CTkFrame(self.contenedor, fg_color="transparent")
        self.frame_contenido.grid(row=1, column=0, sticky="nsew")
        self.frame_contenido.grid_rowconfigure(0, weight=1)
        self.frame_contenido.grid_columnconfigure(0, weight=1)

        # ── VISTA HOME: CHAT DIRECTO CON IA & ACCESO A MÓDULOS ──
        self._crear_vista_home_chat()

    def _abrir_ajustes(self):
        VentanaAjustes(self)

    def _abrir_modal_novedades_ia(self):
        VentanaNovedadesIA(self)

    def _abrir_soporte_e2ee(self):
        if self.email.lower() == "kernossai@support.com":
            VentanaAdminModeracion(self, self.sesion)
        else:
            VentanaSoporteE2EE(self, self.sesion)

    def _abrir_tutoria_profesor(self):
        VentanaTutoriaAlumnoProfesor(self, self.sesion)

    def _al_cambiar_cuenta(self, nueva_sesion):
        self.sesion = nueva_sesion
        self.email = nueva_sesion.get("email", "").lower()
        self.nombre = nueva_sesion.get("nombre", "Usuario")
        self.rol = nueva_sesion.get("rol", "Alumno")
        self.is_premium = nueva_sesion.get("is_premium", False)

        icono = "🎓" if self.rol == "Alumno" else "👨‍🏫"
        if hasattr(self, "lbl_perfil_nombre"):
            self.lbl_perfil_nombre.configure(text=f"{icono} {self.nombre}")
        if hasattr(self, "lbl_perfil_email"):
            self.lbl_perfil_email.configure(text=self.email)
        if hasattr(self, "lbl_perfil_rol"):
            rol_texto = t("lbl_rol_alumno") if self.rol == "Alumno" else t("lbl_rol_profesor")
            badge_color = COLOR_ACCENT_PRIMARY if self.rol == "Alumno" else COLOR_ACCENT_PURPLE
            self.lbl_perfil_rol.configure(text=f"  {rol_texto}  ", fg_color=badge_color)

        if hasattr(self, "btn_tutoria_top"):
            txt_t = t("btn_tutoria_alumno") if self.rol == "Alumno" else t("btn_tutoria_profesor")
            fg_t = "#1e1b4b" if self.rol == "Alumno" else "#312e81"
            self.btn_tutoria_top.configure(text=txt_t, fg_color=fg_t)

        # Comprobar botón maestro de administración
        es_admin = (self.email in ["kernossai@support.com", "admin@kernosai.com", "soporte@kernosai.com"]) or self.is_premium
        if hasattr(self, "btn_admin_top"):
            if es_admin:
                self.btn_admin_top.pack(side="right", padx=(6, 0))
            else:
                self.btn_admin_top.pack_forget()

        self._mostrar_home_chat()

    def _iniciar_vigilante_baneos(self):
        def _loop():
            while self._vigilante_activo:
                threading.Event().wait(3.5)
                if not self._vigilante_activo:
                    break
                try:
                    ok_clean, motivo_ban, tipo_ban = verificar_estado_baneo(self.email)
                    if not ok_clean:
                        self._vigilante_activo = False
                        self.after(0, lambda t=tipo_ban, m=motivo_ban: self._mostrar_pantalla_baneado(t, m))
                        break
                except Exception:
                    pass
        
        threading.Thread(target=_loop, daemon=True).start()

    def _mostrar_pantalla_baneado(self, tipo_ban: str, motivo_ban: str):
        # 1. Borrar token y sesión local
        borrar_token()
        
        # 2. Deshabilitar botones del sidebar
        for w in self.sidebar.winfo_children():
            try:
                w.configure(state="disabled")
            except Exception:
                pass

        # 3. Despejar contenedor y mostrar pantalla de bloqueo roja gigante
        for w in self.contenedor.winfo_children():
            w.destroy()

        frame_lockout = ctk.CTkFrame(self.contenedor, fg_color="#3b0707", corner_radius=0)
        frame_lockout.pack(fill="both", expand=True)

        card_bloqueo = ctk.CTkFrame(
            frame_lockout, fg_color="#180404", corner_radius=20,
            border_width=2, border_color="#ef4444"
        )
        card_bloqueo.place(relx=0.5, rely=0.5, anchor="center")

        # Icono grande
        ctk.CTkLabel(card_bloqueo, text="⛔", font=("Segoe UI", 56)).pack(padx=50, pady=(26, 4))

        # Título en rojo grande
        ctk.CTkLabel(
            card_bloqueo, text="ACCESO BLOQUEADO — CUENTA BANEADA",
            font=("Segoe UI", 18, "bold"), text_color="#ef4444"
        ).pack(padx=50, pady=(0, 4))

        ctk.CTkLabel(
            card_bloqueo, text="Has sido sancionado por la administración de KernossAI.\nTu sesión ha sido revocada y tu acceso a la IA queda suspendido.",
            font=("Segoe UI", 11), text_color="#fca5a5", justify="center"
        ).pack(padx=50, pady=(0, 16))

        # Caja de detalles de la sanción
        box_detalles = ctk.CTkFrame(card_bloqueo, fg_color="#260505", corner_radius=10, border_width=1, border_color="#7f1d1d")
        box_detalles.pack(fill="x", padx=30, pady=(0, 18))

        tipo_str = tipo_ban or "Baneo de Cuenta Permanente"
        motivo_str = motivo_ban or "Infracción de las normas de la comunidad"

        ctk.CTkLabel(box_detalles, text=f"🔒 Tipo de Sanción: {tipo_str}", font=("Segoe UI", 11, "bold"), text_color="#fca5a5").pack(anchor="w", padx=16, pady=(12, 3))
        ctk.CTkLabel(box_detalles, text=f"📋 Motivo: {motivo_str}", font=("Segoe UI", 11), text_color="#fee2e2", wraplength=440, justify="left").pack(anchor="w", padx=16, pady=(0, 3))
        ctk.CTkLabel(box_detalles, text="⚡ Estado: Sesión Cancelada • Inteligencia Artificial Bloqueada.", font=("Segoe UI", 10, "bold"), text_color="#f87171").pack(anchor="w", padx=16, pady=(0, 12))

        # Botón de Cerrar Sesión
        btn_cerrar = ctk.CTkButton(
            card_bloqueo, text="🚪 Cerrar Sesión y Salir", height=42,
            font=("Segoe UI", 12, "bold"), fg_color="#dc2626", hover_color="#b91c1c",
            command=self._ejecutar_cierre_por_ban
        )
        btn_cerrar.pack(fill="x", padx=30, pady=(0, 26))

    def _ejecutar_cierre_por_ban(self):
        borrar_token()
        self.quit()
        self.destroy()
        os._exit(0)

    def _abrir_admin_moderacion(self):
        VentanaAdminModeracion(self, self.sesion)

    def _btn(self, texto, modulo_id, color=COLOR_BG_SURFACE):
        btn = ctk.CTkButton(
            self.sidebar, text=texto,
            font=("Segoe UI", 12, "bold"), height=36, anchor="w",
            fg_color=color, hover_color=COLOR_ACCENT_HOVER,
            command=lambda mid=modulo_id: self._abrir_modulo(mid)
        )
        btn.pack(fill="x", padx=15, pady=2)
        self._botones_modulos[modulo_id] = btn

    def _crear_vista_home_chat(self):
        self.frame_home = ctk.CTkFrame(self.frame_contenido, fg_color="transparent")
        self.frame_home.grid_rowconfigure(2, weight=1)
        self.frame_home.grid_columnconfigure(0, weight=1)

        # 1. Header de Bienvenida
        header_home = ctk.CTkFrame(self.frame_home, fg_color="transparent")
        header_home.grid(row=0, column=0, sticky="ew", padx=25, pady=(10, 10))

        rol_nombre = t("lbl_rol_alumno") if self.rol == "Alumno" else t("lbl_rol_profesor")
        saludo = t("home_bienvenida", rol=rol_nombre, nombre=self.nombre)
        ctk.CTkLabel(header_home, text=saludo, font=("Segoe UI", 26, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")

        sub_header = ctk.CTkFrame(header_home, fg_color="transparent")
        sub_header.pack(fill="x", pady=(2, 0))

        self.lbl_estado_ia_home = ctk.CTkLabel(sub_header, text=t("home_estado_ia"),
                                               font=("Segoe UI", 12), text_color=COLOR_SUCCESS)
        self.lbl_estado_ia_home.pack(side="left")

        # Selector de motor en el home
        frame_selector_ia = ctk.CTkFrame(sub_header, fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER, corner_radius=8)
        frame_selector_ia.pack(side="right")

        self.btn_home_groq = ctk.CTkButton(frame_selector_ia, text=t("opt_groq"), height=28, width=115,
                                           font=("Segoe UI", 11, "bold"),
                                           fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                           command=lambda: self._set_home_modelo("groq"))
        self.btn_home_groq.pack(side="left", padx=2, pady=2)

        self.btn_home_gemini = ctk.CTkButton(frame_selector_ia, text=t("opt_gemini"), height=28, width=125,
                                             font=("Segoe UI", 11, "bold"),
                                             fg_color="transparent", hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                             command=lambda: self._set_home_modelo("gemini"))
        self.btn_home_gemini.pack(side="left", padx=2, pady=2)

        # 2. Barra de sugerencias y cambio de módulo rápido
        bar_modulos = ctk.CTkFrame(self.frame_home, fg_color=COLOR_BG_CARD, corner_radius=12,
                                   border_width=1, border_color=COLOR_BORDER)
        bar_modulos.grid(row=1, column=0, sticky="ew", padx=25, pady=(0, 12))

        ctk.CTkLabel(bar_modulos, text="⚡", font=("Segoe UI", 11, "bold"),
                     text_color=COLOR_TEXT_MUTED).pack(side="left", padx=(14, 4), pady=8)

        modulos_rapidos = [
            (t("mod_mapas"), "mapa_mental"),
            (t("mod_calculador"), "calculador"),
            (t("mod_apuntador"), "apuntador"),
            (t("mod_resumidor"), "resumidor"),
            (t("mod_examenes"), "examen"),
            (t("mod_ayudante"), "ayudador"),
            (t("mod_agenda"), "calendario")
        ]
        if self.rol == "Profesor":
            modulos_rapidos.extend([(t("mod_ejercicios"), "creador"), (t("mod_corrector"), "corrector")])

        for label, mid in modulos_rapidos:
            ctk.CTkButton(bar_modulos, text=label, height=28, font=("Segoe UI", 11, "bold"),
                          fg_color=COLOR_BG_CARD_LIGHT, hover_color=COLOR_ACCENT_HOVER,
                          border_width=1, border_color=COLOR_BORDER,
                          command=lambda m=mid: self._abrir_modulo(m)).pack(side="left", padx=3, pady=6)

        # 3. Ventana del Chat de Bienvenida
        self.txt_home_chat = ctk.CTkTextbox(self.frame_home, font=("Segoe UI", 14), state="disabled", wrap="word",
                                            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER,
                                            corner_radius=14)
        self.txt_home_chat.grid(row=2, column=0, sticky="nsew", padx=25, pady=(0, 12))

        # Mensaje de bienvenida inicial en el chat
        self._inicializar_chat_bienvenida_texto()

        # 4. Input y Botones de Envío
        input_container = ctk.CTkFrame(self.frame_home, fg_color="transparent")
        input_container.grid(row=3, column=0, sticky="ew", padx=25, pady=(0, 20))
        input_container.grid_columnconfigure(0, weight=1)

        self.entry_home_pregunta = ctk.CTkEntry(input_container, placeholder_text=t("home_placeholder_input"),
                                                height=46, font=("Segoe UI", 13),
                                                fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_home_pregunta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_home_pregunta.bind("<Return>", lambda e: self._enviar_chat_home())

        btn_enviar_home = ctk.CTkButton(input_container, text=t("btn_consultar_ia"), width=130, height=46,
                                       font=("Segoe UI", 13, "bold"),
                                       fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                       command=self._enviar_chat_home)
        btn_enviar_home.grid(row=0, column=1, padx=(0, 6))

        self.btn_tts_home = ctk.CTkButton(input_container, text=t("btn_escuchar"), width=105, height=46,
                                          font=("Segoe UI", 12, "bold"),
                                          fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                          hover_color=COLOR_ACCENT_HOVER,
                                          command=self._toggle_tts_home)
        self.btn_tts_home.grid(row=0, column=2, padx=(0, 6))

        btn_word_home = ctk.CTkButton(input_container, text=t("btn_word"), width=75, height=46,
                                      font=("Segoe UI", 12, "bold"),
                                      fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER,
                                      hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                      command=self._exportar_home_word)
        btn_word_home.grid(row=0, column=3)

    def _inicializar_chat_bienvenida_texto(self):
        self.txt_home_chat.configure(state="normal")
        self.txt_home_chat.delete("1.0", "end")
        texto_intro = (
            f"👋 ¡Hola, {self.nombre}! Soy tu asistente académico inteligente en KernossIA.\n\n"
            f"• Pregúntame dudas de cualquier asignatura (Matemáticas, Historia, Ciencias, Idiomas...).\n"
            f"• Pídeme resúmenes, esquemas explicativos o resolución paso a paso de problemas.\n"
            f"• O cambia a cualquiera de los módulos dedicados en el menú lateral o en la barra superior.\n"
            f"──────────────────────────────────────────────────────────────────────────\n"
        )
        self.txt_home_chat.insert("end", texto_intro)
        self.txt_home_chat.configure(state="disabled")

    def _set_home_modelo(self, m):
        self.modelo_chat_home = m
        if m == "groq":
            self.btn_home_groq.configure(fg_color=COLOR_ACCENT_PRIMARY)
            self.btn_home_gemini.configure(fg_color="transparent")
            self.lbl_estado_ia_home.configure(text="🟢 Asistente IA Activo (Groq - LLaMA 3.3 70B Ultrarrápido)", text_color=COLOR_SUCCESS)
        else:
            self.btn_home_groq.configure(fg_color="transparent")
            self.btn_home_gemini.configure(fg_color=COLOR_ACCENT_PURPLE)
            self.lbl_estado_ia_home.configure(text="🟣 Asistente IA Activo (Google Gemini - Análisis Avanzado)", text_color="#a5b4fc")

    def _enviar_chat_home(self):
        pregunta = self.entry_home_pregunta.get().strip()
        if not pregunta:
            return

        self.txt_home_chat.configure(state="normal")
        self.txt_home_chat.insert("end", f"\n\n👤 {self.nombre}:\n{pregunta}\n")
        self.txt_home_chat.see("end")
        self.txt_home_chat.configure(state="disabled")

        self.entry_home_pregunta.delete(0, "end")
        self.lbl_estado_ia_home.configure(text="🟡 Pensando y analizando respuesta...", text_color=COLOR_WARNING)
        self.historial_chat_home.append({"role": "user", "content": pregunta})

        threading.Thread(target=self._thread_proceso_home, daemon=True).start()

    def _thread_proceso_home(self):
        try:
            self.after(0, lambda: self._stream_home_texto(f"\n🤖 KernossIA ({self.modelo_chat_home.upper()}):\n"))
            instrucciones = (
                "Eres el tutor y asistente de estudio principal de la plataforma KernossIA. "
                "Responde de forma clara, didáctica, amigable y rigurosa a las consultas académicas del usuario."
            )
            prompt = construir_prompt(instrucciones, self.historial_chat_home)
            if self.modelo_chat_home == "groq":
                respuesta = llamar_groq(prompt)
            else:
                respuesta = llamar_gemini(prompt)

            self.ultima_respuesta_ia_home = respuesta
            self.historial_chat_home.append({"role": "assistant", "content": respuesta})
            self.after(0, self._stream_home_texto, respuesta + "\n")
            self._guardar_historial_home()

            color = COLOR_SUCCESS if self.modelo_chat_home == "groq" else "#a5b4fc"
            self.after(0, lambda: self.lbl_estado_ia_home.configure(
                text=f"🟢 Listo • {self.modelo_chat_home.upper()} respondió correctamente", text_color=color))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            self.after(0, lambda: self.lbl_estado_ia_home.configure(text="🔴 Error al procesar respuesta", text_color=COLOR_DANGER))
        finally:
            self.after(0, lambda: self.txt_home_chat.configure(state="disabled"))

    def _toggle_tts_home(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_home.configure(text="🔊 Escuchar", fg_color=COLOR_BG_CARD)
        else:
            texto = getattr(self, "ultima_respuesta_ia_home", "")
            if not texto:
                texto = self.txt_home_chat.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin texto", "No hay respuesta para reproducir en voz alta.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_home.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_home.configure(text="🔊 Escuchar", fg_color=COLOR_BG_CARD)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def _stream_home_texto(self, texto):
        self.txt_home_chat.configure(state="normal")
        self.txt_home_chat.insert("end", texto)
        self.txt_home_chat.see("end")
        self.txt_home_chat.configure(state="disabled")

    def _nuevo_chat_home(self):
        self._mostrar_home_chat()
        self.historial_chat_home = []
        self.chat_home_id_actual = None
        self._inicializar_chat_bienvenida_texto()

    def _exportar_home_word(self):
        if not self.historial_chat_home:
            messagebox.showwarning("Vacío", "No hay mensajes en el chat para exportar.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"Conversacion_KernossIA_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Sesión de Estudio – {self.nombre} (KernossIA)", 0)
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph("")
            for msg in self.historial_chat_home:
                rol = self.nombre if msg["role"] == "user" else "KernossIA Tutor"
                p = doc.add_paragraph()
                p.add_run(f"{rol}: ").bold = True
                p.add_run(msg["content"])
            doc.save(path)
            messagebox.showinfo("Éxito", f"Conversación guardada en:\n{path}")

    def _cargar_historial_home(self):
        if os.path.exists(self.ruta_historial_home):
            try:
                with open(self.ruta_historial_home, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def _sincronizar_chats_cloud(self):
        """Sincroniza en segundo plano las conversaciones guardadas en la nube para acceso multi-dispositivo."""
        try:
            chats_remotos = obtener_chats_cloud()
            if chats_remotos and isinstance(chats_remotos, dict):
                cambios = False
                for chat_id, data in chats_remotos.items():
                    mensajes = data.get("mensajes", [])
                    if chat_id not in self.todo_el_historial_home and mensajes:
                        self.todo_el_historial_home[chat_id] = mensajes
                        cambios = True
                if cambios:
                    with open(self.ruta_historial_home, "w", encoding="utf-8") as f:
                        json.dump(self.todo_el_historial_home, f, ensure_ascii=False, indent=2)
                    self.after(0, self._actualizar_historial_home_ui)
        except Exception as e:
            print(f"[Sync Cloud] {e}")

    def _guardar_historial_home(self):
        if not self.historial_chat_home:
            return
        if self.chat_home_id_actual is None:
            primer = next((m["content"] for m in self.historial_chat_home if m["role"] == "user"), "Consulta")
            resumen = primer[:20] + "..." if len(primer) > 20 else primer
            self.chat_home_id_actual = f"[{datetime.now().strftime('%H:%M')}] {resumen}"
        self.todo_el_historial_home[self.chat_home_id_actual] = self.historial_chat_home
        
        # 1. Guardar localmente
        try:
            with open(self.ruta_historial_home, "w", encoding="utf-8") as f:
                json.dump(self.todo_el_historial_home, f, ensure_ascii=False, indent=2)
            self.after(0, self._actualizar_historial_home_ui)
        except Exception as e:
            print(f"Error guardando historial local: {e}")

        # 2. Sincronizar en la nube en segundo plano (Multi-dispositivo)
        chat_id_actual = self.chat_home_id_actual
        mensajes_actuales = list(self.historial_chat_home)
        threading.Thread(
            target=lambda: guardar_chat_cloud(chat_id_actual, chat_id_actual, mensajes_actuales),
            daemon=True
        ).start()

    def _actualizar_historial_home_ui(self):
        for w in self.scroll_historial_home.winfo_children():
            w.destroy()
        for chat_id in reversed(list(self.todo_el_historial_home.keys())):
            fila = ctk.CTkFrame(self.scroll_historial_home, fg_color="transparent")
            fila.pack(fill="x", pady=2)
            ctk.CTkButton(fila, text=chat_id, fg_color="transparent",
                          text_color=COLOR_TEXT_MAIN, anchor="w", hover_color=COLOR_BG_SURFACE, height=26,
                          font=("Segoe UI", 10),
                          command=lambda cid=chat_id: self._cargar_chat_home(cid)).pack(side="left", fill="x", expand=True)
            ctk.CTkButton(fila, text="✕", width=22, height=22,
                          fg_color="transparent", hover_color=COLOR_DANGER_HOVER,
                          text_color=COLOR_TEXT_MUTED, font=("Segoe UI", 9, "bold"),
                          command=lambda cid=chat_id: self._borrar_chat_home(cid)).pack(side="right")

    def _cargar_chat_home(self, chat_id):
        self._mostrar_home_chat()
        self.chat_home_id_actual = chat_id
        self.historial_chat_home = self.todo_el_historial_home.get(chat_id, [])
        self.txt_home_chat.configure(state="normal")
        self.txt_home_chat.delete("1.0", "end")
        self.txt_home_chat.insert("end", f"💬 Cargado: {chat_id}\n{'─'*55}\n")
        for msg in self.historial_chat_home:
            if msg["role"] == "user":
                self.txt_home_chat.insert("end", f"\n👤 {self.nombre}:\n{msg['content']}\n")
            else:
                self.txt_home_chat.insert("end", f"\n🤖 KernossIA:\n{msg['content']}\n")
        self.txt_home_chat.configure(state="disabled")
        self.txt_home_chat.see("end")

    def _borrar_chat_home(self, chat_id):
        if messagebox.askyesno("Confirmar", f"¿Eliminar este chat?\n'{chat_id}'"):
            if chat_id in self.todo_el_historial_home:
                del self.todo_el_historial_home[chat_id]
                try:
                    with open(self.ruta_historial_home, "w", encoding="utf-8") as f:
                        json.dump(self.todo_el_historial_home, f, ensure_ascii=False, indent=2)
                except Exception:
                    pass
                if self.chat_home_id_actual == chat_id:
                    self._nuevo_chat_home()
                self._actualizar_historial_home_ui()

                # Eliminar de la nube en segundo plano
                threading.Thread(
                    target=lambda cid=chat_id: borrar_chat_cloud(cid),
                    daemon=True
                ).start()

    def _mostrar_home_chat(self):
        # Ocultar cualquier módulo activo
        if self._modulo_activo:
            self._modulo_activo.grid_forget()
            self._modulo_activo = None

        # Resetear colores de botones del sidebar
        self.btn_home.configure(fg_color=COLOR_ACCENT_PRIMARY)
        for btn in self._botones_modulos.values():
            btn.configure(fg_color=COLOR_BG_SURFACE)

        if hasattr(self, "lbl_seccion_actual"):
            self.lbl_seccion_actual.configure(text="🏠 Inicio / Chat Asistente IA")

        self.frame_home.grid(row=0, column=0, sticky="nsew")

    def _abrir_modulo(self, modulo_id):
        # Ocultar la pantalla home
        self.frame_home.grid_forget()

        # Ocultar el módulo activo anterior
        if self._modulo_activo:
            self._modulo_activo.grid_forget()

        # Actualizar colores en sidebar
        self.btn_home.configure(fg_color=COLOR_BG_SURFACE)
        for mid, btn in self._botones_modulos.items():
            if mid == modulo_id:
                btn.configure(fg_color=COLOR_ACCENT_PRIMARY)
            else:
                btn.configure(fg_color=COLOR_BG_SURFACE)

        # Actualizar indicador de sección activa en la barra superior
        nombres_secciones = {
            "mapa_mental": "🧠 Generador de Mapas Mentales con IA",
            "calculador": "📊 Calculador de Medias y Ponderaciones",
            "apuntador": "📝 Apuntador de Notas de Clase",
            "resumidor": "🔍 Resumidor Inteligente de Textos",
            "examen": "🎯 Generador y Práctica de Exámenes",
            "ayudador": "🤖 Tutor y Ayudante de Dudas Académicas",
            "calendario": "📅 Agenda y Planificador de Estudios",
            "creador": "✏️ Creador de Ejercicios Didácticos",
            "corrector": "📋 Corrector Inteligente de Exámenes"
        }
        if hasattr(self, "lbl_seccion_actual"):
            self.lbl_seccion_actual.configure(text=nombres_secciones.get(modulo_id, "📂 Módulo de Estudio"))

        # Crear el módulo si no existe todavía (lazy loading en frame_contenido)
        if modulo_id not in self._modulos:
            if modulo_id == "mapa_mental":
                self._modulos[modulo_id] = ModuloMapaMental(self.frame_contenido)
            elif modulo_id == "calculador":
                self._modulos[modulo_id] = ModuloCalculador(self.frame_contenido)
            elif modulo_id == "apuntador":
                self._modulos[modulo_id] = ModuloApuntador(self.frame_contenido)
            elif modulo_id == "resumidor":
                self._modulos[modulo_id] = ModuloResumidor(self.frame_contenido)
            elif modulo_id == "examen":
                self._modulos[modulo_id] = ModuloExamen(self.frame_contenido)
            elif modulo_id == "ayudador":
                self._modulos[modulo_id] = ModuloAyudador(self.frame_contenido, sesion=self.sesion)
            elif modulo_id == "calendario":
                self._modulos[modulo_id] = ModuloCalendario(self.frame_contenido)
            elif modulo_id == "creador":
                self._modulos[modulo_id] = ModuloCreadorEjercicios(self.frame_contenido)
            elif modulo_id == "corrector":
                self._modulos[modulo_id] = ModuloCorrectorExamenes(self.frame_contenido)

        modulo = self._modulos[modulo_id]
        modulo.grid(row=0, column=0, sticky="nsew")
        self._modulo_activo = modulo

    def _cerrar_sesion(self):
        if messagebox.askyesno("Cerrar Sesión", "¿Seguro que quieres cerrar sesión?"):
            borrar_token()
            self.quit()
            self.destroy()
            os._exit(0)

    def _comprobar_actualizaciones(self):
        """Consulta GitHub Releases en segundo plano para comprobar si hay una nueva versión."""
        try:
            url = "https://api.github.com/repos/LDPCICNM2012/KernossAI/releases/latest"
            headers = {"User-Agent": "KernossAI-Desktop-App"}
            resp = requests.get(url, headers=headers, timeout=6)
            if resp.status_code == 200:
                data = resp.json()
                tag_remoto = data.get("tag_name", "")
                if tag_remoto and es_version_superior(tag_remoto, VERSION_APP):
                    assets = data.get("assets", [])
                    download_url = data.get("html_url", "https://github.com/LDPCICNM2012/KernossAI/releases/latest")
                    
                    if sys.platform == "win32":
                        for a in assets:
                            if a.get("name", "").lower().endswith(".exe"):
                                download_url = a.get("browser_download_url", download_url)
                                break
                    elif sys.platform == "darwin":
                        for a in assets:
                            if a.get("name", "").lower().endswith((".zip", ".dmg")):
                                download_url = a.get("browser_download_url", download_url)
                                break

                    info = {
                        "tag": tag_remoto,
                        "url": download_url,
                        "html_url": data.get("html_url", ""),
                        "notas": data.get("body", "Mejoras de rendimiento, corrección de errores y nuevas funciones.")
                    }
                    self._datos_actualizacion = info
                    self.after(500, lambda: self._mostrar_notificacion_actualizacion(info))
        except Exception:
            pass

    def _mostrar_notificacion_actualizacion(self, info):
        """Muestra un banner en la pantalla Home y un botón en el sidebar."""
        tag = info["tag"]
        dl_url = info["url"]

        # 1. Botón destacado en el sidebar
        if hasattr(self, "sidebar") and not hasattr(self, "btn_update_sidebar"):
            self.btn_update_sidebar = ctk.CTkButton(
                self.sidebar,
                text=f"🚀 Actualizar a {tag}",
                font=("Segoe UI", 11, "bold"),
                height=32,
                fg_color="#0284c7",
                hover_color="#0ea5e9",
                command=lambda: self._abrir_modal_actualizacion(info)
            )
            self.btn_update_sidebar.pack(fill="x", padx=15, pady=(2, 4), after=self.btn_home)

        # 2. Banner de actualización en la vista Home
        if hasattr(self, "frame_home") and self.frame_banner_update is None:
            self.frame_banner_update = ctk.CTkFrame(
                self.frame_home,
                fg_color="#0c234a",
                border_width=1,
                border_color="#38bdf8",
                corner_radius=10
            )
            # Insertar en la fila 0 y mover los demás
            self.frame_banner_update.grid(row=0, column=0, sticky="ew", padx=25, pady=(15, 0))

            frame_txt = ctk.CTkFrame(self.frame_banner_update, fg_color="transparent")
            frame_txt.pack(side="left", padx=15, pady=10)

            ctk.CTkLabel(
                frame_txt,
                text=f"🎉 ¡Nueva actualización disponible ({tag})!",
                font=("Segoe UI", 13, "bold"),
                text_color="#38bdf8"
            ).pack(anchor="w")

            ctk.CTkLabel(
                frame_txt,
                text=f"Tu versión actual es v{VERSION_APP}. Haz clic en 'Descargar' para actualizar tu app.",
                font=("Segoe UI", 11),
                text_color=COLOR_TEXT_MUTED
            ).pack(anchor="w")

            frame_btns = ctk.CTkFrame(self.frame_banner_update, fg_color="transparent")
            frame_btns.pack(side="right", padx=15, pady=10)

            ctk.CTkButton(
                frame_btns,
                text="⚡ Descargar",
                font=("Segoe UI", 12, "bold"),
                fg_color=COLOR_ACCENT_PRIMARY,
                hover_color=COLOR_ACCENT_HOVER,
                height=30,
                command=lambda: webbrowser.open(dl_url)
            ).pack(side="left", padx=4)

            ctk.CTkButton(
                frame_btns,
                text="ℹ️ Novedades",
                font=("Segoe UI", 11),
                fg_color=COLOR_BG_SURFACE,
                hover_color=COLOR_BORDER,
                height=30,
                command=lambda: self._abrir_modal_actualizacion(info)
            ).pack(side="left", padx=4)

            ctk.CTkButton(
                frame_btns,
                text="✕",
                width=26,
                height=30,
                fg_color="transparent",
                hover_color=COLOR_DANGER_HOVER,
                command=self._cerrar_banner_update
            ).pack(side="left", padx=2)

    def _cerrar_banner_update(self):
        if self.frame_banner_update:
            self.frame_banner_update.destroy()
            self.frame_banner_update = None

    def _descargar_e_instalar_ota(self, info, progress_bar, lbl_status, btn_ota, btn_web, modal):
        """Descarga e instala en segundo plano (OTA) la nueva versión de KernossAI con reinicio automático."""
        url_descarga = info.get("url", "")
        # Si no es un .exe directo o no estamos en Windows, abrir en navegador como respaldo
        if sys.platform != "win32" or not url_descarga.lower().endswith(".exe"):
            webbrowser.open(info.get("html_url", url_descarga))
            modal.destroy()
            return

        btn_ota.configure(state="disabled", text=t("ota_descargando"))
        btn_web.configure(state="disabled")
        progress_bar.pack(fill="x", padx=30, pady=(6, 4))
        progress_bar.set(0.0)
        lbl_status.pack(anchor="w", padx=30, pady=(0, 10))
        lbl_status.configure(text=t("ota_descargando"), text_color=COLOR_ACCENT_SKY)

        def _worker():
            try:
                tag_limpio = info.get("tag", "latest").replace(" ", "_")
                ruta_temp = os.path.join(tempfile.gettempdir(), f"KernossAI_Setup_OTA_{tag_limpio}.exe")

                # Descargar en bloques (stream) para actualizar la barra de progreso
                resp = requests.get(url_descarga, stream=True, timeout=90)
                total_bytes = int(resp.headers.get("content-length", 0))
                descargados = 0

                with open(ruta_temp, "wb") as f:
                    for chunk in resp.iter_content(chunk_size=65536):
                        if chunk:
                            f.write(chunk)
                            descargados += len(chunk)
                            if total_bytes > 0:
                                p = descargados / total_bytes
                                mb_act = descargados / (1024 * 1024)
                                mb_tot = total_bytes / (1024 * 1024)
                                pct = int(p * 100)
                                self.after(0, lambda p_val=p, txt=f"⬇️ {mb_act:.1f} MB / {mb_tot:.1f} MB ({pct}%)...": (
                                    progress_bar.set(p_val),
                                    lbl_status.configure(text=txt)
                                ))

                # Descarga finalizada
                self.after(0, lambda: (
                    progress_bar.set(1.0),
                    lbl_status.configure(text=t("ota_instalando"), text_color=COLOR_SUCCESS)
                ))

                # Ejecutar Inno Setup en modo silencioso
                cmd = f'"{ruta_temp}" /VERYSILENT /SUPPRESSMSGBOXES /FORCECLOSEAPPLICATIONS'
                subprocess.Popen(cmd, shell=True)

                # Salir de la app para desbloquear archivos en disco
                def _salir():
                    try:
                        self.quit()
                        self.destroy()
                    except Exception:
                        pass
                    os._exit(0)

                self.after(1200, _salir)

            except Exception as e:
                err_msg = str(e)
                self.after(0, lambda: (
                    lbl_status.configure(text=t("ota_error", err=err_msg), text_color=COLOR_DANGER),
                    btn_ota.configure(state="normal", text="🔄 Reintentar OTA"),
                    btn_web.configure(state="normal")
                ))

        threading.Thread(target=_worker, daemon=True).start()

    def _abrir_modal_actualizacion(self, info):
        """Ventana emergente estilizada con el changelog y opciones de actualización OTA / Web."""
        modal = ctk.CTkToplevel(self)
        modal.title(f"Actualización de KernossIA ({info['tag']})")
        modal.geometry("560x490")
        modal.minsize(500, 420)
        modal.configure(fg_color=COLOR_BG_DARK)
        modal.transient(self)
        modal.grab_set()

        ctk.CTkLabel(modal, text="🚀 Nueva Versión Disponible",
                     font=("Segoe UI", 18, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(18, 4))
        ctk.CTkLabel(modal, text=f"Versión instalada: v{VERSION_APP}   ➜   Última versión: {info['tag']}",
                     font=("Segoe UI", 12), text_color=COLOR_TEXT_MAIN).pack(pady=(0, 10))

        ctk.CTkLabel(modal, text="Novedades de esta versión:",
                     font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=30, pady=(4, 2))

        txt_notas = ctk.CTkTextbox(modal, height=140, fg_color=COLOR_BG_CARD,
                                  border_width=1, border_color=COLOR_BORDER,
                                  text_color=COLOR_TEXT_MAIN, font=("Segoe UI", 11))
        txt_notas.pack(fill="both", expand=True, padx=30, pady=(0, 8))
        txt_notas.insert("end", info.get("notas", "Mejoras de rendimiento, estabilidad y nuevas funciones."))
        txt_notas.configure(state="disabled")

        # Barra de progreso y texto de estado de descarga OTA
        progress_bar = ctk.CTkProgressBar(modal, height=12, progress_color=COLOR_ACCENT_CYAN)
        lbl_status = ctk.CTkLabel(modal, text="", font=("Segoe UI", 10, "bold"), text_color=COLOR_TEXT_MUTED)

        frame_modal_btns = ctk.CTkFrame(modal, fg_color="transparent")
        frame_modal_btns.pack(fill="x", padx=30, pady=(8, 20))

        btn_ota = ctk.CTkButton(
            frame_modal_btns,
            text=t("ota_btn_actualizar"),
            font=("Segoe UI", 12, "bold"),
            fg_color=COLOR_ACCENT_PRIMARY,
            hover_color=COLOR_ACCENT_HOVER,
            height=38,
            command=lambda: self._descargar_e_instalar_ota(info, progress_bar, lbl_status, btn_ota, btn_web, modal)
        )
        btn_ota.pack(side="left", fill="x", expand=True, padx=(0, 6))

        btn_web = ctk.CTkButton(
            frame_modal_btns,
            text=t("ota_btn_web"),
            font=("Segoe UI", 11),
            fg_color=COLOR_BG_CARD,
            border_width=1,
            border_color=COLOR_BORDER,
            hover_color=COLOR_BG_SURFACE,
            height=38,
            command=lambda: [webbrowser.open(info.get("html_url", info["url"])), modal.destroy()]
        )
        btn_web.pack(side="left", padx=(0, 6))

        ctk.CTkButton(
            frame_modal_btns,
            text="✕",
            font=("Segoe UI", 12),
            fg_color=COLOR_BG_SURFACE,
            hover_color=COLOR_DANGER_HOVER,
            height=38,
            width=40,
            command=modal.destroy
        ).pack(side="right")

    def _comprobar_changelog_post_actualizacion(self):
        """Muestra una ventana modal con las novedades de la versión actual SOLO UNA VEZ tras actualizar."""
        ruta_version_vista = os.path.expanduser("~/.kernoss_version_seen.json")
        version_vista = ""
        try:
            if os.path.exists(ruta_version_vista):
                with open(ruta_version_vista, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    version_vista = data.get("ultima_version_vista", "")
        except Exception:
            pass

        if version_vista != VERSION_APP:
            # Guardar que ya se vio esta versión
            try:
                with open(ruta_version_vista, "w", encoding="utf-8") as f:
                    json.dump({"ultima_version_vista": VERSION_APP}, f)
            except Exception:
                pass

            # Mostrar modal de bienvenida con novedades
            self._mostrar_modal_bienvenida_changelog()

    def _mostrar_modal_bienvenida_changelog(self):
        """Ventana modal elegante con el changelog de la nueva versión."""
        modal = ctk.CTkToplevel(self)
        modal.title(t("modal_novedades_titulo", version=VERSION_APP))
        modal.geometry("620x540")
        modal.minsize(540, 460)
        modal.configure(fg_color=COLOR_BG_DARK)
        modal.transient(self)
        modal.grab_set()

        # Header
        ctk.CTkLabel(modal, text=t("modal_novedades_titulo", version=VERSION_APP),
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(22, 4))
        ctk.CTkLabel(modal, text=t("modal_novedades_subtitulo"),
                     font=("Segoe UI", 12), text_color=COLOR_TEXT_MUTED).pack(pady=(0, 14))

        # Tarjeta de novedades
        frame_box = ctk.CTkScrollableFrame(modal, fg_color=COLOR_BG_CARD,
                                           border_width=1, border_color=COLOR_BORDER,
                                           corner_radius=12)
        frame_box.pack(fill="both", expand=True, padx=25, pady=(0, 16))

        novedades = [
            (t("nov_item_soporte_tit"), t("nov_item_soporte_desc")),
            (t("nov_item_multi_tit"), t("nov_item_multi_desc")),
            (t("nov_item_rol_tit"), t("nov_item_rol_desc")),
            (t("nov_item_mod_tit"), t("nov_item_mod_desc")),
            (t("nov_item_cloud_tit"), t("nov_item_cloud_desc")),
            (t("nov_item_hwid_tit"), t("nov_item_hwid_desc")),
        ]

        for titulo, desc in novedades:
            item = ctk.CTkFrame(frame_box, fg_color=COLOR_BG_CARD_LIGHT, corner_radius=8,
                                border_width=1, border_color=COLOR_BORDER)
            item.pack(fill="x", pady=5, padx=2)
            ctk.CTkLabel(item, text=titulo, font=("Segoe UI", 12, "bold"),
                         text_color=COLOR_ACCENT_CYAN).pack(anchor="w", padx=12, pady=(8, 2))
            ctk.CTkLabel(item, text=desc, font=("Segoe UI", 11),
                         text_color=COLOR_TEXT_MAIN, wraplength=500, justify="left").pack(anchor="w", padx=12, pady=(0, 8))

        # Botón Empezar
        ctk.CTkButton(modal, text=t("modal_novedades_btn_empezar"),
                      font=("Segoe UI", 13, "bold"), height=42,
                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                      command=modal.destroy).pack(fill="x", padx=25, pady=(0, 20))

    def _al_cerrar(self):
        self._vigilante_activo = False
        self.quit()
        self.destroy()
        os._exit(0)


# ══════════════════════════════════════════════════════════════════════════════
#  PUNTO DE ENTRADA
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    __spec__ = None

    # Intentar recuperar sesión desde token JWT guardado
    _, sesion = token_guardado()

    if not sesion:
        pantalla = PantallaLogin()
        pantalla.mainloop()
        sesion = pantalla.usuario_autenticado

    if sesion:
        app = DashboardEstudios(sesion)
        app.protocol("WM_DELETE_WINDOW", app._al_cerrar)
        app.mainloop() 