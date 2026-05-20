import customtkinter as ctk
from tkinter import messagebox, filedialog
from openai import OpenAI  # <--- Cambiado de 'import ollama' para compatibilidad con Groq
from docx import Document
import threading
import os

# Configuración de apariencia
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class AppExamenIA(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("AI Exam Generator & Evaluator (Groq Cloud Edition)")
        self.geometry("1100x800")

        # ───── CONFIGURACIÓN DEL CLIENTE CLOUD GRATUITO ─────
        # La clave va integrada dentro del paquete ejecutable .app final
        self.cliente_groq = OpenAI(
            base_url="https://api.groq.com/openai/v1",
            api_key="TU API KEY"  # <--- Pega aquí tu gsk_... de console.groq.com
        )

        # Variables de estado
        self.examen_en_memoria = ""
        self.historial_conversacion = []
        self.hilo_activo = None

        # Estructura interna para almacenar múltiples exámenes guardados en el historial
        self.chats_guardados = {}
        self.id_chat_actual = None

        # --- INSTRUCCIONES ---
        self.instrucciones_base = """Eres un evaluador académico profesional. El examen debe tener:
1. Un título relevante.
2. Preguntas de opción múltiple (A-E) o completar huecos (cantidad según usuario).
3. La mitad de preguntas de desarrollo (explicar conceptos).
REGLA CRÍTICA: No des las respuestas hasta que el usuario responda. No pongas las respuestas correctas en el examen. Solo el enunciado de las preguntas. El usuario responderá y tú evaluarás cada respuesta, dando una calificación final al finalizar."""

        self.setup_ui()

    def setup_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # ───── PANEL LATERAL (Configuración) ─────
        self.sidebar = ctk.CTkFrame(self, width=280, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        
        ctk.CTkLabel(self.sidebar, text="Configuración", font=("Segoe UI", 20, "bold")).pack(pady=20)

        self.entry_nombre = ctk.CTkEntry(self.sidebar, placeholder_text="Tu nombre...")
        self.entry_nombre.pack(fill="x", padx=20, pady=10)

        # Estatus del Motor Cloud
        ctk.CTkLabel(self.sidebar, text="Servicio de IA:", font=("Segoe UI", 12)).pack(pady=(10, 0))
        self.status_label = ctk.CTkLabel(self.sidebar, text="🟢 Groq Cloud Listo", text_color="green", font=("Segoe UI", 12, "bold"))
        self.status_label.pack(fill="x", padx=20, pady=5)

        # --- SECCIÓN DE HISTORIAL DE EXÁMENES INTEGRADA ---
        lbl_historial = ctk.CTkLabel(self.sidebar, text="🗂️ Historial de Exámenes", font=("Segoe UI", 12, "bold"), text_color="gray")
        lbl_historial.pack(anchor="w", padx=20, pady=(15, 5))

        self.btn_nuevo_chat = ctk.CTkButton(self.sidebar, text="+ Nuevo Examen", fg_color="#1f6aa5", hover_color="#144d75", command=self.limpiar_pantalla_examen)
        self.btn_nuevo_chat.pack(fill="x", padx=20, pady=5)

        # Contenedor dinámico con scrollbar para listar los exámenes generados
        self.scroll_historial = ctk.CTkScrollableFrame(self.sidebar, height=180, fg_color="transparent")
        self.scroll_historial.pack(fill="x", padx=10, pady=5)
        # --------------------------------------------------

        ctk.CTkLabel(self.sidebar, text="Tema del examen:", font=("Segoe UI", 12)).pack(pady=(15,0))
        self.txt_tema = ctk.CTkTextbox(self.sidebar, height=100)
        self.txt_tema.pack(fill="x", padx=20, pady=10)

        self.btn_generar = ctk.CTkButton(self.sidebar, text="Generar Examen", command=self.iniciar_generacion)
        self.btn_generar.pack(fill="x", padx=20, pady=10)

        self.btn_exportar = ctk.CTkButton(self.sidebar, text="Exportar Word", fg_color="transparent", border_width=2, command=self.exportar_word)
        self.btn_exportar.pack(fill="x", padx=20, pady=10)

        # ───── ÁREA PRINCIPAL ─────
        self.main_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.main_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.main_frame.grid_rowconfigure(0, weight=1)
        self.main_frame.grid_columnconfigure(0, weight=1)

        self.output_text = ctk.CTkTextbox(self.main_frame, font=("Consolas", 13))
        self.output_text.grid(row=0, column=0, sticky="nsew", pady=(0, 10))

        self.input_user_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        self.input_user_frame.grid(row=1, column=0, sticky="ew")
        self.input_user_frame.grid_columnconfigure(0, weight=1)

        self.entry_respuesta = ctk.CTkEntry(self.input_user_frame, placeholder_text="Escribe tus respuestas aquí...")
        self.entry_respuesta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_respuesta.bind("<Return>", lambda e: self.enviar_respuesta())

        self.btn_enviar = ctk.CTkButton(self.input_user_frame, text="Enviar", width=100, command=self.enviar_respuesta)
        self.btn_enviar.grid(row=0, column=1)

    # ───── LÓGICA DE IA ─────

    def iniciar_generacion(self):
        tema = self.txt_tema.get("1.0", "end-1c").strip()
        if not tema:
            messagebox.showwarning("Atención", "Escribe un tema para el examen.")
            return
        
        self.output_text.delete("1.0", "end")
        self.output_text.insert("end", f"Generando examen sobre: {tema}...\n\n")
        self.btn_generar.configure(state="disabled")
        self.status_label.configure(text="🟡 Elaborando Examen...", text_color="orange")
        
        # Guardar identificador del chat basado en el tema elegido
        titulo_examen = tema[:20] + "..." if len(tema) > 20 else tema
        self.id_chat_actual = titulo_examen
        self.historial_conversacion = [
            {'role': 'system', 'content': self.instrucciones_base},
            {'role': 'user', 'content': f"Hazme un examen sobre: {tema}"}
        ]
        self.chats_guardados[self.id_chat_actual] = self.historial_conversacion

        # Ejecutar en hilo separado para no bloquear la UI de CustomTkinter
        thread = threading.Thread(target=self.proceso_groq, args=(tema,))
        thread.start()

    def proceso_groq(self, tema):
        try:
            # Llamada exacta a la API con el modelo inteligente de 70B
            response = self.cliente_groq.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=self.historial_conversacion,
                stream=True,
                temperature=0.3
            )
            
            full_response = ""
            for chunk in response:
                if chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    full_response += content
                    # Actualizar la interfaz gráfica de forma segura desde el hilo secundario
                    self.after(0, self.update_output, content)
            
            self.examen_en_memoria = full_response
            self.historial_conversacion.append({'role': 'assistant', 'content': full_response})
            
            # Guardar el estado final del examen generado en el historial dinámico
            if self.id_chat_actual:
                self.chats_guardados[self.id_chat_actual] = list(self.historial_conversacion)
                self.after(0, self.actualizar_lista_historial)

            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status_label.configure(text="🟢 Groq Cloud Listo", text_color="green"))
            
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error de Conexión Cloud", str(e)))
            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status_label.configure(text="🔴 Error de red", text_color="red"))

    def enviar_respuesta(self):
        msg = self.entry_respuesta.get().strip()
        if not msg: return

        # Si el usuario escribe antes de pulsar "Generar Examen", creamos un entorno dinámico
        if not self.historial_conversacion:
            titulo_examen = msg[:20] + "..." if len(msg) > 20 else msg
            self.id_chat_actual = titulo_examen
            self.chats_guardados[self.id_chat_actual] = self.historial_conversacion

        self.output_text.insert("end", f"\n\n👤 TÚ: {msg}\n\n🤖 IA: ")
        self.entry_respuesta.delete(0, "end")
        self.status_label.configure(text="🟡 Corrigiendo...", text_color="orange")
        
        self.historial_conversacion.append({'role': 'user', 'content': msg})
        
        thread = threading.Thread(target=self.proceso_respuesta_ia)
        thread.start()

    def proceso_respuesta_ia(self):
        try:
            response = self.cliente_groq.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=self.historial_conversacion,
                stream=True,
                temperature=0.2
            )
            
            full_response = ""
            for chunk in response:
                if chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    full_response += content
                    self.after(0, self.update_output, content)
            
            self.historial_conversacion.append({'role': 'assistant', 'content': full_response})
            
            # Actualizar el registro del historial tras la corrección de la IA
            if self.id_chat_actual:
                self.chats_guardados[self.id_chat_actual] = list(self.historial_conversacion)
                self.after(0, self.actualizar_lista_historial)

            self.after(0, lambda: self.status_label.configure(text="🟢 Groq Cloud Listo", text_color="green"))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", str(e)))
            self.after(0, lambda: self.status_label.configure(text="🔴 Error de red", text_color="red"))

    def update_output(self, texto):
        self.output_text.insert("end", texto)
        self.output_text.see("end")

    # ───── NUEVAS FUNCIONES DE SOPORTE PARA EL HISTORIAL ─────
    def actualizar_lista_historial(self):
        # Limpiar los elementos gráficos viejos del panel de scroll
        for widget in self.scroll_historial.winfo_children():
            widget.destroy()

        # Renderizar cada examen guardado con su botón de carga y su papelera de borrado
        for titulo in list(self.chats_guardados.keys()):
            color_btn = "#1f6aa5" if titulo == self.id_chat_actual else "#2b2b2b"
            
            item_frame = ctk.CTkFrame(self.scroll_historial, fg_color="transparent")
            item_frame.pack(fill="x", padx=2, pady=2)
            
            # Botón para restaurar la sesión del examen
            btn_item = ctk.CTkButton(
                item_frame, 
                text=f"📝 {titulo}", 
                fg_color=color_btn,
                hover_color="#3a3a3a",
                anchor="w",
                height=32,
                command=lambda t=titulo: self.cargar_examen_del_historial(t)
            )
            btn_item.pack(side="left", fill="x", expand=True, padx=(0, 4))

            # Botón de papelera independiente para eliminar la sesión
            btn_borrar = ctk.CTkButton(
                item_frame,
                text="🗑️",
                width=32,
                height=32,
                fg_color="#742a2a" if titulo == self.id_chat_actual else "#3a3a3a",
                hover_color="#ba3c3c",
                command=lambda t=titulo: self.eliminar_examen_especifico(t)
            )
            btn_borrar.pack(side="right")

    def cargar_examen_del_historial(self, titulo):
        self.id_chat_actual = titulo
        self.historial_conversacion = list(self.chats_guardados[titulo])
        
        self.output_text.delete("1.0", "end")
        
        # Reconstruir visualmente el hilo de la conversación en la caja de texto central
        for msg in self.historial_conversacion:
            if msg['role'] == 'system':
                continue
            elif msg['role'] == 'user':
                # Si el mensaje contiene el comando inicial de generación, lo saltamos para no afear el chat
                if "Hazme un examen sobre:" in msg['content']:
                    continue
                self.output_text.insert("end", f"\n\n👤 TÚ: {msg['content']}")
            else:
                # Comprobamos si es el cuerpo principal del examen o respuestas siguientes de la IA
                if not self.output_text.get("1.0", "end").strip():
                    self.output_text.insert("end", msg['content'])
                else:
                    self.output_text.insert("end", f"\n\n🤖 IA: {msg['content']}")
                    
        self.output_text.see("end")
        self.actualizar_lista_historial()

    def eliminar_examen_especifico(self, titulo):
        if titulo in self.chats_guardados:
            del self.chats_guardados[titulo]
        
        # Si el usuario ha borrado el examen que está viendo actualmente en pantalla, reseteamos la UI
        if self.id_chat_actual == titulo:
            self.limpiar_pantalla_examen()
        else:
            self.actualizar_lista_historial()

    def limpiar_pantalla_examen(self):
        self.output_text.delete("1.0", "end")
        self.txt_tema.delete("1.0", "end")
        self.historial_conversacion = []
        self.examen_en_memoria = ""
        self.id_chat_actual = None
        self.actualizar_lista_historial()
    # ---------------------------------------------------------

    # ───── ARCHIVOS ─────

    def exportar_word(self):
        texto = self.output_text.get("1.0", "end-1c")
        if not texto.strip():
            messagebox.showwarning("Vacío", "No hay contenido para exportar.")
            return

        nombre = self.entry_nombre.get().strip() or "Usuario"
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=f"Examen_{nombre}.docx"
        )

        if file_path:
            doc = Document()
            doc.add_heading(f"Examen de {nombre}", 0)
            doc.add_paragraph(texto)
            doc.save(file_path)
            messagebox.showinfo("Éxito", f"Archivo guardado en:\n{file_path}")

if __name__ == "__main__":
    app = AppExamenIA()
    app.mainloop()