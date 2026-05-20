import customtkinter as ctk
from tkinter import messagebox, filedialog
from openai import OpenAI  # <--- Cambiado de 'import ollama'
from docx import Document
import threading
import os

# Configuración visual
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class SolucionadorIA(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("AI Problem Solver - Alto Rendimiento Cloud")
        self.geometry("1000x700")

        # ───── CONFIGURACIÓN DEL CLIENTE GROQ ─────
        # RECUERDA: Tus amigos solo necesitan internet, la clave ya va integrada dentro de la app (.app)
        self.cliente_groq = OpenAI(
            base_url="https://api.groq.com/openai/v1",
            api_key="TU API KEY"  # <--- Pega aquí tu gsk_... de console.groq.com
        )

        # Variables de lógica
        self.historial_conversacion = []
        self.instrucciones = (
            "Eres un Solucionador de Problemas de alto rendimiento. "
            "Tu enfoque es: Analizar el problema -> Identificar la causa -> Dar solución técnica/práctica. -> Dar consejos para evitarlo en el futuro. "
            "Responde siempre con estructura de puntos clave y da conversación para acompañar y dar consejos."
        )

        # Estructura interna para almacenar múltiples chats guardados
        self.chats_guardados = {}
        self.id_chat_actual = None

        self.setup_ui()

    def setup_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # ───── PANEL LATERAL ─────
        self.sidebar = ctk.CTkFrame(self, width=250, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        
        ctk.CTkLabel(self.sidebar, text="🛠️ Solver IA", font=("Segoe UI", 24, "bold")).pack(pady=20)

        self.entry_nombre = ctk.CTkEntry(self.sidebar, placeholder_text="¿Cuál es tu nombre?")
        self.entry_nombre.pack(fill="x", padx=20, pady=10)

        # --- SECCIÓN DE HISTORIAL DE CHATS INTEGRADA ---
        lbl_historial = ctk.CTkLabel(self.sidebar, text="🗂️ Historial de Chats", font=("Segoe UI", 12, "bold"), text_color="gray")
        lbl_historial.pack(anchor="w", padx=20, pady=(15, 5))

        self.btn_nuevo_chat = ctk.CTkButton(self.sidebar, text="+ Nuevo Chat", fg_color="#1f6aa5", hover_color="#144d75", command=self.limpiar_chat)
        self.btn_nuevo_chat.pack(fill="x", padx=20, pady=5)

        # Contenedor con scrollbar para listar los chats previos
        self.scroll_historial = ctk.CTkScrollableFrame(self.sidebar, height=180, fg_color="transparent")
        self.scroll_historial.pack(fill="x", padx=10, pady=5)
        # -----------------------------------------------

        self.btn_exportar = ctk.CTkButton(self.sidebar, text="Exportar a Word", 
                                         fg_color="#2c3e50", hover_color="#34495e",
                                         command=self.exportar_a_word)
        self.btn_exportar.pack(fill="x", padx=20, pady=10)

        self.btn_limpiar = ctk.CTkButton(self.sidebar, text="🧹 Limpiar Chat", 
                                        fg_color="transparent", border_width=1,
                                        command=self.limpiar_chat)
        self.btn_limpiar.pack(fill="x", padx=20, pady=10)

        ctk.CTkLabel(self.sidebar, text="Motor de IA:", font=("Segoe UI", 12)).pack(side="bottom", pady=5)
        self.status_label = ctk.CTkLabel(self.sidebar, text="🟢 Groq Cloud Listo", text_color="green")
        self.status_label.pack(side="bottom", pady=(0, 20))

        # ───── ÁREA DE CHAT ─────
        self.chat_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.chat_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.chat_frame.grid_rowconfigure(0, weight=1)
        self.chat_frame.grid_columnconfigure(0, weight=1)

        self.txt_chat = ctk.CTkTextbox(self.chat_frame, font=("Segoe UI", 14), state="disabled", wrap="word")
        self.txt_chat.grid(row=0, column=0, sticky="nsew", pady=(0, 15))

        # Entrada de usuario
        self.input_frame = ctk.CTkFrame(self.chat_frame, fg_color="transparent")
        self.input_frame.grid(row=1, column=0, sticky="ew")
        self.input_frame.grid_columnconfigure(0, weight=1)

        self.entry_pregunta = ctk.CTkEntry(self.input_frame, placeholder_text="Describe el problema aquí...", height=40)
        self.entry_pregunta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_pregunta.bind("<Return>", lambda e: self.enviar_consulta())

        self.btn_enviar = ctk.CTkButton(self.input_frame, text="Analizar", width=120, height=40, command=self.enviar_consulta)
        self.btn_enviar.grid(row=0, column=1)

    # ───── LÓGICA DE INTERACCIÓN ─────

    def agregar_texto(self, emisor, texto):
        self.txt_chat.configure(state="normal")
        if emisor == "IA":
            self.txt_chat.insert("end", f"\n IA:\n{texto}")
        else:
            self.txt_chat.insert("end", f"\n\n {emisor}:\n{texto}\n")
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")

    def enviar_consulta(self):
        pregunta = self.entry_pregunta.get().strip()
        nombre = self.entry_nombre.get().strip() or "Tú"
        
        if not pregunta:
            return

        # Si es el primer mensaje de una sesión limpia, inicializamos su título en el historial
        if not self.historial_conversacion:
            # Recortamos la pregunta inicial para usarla como nombre del botón del historial
            titulo_chat = pregunta[:20] + "..." if len(pregunta) > 20 else pregunta
            self.id_chat_actual = titulo_chat
            self.chats_guardados[self.id_chat_actual] = self.historial_conversacion

        self.agregar_texto(nombre, pregunta)
        self.entry_pregunta.delete(0, "end")
        self.status_label.configure(text="🟡 Pensando (Groq)...", text_color="orange")
        
        # Guardar en memoria (formato compatible con OpenAI/Groq)
        self.historial_conversacion.append({'role': 'user', 'content': pregunta})
        
        # Lanzar hilo para no congelar la UI de CustomTkinter
        threading.Thread(target=self.proceso_ia, daemon=True).start()

    def proceso_ia(self):
        try:
            self.txt_chat.configure(state="normal")
            self.txt_chat.insert("end", "\n IA:\n")
            
            # Llamada exacta a la API Cloud de Groq utilizando streaming
            response = self.cliente_groq.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{'role': 'system', 'content': self.instrucciones}] + self.historial_conversacion,
                stream=True,
                temperature=0.2
            )

            respuesta_completa = ""
            for chunk in response:
                # Comprobar que el fragmento de texto tiene contenido
                if chunk.choices[0].delta.content:
                    contenido = chunk.choices[0].delta.content
                    respuesta_completa += contenido
                    # Actualizar texto en la UI de CustomTkinter
                    self.after(0, self.stream_update, contenido)
            
            self.historial_conversacion.append({'role': 'assistant', 'content': respuesta_completa})
            
            # Almacenar la conversación actualizada en nuestra matriz y regenerar los botones de la barra lateral
            if self.id_chat_actual:
                self.chats_guardados[self.id_chat_actual] = list(self.historial_conversacion)
                self.after(0, self.actualizar_lista_historial)

            self.after(0, lambda: self.status_label.configure(text="🟢 Groq Cloud Listo", text_color="green"))
            
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", f"Fallo de conexión Cloud: {str(e)}"))
            self.after(0, lambda: self.status_label.configure(text="🔴 Error de red", text_color="red"))
        finally:
            self.after(0, lambda: self.txt_chat.configure(state="disabled"))

    def stream_update(self, contenido):
        self.txt_chat.configure(state="normal")
        self.txt_chat.insert("end", contenido)
        self.txt_chat.see("end")
        self.txt_chat.configure(state="disabled")

    # --- NUEVAS FUNCIONES DE SOPORTE PARA EL HISTORIAL ---
    def actualizar_lista_historial(self):
        # Limpiar los widgets antiguos del panel de scroll
        for widget in self.scroll_historial.winfo_children():
            widget.destroy()

        # Crear un contenedor por cada fila (botón de chat + botón borrar)
        for titulo in list(self.chats_guardados.keys()):
            color_btn = "#1f6aa5" if titulo == self.id_chat_actual else "#2b2b2b"
            
            # Frame contenedor horizontal fino para la línea del historial
            item_frame = ctk.CTkFrame(self.scroll_historial, fg_color="transparent")
            item_frame.pack(fill="x", padx=2, pady=2)
            
            # Botón principal para cargar el chat
            btn_item = ctk.CTkButton(
                item_frame, 
                text=f"💬 {titulo}", 
                fg_color=color_btn,
                hover_color="#3a3a3a",
                anchor="w",
                height=32,
                command=lambda t=titulo: self.cargar_chat_del_historial(t)
            )
            btn_item.pack(side="left", fill="x", expand=True, padx=(0, 4))

            # Botón independiente con la papelera para borrar este chat específico
            btn_borrar = ctk.CTkButton(
                item_frame,
                text="🗑️",
                width=32,
                height=32,
                fg_color="#742a2a" if titulo == self.id_chat_actual else "#3a3a3a",
                hover_color="#ba3c3c",
                command=lambda t=titulo: self.eliminar_chat_especifico(t)
            )
            btn_borrar.pack(side="right")

    def cargar_chat_del_historial(self, titulo):
        self.id_chat_actual = titulo
        self.historial_conversacion = list(self.chats_guardados[titulo])
        
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        
        nombre_usuario = self.entry_nombre.get().strip() or "Tú"
        for msg in self.historial_conversacion:
            if msg['role'] == 'user':
                self.txt_chat.insert("end", f"\n\n {nombre_usuario}:\n{msg['content']}\n")
            else:
                self.txt_chat.insert("end", f"\n IA:\n{msg['content']}")
                
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")
        self.actualizar_lista_historial()

    def eliminar_chat_especifico(self, titulo):
        # Elimina el elemento seleccionado del diccionario en memoria
        if titulo in self.chats_guardados:
            del self.chats_guardados[titulo]
        
        # Si el chat que se ha eliminado coincide con el que el usuario estaba leyendo actualmente, reseteamos la UI
        if self.id_chat_actual == titulo:
            self.id_chat_actual = None
            self.historial_conversacion = []
            self.txt_chat.configure(state="normal")
            self.txt_chat.delete("1.0", "end")
            self.txt_chat.configure(state="disabled")

        self.actualizar_lista_historial()
    # -----------------------------------------------------

    def limpiar_chat(self):
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        self.txt_chat.configure(state="disabled")
        self.historial_conversacion = []
        self.id_chat_actual = None
        self.actualizar_lista_historial()

    def exportar_a_word(self):
        if not self.historial_conversacion:
            messagebox.showwarning("Atención", "No hay nada que exportar aún.")
            return

        nombre = self.entry_nombre.get().strip() or "Usuario"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Solucion_Problemas_{nombre}.docx",
            title="Guardar informe técnico"
        )

        if path:
            try:
                doc = Document()
                doc.add_heading(f'Informe de Solución - {nombre}', 0)
                
                for msg in self.historial_conversacion:
                    rol = "Tú" if msg['role'] == 'user' else "IA"
                    p = doc.add_paragraph()
                    p.add_run(f"{rol}: ").bold = True
                    p.add_run(msg['content'])
                
                doc.save(path)
                messagebox.showinfo("Éxito", f"Informe guardado en:\n{path}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar el archivo: {e}")

if __name__ == "__main__":
    app = SolucionadorIA()
    app.mainloop()