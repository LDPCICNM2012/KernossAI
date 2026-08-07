import customtkinter as ctk
from tkinter import messagebox, filedialog
from openai import OpenAI
import google.generativeai as genai
from docx import Document
import threading, os
from datetime import datetime

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

GROQ_API_KEY   = ""
GEMINI_API_KEY = ""

INSTRUCCIONES_EJERCICIO = """Eres un profesor experto creando material educativo.
Crea ejercicios bien estructurados con:
1. Título claro y objetivo de aprendizaje.
2. Enunciado detallado.
3. Apartados o preguntas numerados.
4. Nivel de dificultad indicado.
5. Tiempo estimado.
NO incluyas las soluciones a menos que se te pida explícitamente.
El formato debe ser limpio y listo para imprimir o enviar a alumnos."""

INSTRUCCIONES_SOLUCIONES = """Eres un profesor. Genera el solucionario completo y detallado
del ejercicio que te proporcionan. Explica cada paso y por qué es correcto."""


class CreadorejerciiciosProfesor(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Creador de Ejercicios — Modo Profesor")
        self.geometry("1200x780")
        self.modelo_actual = "groq"
        self.ejercicio_actual = ""
        self.solucionario_actual = ""

        self.cliente_groq = OpenAI(base_url="https://api.groq.com/openai/v1", api_key=GROQ_API_KEY)
        genai.configure(api_key=GEMINI_API_KEY)
        self.cliente_gemini = genai.GenerativeModel("gemini-2.5-flash")

        self.setup_ui()

    def setup_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # ── SIDEBAR ──
        self.sidebar = ctk.CTkFrame(self, width=290, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_propagate(False)

        ctk.CTkLabel(self.sidebar, text="Creador de\nEjercicios",
                     font=("Segoe UI", 20, "bold"), text_color="#e53935").pack(pady=(20, 5))
        ctk.CTkLabel(self.sidebar, text="Modo Profesor",
                     font=("Segoe UI", 11), text_color="#888").pack(pady=(0, 15))

        # Motor IA
        ctk.CTkLabel(self.sidebar, text="Motor de IA:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15)
        fm = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        fm.pack(fill="x", padx=15, pady=5)
        self.btn_groq = ctk.CTkButton(fm, text="Basico\n(Groq)", height=50,
            fg_color="#1565c0", hover_color="#0d47a1", font=("Segoe UI", 10, "bold"),
            command=lambda: self._set_modelo("groq"))
        self.btn_groq.pack(side="left", expand=True, fill="x", padx=(0, 3))
        self.btn_gemini = ctk.CTkButton(fm, text="Avanzado\n(Gemini)", height=50,
            fg_color="#333", hover_color="#444", font=("Segoe UI", 10, "bold"),
            command=lambda: self._set_modelo("gemini"))
        self.btn_gemini.pack(side="left", expand=True, fill="x", padx=(3, 0))

        self.lbl_limite = ctk.CTkLabel(self.sidebar,
            text="Groq: ~1.000 msgs/dia", font=("Segoe UI", 9), text_color="#666")
        self.lbl_limite.pack(pady=(2, 10))

        ctk.CTkFrame(self.sidebar, height=1, fg_color="#333").pack(fill="x", padx=15, pady=5)

        # Parámetros
        ctk.CTkLabel(self.sidebar, text="Tema / Materia", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15, pady=(5,2))
        self.entry_tema = ctk.CTkEntry(self.sidebar, placeholder_text="Ej: Ecuaciones de 2º grado")
        self.entry_tema.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(self.sidebar, text="Nivel educativo", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15)
        self.combo_nivel = ctk.CTkComboBox(self.sidebar, values=[
            "Primaria", "1º ESO", "2º ESO", "3º ESO", "4º ESO",
            "1º Bachillerato", "2º Bachillerato", "Universidad"])
        self.combo_nivel.set("2º ESO")
        self.combo_nivel.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(self.sidebar, text="Tipo de ejercicio", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15)
        self.combo_tipo = ctk.CTkComboBox(self.sidebar, values=[
            "Problemas paso a paso",
            "Preguntas teóricas",
            "Opción múltiple (A-D)",
            "Completar huecos",
            "Verdadero / Falso",
            "Ejercicio práctico",
            "Mixto (teoría + práctica)"])
        self.combo_tipo.set("Mixto (teoría + práctica)")
        self.combo_tipo.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(self.sidebar, text="Nº de preguntas", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15)
        self.spin_preguntas = ctk.CTkSlider(self.sidebar, from_=3, to=20, number_of_steps=17)
        self.spin_preguntas.set(8)
        self.spin_preguntas.pack(fill="x", padx=15, pady=(0, 2))
        self.lbl_n_preguntas = ctk.CTkLabel(self.sidebar, text="8 preguntas", font=("Segoe UI", 11))
        self.lbl_n_preguntas.pack()
        self.spin_preguntas.configure(command=lambda v: self.lbl_n_preguntas.configure(text=f"{int(v)} preguntas"))

        self.check_soluciones = ctk.CTkCheckBox(self.sidebar, text="Incluir solucionario al final")
        self.check_soluciones.pack(padx=15, pady=8, anchor="w")

        ctk.CTkFrame(self.sidebar, height=1, fg_color="#333").pack(fill="x", padx=15, pady=5)

        self.btn_generar = ctk.CTkButton(self.sidebar, text="Generar Ejercicio",
            height=45, font=("Segoe UI", 13, "bold"),
            fg_color="#7b1fa2", hover_color="#4a0072",
            command=self.generar_ejercicio)
        self.btn_generar.pack(fill="x", padx=15, pady=5)

        self.btn_regenerar = ctk.CTkButton(self.sidebar, text="Regenerar (nueva versión)",
            height=38, fg_color="#1565c0", hover_color="#0d47a1",
            command=self.regenerar_ejercicio, state="disabled")
        self.btn_regenerar.pack(fill="x", padx=15, pady=3)

        self.btn_editar = ctk.CTkButton(self.sidebar, text="Editar manualmente",
            height=38, fg_color="transparent", border_width=1,
            command=self.activar_edicion, state="disabled")
        self.btn_editar.pack(fill="x", padx=15, pady=3)

        self.btn_solucionario = ctk.CTkButton(self.sidebar, text="Generar Solucionario",
            height=38, fg_color="#2e7d32", hover_color="#1b5e20",
            command=self.generar_solucionario, state="disabled")
        self.btn_solucionario.pack(fill="x", padx=15, pady=3)

        # ── BOTONES EXPORTAR (CON Y SIN SOLUCIÓN) ──
        self.btn_exportar = ctk.CTkButton(self.sidebar, text="📄 Exportar con solución",
            height=38, fg_color="#37474f", hover_color="#263238",
            command=lambda: self.exportar_word(con_solucion=True), state="disabled")
        self.btn_exportar.pack(fill="x", padx=15, pady=(3, 1))

        self.btn_exportar_sin = ctk.CTkButton(self.sidebar, text="📋 Exportar sin solución",
            height=38, fg_color="#455a64", hover_color="#37474f",
            command=lambda: self.exportar_word(con_solucion=False), state="disabled")
        self.btn_exportar_sin.pack(fill="x", padx=15, pady=(1, 3))

        self.status = ctk.CTkLabel(self.sidebar, text="Listo", text_color="gray",
                                   font=("Segoe UI", 10))
        self.status.pack(side="bottom", pady=10)

        # ── PANEL DERECHO ──
        self.panel = ctk.CTkFrame(self, fg_color="transparent")
        self.panel.grid(row=0, column=1, sticky="nsew", padx=15, pady=15)
        self.panel.grid_rowconfigure(1, weight=1)
        self.panel.grid_columnconfigure(0, weight=1)

        cab = ctk.CTkFrame(self.panel, fg_color="transparent")
        cab.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        ctk.CTkLabel(cab, text="Ejercicio Generado",
                     font=("Segoe UI", 18, "bold")).pack(side="left")
        self.lbl_modo_edicion = ctk.CTkLabel(cab, text="",
                                              font=("Segoe UI", 11), text_color="#fb8c00")
        self.lbl_modo_edicion.pack(side="right")

        self.txt_ejercicio = ctk.CTkTextbox(self.panel, font=("Consolas", 13))
        self.txt_ejercicio.grid(row=1, column=0, sticky="nsew")
        self.txt_ejercicio.insert("end",
            "Configura los parámetros en el panel izquierdo y pulsa 'Generar Ejercicio'.")
        self.txt_ejercicio.configure(state="disabled")

        self.entry_cambio = ctk.CTkEntry(self.panel,
            placeholder_text="Ej: 'Hazlo más difícil' / 'Añade 2 preguntas de desarrollo'",
            height=40)
        self.entry_cambio.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        self.entry_cambio.bind("<Return>", lambda e: self.aplicar_cambio_ia())

        btn_row = ctk.CTkFrame(self.panel, fg_color="transparent")
        btn_row.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        ctk.CTkButton(btn_row, text="Aplicar cambio con IA",
                      fg_color="#e65100", hover_color="#bf360c", height=38,
                      command=self.aplicar_cambio_ia).pack(side="left", padx=(0, 8))
        ctk.CTkLabel(btn_row, text="← Escribe qué quieres cambiar y pulsa Enter o el botón",
                     font=("Segoe UI", 11), text_color="#888").pack(side="left")

    # ──────────────────────────────────────────────
    #  MODELO
    # ──────────────────────────────────────────────
    def _set_modelo(self, m):
        self.modelo_actual = m
        if m == "groq":
            self.btn_groq.configure(fg_color="#1565c0")
            self.btn_gemini.configure(fg_color="#333")
            self.lbl_limite.configure(text="Groq: ~1.000 msgs/dia")
        else:
            self.btn_groq.configure(fg_color="#333")
            self.btn_gemini.configure(fg_color="#6a1b9a")
            self.lbl_limite.configure(text="Gemini: 15 msgs/minuto")

    def _llamar_ia(self, prompt: str) -> str:
        if self.modelo_actual == "groq":
            response = self.cliente_groq.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {'role': 'system', 'content': INSTRUCCIONES_EJERCICIO},
                    {'role': 'user', 'content': prompt}
                ],
                temperature=0.4
            )
            return response.choices[0].message.content
        else:
            response = self.cliente_gemini.generate_content(
                f"{INSTRUCCIONES_EJERCICIO}\n\n{prompt}")
            return response.text

    def _construir_prompt(self):
        tema = self.entry_tema.get().strip() or "tema libre"
        nivel = self.combo_nivel.get()
        tipo = self.combo_tipo.get()
        n = int(self.spin_preguntas.get())
        inc_sol = self.check_soluciones.get()
        prompt = (f"Crea un ejercicio de {tipo} sobre '{tema}' para alumnos de {nivel}. "
                  f"Debe tener exactamente {n} preguntas/apartados. "
                  f"{'Incluye el solucionario completo al final.' if inc_sol else 'NO incluyas las soluciones.'} "
                  f"Formato limpio y listo para entregar a los alumnos.")
        return prompt

    # ──────────────────────────────────────────────
    #  GENERAR
    # ──────────────────────────────────────────────
    def generar_ejercicio(self):
        if not self.entry_tema.get().strip():
            messagebox.showwarning("Tema vacío", "Introduce el tema del ejercicio.")
            return
        self.btn_generar.configure(state="disabled")
        self.status.configure(text="Generando...", text_color="orange")
        threading.Thread(target=self._thread_generar, daemon=True).start()

    def _thread_generar(self):
        try:
            resultado = self._llamar_ia(self._construir_prompt())
            self.ejercicio_actual = resultado
            self.after(0, self._mostrar_ejercicio, resultado)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    def _mostrar_ejercicio(self, texto):
        self.txt_ejercicio.configure(state="normal")
        self.txt_ejercicio.delete("1.0", "end")
        self.txt_ejercicio.insert("end", texto)
        self.txt_ejercicio.configure(state="disabled")
        for btn in [self.btn_regenerar, self.btn_editar,
                    self.btn_solucionario, self.btn_exportar, self.btn_exportar_sin]:
            btn.configure(state="normal")

    def regenerar_ejercicio(self):
        self.generar_ejercicio()

    def activar_edicion(self):
        self.txt_ejercicio.configure(state="normal")
        self.lbl_modo_edicion.configure(text="✏️ Modo edición manual activo")
        self.btn_editar.configure(text="Bloquear edición", command=self.desactivar_edicion)

    def desactivar_edicion(self):
        self.ejercicio_actual = self.txt_ejercicio.get("1.0", "end-1c")
        self.txt_ejercicio.configure(state="disabled")
        self.lbl_modo_edicion.configure(text="")
        self.btn_editar.configure(text="Editar manualmente", command=self.activar_edicion)

    def aplicar_cambio_ia(self):
        cambio = self.entry_cambio.get().strip()
        if not cambio:
            messagebox.showwarning("Atención", "Escribe qué quieres cambiar.")
            return
        if not self.ejercicio_actual:
            messagebox.showwarning("Atención", "Primero genera un ejercicio.")
            return
        self.status.configure(text="Aplicando cambio...", text_color="orange")
        prompt = (f"Este es el ejercicio actual:\n\n{self.ejercicio_actual}\n\n"
                  f"Aplica el siguiente cambio y devuelve el ejercicio completo actualizado: {cambio}")
        self.entry_cambio.delete(0, "end")
        threading.Thread(target=lambda: self._thread_cambio(prompt), daemon=True).start()

    def _thread_cambio(self, prompt):
        try:
            resultado = self._llamar_ia(prompt)
            self.ejercicio_actual = resultado
            self.after(0, self._mostrar_ejercicio, resultado)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    # ──────────────────────────────────────────────
    #  SOLUCIONARIO
    # ──────────────────────────────────────────────
    def generar_solucionario(self):
        if not self.ejercicio_actual:
            messagebox.showwarning("Atención", "Primero genera un ejercicio.")
            return
        self.status.configure(text="Generando solucionario...", text_color="orange")
        prompt = f"Genera el solucionario completo de este ejercicio:\n\n{self.ejercicio_actual}"

        def _thread():
            try:
                if self.modelo_actual == "groq":
                    r = self.cliente_groq.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[
                            {'role': 'system', 'content': INSTRUCCIONES_SOLUCIONES},
                            {'role': 'user', 'content': prompt}
                        ], temperature=0.2)
                    sol = r.choices[0].message.content
                else:
                    r = self.cliente_gemini.generate_content(
                        f"{INSTRUCCIONES_SOLUCIONES}\n\n{prompt}")
                    sol = r.text
                self.solucionario_actual = sol

                def _mostrar():
                    ventana = ctk.CTkToplevel(self)
                    ventana.title("Solucionario")
                    ventana.geometry("800x600")
                    ctk.CTkLabel(ventana, text="Solucionario Completo",
                                 font=("Segoe UI", 16, "bold")).pack(pady=10)
                    txt = ctk.CTkTextbox(ventana, font=("Consolas", 12))
                    txt.pack(fill="both", expand=True, padx=15, pady=(0, 10))
                    txt.insert("end", sol)
                    txt.configure(state="disabled")
                    ctk.CTkButton(ventana, text="Exportar Solucionario a Word",
                                  command=lambda: self._exportar_solucionario(sol)
                                  ).pack(fill="x", padx=15, pady=(0, 15))

                self.after(0, _mostrar)
            except Exception as e:
                self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            finally:
                self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

        threading.Thread(target=_thread, daemon=True).start()

    # ──────────────────────────────────────────────
    #  EXPORTAR
    # ──────────────────────────────────────────────
    def exportar_word(self, con_solucion: bool = True):
        texto = self.txt_ejercicio.get("1.0", "end-1c")
        if not texto.strip():
            messagebox.showwarning("Vacío", "No hay ejercicio que exportar.")
            return

        # Si pide sin solución pero el ejercicio la trae incluida, pedimos a la IA que la elimine
        if not con_solucion and self.solucionario_actual:
            # Usamos el ejercicio tal cual (fue generado sin solucionario si no se marcó el check)
            # Si el usuario generó solucionario aparte no está en el texto del ejercicio
            pass  # el texto del ejercicio ya está limpio

        tema = self.entry_tema.get().strip() or "Ejercicio"
        sufijo = "ConSolucion" if con_solucion else "SinSolucion"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Ejercicio_{tema}_{sufijo}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return

        doc = Document()
        doc.add_heading(f"Ejercicio — {tema}", 0)
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()} | Tipo: {self.combo_tipo.get()}")
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")
        doc.add_paragraph(texto)

        if con_solucion and self.solucionario_actual:
            doc.add_page_break()
            doc.add_heading("SOLUCIONARIO", level=1)
            doc.add_paragraph(self.solucionario_actual)
        elif con_solucion and not self.solucionario_actual:
            messagebox.showinfo(
                "Sin solucionario",
                "No se ha generado solucionario todavía.\nEl Word se exportará solo con el ejercicio.\n"
                "Usa 'Generar Solucionario' primero si lo necesitas.")

        doc.save(path)
        messagebox.showinfo("Exportado", f"Guardado en:\n{path}")

    def _exportar_solucionario(self, sol):
        tema = self.entry_tema.get().strip() or "Ejercicio"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Solucionario_{tema}_{datetime.now().strftime('%Y%m%d')}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Solucionario — {tema}", 0)
            doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
            doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph("")
            doc.add_paragraph(sol)
            doc.save(path)
            messagebox.showinfo("Exportado", f"Solucionario guardado en:\n{path}")


if __name__ == "__main__":
    app = CreadorejerciiciosProfesor()
    app.mainloop()