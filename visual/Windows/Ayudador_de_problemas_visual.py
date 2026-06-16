import customtkinter as ctk
from tkinter import messagebox, filedialog
from openai import OpenAI
from google import genai
from google.genai import types
from docx import Document
import threading
import os
import json
from datetime import datetime

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

GROQ_API_KEY =""
GEMINI_API_KEY = ""

NOMBRE_SESION = os.environ.get("IMMUNE_NOMBRE", "")
EMAIL_SESION  = os.environ.get("IMMUNE_EMAIL", "")
ROL_SESION    = os.environ.get("IMMUNE_ROL", "Alumno")

# Modelo Gemini más capaz disponible; cambia a "gemini-2.5-flash" si tu key es gratuita
GEMINI_MODEL = "gemini-2.5-flash"


class SolucionadorIA(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("🤖 Ayudante de Problemas IA")
        self.geometry("1050x780")

        # Clientes IA
        self.cliente_groq = OpenAI(
            base_url="https://api.groq.com/openai/v1",
            api_key=GROQ_API_KEY
        )
        self.cliente_gemini = genai.Client(api_key=GEMINI_API_KEY)

        # Estado
        self.modelo_actual = "groq"
        self.historial_conversacion = []

        # Instrucciones independientes por modelo
        self.instrucciones_groq = (
            "Eres un asistente BÁSICO y RÁPIDO (Groq). "
            "Tu objetivo es ayudar mediar y resolver de la mejor manera posible lo que te pida el usuario"
            "Sé directo pero si hace falta explicar cualquier cosa hazlo"
        )
        self.instrucciones_gemini = (
            "Eres un asistente AVANZADO y PROFUNDO (Gemini). "
            "Tu función es resolver de la forma más inteligente cualquier cosa que te pidan. Si crees que la respuesta es erronea especifica que puede ser erronea"
            "Siempre explica el desarrollo y el porque de las cosas(si es por ejemplo un problema de matematicas o un acertijo di como has llegado ahí"

)
        # Persistencia historial
        suffix = EMAIL_SESION.replace("@", "_").replace(".", "_") if EMAIL_SESION else "default"
        self.ruta_historial = os.path.expanduser(f"~/.historial_solver_{suffix}.json")
        self.chat_actual_id = None
        self.todo_el_historial = self._cargar_historial()

        self._build_ui()

    # ──────────────────────────────────────────────
    #  UI
    # ──────────────────────────────────────────────
    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # ── SIDEBAR ──
        self.sidebar = ctk.CTkFrame(self, width=280, corner_radius=0, fg_color="#0f0f1a")
        self.sidebar.grid(row=0, column=0, sticky="nsew")

        ctk.CTkLabel(self.sidebar, text="🛠️ Solver IA",
                     font=("Segoe UI", 22, "bold")).pack(pady=(20, 5))

        nombre_default = NOMBRE_SESION or ""
        self.entry_nombre = ctk.CTkEntry(self.sidebar, placeholder_text="¿Cuál es tu nombre?")
        self.entry_nombre.pack(fill="x", padx=15, pady=8)
        if nombre_default:
            self.entry_nombre.insert(0, nombre_default)

        # Instrucciones dinámicas
        ctk.CTkLabel(self.sidebar, text="⚙️ Instrucciones del Sistema",
                     font=("Segoe UI", 12, "bold")).pack(pady=(10, 2))

        self.txt_instrucciones = ctk.CTkTextbox(self.sidebar, height=120, wrap="word",
                                                font=("Segoe UI", 11))
        self.txt_instrucciones.pack(fill="x", padx=15, pady=5)
        self.txt_instrucciones.insert("1.0", self.instrucciones_groq)

        # Selector de modelo
        ctk.CTkLabel(self.sidebar, text="Modelo de IA",
                     font=("Segoe UI", 12, "bold")).pack(pady=(10, 2))

        frame_modelo = ctk.CTkFrame(self.sidebar, fg_color="#1a1a2e", corner_radius=10)
        frame_modelo.pack(fill="x", padx=15, pady=5)

        self.btn_basico = ctk.CTkButton(
            frame_modelo, text="🔵 Básico (Groq)", height=36,
            fg_color="#1565c0", hover_color="#0d47a1",
            command=lambda: self._cambiar_modelo("groq")
        )
        self.btn_basico.pack(fill="x", padx=8, pady=(8, 4))

        self.btn_avanzado = ctk.CTkButton(
            frame_modelo, text="🟣 Avanzado (Gemini)", height=36,
            fg_color="transparent", border_width=1, border_color="#7b1fa2",
            hover_color="#4a1a7a",
            command=lambda: self._cambiar_modelo("gemini")
        )
        self.btn_avanzado.pack(fill="x", padx=8, pady=(4, 8))

        self.lbl_limite = ctk.CTkLabel(
            self.sidebar, text="ℹ️ Groq: ~1.000 msgs/día",
            font=("Segoe UI", 10), text_color="#888"
        )
        self.lbl_limite.pack(pady=(0, 5))

        ctk.CTkFrame(self.sidebar, height=1, fg_color="#2a2a3a").pack(fill="x", padx=10, pady=8)

        ctk.CTkButton(self.sidebar, text="📄 Exportar a Word",
                      fg_color="#2c3e50", hover_color="#34495e",
                      command=self._exportar_word).pack(fill="x", padx=15, pady=4)
        ctk.CTkButton(self.sidebar, text="🧹 Nuevo Chat",
                      fg_color="transparent", border_width=1,
                      command=self._limpiar_chat).pack(fill="x", padx=15, pady=4)

        ctk.CTkLabel(self.sidebar, text="🕒 Historial",
                     font=("Segoe UI", 12, "bold")).pack(pady=(12, 3))
        self.frame_historial = ctk.CTkScrollableFrame(self.sidebar, fg_color="transparent")
        self.frame_historial.pack(fill="both", expand=True, padx=10, pady=5)

        self.status_label = ctk.CTkLabel(self.sidebar, text="🟢 Groq Listo",
                                         text_color="green", font=("Segoe UI", 11))
        self.status_label.pack(side="bottom", pady=(0, 15))

        self._actualizar_historial_ui()

        # ── ÁREA DE CHAT ──
        self.chat_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.chat_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.chat_frame.grid_rowconfigure(1, weight=1)
        self.chat_frame.grid_columnconfigure(0, weight=1)

        self.lbl_banner = ctk.CTkLabel(
            self.chat_frame,
            text="Modelo activo: Groq  •  Límite aprox: 1.000 mensajes / día",
            font=("Segoe UI", 10), text_color="#555"
        )
        self.lbl_banner.grid(row=0, column=0, sticky="w", pady=(0, 6))

        self.txt_chat = ctk.CTkTextbox(self.chat_frame, font=("Segoe UI", 14),
                                       state="disabled", wrap="word")
        self.txt_chat.grid(row=1, column=0, sticky="nsew", pady=(0, 12))

        input_frame = ctk.CTkFrame(self.chat_frame, fg_color="transparent")
        input_frame.grid(row=2, column=0, sticky="ew")
        input_frame.grid_columnconfigure(0, weight=1)

        self.entry_pregunta = ctk.CTkEntry(input_frame,
                                           placeholder_text="Describe el problema aquí...",
                                           height=42)
        self.entry_pregunta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_pregunta.bind("<Return>", lambda e: self._enviar())

        ctk.CTkButton(input_frame, text="Analizar", width=120, height=42,
                      command=self._enviar).grid(row=0, column=1)

    # ──────────────────────────────────────────────
    #  CAMBIO DE MODELO
    # ──────────────────────────────────────────────
    def _cambiar_modelo(self, modelo):
        # Guardar instrucciones editadas del modelo actual
        texto_actual = self.txt_instrucciones.get("1.0", "end-1c").strip()
        if self.modelo_actual == "groq":
            self.instrucciones_groq = texto_actual
        else:
            self.instrucciones_gemini = texto_actual

        self.modelo_actual = modelo
        self.txt_instrucciones.delete("1.0", "end")

        if modelo == "groq":
            self.btn_basico.configure(fg_color="#1565c0")
            self.btn_avanzado.configure(fg_color="transparent")
            self.lbl_limite.configure(text="ℹ️ Groq: ~1.000 msgs/día")
            self.lbl_banner.configure(
                text="Modelo activo: Groq (Básico)  •  Límite aprox: 1.000 mensajes / día")
            self.status_label.configure(text="🟢 Groq Listo", text_color="green")
            self.txt_instrucciones.insert("1.0", self.instrucciones_groq)
        else:
            self.btn_avanzado.configure(fg_color="#7b1fa2")
            self.btn_basico.configure(fg_color="transparent")
            self.lbl_limite.configure(text="ℹ️ Gemini: 5 msgs/minuto (2.5 Pro)")
            self.lbl_banner.configure(
                text=f"Modelo activo: Gemini ({GEMINI_MODEL})  •  Límite: 5 mensajes / minuto")
            self.status_label.configure(text="🟣 Gemini Listo", text_color="#bb86fc")
            self.txt_instrucciones.insert("1.0", self.instrucciones_gemini)

    # ──────────────────────────────────────────────
    #  CHAT
    # ──────────────────────────────────────────────
    def _agregar_texto(self, emisor, texto):
        self.txt_chat.configure(state="normal")
        if emisor == "IA":
            self.txt_chat.insert("end", f"\n IA:\n{texto}")
        else:
            self.txt_chat.insert("end", f"\n\n {emisor}:\n{texto}\n")
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")

    def _enviar(self):
        pregunta = self.entry_pregunta.get().strip()
        nombre = self.entry_nombre.get().strip() or "Tú"
        if not pregunta:
            return
        self._agregar_texto(nombre, pregunta)
        self.entry_pregunta.delete(0, "end")
        self.status_label.configure(text="🟡 Pensando...", text_color="orange")
        self.historial_conversacion.append({"role": "user", "content": pregunta})
        threading.Thread(target=self._proceso_ia, daemon=True).start()

    def _proceso_ia(self):
        try:
            self.after(0, lambda: self._stream_update("\n IA:\n"))

            instrucciones_actuales = self.txt_instrucciones.get("1.0", "end-1c").strip()

            if self.modelo_actual == "groq":
                response = self.cliente_groq.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "system", "content": instrucciones_actuales}]
                             + self.historial_conversacion,
                    stream=True,
                    temperature=0.2
                )
                respuesta_completa = ""
                for chunk in response:
                    delta = chunk.choices[0].delta.content
                    if delta:
                        respuesta_completa += delta
                        self.after(0, self._stream_update, delta)

            else:  # Gemini
                gemini_contents = []
                for msg in self.historial_conversacion:
                    rol_gemini = "user" if msg["role"] == "user" else "model"
                    gemini_contents.append(
                        types.Content(
                            role=rol_gemini,
                            parts=[types.Part.from_text(text=msg["content"])]
                        )
                    )

                config_gemini = types.GenerateContentConfig(
                    system_instruction=instrucciones_actuales,
                    temperature=0.2,
                )

                response = self.cliente_gemini.models.generate_content_stream(
                    model=GEMINI_MODEL,
                    contents=gemini_contents,
                    config=config_gemini,
                )

                respuesta_completa = ""
                for chunk in response:
                    if chunk.text:
                        respuesta_completa += chunk.text
                        self.after(0, self._stream_update, chunk.text)

            self.historial_conversacion.append(
                {"role": "assistant", "content": respuesta_completa}
            )
            self._guardar_historial()

            color = "green" if self.modelo_actual == "groq" else "#bb86fc"
            texto = "🟢 Groq Listo" if self.modelo_actual == "groq" else "🟣 Gemini Listo"
            self.after(0, lambda: self.status_label.configure(text=texto, text_color=color))

        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            self.after(0, lambda: self.status_label.configure(text="🔴 Error", text_color="red"))
        finally:
            self.after(0, lambda: self.txt_chat.configure(state="disabled"))

    def _stream_update(self, contenido):
        self.txt_chat.configure(state="normal")
        self.txt_chat.insert("end", contenido)
        self.txt_chat.see("end")
        self.txt_chat.configure(state="disabled")

    def _limpiar_chat(self):
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        self.txt_chat.configure(state="disabled")
        self.historial_conversacion = []
        self.chat_actual_id = None

    # ──────────────────────────────────────────────
    #  EXPORTAR
    # ──────────────────────────────────────────────
    def _exportar_word(self):
        if not self.historial_conversacion:
            messagebox.showwarning("Vacío", "No hay nada que exportar.")
            return
        nombre = self.entry_nombre.get().strip() or "Usuario"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Solucion_{nombre}.docx"
        )
        if path:
            doc = Document()
            doc.add_heading(f"Informe de Solución – {nombre}", 0)
            for msg in self.historial_conversacion:
                rol = "Tú" if msg["role"] == "user" else "IA"
                p = doc.add_paragraph()
                p.add_run(f"{rol}: ").bold = True
                p.add_run(msg["content"])
            doc.save(path)
            messagebox.showinfo("Éxito", f"Guardado en:\n{path}")

    # ──────────────────────────────────────────────
    #  HISTORIAL PERSISTENTE
    # ──────────────────────────────────────────────
    def _cargar_historial(self):
        if os.path.exists(self.ruta_historial):
            try:
                with open(self.ruta_historial, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def _guardar_historial(self):
        if not self.historial_conversacion:
            return
        if self.chat_actual_id is None:
            primer = next(
                (m["content"] for m in self.historial_conversacion if m["role"] == "user"),
                "Chat"
            )
            resumen = primer[:25] + "..." if len(primer) > 25 else primer
            hora = datetime.now().strftime("%H:%M")
            self.chat_actual_id = f"[{hora}] {resumen}"

        self.todo_el_historial[self.chat_actual_id] = self.historial_conversacion
        try:
            with open(self.ruta_historial, "w", encoding="utf-8") as f:
                json.dump(self.todo_el_historial, f, ensure_ascii=False, indent=2)
            self.after(0, self._actualizar_historial_ui)
        except Exception as e:
            print(f"Error guardando historial: {e}")

    def _actualizar_historial_ui(self):
        for w in self.frame_historial.winfo_children():
            w.destroy()
        for chat_id in reversed(list(self.todo_el_historial.keys())):
            fila = ctk.CTkFrame(self.frame_historial, fg_color="transparent")
            fila.pack(fill="x", pady=3)
            ctk.CTkButton(
                fila, text=chat_id, fg_color="transparent",
                text_color="#ddd", anchor="w", hover_color="#2a2a4a", height=32,
                command=lambda cid=chat_id: self._cargar_chat(cid)
            ).pack(side="left", fill="x", expand=True, padx=(0, 4))
            ctk.CTkButton(
                fila, text="❌", width=32, height=32,
                fg_color="#7a0000", hover_color="#500000",
                command=lambda cid=chat_id: self._borrar_chat(cid)
            ).pack(side="right")

    def _cargar_chat(self, chat_id):
        self.chat_actual_id = chat_id
        self.historial_conversacion = self.todo_el_historial[chat_id]
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        nombre = self.entry_nombre.get().strip() or "Tú"
        for msg in self.historial_conversacion:
            if msg["role"] == "user":
                self.txt_chat.insert("end", f"\n\n {nombre}:\n{msg['content']}\n")
            else:
                self.txt_chat.insert("end", f"\n IA:\n{msg['content']}")
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")

    def _borrar_chat(self, chat_id):
        if messagebox.askyesno("Confirmar", f"¿Borrar este chat?\n'{chat_id}'"):
            del self.todo_el_historial[chat_id]
            try:
                with open(self.ruta_historial, "w", encoding="utf-8") as f:
                    json.dump(self.todo_el_historial, f, ensure_ascii=False, indent=2)
            except Exception:
                pass
            if self.chat_actual_id == chat_id:
                self._limpiar_chat()
            self._actualizar_historial_ui()


if __name__ == "__main__":
    app = SolucionadorIA()
    app.mainloop()