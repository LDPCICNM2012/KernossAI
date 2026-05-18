import matplotlib
matplotlib.use("TkAgg")  # Forzar backend antes de cualquier import de matplotlib
import customtkinter as ctk
from tkinter import messagebox
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
 
# Configurar el tema oscuro/claro según el sistema
ctk.set_appearance_mode("dark")
# Usar un tema de color moderno y profesional (azul)
ctk.set_default_color_theme("blue")
 
class AppCalculadora(ctk.CTk):
    def __init__(self):
        # Heredar de CTk para crear una ventana moderna
        super().__init__()
 
        # Establecer el título de la ventana
        self.title("Calculadora de Medias")
        # Definir dimensiones de la ventana (ancho x alto)
        self.geometry("1300x1000")
        # Agregar un poco de padding para que no se vea pegado a los bordes
        self.resizable(False, False)
        
        # Configurar el grid principal para centrar contenido horizontalmente
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=0)
 
        # ───── SECCIÓN DE TÍTULO ─────
        # Frame superior para el título con un color de fondo sutil
        frame_titulo = ctk.CTkFrame(self, fg_color="transparent")
        frame_titulo.pack(fill="x", padx=20, pady=(20, 15))
        
        # Label con el título principal
        label_titulo = ctk.CTkLabel(
            frame_titulo, 
            # Texto del encabezado
            text="Calculadora de Medias", 
            # Fuente: Segoe UI, tamaño 32, negrita
            font=("Segoe UI", 32, "bold"),
            # Color dinámico según tema claro/oscuro
            text_color=["#0d47a1", "#64b5f6"]
        )
        label_titulo.pack()
 
        # Subtítulo descriptivo con fuente más pequeña
        label_subtitulo = ctk.CTkLabel(
            frame_titulo,
            # Texto informativo
            text="Organiza y calcula el promedio de tus calificaciones",
            # Fuente más pequeña y ligera
            font=("Segoe UI", 12),
            # Color gris para menos énfasis
            text_color=["#424242", "#bdbdbd"]
        )
        label_subtitulo.pack(pady=(3, 0))
 
        # ───── SECCIÓN DE ENTRADA PRINCIPAL ─────
        # Frame contenedor principal con esquinas redondeadas
        self.frame_entrada = ctk.CTkFrame(
            self, 
            # Radio de las esquinas redondeadas (20px)
            corner_radius=12,
            # Borde visible de 1px
            border_width=1,
            # Color de borde oscuro/claro según tema
            border_color=["#e0e0e0", "#333333"]
        )
        self.frame_entrada.pack(padx=25, pady=15, fill="both", expand=False)
        
        # Configurar dos columnas iguales dentro del frame
        self.frame_entrada.grid_columnconfigure((0, 1), weight=1)
 
        # Label para el campo de nombre de materia
        label_materia = ctk.CTkLabel(
            self.frame_entrada, 
            # Texto descriptivo
            text="Materia o Asignatura", 
            # Fuente más pequeña que el título
            font=("Segoe UI", 13, "bold"),
            # Color oscuro
            text_color=["#212121", "#f5f5f5"]
        )
        label_materia.grid(row=0, column=0, padx=15, pady=(15, 5), sticky="w")
        
        # Campo de entrada para el nombre de la materia
        self.entrada_nombre = ctk.CTkEntry(
            self.frame_entrada, 
            # Texto de ayuda que desaparece al escribir
            placeholder_text="Ej: Matemáticas, Física...",
            # Alto del campo de entrada
            height=40,
            # Fuente proporcional
            font=("Segoe UI", 12),
            # Borde sutil
            border_width=1
        )
        self.entrada_nombre.grid(row=1, column=0, padx=15, pady=(0, 12), sticky="ew")
 
        # Label para nota directa (sin subnotas)
        label_nota = ctk.CTkLabel(
            self.frame_entrada, 
            # Texto descriptivo
            text="Nota Directa", 
            # Fuente igual al anterior
            font=("Segoe UI", 13, "bold"),
            # Color oscuro
            text_color=["#212121", "#f5f5f5"]
        )
        label_nota.grid(row=0, column=1, padx=15, pady=(15, 5), sticky="w")
 
        # Campo de entrada para la nota numérica
        self.entrada_nota_directa = ctk.CTkEntry(
            self.frame_entrada, 
            # Texto de ayuda
            placeholder_text="Ej: 9.5",
            # Alto similar a la entrada anterior
            height=40,
            # Fuente proporcional
            font=("Segoe UI", 12),
            # Borde sutil
            border_width=1
        )
        self.entrada_nota_directa.grid(row=1, column=1, padx=15, pady=(0, 12), sticky="ew")
 
        # Label para porcentaje sobre el total
        label_porcentaje = ctk.CTkLabel(
            self.frame_entrada, 
            text="% del Total (según criterios: 60% Pruebas / 30% Proyectos / 10% Observación)", 
            font=("Segoe UI", 13, "bold"),
            text_color=["#212121", "#f5f5f5"]
        )
        label_porcentaje.grid(row=2, column=0, padx=15, pady=(0, 5), sticky="w")
 
        # Campo de entrada para porcentaje
        self.entrada_porcentaje = ctk.CTkEntry(
            self.frame_entrada, 
            placeholder_text="Ej: 60 (Pruebas Obj.) / 30 (Proyectos) / 10 (Observación)",
            height=40,
            font=("Segoe UI", 12),
            border_width=1
        )
        self.entrada_porcentaje.grid(row=3, column=0, columnspan=2, padx=15, pady=(0, 15), sticky="ew")
 
        # ───── SECCIÓN DE BOTONES DE ACCIÓN ─────
        # Frame para agrupar los botones de manera ordenada
        frame_botones = ctk.CTkFrame(self, fg_color="transparent")
        frame_botones.pack(fill="x", padx=25, pady=12)
        
        # Distribuir dos columnas iguales para los botones
        frame_botones.grid_columnconfigure((0, 1), weight=1)
 
        # Botón para guardar una nota simple
        self.btn_agregar = ctk.CTkButton(
            frame_botones, 
            # Texto del botón con icono
            text="Guardar Nota",
            # Función que se ejecuta al presionar
            command=self.agregar_nota_principal,
            # Color verde para acción positiva
            fg_color="#4caf50",
            # Color más oscuro al pasar el mouse
            hover_color="#388e3c",
            # Alto del botón
            height=40,
            # Fuente legible
            font=("Segoe UI", 12, "bold"),
            # Color del texto blanco
            text_color="white"
        )
        self.btn_agregar.grid(row=0, column=0, padx=(0, 6), sticky="ew")
 
        # Botón para agregar subnotas (notas derivadas)
        self.btn_subnotas = ctk.CTkButton(
            frame_botones, 
            # Texto descriptivo
            text="Agregar Subnotas",
            # Función cuando se presiona
            command=self.gestionar_subnotas,
            # Fondo transparente con borde
            fg_color="transparent",
            # Borde visible de 2px
            border_width=2,
            # Color del borde (cian/azul)
            border_color=["#0288d1", "#4dd0e1"],
            # Alto consistente
            height=40,
            # Fuente similar
            font=("Segoe UI", 12, "bold")
        )
        self.btn_subnotas.grid(row=0, column=1, padx=(6, 0), sticky="ew")
 
        # ───── SECCIÓN PRINCIPAL CON DOS COLUMNAS ─────
        # Frame contenedor para dividir en dos columnas (texto y gráfico)
        frame_principal = ctk.CTkFrame(self, fg_color="transparent")
        frame_principal.pack(padx=25, pady=15, fill="both", expand=True)
        
        # Configurar dos columnas: izquierda para texto, derecha para gráfico
        frame_principal.grid_columnconfigure((0, 1), weight=1)
        
        # ───── COLUMNA IZQUIERDA: HISTORIAL ─────
        # Label para la sección de notas guardadas
        label_historial = ctk.CTkLabel(
            frame_principal, 
            # Título de la sección
            text="Registro de Notas",
            # Fuente grande pero no excesiva
            font=("Segoe UI", 16, "bold"),
            # Color oscuro
            text_color=["#212121", "#f5f5f5"]
        )
        label_historial.grid(row=0, column=0, pady=(0, 8), padx=(0, 10), sticky="w")
 
        # Caja de texto para mostrar todas las notas guardadas
        self.salida_texto = ctk.CTkTextbox(
            frame_principal, 
            # Ancho del área de texto
            width=500,
            # Altura del área de texto
            height=350,
            # Fuente monoespaciada para alineación
            font=("Courier New", 11),
            # Esquinas redondeadas suaves
            corner_radius=10,
            # Borde visible
            border_width=1,
            # No editable por el usuario
            state="disabled"
        )
        self.salida_texto.grid(row=1, column=0, padx=(0, 10), sticky="nsew")
 
        # ───── COLUMNA DERECHA: GRÁFICO ─────
        # Label para la sección de gráfico
        label_grafico = ctk.CTkLabel(
            frame_principal, 
            # Título de la sección
            text="Visualización de Notas",
            # Fuente grande pero no excesiva
            font=("Segoe UI", 16, "bold"),
            # Color oscuro
            text_color=["#212121", "#f5f5f5"]
        )
        label_grafico.grid(row=0, column=1, pady=(0, 8), padx=(10, 0), sticky="w")
 
        # Frame para el canvas del gráfico
        self.frame_grafico = ctk.CTkFrame(
            frame_principal,
            # Esquinas redondeadas
            corner_radius=10,
            # Borde visible
            border_width=1
        )
        self.frame_grafico.grid(row=1, column=1, padx=(10, 0), sticky="nsew")
        
        # Canvas vacío que será reemplazado por el gráfico
        self.canvas_grafico = None
 
        # ───── SECCIÓN DE BOTONES FINALES ─────
        # Frame para los botones de acción final
        frame_acciones = ctk.CTkFrame(self, fg_color="transparent")
        frame_acciones.pack(padx=25, pady=15, fill="x")
        
        # Configurar tres columnas iguales
        frame_acciones.grid_columnconfigure((0, 1, 2), weight=1)
 
        # Botón para calcular promedio final
        self.btn_calcular = ctk.CTkButton(
            frame_acciones, 
            # Texto del botón principal
            text="Calcular Promedio Final",
            # Función para calcular resultado
            command=self.calcular_total_final,
            # Alto visual
            height=45,
            # Fuente grande y bold
            font=("Segoe UI", 12, "bold"),
            # Color azul principal
            fg_color="#0277bd",
            # Azul más oscuro al hover
            hover_color="#01579b",
            # Esquinas redondeadas
            corner_radius=10,
            # Texto blanco
            text_color="white"
        )
        self.btn_calcular.grid(row=0, column=0, padx=(0, 6), sticky="ew")
 
        # Botón para limpiar datos
        self.btn_limpiar = ctk.CTkButton(
            frame_acciones, 
            # Texto descriptivo
            text="Limpiar Todo",
            # Función para limpiar
            command=self.limpiar_datos,
            # Color naranja para advertencia
            fg_color="#ff9800",
            # Naranja más oscuro al hover
            hover_color="#f57c00",
            # Alto consistente
            height=45,
            # Fuente similar
            font=("Segoe UI", 12, "bold"),
            # Texto blanco
            text_color="white"
        )
        self.btn_limpiar.grid(row=0, column=1, padx=(3, 3), sticky="ew")
 
        # Botón para exportar a Word
        self.btn_exportar = ctk.CTkButton(
            frame_acciones, 
            # Texto descriptivo
            text="Exportar a Word",
            # Función para exportar
            command=self.exportar_a_word,
            # Color púrpura para exportación
            fg_color="#7c3aed",
            # Púrpura más oscuro al hover
            hover_color="#5b21b6",
            # Alto consistente
            height=45,
            # Fuente similar
            font=("Segoe UI", 12, "bold"),
            # Texto blanco
            text_color="white"
        )
        self.btn_exportar.grid(row=0, column=2, padx=(6, 0), sticky="ew")
 
        # Listas para almacenar los datos de notas
        # Lista que guarda los valores numéricos de todas las notas
        self.notas_finales = []
        # Lista que guarda los nombres correspondientes a cada nota
        self.nombres_notas = []
        # Lista que guarda los porcentajes de cada nota sobre el total
        self.porcentajes = []
 
    def calcular_media(self, lista):
        # Retornar 0 si la lista está vacía (evitar división por cero)
        if not lista: 
            return 0
        # Sumar todos los elementos y dividir entre la cantidad
        return sum(lista) / len(lista)
 
    def actualizar_grafico(self):
        # Validar que hay notas para graficar
        if not self.notas_finales:
            return
        
        # Destruir gráfico anterior si existe
        if self.canvas_grafico:
            self.canvas_grafico.get_tk_widget().destroy()
        
        # Crear figura con estilo moderno
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
        # Establecer color de fondo oscuro
        fig.patch.set_facecolor('#1a1a1a')
        
        # ─── Gráfico 1: Barras de notas individuales ───
        # Colores para las barras basado en rango de notas
        colores = ['#4caf50' if nota >= 7 else '#ff9800' if nota >= 5 else '#f44336' 
                   for nota in self.notas_finales]
        # Crear gráfico de barras
        ax1.bar(range(len(self.nombres_notas)), self.notas_finales, color=colores, alpha=0.8, edgecolor='white')
        # Establecer etiquetas en X
        ax1.set_xticks(range(len(self.nombres_notas)))
        # Rotar etiquetas para mejor legibilidad
        ax1.set_xticklabels(self.nombres_notas, rotation=45, ha='right', fontsize=9)
        # Etiqueta del eje Y
        ax1.set_ylabel('Calificación', color='white', fontsize=10)
        # Título del gráfico
        ax1.set_title('Notas por Asignatura', color='white', fontsize=12, fontweight='bold')
        # Establecer límites del eje Y
        ax1.set_ylim(0, 10)
        # Color de fondo del gráfico oscuro
        ax1.set_facecolor('#2a2a2a')
        # Color de los ticks y etiquetas
        ax1.tick_params(colors='white')
        # Línea de grid horizontal
        ax1.grid(axis='y', alpha=0.3, color='white')
        
        # ─── Gráfico 2: Pie de distribución ───
        # Calcular el promedio general
        promedio = self.calcular_media(self.notas_finales)
        # Datos para el pie (notas por encima y por debajo del promedio)
        por_encima = sum(1 for n in self.notas_finales if n >= promedio)
        por_debajo = sum(1 for n in self.notas_finales if n < promedio)
        # Crear gráfico tipo pie
        ax2.pie([por_encima, por_debajo], 
                labels=[f'Arriba del promedio\n({por_encima})', f'Debajo del promedio\n({por_debajo})'],
                colors=['#4caf50', '#f44336'],
                autopct='%1.1f%%',
                startangle=90,
                textprops={'color': 'white', 'fontsize': 10})
        # Título del gráfico
        ax2.set_title(f'Distribución (Promedio: {promedio:.2f})', color='white', fontsize=12, fontweight='bold')
        
        # Ajustar espaciado
        plt.tight_layout()
        
        # Insertar gráfico en el canvas de tkinter
        self.canvas_grafico = FigureCanvasTkAgg(fig, master=self.frame_grafico)
        # Dibujar el gráfico
        self.canvas_grafico.draw()
        # Mostrar el widget
        self.canvas_grafico.get_tk_widget().pack(fill="both", expand=True)
 
    def gestionar_subnotas(self):
        # Obtener el nombre de la materia del campo de entrada
        nombre_principal = self.entrada_nombre.get().strip()
        if not nombre_principal:
            messagebox.showwarning("Atención", "Escribe el nombre de la asignatura primero.")
            return

        # ── Preguntar cuántos BLOQUES tiene la asignatura ──
        d = ctk.CTkInputDialog(text=f"¿Cuántos bloques tiene {nombre_principal}?\n(Ej: 3 → Exámenes, Clase, Deberes)", title="Bloques de calificación")
        res = d.get_input()
        self.focus_force(); self.lift()
        if not res or not res.isdigit():
            return

        num_bloques = int(res)
        self.salida_texto.configure(state="normal")
        self.salida_texto.insert("end", f"\n{'─'*60}\n")
        self.salida_texto.insert("end", f"  📚 {nombre_principal.upper()}\n")
        self.salida_texto.insert("end", f"{'─'*60}\n")

        nota_final_asignatura = 0.0
        suma_pesos_bloques = 0.0

        for i in range(num_bloques):
            # ── Nombre del bloque ──
            d = ctk.CTkInputDialog(text=f"Nombre del bloque {i+1}:\n(Ej: Exámenes)", title="Nombre del bloque")
            nombre_bloque = d.get_input()
            self.focus_force(); self.lift()
            if nombre_bloque is None:
                break

            # ── Peso del bloque sobre el TOTAL de la asignatura ──
            d = ctk.CTkInputDialog(text=f"¿Cuánto vale '{nombre_bloque}' en % sobre el total de {nombre_principal}?\n(Ej: 60)", title="% del bloque")
            peso_bloque_str = d.get_input()
            self.focus_force(); self.lift()
            if peso_bloque_str is None:
                break
            try:
                peso_bloque = float(peso_bloque_str)
            except ValueError:
                messagebox.showerror("Error", "Introduce un número válido para el porcentaje.")
                continue

            # ── Cuántas notas hay dentro del bloque ──
            d = ctk.CTkInputDialog(text=f"¿Cuántas notas hay dentro de '{nombre_bloque}'?", title="Notas del bloque")
            res_sub = d.get_input()
            self.focus_force(); self.lift()
            if not res_sub or not res_sub.isdigit():
                continue
            num_sub = int(res_sub)

            self.salida_texto.insert("end", f"\n  📂 {nombre_bloque} ({peso_bloque:.0f}% del total)\n")

            nota_bloque_ponderada = 0.0
            suma_pesos_sub = 0.0

            for j in range(num_sub):
                # ── Nombre de la nota dentro del bloque ──
                d = ctk.CTkInputDialog(text=f"Nombre de la nota {j+1} dentro de '{nombre_bloque}':", title="Nota")
                nombre_sub = d.get_input()
                self.focus_force(); self.lift()
                if nombre_sub is None:
                    break

                # ── Peso de esta nota dentro del bloque ──
                d = ctk.CTkInputDialog(text=f"¿Cuánto vale '{nombre_sub}' en % dentro de '{nombre_bloque}'?\n(Ej: 33.3)", title=f"% dentro de {nombre_bloque}")
                peso_sub_str = d.get_input()
                self.focus_force(); self.lift()
                if peso_sub_str is None:
                    break
                try:
                    peso_sub = float(peso_sub_str)
                except ValueError:
                    messagebox.showerror("Error", "Introduce un número válido.")
                    continue

                # ── Valor de la nota ──
                d = ctk.CTkInputDialog(text=f"Calificación de '{nombre_sub}':", title="Calificación")
                valor_sub_str = d.get_input()
                self.focus_force(); self.lift()
                if valor_sub_str is None:
                    break
                try:
                    valor_sub = float(valor_sub_str)
                except ValueError:
                    messagebox.showerror("Error", "Introduce un número válido.")
                    continue

                nota_bloque_ponderada += valor_sub * (peso_sub / 100)
                suma_pesos_sub += peso_sub
                self.salida_texto.insert("end", f"      • {nombre_sub:20} {peso_sub:.0f}% dentro del bloque → {valor_sub:.2f}\n")

            # Normalizar si los pesos del bloque no suman 100
            if suma_pesos_sub > 0 and suma_pesos_sub != 100:
                nota_bloque_ponderada = nota_bloque_ponderada / (suma_pesos_sub / 100)

            self.salida_texto.insert("end", f"    ✓ Nota del bloque '{nombre_bloque}': {nota_bloque_ponderada:.2f}\n")

            nota_final_asignatura += nota_bloque_ponderada * (peso_bloque / 100)
            suma_pesos_bloques += peso_bloque

        # Normalizar si los pesos de los bloques no suman 100
        if suma_pesos_bloques > 0 and suma_pesos_bloques != 100:
            nota_final_asignatura = nota_final_asignatura / (suma_pesos_bloques / 100)

        # Guardar resultado final
        self.nombres_notas.append(nombre_principal)
        self.notas_finales.append(nota_final_asignatura)
        self.porcentajes.append(100.0)  # ya viene ponderada internamente

        self.salida_texto.insert("end", f"\n  {'═'*50}\n")
        self.salida_texto.insert("end", f"  ✅ NOTA FINAL {nombre_principal.upper()}: {nota_final_asignatura:.2f}\n")
        self.salida_texto.insert("end", f"  {'═'*50}\n\n")
        self.salida_texto.configure(state="disabled")
        self.entrada_nombre.delete(0, "end")
        self.actualizar_grafico()
 
    def agregar_nota_principal(self):
        # Obtener el nombre de la materia del campo de entrada
        nombre = self.entrada_nombre.get().strip()
        # Obtener la nota numérica del campo de entrada
        nota_str = self.entrada_nota_directa.get().strip()
        # Obtener el porcentaje personalizado (opcional)
        porcentaje_str = self.entrada_porcentaje.get().strip()
 
        # Validar que al menos nombre y nota estén presentes
        if nombre and nota_str:
            # Intentar convertir a número decimal
            try:
                # Convertir string a float
                nota = float(nota_str)
                
                # Manejar porcentaje personalizado
                if porcentaje_str:
                    porcentaje_sobre_total = float(porcentaje_str)
                else:
                    # Si no hay porcentaje, distribuir equitativamente
                    total_notas_actuales = len(self.notas_finales) + 1
                    porcentaje_sobre_total = (1 / total_notas_actuales) * 100
                
                # Guardar el nombre en la lista
                self.nombres_notas.append(nombre)
                # Guardar la nota en la lista
                self.notas_finales.append(nota)
                # Guardar el porcentaje en la lista
                self.porcentajes.append(porcentaje_sobre_total)
                
                # Habilitar edición del textbox
                self.salida_texto.configure(state="normal")
                # Mostrar la nota guardada en el historial
                self.salida_texto.insert("end", f"  ✓ {nombre:25} → {nota:.2f}")
                # Mostrar el porcentaje que representa sobre el total
                self.salida_texto.insert("end", f" ({porcentaje_sobre_total:.1f}%)\n")
                # Deshabilitar edición nuevamente
                self.salida_texto.configure(state="disabled")
                
                # Limpiar campos de entrada
                self.entrada_nombre.delete(0, "end")
                self.entrada_nota_directa.delete(0, "end")
                self.entrada_porcentaje.delete(0, "end")
                # Actualizar gráfico
                self.actualizar_grafico()
            # Si la conversión falla
            except ValueError:
                # Mostrar error de formato
                messagebox.showerror("Error", "La calificación y el porcentaje deben ser números válidos.")
        # Si falta algún campo
        else:
            # Advertir al usuario que complete ambos campos
            messagebox.showwarning("Campos incompletos", "Por favor completa materia y calificación.")
 
    def calcular_total_final(self):
        # Validar que haya al menos una nota guardada
        if not self.notas_finales:
            # Informar que no hay datos
            messagebox.showinfo("Sin datos", "No hay calificaciones guardadas para calcular.")
            # Salir sin hacer nada
            return
 
        # Calcular la media ponderada usando los porcentajes
        suma_ponderada = sum(nota * (peso / 100) for nota, peso in zip(self.notas_finales, self.porcentajes))
        total_porcentaje = sum(self.porcentajes)
        
        # Si el total de porcentajes no suma 100%, normalizar
        if total_porcentaje != 100:
            media_total = suma_ponderada / (total_porcentaje / 100)
        else:
            media_total = suma_ponderada
        
        # Habilitar edición del textbox para agregar resultado
        self.salida_texto.configure(state="normal")
        # Línea decorativa superior
        self.salida_texto.insert("end", f"\n{'═'*60}\n")
        # Línea de título del resultado
        self.salida_texto.insert("end", f"  📊 RESULTADO FINAL\n")
        # Mostrar cantidad de asignaturas
        self.salida_texto.insert("end", f"  Asignaturas: {len(self.nombres_notas)}\n")
        # Mostrar suma total de porcentajes
        self.salida_texto.insert("end", f"  Peso total asignado: {total_porcentaje:.1f}%\n")
        # Mostrar promedio ponderado con dos decimales
        self.salida_texto.insert("end", f"  Promedio Ponderado: {media_total:.2f}\n")
        # Línea decorativa inferior
        self.salida_texto.insert("end", f"{'═'*60}\n\n")
        # Deshabilitar edición nuevamente
        self.salida_texto.configure(state="disabled")
        # Desplazar vista al final del contenido
        self.salida_texto.see("end")
 
    def limpiar_datos(self):
        # Preguntar al usuario si está seguro
        respuesta = messagebox.askyesno("Confirmar", "¿Deseas limpiar todas las notas?")
        
        # Si el usuario confirma
        if respuesta:
            # Limpiar lista de notas
            self.notas_finales = []
            # Limpiar lista de nombres
            self.nombres_notas = []
            # Limpiar lista de porcentajes
            self.porcentajes = []
            
            # Limpiar textbox
            self.salida_texto.configure(state="normal")
            self.salida_texto.delete("1.0", "end")
            self.salida_texto.configure(state="disabled")
            
            # Limpiar campos de entrada
            self.entrada_nombre.delete(0, "end")
            self.entrada_nota_directa.delete(0, "end")
            self.entrada_porcentaje.delete(0, "end")
            
            # Destruir gráfico
            if self.canvas_grafico:
                self.canvas_grafico.get_tk_widget().destroy()
                self.canvas_grafico = None
            
            # Notificar éxito
            messagebox.showinfo("Éxito", "Todos los datos han sido limpiados.")
 
    def exportar_a_word(self):
        # Validar que haya notas para exportar
        if not self.notas_finales:
            # Informar que no hay datos
            messagebox.showwarning("Sin datos", "No hay calificaciones para exportar.")
            # Salir sin hacer nada
            return
 
        # Crear nuevo documento Word
        doc = Document()
        
        # Agregar título al documento
        titulo = doc.add_heading('Reporte de Calificaciones', 0)
        # Centrar el título
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar fecha del reporte
        fecha = doc.add_paragraph(f'Generado el: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        # Centrar la fecha
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar salto de línea
        doc.add_paragraph()
        
        # Agregar tabla de notas
        tabla = doc.add_table(rows=len(self.nombres_notas) + 1, cols=3)
        # Establecer estilo de tabla
        tabla.style = 'Light Grid Accent 1'
        
        # Encabezado de la tabla
        encabezado = tabla.rows[0].cells
        # Columna 1: Número
        encabezado[0].text = '#'
        # Columna 2: Asignatura
        encabezado[1].text = 'Asignatura'
        # Columna 3: Calificación
        encabezado[2].text = 'Calificación'
        
        # Llenar tabla con datos
        for i, (nombre, nota) in enumerate(zip(self.nombres_notas, self.notas_finales), start=1):
            # Fila actual
            fila = tabla.rows[i].cells
            # Número de fila
            fila[0].text = str(i)
            # Nombre de asignatura
            fila[1].text = nombre
            # Nota con 2 decimales
            fila[2].text = f"{nota:.2f}"
        
        # Agregar salto de línea
        doc.add_paragraph()
        
        # Agregar promedio final
        promedio = self.calcular_media(self.notas_finales)
        parrafo_promedio = doc.add_paragraph(f'Promedio General: {promedio:.2f}')
        # Aplicar negrita al promedio
        parrafo_promedio.runs[0].font.bold = True
        parrafo_promedio.runs[0].font.size = Pt(14)
        
        # Agregar análisis de desempeño
        doc.add_paragraph()
        doc.add_heading('Análisis de Desempeño', level=2)
        
        # Contar notas por rango
        excelente = sum(1 for n in self.notas_finales if n >= 9)
        muy_bien = sum(1 for n in self.notas_finales if 8 <= n < 9)
        bien = sum(1 for n in self.notas_finales if 7 <= n < 8)
        regular = sum(1 for n in self.notas_finales if 6 <= n < 7)
        insuficiente = sum(1 for n in self.notas_finales if n < 6)
        
        # Agregar estadísticas
        doc.add_paragraph(f'Excelente (≥9): {excelente} asignatura(s)', style='List Bullet')
        doc.add_paragraph(f'Muy Bien (8-8.9): {muy_bien} asignatura(s)', style='List Bullet')
        doc.add_paragraph(f'Bien (7-7.9): {bien} asignatura(s)', style='List Bullet')
        doc.add_paragraph(f'Regular (6-6.9): {regular} asignatura(s)', style='List Bullet')
        doc.add_paragraph(f'Insuficiente (<6): {insuficiente} asignatura(s)', style='List Bullet')
        
        # Obtener ruta de usuario
        ruta_documentos = os.path.expanduser("~/Documents")
        # Crear nombre de archivo con timestamp
        nombre_archivo = f"Reporte_Calificaciones_{datetime.now().strftime('%d%m%Y_%H%M%S')}.docx"
        # Ruta completa del archivo
        ruta_completa = os.path.join(ruta_documentos, nombre_archivo)
        
        # Guardar documento
        doc.save(ruta_completa)
        
        # Notificar éxito
        messagebox.showinfo("Éxito", f"Documento exportado a:\n{ruta_completa}")
 
# Punto de entrada del programa
# Asegúrate de que el final de tus archivos use estrictamente esta condición:
if __name__ == "__main__":
    app = AppCalculadora() # (O el nombre de la clase de ese archivo)
    app.mainloop()