import customtkinter as ctk
from tkinter import messagebox, filedialog
from openai import OpenAI
import google.generativeai as genai
from docx import Document
import threading, os
from datetime import datetime

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

GROQ_API_KEY   = ""
GEMINI_API_KEY = ""

INSTRUCCIONES_CORRECTOR = """Eres un profesor corrector experto y riguroso.
Tu tarea es evaluar las respuestas de un alumno comparándolas con el ejercicio/criterios dados.
Para cada pregunta debes:
1. Indicar si está CORRECTA, PARCIALMENTE CORRECTA o INCORRECTA.
2. Puntuación obtenida sobre la puntuación máxima de esa pregunta.
3. Comentario breve explicando el acierto o el error.
4. Sugerencia de mejora si aplica.

Al final incluye:
- NOTA TOTAL (sobre 10).
- Resumen general de puntos fuertes y débiles del alumno.
- Recomendaciones de estudio personalizadas.

Sé justo, constructivo y motivador en el tono."""


class CorrectorExamenesProfesor(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Corrector de Exámenes — Modo Profesor")
        self.geometry("1350x820")
        self.modelo_actual = "groq"
        self.correccion_actual = ""
        self.alumnos = {}           # nombre -> {respuestas, correccion}
        self.alumno_actual = None

        self.cliente_groq = OpenAI(base_url="https://api.groq.com/openai/v1", api_key=GROQ_API_KEY)
        genai.configure(api_key=GEMINI_API_KEY)
        self.cliente_gemini = genai.GenerativeModel("gemini-2.5-flash")

        self.setup_ui()

    # ──────────────────────────────────────────────
    #  UI
    # ──────────────────────────────────────────────
    def setup_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # ══ SIDEBAR ══
        self.sidebar = ctk.CTkFrame(self, width=270, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_propagate(False)

        ctk.CTkLabel(self.sidebar, text="Corrector de\nExámenes",
                     font=("Segoe UI", 20, "bold"), text_color="#1e88e5").pack(pady=(20, 5))
        ctk.CTkLabel(self.sidebar, text="Modo Profesor",
                     font=("Segoe UI", 11), text_color="#888").pack(pady=(0, 15))

        # Motor IA
        ctk.CTkLabel(self.sidebar, text="Motor de IA:", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15)
        fm = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        fm.pack(fill="x", padx=15, pady=5)
        self.btn_groq = ctk.CTkButton(fm, text="Básico\n(Groq)", height=50,
            fg_color="#1565c0", hover_color="#0d47a1", font=("Segoe UI", 10, "bold"),
            command=lambda: self._set_modelo("groq"))
        self.btn_groq.pack(side="left", expand=True, fill="x", padx=(0, 3))
        self.btn_gemini = ctk.CTkButton(fm, text="Avanzado\n(Gemini)", height=50,
            fg_color="#333", hover_color="#444", font=("Segoe UI", 10, "bold"),
            command=lambda: self._set_modelo("gemini"))
        self.btn_gemini.pack(side="left", expand=True, fill="x", padx=(3, 0))

        self.lbl_limite = ctk.CTkLabel(self.sidebar,
            text="Groq: ~1.000 msgs/día", font=("Segoe UI", 9), text_color="#666")
        self.lbl_limite.pack(pady=(2, 8))

        ctk.CTkFrame(self.sidebar, height=1, fg_color="#333").pack(fill="x", padx=15, pady=5)

        # Parámetros de corrección
        ctk.CTkLabel(self.sidebar, text="Materia / Examen", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15, pady=(5,2))
        self.entry_materia = ctk.CTkEntry(self.sidebar, placeholder_text="Ej: Matemáticas — Tema 3")
        self.entry_materia.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(self.sidebar, text="Nivel educativo", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15)
        self.combo_nivel = ctk.CTkComboBox(self.sidebar, values=[
            "Primaria", "1º ESO", "2º ESO", "3º ESO", "4º ESO",
            "1º Bachillerato", "2º Bachillerato", "Universidad"])
        self.combo_nivel.set("2º ESO")
        self.combo_nivel.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(self.sidebar, text="Nombre del alumno", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15)
        self.entry_alumno = ctk.CTkEntry(self.sidebar, placeholder_text="Nombre y apellidos")
        self.entry_alumno.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(self.sidebar, text="Criterios de puntuación", font=("Segoe UI", 12, "bold")).pack(anchor="w", padx=15)
        self.entry_criterios = ctk.CTkEntry(self.sidebar,
            placeholder_text="Ej: P1=2pts, P2=3pts, P3=5pts")
        self.entry_criterios.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkFrame(self.sidebar, height=1, fg_color="#333").pack(fill="x", padx=15, pady=5)

        # Botones principales
        self.btn_corregir = ctk.CTkButton(self.sidebar, text="✅ Corregir Examen",
            height=45, font=("Segoe UI", 13, "bold"),
            fg_color="#1b5e20", hover_color="#003300",
            command=self.corregir_examen)
        self.btn_corregir.pack(fill="x", padx=15, pady=5)

        self.btn_guardar_alumno = ctk.CTkButton(self.sidebar, text="💾 Guardar resultado alumno",
            height=38, fg_color="#37474f", hover_color="#263238",
            command=self.guardar_alumno, state="disabled")
        self.btn_guardar_alumno.pack(fill="x", padx=15, pady=3)

        self.btn_exportar_uno = ctk.CTkButton(self.sidebar, text="📄 Exportar corrección (Word)",
            height=38, fg_color="#4a148c", hover_color="#2d0065",
            command=self.exportar_correccion, state="disabled")
        self.btn_exportar_uno.pack(fill="x", padx=15, pady=3)

        self.btn_exportar_clase = ctk.CTkButton(self.sidebar, text="📊 Exportar informe de clase",
            height=38, fg_color="#0d47a1", hover_color="#002171",
            command=self.exportar_informe_clase, state="disabled")
        self.btn_exportar_clase.pack(fill="x", padx=15, pady=3)

        self.btn_limpiar = ctk.CTkButton(self.sidebar, text="🧹 Nuevo examen",
            height=35, fg_color="transparent", border_width=1,
            command=self.limpiar_todo)
        self.btn_limpiar.pack(fill="x", padx=15, pady=3)

        # Lista de alumnos corregidos
        ctk.CTkLabel(self.sidebar, text="Alumnos corregidos",
                     font=("Segoe UI", 11, "bold"), text_color="#aaa").pack(pady=(10, 3))
        self.frame_alumnos = ctk.CTkScrollableFrame(self.sidebar, fg_color="transparent", height=120)
        self.frame_alumnos.pack(fill="x", padx=10, pady=3)

        self.status = ctk.CTkLabel(self.sidebar, text="Listo", text_color="gray",
                                   font=("Segoe UI", 10))
        self.status.pack(side="bottom", pady=10)

        # ══ PANEL CENTRAL — Enunciado / Ejercicio ══
        panel_izq = ctk.CTkFrame(self, fg_color="transparent")
        panel_izq.grid(row=0, column=1, sticky="nsew", padx=(15, 7), pady=15)
        panel_izq.grid_rowconfigure(1, weight=1)
        panel_izq.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(panel_izq, text="📝 Enunciado / Criterios del examen",
                     font=("Segoe UI", 15, "bold")).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.txt_enunciado = ctk.CTkTextbox(panel_izq, font=("Consolas", 12))
        self.txt_enunciado.grid(row=1, column=0, sticky="nsew")
        self.txt_enunciado.insert("end",
            "Pega aquí el enunciado del examen y/o los criterios de corrección...\n\n"
            "Ejemplo:\n"
            "Pregunta 1 (2 pts): ¿Qué es la fotosíntesis? Explícala.\n"
            "Pregunta 2 (3 pts): Nombra 3 diferencias entre células animales y vegetales.\n"
            "Pregunta 3 (5 pts): Resuelve los problemas adjuntos.")

        # ══ PANEL DERECHO — Respuestas y corrección ══
        panel_der = ctk.CTkFrame(self, fg_color="transparent")
        panel_der.grid(row=0, column=2, sticky="nsew", padx=(7, 15), pady=15)
        panel_der.grid_rowconfigure(1, weight=2)
        panel_der.grid_rowconfigure(3, weight=3)
        panel_der.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(panel_der, text="✍️ Respuestas del alumno",
                     font=("Segoe UI", 15, "bold")).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.txt_respuestas = ctk.CTkTextbox(panel_der, font=("Consolas", 12))
        self.txt_respuestas.grid(row=1, column=0, sticky="nsew", pady=(0, 10))
        self.txt_respuestas.insert("end",
            "Pega o escribe aquí las respuestas del alumno...\n\n"
            "Ejemplo:\n"
            "Respuesta 1: La fotosíntesis es el proceso por el cual las plantas...\n"
            "Respuesta 2: 1) Las células vegetales tienen cloroplastos...\n"
            "Respuesta 3: Ejercicio a) x=5  b) y=3...")

        # Separador con etiqueta
        sep_frame = ctk.CTkFrame(panel_der, fg_color="transparent")
        sep_frame.grid(row=2, column=0, sticky="ew", pady=4)
        ctk.CTkLabel(sep_frame, text="📋 Corrección de la IA",
                     font=("Segoe UI", 15, "bold")).pack(side="left")
        self.lbl_nota = ctk.CTkLabel(sep_frame, text="",
                                      font=("Segoe UI", 16, "bold"), text_color="#4caf50")
        self.lbl_nota.pack(side="right")

        self.txt_correccion = ctk.CTkTextbox(panel_der, font=("Consolas", 12), state="disabled")
        self.txt_correccion.grid(row=3, column=0, sticky="nsew")

    # ──────────────────────────────────────────────
    #  MODELO
    # ──────────────────────────────────────────────
    def _set_modelo(self, m):
        self.modelo_actual = m
        if m == "groq":
            self.btn_groq.configure(fg_color="#1565c0")
            self.btn_gemini.configure(fg_color="#333")
            self.lbl_limite.configure(text="Groq: ~1.000 msgs/día")
        else:
            self.btn_groq.configure(fg_color="#333")
            self.btn_gemini.configure(fg_color="#6a1b9a")
            self.lbl_limite.configure(text="Gemini: 15 msgs/minuto")

    def _llamar_ia(self, prompt: str) -> str:
        if self.modelo_actual == "groq":
            response = self.cliente_groq.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {'role': 'system', 'content': INSTRUCCIONES_CORRECTOR},
                    {'role': 'user', 'content': prompt}
                ],
                temperature=0.2
            )
            return response.choices[0].message.content
        else:
            response = self.cliente_gemini.generate_content(
                f"{INSTRUCCIONES_CORRECTOR}\n\n{prompt}")
            return response.text

    # ──────────────────────────────────────────────
    #  CORREGIR
    # ──────────────────────────────────────────────
    def corregir_examen(self):
        enunciado = self.txt_enunciado.get("1.0", "end-1c").strip()
        respuestas = self.txt_respuestas.get("1.0", "end-1c").strip()
        alumno = self.entry_alumno.get().strip() or "Alumno sin nombre"
        materia = self.entry_materia.get().strip() or "Examen"
        criterios = self.entry_criterios.get().strip()

        if not enunciado or "Pega aquí" in enunciado:
            messagebox.showwarning("Falta enunciado", "Introduce el enunciado del examen.")
            return
        if not respuestas or "Pega o escribe" in respuestas:
            messagebox.showwarning("Faltan respuestas", "Introduce las respuestas del alumno.")
            return

        self.btn_corregir.configure(state="disabled")
        self.status.configure(text="Corrigiendo...", text_color="orange")
        self.lbl_nota.configure(text="")

        prompt = (
            f"MATERIA: {materia}\n"
            f"NIVEL: {self.combo_nivel.get()}\n"
            f"ALUMNO: {alumno}\n"
            + (f"CRITERIOS DE PUNTUACIÓN: {criterios}\n" if criterios else "")
            + f"\nENUNCIADO DEL EXAMEN:\n{enunciado}\n\n"
            f"RESPUESTAS DEL ALUMNO:\n{respuestas}"
        )

        threading.Thread(target=self._thread_corregir,
                         args=(prompt, alumno, respuestas), daemon=True).start()

    def _thread_corregir(self, prompt, alumno, respuestas):
        try:
            resultado = self._llamar_ia(prompt)
            self.correccion_actual = resultado

            # Extraer nota del texto si aparece "NOTA TOTAL"
            nota_str = ""
            for linea in resultado.splitlines():
                if "NOTA TOTAL" in linea.upper() or "NOTA FINAL" in linea.upper():
                    partes = linea.split(":")
                    if len(partes) > 1:
                        nota_str = partes[-1].strip().split()[0]
                    break

            def _mostrar():
                self.txt_correccion.configure(state="normal")
                self.txt_correccion.delete("1.0", "end")
                self.txt_correccion.insert("end", resultado)
                self.txt_correccion.configure(state="disabled")
                if nota_str:
                    self.lbl_nota.configure(text=f"Nota: {nota_str}")
                for btn in [self.btn_guardar_alumno, self.btn_exportar_uno]:
                    btn.configure(state="normal")

            self.after(0, _mostrar)

        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.btn_corregir.configure(state="normal"))
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    # ──────────────────────────────────────────────
    #  GESTIÓN DE ALUMNOS
    # ──────────────────────────────────────────────
    def guardar_alumno(self):
        alumno = self.entry_alumno.get().strip() or "Alumno sin nombre"
        if not self.correccion_actual:
            messagebox.showwarning("Sin corrección", "Primero corrige un examen.")
            return
        self.alumnos[alumno] = {
            "respuestas": self.txt_respuestas.get("1.0", "end-1c"),
            "correccion": self.correccion_actual,
            "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M")
        }
        self._actualizar_lista_alumnos()
        self.btn_exportar_clase.configure(state="normal")
        messagebox.showinfo("Guardado", f"Resultado de '{alumno}' guardado.\nTotal alumnos: {len(self.alumnos)}")

    def _actualizar_lista_alumnos(self):
        for w in self.frame_alumnos.winfo_children():
            w.destroy()
        for nombre in self.alumnos:
            fila = ctk.CTkFrame(self.frame_alumnos, fg_color="transparent")
            fila.pack(fill="x", pady=2)
            ctk.CTkButton(fila, text=nombre, fg_color="transparent",
                          text_color="#ccc", anchor="w", hover_color="#2a2a4a", height=28,
                          command=lambda n=nombre: self._cargar_alumno(n)
                          ).pack(side="left", fill="x", expand=True)
            ctk.CTkButton(fila, text="❌", width=28, height=28,
                          fg_color="#7a0000", hover_color="#500000",
                          command=lambda n=nombre: self._borrar_alumno(n)
                          ).pack(side="right")

    def _cargar_alumno(self, nombre):
        datos = self.alumnos[nombre]
        self.entry_alumno.delete(0, "end")
        self.entry_alumno.insert(0, nombre)
        self.txt_respuestas.delete("1.0", "end")
        self.txt_respuestas.insert("end", datos["respuestas"])
        self.txt_correccion.configure(state="normal")
        self.txt_correccion.delete("1.0", "end")
        self.txt_correccion.insert("end", datos["correccion"])
        self.txt_correccion.configure(state="disabled")
        self.correccion_actual = datos["correccion"]
        for btn in [self.btn_guardar_alumno, self.btn_exportar_uno]:
            btn.configure(state="normal")

    def _borrar_alumno(self, nombre):
        if messagebox.askyesno("Confirmar", f"¿Borrar resultado de '{nombre}'?"):
            del self.alumnos[nombre]
            self._actualizar_lista_alumnos()
            if not self.alumnos:
                self.btn_exportar_clase.configure(state="disabled")

    def limpiar_todo(self):
        self.txt_respuestas.delete("1.0", "end")
        self.txt_correccion.configure(state="normal")
        self.txt_correccion.delete("1.0", "end")
        self.txt_correccion.configure(state="disabled")
        self.entry_alumno.delete(0, "end")
        self.correccion_actual = ""
        self.lbl_nota.configure(text="")
        for btn in [self.btn_guardar_alumno, self.btn_exportar_uno]:
            btn.configure(state="disabled")

    # ──────────────────────────────────────────────
    #  EXPORTAR
    # ──────────────────────────────────────────────
    def exportar_correccion(self):
        if not self.correccion_actual:
            messagebox.showwarning("Vacío", "No hay corrección que exportar.")
            return
        alumno = self.entry_alumno.get().strip() or "Alumno"
        materia = self.entry_materia.get().strip() or "Examen"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Correccion_{alumno}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return

        doc = Document()
        doc.add_heading(f"Corrección — {materia}", 0)
        doc.add_paragraph(f"Alumno: {alumno}")
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
        doc.add_paragraph(f"Fecha corrección: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")
        doc.add_heading("Respuestas del alumno", level=1)
        doc.add_paragraph(self.txt_respuestas.get("1.0", "end-1c"))
        doc.add_page_break()
        doc.add_heading("Corrección y calificación", level=1)
        doc.add_paragraph(self.correccion_actual)
        doc.save(path)
        messagebox.showinfo("Exportado", f"Corrección guardada en:\n{path}")

    def exportar_informe_clase(self):
        if not self.alumnos:
            messagebox.showwarning("Sin datos", "No hay alumnos guardados todavía.")
            return
        materia = self.entry_materia.get().strip() or "Examen"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"InformeClase_{materia}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return

        doc = Document()
        doc.add_heading(f"Informe de Clase — {materia}", 0)
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
        doc.add_paragraph(f"Total alumnos corregidos: {len(self.alumnos)}")
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")

        for nombre, datos in self.alumnos.items():
            doc.add_heading(nombre, level=1)
            doc.add_paragraph(f"Corregido: {datos['timestamp']}")
            doc.add_paragraph("")
            doc.add_heading("Corrección", level=2)
            doc.add_paragraph(datos["correccion"])
            doc.add_page_break()

        doc.save(path)
        messagebox.showinfo("Exportado", f"Informe de clase guardado en:\n{path}")


if __name__ == "__main__":
    app = CorrectorExamenesProfesor()
    app.mainloop()