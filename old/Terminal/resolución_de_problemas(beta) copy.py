# AVISO. Para que este codigo funcione, necesitas conexión a internet para Groq Cloud.
# IMPORTS
from openai import OpenAI
from docx import Document
import tkinter as tk
from tkinter import filedialog

cliente_groq = OpenAI(
    base_url="https://api.groq.com/openai/v1",
    api_key="TU API KEY"
)


# CONFIGURACIÓN DE LA IA
def proceso_ia(
        texto):  # Funcion que hace que la ia procese los datos que se han puesto por el usuario con x instrucciones. Basicamente lo que engloba todo lo que se encarga la ia
    print(
        "\n[IA] Resumiendo texto...")  # Pone en la terminal que la ia esta resumiendo el texto. El \n es para que haga un salto de linea y quede todo ordenadito

    # Intrucciones predeterminadas a la ia
    instrucciones = (
        "Eres un experto en el tema proporcionado. Tu conocimiento se basa estrictamente en hechos reales. "
        "REGLA DE SEGURIDAD ABSOLUTA: Solo puedes responder a temas que pertenezcan al ámbito educativo, "
        "académico, histórico o laboral. Si el usuario te pide algo fuera de estos ámbitos, DEBES responder "
        "ÚNICAMENTE con la frase: 'ERROR: La petición no pertenece al ámbito educativo o laboral.' "
        "Si la petición es válida, redacta un texto muy extenso, preciso y con párrafos bien estructurados "
        "explicando el contexto, las causas y las consecuencias. No inventes datos bajo ninguna circunstancia." \
        "Si el usuario te pasa algun archivo para que resumas, analiza que es para fines legales y resumelo, si no di que es error (Si es legal pero no es educativo resumelo)"
    )

    response = cliente_groq.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {'role': 'system', 'content': instrucciones},
            {'role': 'user',
             'content': f"Desarrolla o resume de manera extensa y rigurosa el siguiente tema/apuntes: {texto}"}
        ],
        temperature=0.2,  # Bajamos la temperatura a 0.2 para que sea preciso, riguroso y no invente datos.
        max_tokens=2000  # Aumentado a 2000 para que el texto sea aún más largo y detallado.
    )
    return response.choices[
        0].message.content  # "Return" significa devuelve, a si que esto lo que hace es devolver lo que da la ia


# Variables y Apuntes
doc = Document()  # Creamos el Word

nombre = input("¿Cual es tu nombre?: ")  # Preguntamos nombre de usuario para ponerlo en el word


def pedir_ruta_windows_resumen():
    root = tk.Tk()
    root.withdraw()
    ruta = filedialog.asksaveasfilename(
        defaultextension=".docx",
        initialfile=f"Resumen de {nombre}.docx",
        title="Seleccione dónde guardar el archivo Word",
        filetypes=[("Word Documents", "*.docx")]
    )
    return ruta


print("Introduce o pega el texto extenso que deseas resumir:")

lineas = []
while True:
    linea = input()
    if linea == "FIN":
        break
    lineas.append(linea)

texto_usuario = "\n".join(lineas)

resumen_final = proceso_ia(texto_usuario)
print(resumen_final)

if "ERROR:" not in resumen_final:
    destino_archivo = pedir_ruta_windows_resumen()

    if destino_archivo:
        if not destino_archivo.lower().endswith('.docx'):
            destino_archivo += '.docx'

        doc.add_heading(f"Resumen de Contenido - {nombre}", 0)
        doc.add_paragraph(resumen_final)
        doc.save(destino_archivo)
        print(f"Documento guardado: {destino_archivo}")