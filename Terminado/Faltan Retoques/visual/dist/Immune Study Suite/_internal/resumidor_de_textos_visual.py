import customtkinter as ctk
from tkinter import messagebox, filedialog
import ollama
from docx import Document
import threading
import os

# Configuración de tema
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class LanderResumidor(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Lander_Resumidor 2.0 Pro Plus Ultra Deluxe")
        self.geometry("1100x850")

        # Configuración de la IA
        self.modelo = "llama3.2"
        self.instrucciones = (
            "Eres un experto en el tema proporcionado. Tu conocimiento se basa estrictamente en hechos reales. "
            "REGLA DE SEGURIDAD ABSOLUTA: Solo puedes responder a temas que pertenezcan al ámbito educativo, "
            "académico, histórico o laboral. Si el usuario te pide algo fuera de estos ámbitos, DEBES responder "
            "ÚNICAMENTE con la frase: 'ERROR: La petición no pertenece al ámbito educativo o laboral.' "
            "Si la petición es válida, redacta un texto muy extenso, preciso y con párrafos bien estructurados "
            "explicando el contexto, las causas y las consecuencias. No inventes datos bajo ninguna circunstancia."
        )

        self.setup_ui()

    def setup_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # ───── CABECERA ─────
        self.header_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.header_frame.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        
        self.label_titulo = ctk.CTkLabel(
            self.header_frame, 
            text="🎓 Lander_Resumidor 2.0", 
            font=("Segoe UI", 28, "bold")
        )
        self.label_titulo.pack(side="left")

        self.entry_nombre = ctk.CTkEntry(self.header_frame, placeholder_text="Tu nombre...", width=200)
        self.entry_nombre.pack(side="right", padx=10)

        # ───── CUERPO (ENTRADA Y SALIDA) ─────
        self.main_container = ctk.CTkFrame(self, fg_color="transparent")
        self.main_container.grid(row=1, column=0, padx=20, pady=0, sticky="nsew")
        self.main_container.grid_columnconfigure((0, 1), weight=1)
        self.main_container.grid_rowconfigure(0, weight=1)

        # Izquierda: Entrada de apuntes
        self.input_frame = ctk.CTkFrame(self.main_container)
        self.input_frame.grid(row=0, column=0, padx=(0, 10), sticky="nsew")
        
        ctk.CTkLabel(self.input_frame, text="Pega tus apuntes o tema aquí:", font=("Segoe UI", 13, "bold")).pack(pady=10)
        self.txt_input = ctk.CTkTextbox(self.input_frame, font=("Segoe UI", 12))
        self.txt_input.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        # Derecha: Resultado de la IA
        self.output_frame = ctk.CTkFrame(self.main_container)
        self.output_frame.grid(row=0, column=1, padx=(10, 0), sticky="nsew")
        
        ctk.CTkLabel(self.output_frame, text="Resumen generado por IA:", font=("Segoe UI", 13, "bold")).pack(pady=10)
        self.txt_output = ctk.CTkTextbox(self.output_frame, font=("Segoe UI", 13), fg_color=("#ebebeb", "#1a1a1a"))
        self.txt_output.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        # ───── BOTONES Y PROGRESO ─────
        self.footer_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.footer_frame.grid(row=2, column=0, padx=20, pady=20, sticky="ew")

        self.progress_bar = ctk.CTkProgressBar(self.footer_frame, orientation="horizontal")
        self.progress_bar.pack(fill="x", pady=(0, 15))
        self.progress_bar.set(0)

        self.btn_procesar = ctk.CTkButton(
            self.footer_frame, 
            text="🔍 Generar Resumen Riguroso", 
            height=45, 
            font=("Segoe UI", 14, "bold"),
            command=self.iniciar_proceso
        )
        self.btn_procesar.pack(side="left", fill="x", expand=True, padx=(0, 10))

        self.btn_word = ctk.CTkButton(
            self.footer_frame, 
            text="📝 Guardar en Word", 
            height=45,
            width=200, # El tamaño se define aquí, no en el pack
            fg_color="#27ae60",
            hover_color="#219150",
            command=self.exportar_word
        )
        self.btn_word.pack(side="right")

    # ───── LÓGICA ─────

    def iniciar_proceso(self):
        texto = self.txt_input.get("1.0", "end-1c").strip()
        if not texto:
            messagebox.showwarning("Atención", "Por favor, introduce el texto que deseas resumir.")
            return

        self.txt_output.delete("1.0", "end")
        self.btn_procesar.configure(state="disabled")
        self.progress_bar.configure(mode="indeterminate")
        self.progress_bar.start()

        # Hilo para Ollama
        threading.Thread(target=self.ejecutar_ia, args=(texto,), daemon=True).start()

    def ejecutar_ia(self, texto_usuario):
        try:
            response = ollama.chat(
                model=self.modelo,
                messages=[
                    {'role': 'system', 'content': self.instrucciones},
                    {'role': 'user', 'content': f"Desarrolla o resume de manera extensa y rigurosa: {texto_usuario}"}
                ],
                options={"temperature": 0.2, "num_predict": 2000},
                stream=True
            )

            for chunk in response:
                contenido = chunk['message']['content']
                self.after(0, self.escribir_en_output, contenido)

        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", f"Fallo al conectar con Ollama: {e}"))
        finally:
            self.after(0, self.finalizar_proceso)

    def escribir_en_output(self, char):
        self.txt_output.insert("end", char)
        self.txt_output.see("end")

    def finalizar_proceso(self):
        self.progress_bar.stop()
        self.progress_bar.set(1)
        self.btn_procesar.configure(state="normal")

    def exportar_word(self):
        contenido = self.txt_output.get("1.0", "end-1c").strip()
        if not contenido or "ERROR:" in contenido:
            messagebox.showwarning("No se puede guardar", "No hay un resumen válido para exportar.")
            return

        nombre = self.entry_nombre.get().strip() or "Usuario"
        ruta = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Resumen_Lander_{nombre}.docx",
            title="Seleccionar destino del Word"
        )

        if ruta:
            try:
                doc = Document()
                doc.add_heading(f'Resumen Académico de {nombre}', 0)
                doc.add_paragraph(contenido)
                doc.save(ruta)
                messagebox.showinfo("Éxito", f"Archivo guardado en:\n{ruta}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar: {e}")

if __name__ == "__main__":
    app = LanderResumidor()
    app.mainloop()