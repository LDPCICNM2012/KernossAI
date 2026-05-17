import customtkinter as ctk
from tkinter import messagebox, filedialog
import ollama
from docx import Document
import threading
import os

# Configuración visual
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class SolucionadorIA(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("AI Problem Solver - Alto Rendimiento")
        self.geometry("1000x700")

        # Variables de lógica
        self.historial_conversacion = []
        self.instrucciones = (
            "Eres un Solucionador de Problemas de alto rendimiento. "
            "Tu enfoque es: Analizar el problema -> Identificar la causa -> Dar solución técnica/práctica. -> Dar consejos para evitarlo en el futuro. "
            "Responde siempre con estructura de puntos clave y da conversación para acompañar y dar consejos."
        )

        self.setup_ui()

    def setup_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # ───── PANEL LATERAL ─────
        self.sidebar = ctk.CTkFrame(self, width=250, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        
        ctk.CTkLabel(self.sidebar, text="🛠️ Solver IA", font=("Segoe UI", 24, "bold")).pack(pady=20)

        self.entry_nombre = ctk.CTkEntry(self.sidebar, placeholder_text="¿Cuál es tu nombre?")
        self.entry_nombre.pack(fill="x", padx=20, pady=10)

        self.btn_exportar = ctk.CTkButton(self.sidebar, text="💾 Exportar a Word", 
                                         fg_color="#2c3e50", hover_color="#34495e",
                                         command=self.exportar_a_word)
        self.btn_exportar.pack(fill="x", padx=20, pady=10)

        self.btn_limpiar = ctk.CTkButton(self.sidebar, text="🧹 Limpiar Chat", 
                                        fg_color="transparent", border_width=1,
                                        command=self.limpiar_chat)
        self.btn_limpiar.pack(fill="x", padx=20, pady=10)

        ctk.CTkLabel(self.sidebar, text="Estatus de Ollama:", font=("Segoe UI", 12)).pack(side="bottom", pady=5)
        self.status_label = ctk.CTkLabel(self.sidebar, text="🟢 Listo", text_color="green")
        self.status_label.pack(side="bottom", pady=(0, 20))

        # ───── ÁREA DE CHAT ─────
        self.chat_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.chat_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.chat_frame.grid_rowconfigure(0, weight=1)
        self.chat_frame.grid_columnconfigure(0, weight=1)

        self.txt_chat = ctk.CTkTextbox(self.chat_frame, font=("Segoe UI", 14), state="disabled", wrap="word")
        self.txt_chat.grid(row=0, column=0, sticky="nsew", pady=(0, 15))

        # Entrada de usuario
        self.input_frame = ctk.CTkFrame(self.chat_frame, fg_color="transparent")
        self.input_frame.grid(row=1, column=0, sticky="ew")
        self.input_frame.grid_columnconfigure(0, weight=1)

        self.entry_pregunta = ctk.CTkEntry(self.input_frame, placeholder_text="Describe el problema aquí...", height=40)
        self.entry_pregunta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_pregunta.bind("<Return>", lambda e: self.enviar_consulta())

        self.btn_enviar = ctk.CTkButton(self.input_frame, text="Analizar", width=120, height=40, command=self.enviar_consulta)
        self.btn_enviar.grid(row=0, column=1)

    # ───── LÓGICA DE INTERACCIÓN ─────

    def agregar_texto(self, emisor, texto):
        self.txt_chat.configure(state="normal")
        if emisor == "IA":
            self.txt_chat.insert("end", f"\n IA:\n{texto}")
        else:
            self.txt_chat.insert("end", f"\n\n {emisor}:\n{texto}\n")
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")

    def enviar_consulta(self):
        pregunta = self.entry_pregunta.get().strip()
        nombre = self.entry_nombre.get().strip() or "Tú"
        
        if not pregunta:
            return

        self.agregar_texto(nombre, pregunta)
        self.entry_pregunta.delete(0, "end")
        self.status_label.configure(text="🟡 Analizando...", text_color="orange")
        
        # Guardar en memoria
        self.historial_conversacion.append({'role': 'user', 'content': pregunta})
        
        # Lanzar hilo para no congelar la UI
        threading.Thread(target=self.proceso_ia, daemon=True).start()

    def proceso_ia(self):
        try:
            self.txt_chat.configure(state="normal")
            self.txt_chat.insert("end", "\n IA:\n")
            
            response = ollama.chat(
                model="llama3.2",
                messages=[{'role': 'system', 'content': self.instrucciones}] + self.historial_conversacion,
                stream=True,
                options={"temperature": 0.2}
            )

            respuesta_completa = ""
            for chunk in response:
                contenido = chunk['message']['content']
                respuesta_completa += contenido
                # Actualizar texto en tiempo real
                self.after(0, self.stream_update, contenido)
            
            self.historial_conversacion.append({'role': 'assistant', 'content': respuesta_completa})
            self.after(0, lambda: self.status_label.configure(text="🟢 Listo", text_color="green"))
            
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", f"Ollama no responde: {str(e)}"))
        finally:
            self.after(0, lambda: self.txt_chat.configure(state="disabled"))

    def stream_update(self, contenido):
        self.txt_chat.configure(state="normal")
        self.txt_chat.insert("end", contenido)
        self.txt_chat.see("end")
        self.txt_chat.configure(state="disabled")

    def limpiar_chat(self):
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        self.txt_chat.configure(state="disabled")
        self.historial_conversacion = []

    def exportar_a_word(self):
        if not self.historial_conversacion:
            messagebox.showwarning("Atención", "No hay nada que exportar aún.")
            return

        nombre = self.entry_nombre.get().strip() or "Usuario"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Solucion_Problemas_{nombre}.docx",
            title="Guardar informe técnico"
        )

        if path:
            try:
                doc = Document()
                doc.add_heading(f'Informe de Solución - {nombre}', 0)
                
                for msg in self.historial_conversacion:
                    rol = "Tú" if msg['role'] == 'user' else "IA"
                    p = doc.add_paragraph()
                    p.add_run(f"{rol}: ").bold = True
                    p.add_run(msg['content'])
                
                doc.save(path)
                messagebox.showinfo("Éxito", f"Informe guardado en:\n{path}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar el archivo: {e}")

if __name__ == "__main__":
    app = SolucionadorIA()
    app.mainloop()