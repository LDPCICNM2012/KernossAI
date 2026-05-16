import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import json

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

class ApuntadorNotas(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Apuntador de Notas")
        self.geometry("1200x800")
        
        self.archivo_actual = None
        self.notas_guardadas = {}
        self.cargar_notas_guardadas()
        
        # Grid layout
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # ───── PANEL LATERAL ─────
        self.frame_lateral = ctk.CTkFrame(self, width=250, corner_radius=0)
        self.frame_lateral.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        
        self.label_titulo = ctk.CTkLabel(self.frame_lateral, text="📚 Mis Notas", font=("Segoe UI", 20, "bold"))
        self.label_titulo.pack(pady=20, padx=10)

        self.entrada_nombre_nota = ctk.CTkEntry(self.frame_lateral, placeholder_text="Nueva nota...")
        self.entrada_nombre_nota.pack(fill="x", padx=10, pady=5)

        self.btn_nueva = ctk.CTkButton(self.frame_lateral, text="Crear Nota", command=self.crear_nueva_nota, fg_color="#28a745", hover_color="#218838")
        self.btn_nueva.pack(fill="x", padx=10, pady=5)

        # Reemplazo del Textbox por un ScrollableFrame
        self.lista_notas_frame = ctk.CTkScrollableFrame(self.frame_lateral, label_text="Notas Guardadas")
        self.lista_notas_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # ───── ÁREA DE EDICIÓN ─────
        self.frame_editor = ctk.CTkFrame(self, fg_color="transparent")
        self.frame_editor.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.frame_editor.grid_rowconfigure(1, weight=1)
        self.frame_editor.grid_columnconfigure(0, weight=1)

        self.label_nota_abierta = ctk.CTkLabel(self.frame_editor, text="Seleccione una nota", font=("Segoe UI", 16, "italic"))
        self.label_nota_abierta.grid(row=0, column=0, pady=(0, 10), sticky="w")

        self.editor_texto = ctk.CTkTextbox(self.frame_editor, font=("Consolas", 13), border_width=1)
        self.editor_texto.grid(row=1, column=0, sticky="nsew")

        # Botones Inferiores
        self.frame_btns = ctk.CTkFrame(self.frame_editor, fg_color="transparent")
        self.frame_btns.grid(row=2, column=0, pady=(15, 0), sticky="ew")
        
        ctk.CTkButton(self.frame_btns, text="💾 Guardar", command=self.guardar_nota).pack(side="left", padx=5)
        ctk.CTkButton(self.frame_btns, text="🗑️ Eliminar", fg_color="#dc3545", hover_color="#c82333", command=self.eliminar_nota).pack(side="left", padx=5)
        ctk.CTkButton(self.frame_btns, text="📥 Exportar Word", command=self.exportar_nota_word).pack(side="right", padx=5)

        self.actualizar_listbox()

    def crear_nueva_nota(self):
        nombre = self.entrada_nombre_nota.get().strip()
        if nombre:
            if nombre not in self.notas_guardadas:
                self.notas_guardadas[nombre] = ""
                self.guardar_notas_persistente()
                self.abrir_nota(nombre)
                self.actualizar_listbox()
                self.entrada_nombre_nota.delete(0, "end")
            else:
                messagebox.showwarning("Error", "La nota ya existe")

    def abrir_nota(self, nombre):
        self.archivo_actual = nombre
        self.label_nota_abierta.configure(text=f"Editando: {nombre}", font=("Segoe UI", 16, "bold"))
        self.editor_texto.delete("1.0", "end")
        self.editor_texto.insert("1.0", self.notas_guardadas[nombre])

    def actualizar_listbox(self):
        # Limpiar botones anteriores
        for widget in self.lista_notas_frame.winfo_children():
            widget.destroy()

        # Crear un botón por cada nota
        for nombre in self.notas_guardadas.keys():
            btn = ctk.CTkButton(
                self.lista_notas_frame, 
                text=f"• {nombre}", 
                anchor="w",
                fg_color="transparent",
                text_color=("black", "white"),
                hover_color=("#dbdbdb", "#2b2b2b"),
                command=lambda n=nombre: self.abrir_nota(n)
            )
            btn.pack(fill="x", pady=2)

    def guardar_nota(self):
        if self.archivo_actual:
            contenido = self.editor_texto.get("1.0", "end-1c")
            self.notas_guardadas[self.archivo_actual] = contenido
            self.guardar_notas_persistente()
            messagebox.showinfo("Guardado", "Nota guardada correctamente")

    def eliminar_nota(self):
        if self.archivo_actual and messagebox.askyesno("Confirmar", f"¿Eliminar '{self.archivo_actual}'?"):
            del self.notas_guardadas[self.archivo_actual]
            self.guardar_notas_persistente()
            self.archivo_actual = None
            self.editor_texto.delete("1.0", "end")
            self.label_nota_abierta.configure(text="Seleccione una nota")
            self.actualizar_listbox()

    def exportar_nota_word(self):
        if not self.archivo_actual: return
        doc = Document()
        doc.add_heading(self.archivo_actual, 0)
        doc.add_paragraph(self.editor_texto.get("1.0", "end-1c"))
        
        path = filedialog.asksaveasfilename(defaultextension=".docx", initialfile=f"{self.archivo_actual}.docx")
        if path:
            doc.save(path)
            messagebox.showinfo("Éxito", "Exportado correctamente")

    def guardar_notas_persistente(self):
        ruta = os.path.expanduser("~/.apuntador_notas.json")
        with open(ruta, 'w', encoding='utf-8') as f:
            json.dump(self.notas_guardadas, f, ensure_ascii=False, indent=2)

    def cargar_notas_guardadas(self):
        ruta = os.path.expanduser("~/.apuntador_notas.json")
        if os.path.exists(ruta):
            with open(ruta, 'r', encoding='utf-8') as f:
                self.notas_guardadas = json.load(f)

if __name__ == "__main__":
    app = ApuntadorNotas()
    app.mainloop()