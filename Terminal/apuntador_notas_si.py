# AVISO. Para que este codigo funcione, necesitas conexión a internet para Groq Cloud.
from docx import Document
import time
from openai import OpenAI
import tkinter as tk
from tkinter import filedialog

#                                               Imports
# --------------------------------------------------------------------------------------------------------------
#                                               Variables
cliente_groq = OpenAI(
    base_url="https://api.groq.com/openai/v1",
    api_key="TU API KEY"
)

doc = Document()  # Creamos el Word

nombre = input("¿Cual es tu nombre?: ")  # Preguntamos nombre de usuario para ponerlo en el word


# Esto sustituye al applescript de mac. Llama a un menú nativo de Windows para guardar archivos.
def pedir_ruta_windows():
    root = tk.Tk()
    root.withdraw()  # Oculta la ventana principal
    ruta = filedialog.asksaveasfilename(
        defaultextension=".docx",
        initialfile=f"notas_de_{nombre}.docx",
        title="Seleccione dónde guardar el archivo Word",
        filetypes=[("Word Documents", "*.docx")]
    )
    return ruta


# ----------------------------------------------------------------------------------------------------------------
#                                               Codigo
print(
    "Bienvenido a el apuntador que apunta los apuntes apuntados por el apuntador.")  # Bienvenido a el apuntador que apunta los apuntes apuntador por el apuntador.

time.sleep(2)

print(
    "Escribe las notas que quieras y cuando termines escribe FIN")  # Escribe las notas que quieras y cuando termines escribe FIN

time.sleep(1)

print("Recuerda que para cambiar de linea es Enter")  # Recuerda que para cambiar de linea es Enter

time.sleep(1)

print("Ya puedes escribir:")  # Ya puedes escribir:

lineas = []  # [] Es una lista que almacena las lineas que escribas
while True:
    linea = input()  # Aqui es donde tu pones las lineas y lo que quieras escribir en ellas
    if linea == "FIN":  # Si pones FIN, acaba la nota y terminas de escribir para llevarlo al proceso de exportación al word
        break
    lineas.append(linea)  # Mete en la lista de lineas la linea que acabas de escribir

notas = "\n".join(lineas)  # Esto junta TODAS las lineas escritas en un solo bloque
time.sleep(1)

print("Se ha registrado correctamente tus notas")

# ---------------------------------------------------------------------------------------------------------------
#                                               Parte Word

destino_archivo = pedir_ruta_windows()  # Esto es como el script de arriba, es el menú que sale ya llamado para guardar el archivo .word

doc.add_heading(0)  # Ponemos los datos del word que se han preguntado anteriormente
if destino_archivo:
    if not destino_archivo.lower().endswith('.docx'):  # Esto es para que tenga la extensión de Word
        destino_archivo += '.docx'

    # Aquí llamamos a la IA de Groq para que mejore las notas antes de guardarlas (como hacías en el apuntador)
    respuesta = cliente_groq.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system",
             "content": "Eres un asistente académico experto. Tu tarea es tomar los apuntes en bruto del usuario, corregir la ortografía, estructurarlos de manera limpia con títulos o viñetas claros, manteniendo toda la información técnica original."},
            {"role": "user", "content": notas}
        ]
    )
    notas_procesadas = respuesta.choices[0].message.content

    doc.add_paragraph(notas_procesadas, style='List Bullet')  # Font y añadimos el texto
    doc.save(destino_archivo)  # Guardamos
    print(f"Éxito: Archivo guardado en {destino_archivo}")  # Confirmación

nombre_archivo = f"Apuntes de {nombre} "  # Guardamos el word como Notas de (y el nombre del usuario)
try:
    doc.save(nombre_archivo)  # Guardamos el Word
except:
    pass

print(f"Se ha exportado correctamente '{nombre_archivo}', Gracias por utilizar mi programa :)")  # Fin del programa