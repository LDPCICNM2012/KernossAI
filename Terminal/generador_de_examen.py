# AVISO. Para que este codigo funcione, necesitas conexión a internet para Groq Cloud.
# IMPORTS
from openai import OpenAI
from docx import Document
import time
import tkinter as tk
from tkinter import filedialog

cliente_groq = OpenAI(
    base_url="https://api.groq.com/openai/v1",
    api_key="TU API KEY"
)


# Ventana nativa para reemplazar osascript
def pedir_ruta_windows_examen(nombre):
    root = tk.Tk()
    root.withdraw()
    ruta = filedialog.asksaveasfilename(
        defaultextension=".docx",
        initialfile=f"Apuntes de {nombre}.docx",
        title="Seleccione dónde guardar el archivo Word",
        filetypes=[("Word Documents", "*.docx")]
    )
    return ruta


# -----------------------------------------------FUNCIONES-------------------------------------------------
# INSTRUCCIONES PARA LA IA SOBRE CÓMO DEBE SER EL EXAMEN Y LAS EXPORTE A WORD
instrucciones1 = """El examen debe tener:
        1. Un título que tenga que ver con el tema.
        2. El numero de preguntas que elija el usuario de opción múltiple (A, B, C, D, E) o rellena la respuesta dejando un hueco en blanco.
        3. La mitad de las preguntas que haya dicho el usuario que sean de desarrollo (explicar conceptos)(Si el numero no es par, se redondea hacia arriba).
        4. Cuando el usuario responda a las preguntas, debes evaluar sus respuestar sobre 10 dandole la calificacion y explicando porque el fallo o el acierto.
        5. Si el usuario dice que otro examen, generas otro examen diferente al anterior pero con las mismas características. Si el usuario dice que no, le dices que hasta luego y se acaba el programa."""

# INSTRUCCIONES PARA LA IA SOBRE CÓMO DEBE SER EL EXAMEN Y LAS RESPUESTA EN EL PROGRAMA
instrucciones2 = """"Eres un evaluador académico profesional. Tu tarea es crear un examen sobre el tema proporcionado."
        "El examen debe tener:"
        "1. Un título que tenga que ver con el tema."
        "2. El numero de preguntas que elija el usuario de opción múltiple (A, B, C, D, E) o rellena la respuesta dejando un hueco en blanco."
        "3. La mitad de las preguntas que haya dicho el usuario que sean de desarrollo (explicar conceptos)(Si el numero no es par, se redondea hacia arriba)."
        "4. Cuando el usuario responda a las preguntas, debes evaluar sus respuestar sobre 10 dandole la calificacion y explicando porque el fallo o el acierto."
        "5. Si el usuario dice que otro examen, generas otro examen diferente al anterior pero con las mismas características. Si el usuario dice que no, le dices que hasta luego y se acaba el programa."
        "Solo pon la explicación y la respuesta cuando el usuario haya respondido a todas las preguntas, no antes."
        "NO PONGAS LAS RESPUESTAS HASTA QUE EL USUARIO TE DIGA LAS RESPUESTAR (EJ. 1.A 2.C 3.B etc...)"""

# INSTRUCCIONES PARA LA IA SOBRE CÓMO DEBE SER EL EXAMEN Y LAS RESPUESTA EN EL PROGRAMA
instrucciones_corregir = """Eres un evaluador académico profesional. Tu tarea es corregir las respuestas que el USUARIO te dé sobre el examen proporcionado.
REGLAS ESTRICTAS:
- Cuando recibas el examen, ÚNICAMENTE responde con: "Examen recibido. Escribe tus respuestas cuando quieras y las corregiré."
- NO respondas las preguntas tú mismo bajo ningún concepto.
- NO des las respuestas correctas hasta que el usuario haya dado las suyas.
- Espera SIEMPRE a que el usuario escriba sus respuestas antes de corregir nada.
- Solo cuando el usuario envíe sus respuestas, evalúalas sobre 10 explicando cada acierto y fallo."""


# CONFIGURACIÓN DE LA IA PARA QUE EL USUARIO ELIJA EXPORTAR A WORD
def proceso_ia_exportar(
        texto):  # Funcion que hace que la ia procese los datos que se han puesto por el usuario con x instrucciones. Basicamente lo que engloba todo lo que se encarga la ia
    print(
        "\n[IA] Generando examen...")  # Pone en la terminal que la ia esta resumiendo el texto. El \n es para que haga un salto de linea y quede todo ordenadito

    # Intrucciones predeterminadas a la ia
    instrucciones = (instrucciones1)

    response = cliente_groq.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {'role': 'system', 'content': instrucciones},
            {'role': 'user',
             'content': f"Hazme un examen sobre: {texto} con las características mencionadas anteriormente."}
        ],
        stream=True,
        temperature=0.2,
        max_tokens=2000
    )

    respuesta_completa = ""
    for chunk in response:
        if chunk.choices[0].delta.content:
            contenido = chunk.choices[0].delta.content
            print(contenido, end='', flush=True)  # El truco del tiempo real
            respuesta_completa += contenido  # Vamos acumulando la respuesta completa para devolverla al final
    print("\n")  # Espacio al terminar
    return respuesta_completa  # Devolvemos la respuesta completa para que se guarde en la memoria y se pueda exportar luego a Word.


# CONFIGURACIÓN DE LA IA PARA QUE EL USUARIO ELIJA RESPONDER EN EL POROGRAMA
def proceso_ia_responder(
        texto):  # Funcion que hace que la ia procese los datos que se han puesto por el usuario con x instrucciones. Basicamente lo que engloba todo lo que se encarga la ia
    print(
        "\n[IA] Generando examen...")  # Pone en la terminal que la ia esta resumiendo el texto. El \n es para que haga un salto de linea y quede todo ordenadito

    # Intrucciones predeterminadas a la ia
    instrucciones = (instrucciones2)

    conversacionales = [
        {'role': 'system', 'content': instrucciones},
        {'role': 'user',
         'content': f"Hazme un examen sobre: {texto} con las características mencionadas anteriormente."}
    ]

    response = cliente_groq.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=conversacionales,
        stream=True,
        temperature=0.2,
        max_tokens=2000
    )
    respuesta_completa = ""
    for chunk in response:
        if chunk.choices[0].delta.content:
            contenido = chunk.choices[0].delta.content
            print(contenido, end='', flush=True)  # El truco del tiempo real
            respuesta_completa += contenido  # Vamos acumulando la respuesta completa para devolverla al final
    print("\n")  # Espacio al terminar

    conversacionales.append({'role': 'assistant', 'content': respuesta_completa})

    # Bucle para que el usuario pueda mandar sus respuestas y la IA las corrija
    while True:
        respuesta = input("\nEscribe tus respuestas (o 'FIN' para terminar): ")
        if respuesta.upper() == "FIN":
            break
        conversacionales.append({"role": "user", "content": respuesta})

        print("\n[IA]: ", end="")
        stream_resp = cliente_groq.chat.completions.create(model="llama-3.3-70b-versatile", messages=conversacionales,
                                                           stream=True)
        feedback = ""
        for chunk in stream_resp:
            if chunk.choices[0].delta.content:
                content = chunk.choices[0].delta.content
                print(content, end='', flush=True)
                feedback += content
        print("\n")
        conversacionales.append({"role": "assistant", "content": feedback})

    return respuesta_completa  # Devolvemos la respuesta completa para que se guarde en la memoria y se pueda exportar luego a Word.


def generar_examen():
    print(
        "Escribe el tema sobre el que quieres hacer el examen. Escribe 'FIN' en una línea nueva para terminar:")  # Aqui esta lo de mi codigo para que el usuario escriba sus nota/instrucciones y se lo mande a la ia sin limite de lineas.
    lineas = []  # Lista donde se guardan las lineas que el usuario va escribiendo
    while True:
        linea = input(
            "> ")  # Esto es para que salga ete logo > y sepas que tienes que escribir ahí, detallitos importantes.
        if linea.upper() == "FIN":  # Si pones fin en mayusculas sales de escribir a la ia y se manda
            break
        lineas.append(linea)
    return "\n".join(lineas)  # Copila todas las lineas y lo hace en un solo bloque para mandarselo a la ia.


def leer_docx(ruta):
    """Esta función abre un archivo Word y extrae todo su texto"""
    try:
        # Intentamos abrir el documento con la ruta que diste
        doc_lectura = Document(ruta)
        texto_extraido = []
        for parrafo in doc_lectura.paragraphs:
            texto_extraido.append(parrafo.text)

        # Unimos todos los párrafos con saltos de línea
        return "\n".join(texto_extraido)
    except Exception as e:
        return f"Error: No se pudo leer el archivo. Asegúrate de que la ruta sea correcta. {e}"


def nombre_usuario():
    global nombre
    nombre = input("¿Cuál es tu nombre? (Esto se usará para nombrar el examen y la corrección): ").strip()
    if not nombre:
        nombre = "Usuario"  # Nombre por defecto si no se proporciona uno


def proceso_ia_corregir(examen_memoria):
    print("\n[MODO CORRECCIÓN] Iniciado.")
    print("Si exportaste el examen, pega aquí la RUTA del archivo .docx.")
    print("Si no, presiona ENTER para corregir el examen de la memoria actual.")

    entrada = input("Ruta o ENTER: ").strip()

    # Decidimos qué texto enviar basándonos en la entrada
    if entrada.lower().endswith(".docx"):
        print(f"[SISTEMA] Leyendo archivo en: {entrada}...")
        contenido_examen = leer_docx(entrada)
    else:
        print("[SISTEMA] Usando examen almacenado en memoria...")
        contenido_examen = examen_memoria

    conversacionales = [
        {"role": "system", "content": instrucciones_corregir},
        {"role": "user",
         "content": f"Este es el examen. NO lo corrijas tú. NO des las respuestas correctas. Solo acusa recibo y espera a que yo te dé MIS respuestas:\n\n{contenido_examen}"}
    ]

    # Llamada inicial para que la IA acuse recibo del examen y pida las respuestas
    print("\n[IA]: ", end="")
    stream_inicial = cliente_groq.chat.completions.create(model="llama-3.3-70b-versatile", messages=conversacionales,
                                                          stream=True)
    respuesta_inicial = ""
    for chunk in stream_inicial:
        if chunk.choices[0].delta.content:
            content = chunk.choices[0].delta.content
            print(content, end='', flush=True)
            respuesta_inicial += content
    print("\n")
    conversacionales.append({"role": "assistant", "content": respuesta_inicial})

    while True:
        respuesta = input("\nEscribe tus respuestas (o 'FIN' para terminar y exportar corrección): ")

        if respuesta.upper() == "FIN":
            # Opción de exportar la corrección a Word
            exportar_correccion = input("¿Quieres exportar esta corrección a Word? (S/N): ")
            if exportar_correccion.upper() == "S":
                doc_corr = Document()
                doc_corr.add_heading(f"Corrección de Examen - {nombre}", 0)
                for msg in conversacionales:
                    if msg['role'] == 'system': continue
                    rol = "IA" if msg['role'] == 'assistant' else "Usuario"
                    p = doc_corr.add_paragraph()
                    p.add_run(f"{rol}: ").bold = True
                    p.add_run(msg['content'])

                # Usamos el Dialog de Windows para guardar
                try:
                    ruta_guardar = pedir_ruta_windows_examen(nombre)
                    if ruta_guardar:
                        if not ruta_guardar.endswith(".docx"): ruta_guardar += ".docx"
                        doc_corr.save(ruta_guardar)
                        print(f"Corrección guardada en: {ruta_guardar}")
                except:
                    print("No se pudo guardar el archivo.")
            break

        conversacionales.append({"role": "user", "content": respuesta})

        print("\n[IA Corrigiendo]: ", end="")
        stream = cliente_groq.chat.completions.create(model="llama-3.3-70b-versatile", messages=conversacionales,
                                                      stream=True)

        feedback = ""
        for chunk in stream:
            if chunk.choices[0].delta.content:
                content = chunk.choices[0].delta.content
                print(content, end='', flush=True)
                feedback += content
        print("\n")
        conversacionales.append({"role": "assistant", "content": feedback})


# -----------------------------------------------PROGRAMA PRINCIPAL-------------------------------------------------

Apuntes_finales = ""

while True:

    nombre_usuario()

    # ----------------------------------------------Condicionales-----------------------------------------------------------
    destino_examen = input(
        "¿Que prefieres?: 1. Exportar el examen a Word y responderlo ahí(luego tendrás que copiar la ruta del archivo y pasarla al otro proyecto), 2. Responder aquí mismo en el programa? 3: Corregir el examen almacenado en la memoria(tienes que haber generado un examen antes) O 4. Salir: ")  # Aqui le das al usuario la opción de exportar a Word o responder en el programa.
    doc = Document()

    if destino_examen == "1":
        Apuntes_puras = generar_examen()
        Apuntes_finales = proceso_ia_exportar(Apuntes_puras)

        try:  # El try es para que haga lo que hay abajo y si no va que no pete como un tonto si no que te diga pq.
            print("Abriendo menú para que guardes tu archivo...")
            destino_archivo = pedir_ruta_windows_examen(nombre)  # Aqui llamamos al menu de guardado.

            if destino_archivo:
                if not destino_archivo.lower().endswith(
                        '.docx'):  # Basicamente (que esto no lo hago ni yo) es para que el usuario si no pone la terminacion de word .docx la pone automaticamente. Si no sería horrible.
                    destino_archivo += '.docx'  # La extensión.

                doc.add_heading(f'Examen de {nombre}', 0)  # Pone titulo al Word con el nombre del usuario.
            doc.add_paragraph(Apuntes_finales)  # Pone los Apuntes ya procesados por la ia en el Word.

            doc.save(destino_archivo)  # Destino del archivo que se guardará donde haya escogido el usuario.
            print(
                f"\n Éxito: Archivo guardado en {destino_archivo}")  # Confirmación de que se ha guardado el archivo.")
        except Exception as e:
            print(f"Error al guardar: {e}")  # Si hay un error, que te diga pq y donde

    elif destino_examen == "2":
        Apuntes_puras = generar_examen()
        Apuntes_finales = proceso_ia_responder(Apuntes_puras)

    elif destino_examen == "3":
        proceso_ia_corregir(Apuntes_finales)

    elif destino_examen == "4":
        break

# COSAS QUE FALTAN:
# -QUE CORRIGA BIEN EL CBRN
# -PROBAR EL EXPORTE A WORD
