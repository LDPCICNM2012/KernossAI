import os
from docx import Document
from openai import OpenAI

cliente_groq = OpenAI(
    base_url="https://api.groq.com/openai/v1",
    api_key="TU API KEY"
)

doc = Document()

nombre = input("¿Cuál es tu nombre?: ")
texto_usuario = input("Introduce o pega el texto extenso que deseas resumir:\n")

instrucciones = "Eres un experto en el tema proporcionado. Tu conocimiento se basa estrictamente en hechos reales. REGLA DE SEGURIDAD ABSOLUTA: Solo puedes responder a temas que pertenezcan al ámbito educativo, académico, histórico o laboral. Si el usuario te pide algo fuera de estos ámbitos, DEBES responder ÚNICAMENTE con la frase: 'ERROR: La petición no pertenece al ámbito educativo o laboral.' Si la petición es válida, redacta un texto muy extenso, preciso y con párrafos bien estructurados explicando el contexto, las causas y las consecuencias. No inventes datos bajo ninguna circunstancia."

response = cliente_groq.chat.completions.create(
    model="llama-3.3-70b-versatile",
    messages=[
        {"role": "system", "content": instrucciones},
        {"role": "user", "content": texto_usuario}
    ]
)
resumen_final = response.choices[0].message.content

print(resumen_final)

if "ERROR:" not in resumen_final:
    nombre_archivo = f"Resumen_de_{nombre}.docx"
    doc.add_heading(f"Resumen de Contenido - {nombre}", 0)
    doc.add_paragraph(resumen_final)
    doc.save(nombre_archivo)
    print(f"Documento guardado: {nombre_archivo}")