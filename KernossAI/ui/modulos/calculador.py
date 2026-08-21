"""
KernossAI - Módulo: Calculador de Medias y Ponderaciones
Registro de materias, bloques porcentuales, gráficos interactivos y exportación Word.
"""

import os
from datetime import datetime
import customtkinter as ctk
from tkinter import messagebox
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from KernossAI.core.theme import (
    COLOR_BG_CARD,
    COLOR_BG_CARD_LIGHT,
    COLOR_BG_SURFACE,
    COLOR_BORDER,
    COLOR_ACCENT_PRIMARY,
    COLOR_ACCENT_HOVER,
    COLOR_ACCENT_CYAN,
    COLOR_ACCENT_SKY,
    COLOR_ACCENT_PURPLE,
    COLOR_ACCENT_PURPLE_HOVER,
    COLOR_TEXT_MAIN,
    COLOR_TEXT_MUTED,
    COLOR_SUCCESS,
    COLOR_SUCCESS_HOVER,
    COLOR_DANGER_HOVER,
)
from KernossAI.core.i18n import t


class ModuloCalculador(ctk.CTkFrame):
    """Módulo de cálculo de calificaciones ponderadas y visualización analítica."""
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.notas_finales = []
        self.nombres_notas = []
        self.porcentajes = []
        self.canvas_grafico = None
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=1)

        frame_titulo = ctk.CTkFrame(self, fg_color="transparent")
        frame_titulo.pack(fill="x", padx=20, pady=(20, 15))
        ctk.CTkLabel(frame_titulo, text=t("calc_titulo"),
                     font=("Segoe UI", 28, "bold"),
                     text_color=COLOR_ACCENT_SKY).pack()
        ctk.CTkLabel(frame_titulo, text=t("calc_subtitulo"),
                     font=("Segoe UI", 13),
                     text_color=COLOR_TEXT_MUTED).pack(pady=(3, 0))

        self.frame_entrada = ctk.CTkFrame(self, corner_radius=14, border_width=1,
                                          fg_color=COLOR_BG_CARD,
                                          border_color=COLOR_BORDER)
        self.frame_entrada.pack(padx=25, pady=15, fill="both", expand=False)
        self.frame_entrada.grid_columnconfigure((0, 1), weight=1)

        ctk.CTkLabel(self.frame_entrada, text=t("calc_lbl_materia"),
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=0, padx=15, pady=(15, 5), sticky="w")
        self.entrada_nombre = ctk.CTkEntry(self.frame_entrada, placeholder_text=t("calc_plh_materia"),
                                           height=40, font=("Segoe UI", 12),
                                           fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entrada_nombre.grid(row=1, column=0, padx=15, pady=(0, 12), sticky="ew")

        ctk.CTkLabel(self.frame_entrada, text=t("calc_lbl_nota"),
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=1, padx=15, pady=(15, 5), sticky="w")
        self.entrada_nota_directa = ctk.CTkEntry(self.frame_entrada, placeholder_text=t("calc_plh_nota"),
                                                  height=40, font=("Segoe UI", 12),
                                                  fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entrada_nota_directa.grid(row=1, column=1, padx=15, pady=(0, 12), sticky="ew")

        ctk.CTkLabel(self.frame_entrada,
                     text=t("calc_lbl_pct"),
                     font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=2, column=0, padx=15, pady=(0, 5), sticky="w")
        self.entrada_porcentaje = ctk.CTkEntry(self.frame_entrada,
                                               placeholder_text=t("calc_plh_pct"),
                                               height=40, font=("Segoe UI", 12),
                                               fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entrada_porcentaje.grid(row=3, column=0, columnspan=2, padx=15, pady=(0, 15), sticky="ew")

        frame_botones = ctk.CTkFrame(self, fg_color="transparent")
        frame_botones.pack(fill="x", padx=25, pady=8)
        frame_botones.grid_columnconfigure((0, 1), weight=1)

        ctk.CTkButton(frame_botones, text=t("calc_btn_guardar_nota"),
                       command=self.agregar_nota_principal,
                       fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                       height=40, font=("Segoe UI", 12, "bold")).grid(row=0, column=0, padx=(0, 6), sticky="ew")
        ctk.CTkButton(frame_botones, text=t("calc_btn_bloques"),
                       command=self.gestionar_subnotas,
                       fg_color=COLOR_BG_SURFACE, border_width=1,
                       border_color=COLOR_ACCENT_CYAN,
                       hover_color=COLOR_ACCENT_PRIMARY,
                       height=40, font=("Segoe UI", 12, "bold")).grid(row=0, column=1, padx=(6, 0), sticky="ew")

        frame_principal = ctk.CTkFrame(self, fg_color="transparent")
        frame_principal.pack(padx=25, pady=12, fill="both", expand=True)
        frame_principal.grid_columnconfigure((0, 1), weight=1)
        frame_principal.grid_rowconfigure(1, weight=1)

        ctk.CTkLabel(frame_principal, text=t("calc_hdr_registro"),
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=0, pady=(0, 8), padx=(0, 10), sticky="w")
        self.salida_texto = ctk.CTkTextbox(frame_principal, width=500, height=320,
                                           font=("Consolas", 12), corner_radius=12,
                                           fg_color=COLOR_BG_CARD, border_width=1,
                                           border_color=COLOR_BORDER, state="disabled")
        self.salida_texto.grid(row=1, column=0, padx=(0, 10), sticky="nsew")

        ctk.CTkLabel(frame_principal, text=t("calc_hdr_grafica"),
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=1, pady=(0, 8), padx=(10, 0), sticky="w")
        self.frame_grafico = ctk.CTkFrame(frame_principal, corner_radius=12, border_width=1,
                                          fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.frame_grafico.grid(row=1, column=1, padx=(10, 0), sticky="nsew")

        frame_acciones = ctk.CTkFrame(self, fg_color="transparent")
        frame_acciones.pack(padx=25, pady=15, fill="x")
        frame_acciones.grid_columnconfigure((0, 1, 2), weight=1)

        ctk.CTkButton(frame_acciones, text=t("calc_btn_calcular"),
                      command=self.calcular_total_final,
                      height=44, font=("Segoe UI", 13, "bold"),
                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER).grid(row=0, column=0, padx=(0, 6), sticky="ew")
        ctk.CTkButton(frame_acciones, text=t("calc_btn_limpiar"),
                      command=self.limpiar_datos,
                      fg_color=COLOR_BG_SURFACE, hover_color=COLOR_DANGER_HOVER,
                      border_width=1, border_color=COLOR_BORDER,
                      height=44, font=("Segoe UI", 13, "bold")).grid(row=0, column=1, padx=3, sticky="ew")
        ctk.CTkButton(frame_acciones, text=t("calc_btn_exportar"),
                      command=self.exportar_a_word,
                      fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                      height=44, font=("Segoe UI", 13, "bold")).grid(row=0, column=2, padx=(6, 0), sticky="ew")

    def calcular_media(self, lista):
        if not lista:
            return 0
        return sum(lista) / len(lista)

    def actualizar_grafico(self):
        if not self.notas_finales:
            return
        if self.canvas_grafico:
            self.canvas_grafico.get_tk_widget().destroy()
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
        fig.patch.set_facecolor('#0a1124')
        colores = ['#10b981' if n >= 7 else '#f59e0b' if n >= 5 else '#ef4444' for n in self.notas_finales]
        ax1.bar(range(len(self.nombres_notas)), self.notas_finales, color=colores, alpha=0.85, edgecolor='#3b82f6')
        ax1.set_xticks(range(len(self.nombres_notas)))
        ax1.set_xticklabels(self.nombres_notas, rotation=35, ha='right', fontsize=9, color='#cbd5e1')
        ax1.set_ylabel('Calificación', color='#94a3b8', fontsize=10)
        ax1.set_title('Notas por Asignatura', color='#f8fafc', fontsize=12, fontweight='bold')
        ax1.set_ylim(0, 10)
        ax1.set_facecolor('#0f1a35')
        ax1.tick_params(colors='#94a3b8')
        ax1.grid(axis='y', alpha=0.2, color='#3b82f6')

        promedio = self.calcular_media(self.notas_finales)
        por_encima = sum(1 for n in self.notas_finales if n >= promedio)
        por_debajo = sum(1 for n in self.notas_finales if n < promedio)
        ax2.pie([max(por_encima, 0.001), max(por_debajo, 0.001)],
                labels=[f'Arriba media\n({por_encima})', f'Debajo media\n({por_debajo})'],
                colors=['#10b981', '#ef4444'], autopct='%1.1f%%', startangle=90,
                textprops={'color': '#f8fafc', 'fontsize': 10})
        ax2.set_facecolor('#0a1124')
        ax2.set_title(f'Distribución (Promedio: {promedio:.2f})', color='#38bdf8', fontsize=12, fontweight='bold')
        plt.tight_layout()
        self.canvas_grafico = FigureCanvasTkAgg(fig, master=self.frame_grafico)
        self.canvas_grafico.draw()
        self.canvas_grafico.get_tk_widget().pack(fill="both", expand=True)

    def gestionar_subnotas(self):
        nombre_principal = self.entrada_nombre.get().strip()
        if not nombre_principal:
            messagebox.showwarning("Atención", "Escribe el nombre de la asignatura primero.")
            return
        d = ctk.CTkInputDialog(text=f"¿Cuántos bloques tiene {nombre_principal}?\n(Ej: 3 → Exámenes, Prácticas, Trabajo)",
                               title="Bloques de calificación")
        res = d.get_input()
        self.focus_force(); self.lift()
        if not res or not res.isdigit():
            return
        num_bloques = int(res)
        self.salida_texto.configure(state="normal")
        self.salida_texto.insert("end", f"\n{'─'*55}\n  📚 {nombre_principal.upper()}\n{'─'*55}\n")
        nota_final_asignatura = 0.0
        suma_pesos_bloques = 0.0
        for i in range(num_bloques):
            d = ctk.CTkInputDialog(text=f"Nombre del bloque {i+1}:", title="Nombre del bloque")
            nombre_bloque = d.get_input(); self.focus_force(); self.lift()
            if nombre_bloque is None: break
            d = ctk.CTkInputDialog(text=f"¿Cuánto vale '{nombre_bloque}' en % sobre el total?", title="% del bloque")
            peso_bloque_str = d.get_input(); self.focus_force(); self.lift()
            if peso_bloque_str is None: break
            try:
                peso_bloque = float(peso_bloque_str)
            except ValueError:
                messagebox.showerror("Error", "Número inválido."); continue
            d = ctk.CTkInputDialog(text=f"¿Cuántas notas hay dentro de '{nombre_bloque}'?", title="Notas del bloque")
            res_sub = d.get_input(); self.focus_force(); self.lift()
            if not res_sub or not res_sub.isdigit(): continue
            num_sub = int(res_sub)
            self.salida_texto.insert("end", f"\n  📂 {nombre_bloque} ({peso_bloque:.0f}% del total)\n")
            nota_bloque_ponderada = 0.0; suma_pesos_sub = 0.0
            for j in range(num_sub):
                d = ctk.CTkInputDialog(text=f"Nombre de la nota {j+1}:", title="Nota")
                nombre_sub = d.get_input(); self.focus_force(); self.lift()
                if nombre_sub is None: break
                d = ctk.CTkInputDialog(text=f"¿Cuánto vale '{nombre_sub}' en % dentro de '{nombre_bloque}'?", title=f"% dentro de {nombre_bloque}")
                peso_sub_str = d.get_input(); self.focus_force(); self.lift()
                if peso_sub_str is None: break
                try:
                    peso_sub = float(peso_sub_str)
                except ValueError:
                    messagebox.showerror("Error", "Número inválido."); continue
                d = ctk.CTkInputDialog(text=f"Calificación de '{nombre_sub}':", title="Calificación")
                valor_sub_str = d.get_input(); self.focus_force(); self.lift()
                if valor_sub_str is None: break
                try:
                    valor_sub = float(valor_sub_str)
                except ValueError:
                    messagebox.showerror("Error", "Número inválido."); continue
                nota_bloque_ponderada += valor_sub * (peso_sub / 100)
                suma_pesos_sub += peso_sub
                self.salida_texto.insert("end", f"      • {nombre_sub:20} {peso_sub:.0f}% → {valor_sub:.2f}\n")
            if suma_pesos_sub > 0 and suma_pesos_sub != 100:
                nota_bloque_ponderada = nota_bloque_ponderada / (suma_pesos_sub / 100)
            self.salida_texto.insert("end", f"    ✓ Nota bloque '{nombre_bloque}': {nota_bloque_ponderada:.2f}\n")
            nota_final_asignatura += nota_bloque_ponderada * (peso_bloque / 100)
            suma_pesos_bloques += peso_bloque
        if suma_pesos_bloques > 0 and suma_pesos_bloques != 100:
            nota_final_asignatura = nota_final_asignatura / (suma_pesos_bloques / 100)
        self.nombres_notas.append(nombre_principal)
        self.notas_finales.append(nota_final_asignatura)
        self.porcentajes.append(100.0)
        self.salida_texto.insert("end", f"\n  {'═'*50}\n  ✅ NOTA FINAL {nombre_principal.upper()}: {nota_final_asignatura:.2f}\n  {'═'*50}\n\n")
        self.salida_texto.configure(state="disabled")
        self.entrada_nombre.delete(0, "end")
        self.actualizar_grafico()

    def agregar_nota_principal(self):
        nombre = self.entrada_nombre.get().strip()
        nota_str = self.entrada_nota_directa.get().strip()
        porcentaje_str = self.entrada_porcentaje.get().strip()
        if nombre and nota_str:
            try:
                nota = float(nota_str)
                if porcentaje_str:
                    porcentaje_sobre_total = float(porcentaje_str)
                else:
                    total = len(self.notas_finales) + 1
                    porcentaje_sobre_total = (1 / total) * 100
                self.nombres_notas.append(nombre)
                self.notas_finales.append(nota)
                self.porcentajes.append(porcentaje_sobre_total)
                self.salida_texto.configure(state="normal")
                self.salida_texto.insert("end", f"  ✓ {nombre:25} → {nota:.2f} ({porcentaje_sobre_total:.1f}%)\n")
                self.salida_texto.configure(state="disabled")
                self.entrada_nombre.delete(0, "end")
                self.entrada_nota_directa.delete(0, "end")
                self.entrada_porcentaje.delete(0, "end")
                self.actualizar_grafico()
            except ValueError:
                messagebox.showerror("Error", "La calificación y el porcentaje deben ser números válidos.")
        else:
            messagebox.showwarning("Campos incompletos", "Por favor completa materia y calificación.")

    def calcular_total_final(self):
        if not self.notas_finales:
            messagebox.showinfo("Sin datos", "No hay calificaciones guardadas para calcular.")
            return
        suma_ponderada = sum(n * (p / 100) for n, p in zip(self.notas_finales, self.porcentajes))
        total_porcentaje = sum(self.porcentajes)
        media_total = suma_ponderada / (total_porcentaje / 100) if total_porcentaje != 100 else suma_ponderada
        self.salida_texto.configure(state="normal")
        self.salida_texto.insert("end", f"\n{'═'*55}\n  📊 RESULTADO FINAL\n")
        self.salida_texto.insert("end", f"  Asignaturas: {len(self.nombres_notas)}\n")
        self.salida_texto.insert("end", f"  Peso total asignado: {total_porcentaje:.1f}%\n")
        self.salida_texto.insert("end", f"  Promedio Ponderado General: {media_total:.2f}\n{'═'*55}\n\n")
        self.salida_texto.configure(state="disabled")
        self.salida_texto.see("end")

    def limpiar_datos(self):
        if messagebox.askyesno("Confirmar", "¿Deseas limpiar todas las notas?"):
            self.notas_finales.clear()
            self.nombres_notas.clear()
            self.porcentajes.clear()
            self.salida_texto.configure(state="normal")
            self.salida_texto.delete("1.0", "end")
            self.salida_texto.configure(state="disabled")
            self.entrada_nombre.delete(0, "end")
            self.entrada_nota_directa.delete(0, "end")
            self.entrada_porcentaje.delete(0, "end")
            if self.canvas_grafico:
                self.canvas_grafico.get_tk_widget().destroy()
                self.canvas_grafico = None

    def exportar_a_word(self):
        if not self.notas_finales:
            messagebox.showwarning("Sin datos", "No hay calificaciones para exportar.")
            return
        doc = Document()
        titulo = doc.add_heading('Reporte de Calificaciones – KernossAI', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha = doc.add_paragraph(f'Generado el: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        tabla = doc.add_table(rows=len(self.nombres_notas) + 1, cols=3)
        tabla.style = 'Light Grid Accent 1'
        enc = tabla.rows[0].cells
        enc[0].text = '#'; enc[1].text = 'Asignatura'; enc[2].text = 'Calificación'
        for i, (nombre, nota) in enumerate(zip(self.nombres_notas, self.notas_finales), start=1):
            fila = tabla.rows[i].cells
            fila[0].text = str(i); fila[1].text = nombre; fila[2].text = f"{nota:.2f}"
        doc.add_paragraph()
        promedio = self.calcular_media(self.notas_finales)
        p = doc.add_paragraph(f'Promedio General: {promedio:.2f}')
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(14)
        ruta = os.path.join(os.path.expanduser("~/Documents"),
                            f"Reporte_Calificaciones_{datetime.now().strftime('%d%m%Y_%H%M%S')}.docx")
        doc.save(ruta)
        messagebox.showinfo("Éxito", f"Documento exportado a:\n{ruta}")
