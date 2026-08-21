"""
KernossAI - Módulo: Mapas Mentales y Conceptuales con IA
Generador de diagramas conceptuales extensos, visualizador Matplotlib HD y exportación.
"""

import os
import re
import json
import threading
from datetime import datetime
import numpy as np
import customtkinter as ctk
from tkinter import messagebox, filedialog
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
from docx.shared import Inches

from KernossAI.core.theme import (
    COLOR_BG_CARD,
    COLOR_BG_CARD_LIGHT,
    COLOR_BG_SURFACE,
    COLOR_BORDER,
    COLOR_ACCENT_PRIMARY,
    COLOR_ACCENT_HOVER,
    COLOR_ACCENT_CYAN,
    COLOR_ACCENT_SKY,
    COLOR_ACCENT_PURPLE,
    COLOR_ACCENT_PURPLE_HOVER,
    COLOR_TEXT_MAIN,
    COLOR_TEXT_MUTED,
    COLOR_TEXT_DIM,
    COLOR_SUCCESS,
    COLOR_SUCCESS_HOVER,
)
from KernossAI.core.i18n import t
from KernossAI.core.auth import llamar_groq, llamar_gemini

INSTRUCCIONES_MAPA_MENTAL = """
Eres un pedagogo experto en síntesis visual avanzada, mapas conceptuales educativos y pensamiento sistémico.
Tu objetivo es transformar el tema académico en una estructura de Mapa Mental COMPLETA, EXTENSA, PROFUNDA, JERÁRQUICA Y GRANDE.

CRITERIOS OBLIGATORIOS PARA EL MAPA MENTAL:
1. Genera entre 6 y 8 RAMAS PRINCIPALES exhaustivas que cubran todos los ángulos del tema (ej: Origen e Historia, Conceptos Fundamentales, Mecanismos/Funcionamiento, Clasificación y Tipos, Aplicaciones Prácticas, Ventajas y Desafíos, Ejemplos Reales, Futuro/Tendencias).
2. Para CADA rama principal, incluye entre 3 y 5 SUBCONCEPTOS específicos, con nombres concisos y detalles explicativos enriquecedores.
3. El resultado debe ser amplio, formativo y de alto valor pedagógico.

Debes responder ÚNICAMENTE con un objeto JSON válido (sin explicaciones previas ni posteriores, sin bloques de razonamiento):
{
  "tema_central": "Título conciso del tema central",
  "descripcion_general": "Resumen conceptual sintético de 2-3 frases.",
  "ramas": [
    {
      "titulo": "1. Fundamentos & Origen",
      "descripcion": "Base conceptual y antecedentes históricos.",
      "sub_conceptos": [
        { "nombre": "Definición Clave", "detalle": "Explicación exacta y formal del concepto." },
        { "nombre": "Antecedentes", "detalle": "Origen histórico y evolución." },
        { "nombre": "Principios Básicos", "detalle": "Leyes o axiomas que lo rigen." },
        { "nombre": "Contexto", "detalle": "Marco de aplicación." }
      ]
    }
  ]
}
"""


class ModuloMapaMental(ctk.CTkFrame):
    """Módulo interactivo de mapas mentales con editor JSON, visualización gráfica y exportación."""
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.modelo_actual = "groq"
        self.datos_mapa = None
        self.fig = None
        self.ax = None
        self.canvas_grafico = None
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=4, minsize=420)
        self.grid_columnconfigure(1, weight=6, minsize=550)
        self.grid_rowconfigure(0, weight=1)

        # Panel izquierdo
        panel_izq = ctk.CTkFrame(self, corner_radius=14, fg_color=COLOR_BG_CARD,
                                 border_width=1, border_color=COLOR_BORDER)
        panel_izq.grid(row=0, column=0, sticky="nsew", padx=(20, 10), pady=20)
        panel_izq.grid_rowconfigure(7, weight=1)
        panel_izq.grid_columnconfigure(0, weight=1)

        frame_header = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_header.grid(row=0, column=0, sticky="ew", padx=18, pady=(16, 10))

        ctk.CTkLabel(frame_header, text=t("mapa_titulo"),
                     font=("Segoe UI", 18, "bold"), text_color=COLOR_ACCENT_SKY).pack(anchor="w")
        ctk.CTkLabel(frame_header, text="Genera mapas mentales estructurados y profundos con IA.",
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(anchor="w", pady=(2, 0))

        ctk.CTkLabel(panel_izq, text=t("mapa_lbl_tema"),
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=1, column=0, sticky="w", padx=18, pady=(4, 2))
        self.entry_tema = ctk.CTkEntry(panel_izq, placeholder_text="Ej: Fotosíntesis, Revolución Industrial, Redes Neuronales...",
                                       height=38, font=("Segoe UI", 12),
                                       fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_tema.grid(row=2, column=0, sticky="ew", padx=18, pady=(0, 10))

        frame_opts = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_opts.grid(row=3, column=0, sticky="ew", padx=18, pady=(0, 10))
        frame_opts.grid_columnconfigure((0, 1), weight=1)

        ctk.CTkLabel(frame_opts, text="Level / Curso:", font=("Segoe UI", 11, "bold"),
                     text_color=COLOR_TEXT_MUTED).grid(row=0, column=0, sticky="w", padx=(0, 6), pady=(0, 2))
        self.combo_nivel = ctk.CTkComboBox(frame_opts,
                                           values=["Secondary / ESO", "High School", "University", "Primary", "General"],
                                           font=("Segoe UI", 11), height=34,
                                           fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.combo_nivel.set("Secondary / ESO")
        self.combo_nivel.grid(row=1, column=0, sticky="ew", padx=(0, 6))

        ctk.CTkLabel(frame_opts, text="Focus / Enfoque:", font=("Segoe UI", 11, "bold"),
                     text_color=COLOR_TEXT_MUTED).grid(row=0, column=1, sticky="w", padx=(6, 0), pady=(0, 2))
        self.entry_enfoque = ctk.CTkEntry(frame_opts, placeholder_text="Puntos clave opcionales...",
                                          font=("Segoe UI", 11), height=34,
                                          fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_enfoque.grid(row=1, column=1, sticky="ew", padx=(6, 0))

        frame_ia_bar = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_ia_bar.grid(row=4, column=0, sticky="ew", padx=18, pady=(0, 10))
        frame_ia_bar.grid_columnconfigure(0, weight=1)

        frame_model_switch = ctk.CTkFrame(frame_ia_bar, fg_color=COLOR_BG_CARD_LIGHT,
                                          border_width=1, border_color=COLOR_BORDER, corner_radius=8)
        frame_model_switch.pack(side="left")

        self.btn_groq = ctk.CTkButton(frame_model_switch, text="⚡ Groq", height=28, width=75,
                                      font=("Segoe UI", 10, "bold"),
                                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                      command=lambda: self._set_modelo("groq"))
        self.btn_groq.pack(side="left", padx=2, pady=2)

        self.btn_gemini = ctk.CTkButton(frame_model_switch, text="🧠 Gemini", height=28, width=75,
                                        font=("Segoe UI", 10, "bold"),
                                        fg_color="transparent", hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                        command=lambda: self._set_modelo("gemini"))
        self.btn_gemini.pack(side="left", padx=2, pady=2)

        self.btn_generar = ctk.CTkButton(frame_ia_bar, text="🧠 Generar Mapa Extenso", height=34,
                                         font=("Segoe UI", 12, "bold"),
                                         fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                         command=self.generar_mapa_mental)
        self.btn_generar.pack(side="right", fill="x", expand=True, padx=(10, 0))

        self.lbl_status = ctk.CTkLabel(panel_izq, text="Listo para generar mapa conceptual profundo", font=("Segoe UI", 11),
                                       text_color=COLOR_TEXT_DIM)
        self.lbl_status.grid(row=5, column=0, sticky="w", padx=18, pady=(0, 6))

        frame_edit_header = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_edit_header.grid(row=6, column=0, sticky="ew", padx=18, pady=(4, 4))
        ctk.CTkLabel(frame_edit_header, text="📝 Editor de Estructura:",
                     font=("Segoe UI", 11, "bold"), text_color=COLOR_ACCENT_CYAN).pack(side="left")

        self.btn_redibujar = ctk.CTkButton(frame_edit_header, text="🔄 Redibujar Mapa", height=24, width=120,
                                           font=("Segoe UI", 10, "bold"),
                                           fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_HOVER,
                                           border_width=1, border_color=COLOR_BORDER,
                                           command=self.redibujar_desde_editor)
        self.btn_redibujar.pack(side="right")

        self.txt_estructura = ctk.CTkTextbox(panel_izq, font=("Consolas", 11), wrap="word",
                                             fg_color=COLOR_BG_CARD_LIGHT, border_width=1,
                                             border_color=COLOR_BORDER, corner_radius=10)
        self.txt_estructura.grid(row=7, column=0, sticky="nsew", padx=18, pady=(0, 10))

        frame_acciones = ctk.CTkFrame(panel_izq, fg_color="transparent")
        frame_acciones.grid(row=8, column=0, sticky="ew", padx=18, pady=(0, 16))
        frame_acciones.grid_columnconfigure((0, 1), weight=1)

        self.btn_exportar_word = ctk.CTkButton(frame_acciones, text="📄 Exportar a Word (.docx)", height=36,
                                               font=("Segoe UI", 11, "bold"),
                                               fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                               state="disabled",
                                               command=self.exportar_word)
        self.btn_exportar_word.grid(row=0, column=0, sticky="ew", padx=(0, 4))

        self.btn_exportar_img = ctk.CTkButton(frame_acciones, text="🖼️ Guardar Imagen (.png)", height=36,
                                              font=("Segoe UI", 11, "bold"),
                                              fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                                              state="disabled",
                                              command=self.exportar_imagen)
        self.btn_exportar_img.grid(row=0, column=1, sticky="ew", padx=(4, 0))

        # Panel derecho
        self.panel_der = ctk.CTkFrame(self, corner_radius=14, fg_color=COLOR_BG_CARD,
                                      border_width=1, border_color=COLOR_BORDER)
        self.panel_der.grid(row=0, column=1, sticky="nsew", padx=(10, 20), pady=20)
        self.panel_der.grid_rowconfigure(1, weight=1)
        self.panel_der.grid_columnconfigure(0, weight=1)

        header_der = ctk.CTkFrame(self.panel_der, fg_color="transparent")
        header_der.grid(row=0, column=0, sticky="ew", padx=18, pady=(16, 6))

        ctk.CTkLabel(header_der, text="🎨 Vista Gráfica del Mapa Mental",
                     font=("Segoe UI", 16, "bold"), text_color=COLOR_TEXT_MAIN).pack(side="left")

        self.lbl_info_ramas = ctk.CTkLabel(header_der, text="Sin mapa generado",
                                           font=("Segoe UI", 11), text_color=COLOR_TEXT_DIM)
        self.lbl_info_ramas.pack(side="right")

        self.frame_canvas = ctk.CTkFrame(self.panel_der, fg_color="#060a14", corner_radius=10,
                                         border_width=1, border_color=COLOR_BORDER)
        self.frame_canvas.grid(row=1, column=0, sticky="nsew", padx=18, pady=(0, 18))
        self.frame_canvas.grid_rowconfigure(0, weight=1)
        self.frame_canvas.grid_columnconfigure(0, weight=1)

        self._inicializar_canvas_vacio()

    def _set_modelo(self, modelo):
        self.modelo_actual = modelo
        if modelo == "groq":
            self.btn_groq.configure(fg_color=COLOR_ACCENT_PRIMARY)
            self.btn_gemini.configure(fg_color="transparent")
        else:
            self.btn_groq.configure(fg_color="transparent")
            self.btn_gemini.configure(fg_color=COLOR_ACCENT_PURPLE)

    def _inicializar_canvas_vacio(self):
        plt.style.use("dark_background")
        self.fig, self.ax = plt.subplots(figsize=(9, 8), facecolor="#060a14")
        self.ax.set_facecolor("#060a14")
        self.ax.text(0, 0, "🧠 Escribe un tema y haz clic en\n'🧠 Generar Mapa Extenso' para crear un esquema conceptual",
                     color="#94a3b8", fontsize=12, ha="center", va="center",
                     bbox=dict(boxstyle="round,pad=0.9", facecolor="#0b1329", edgecolor="#1e3a6a", lw=1.5))
        self.ax.set_xlim(-8, 8)
        self.ax.set_ylim(-8, 8)
        self.ax.axis("off")

        self.canvas_grafico = FigureCanvasTkAgg(self.fig, master=self.frame_canvas)
        self.canvas_grafico.get_tk_widget().pack(fill="both", expand=True)
        self.canvas_grafico.draw()

    def generar_mapa_mental(self):
        tema = self.entry_tema.get().strip()
        if not tema:
            messagebox.showwarning("Tema requerido", "Introduce el tema o concepto para crear el mapa mental.")
            return

        nivel = self.combo_nivel.get()
        enfoque = self.entry_enfoque.get().strip()

        self.btn_generar.configure(state="disabled")
        self.lbl_status.configure(text="✨ Creando estructura grande con IA...", text_color=COLOR_ACCENT_SKY)

        prompt = f"TEMA PRINCIPAL: {tema}\nNIVEL EDUCATIVO: {nivel}\n"
        if enfoque:
            prompt += f"ENFOQUE Y PUNTOS CLAVE: {enfoque}\n"

        threading.Thread(target=self._thread_generar_mapa, args=(prompt,), daemon=True).start()

    def _thread_generar_mapa(self, prompt):
        try:
            full_prompt = f"{INSTRUCCIONES_MAPA_MENTAL}\n\n{prompt}"
            if self.modelo_actual == "groq":
                respuesta = llamar_groq(full_prompt)
            else:
                respuesta = llamar_gemini(full_prompt)

            datos = self._extraer_json(respuesta)
            if not datos:
                t_nombre = self.entry_tema.get().strip() or "Tema Principal"
                datos = {
                    "tema_central": t_nombre,
                    "descripcion_general": f"Mapa conceptual completo y sistemático sobre {t_nombre}.",
                    "ramas": [
                        {"titulo": "1. Fundamentos & Origen", "descripcion": "Bases teóricas y antecedentes.", "sub_conceptos": [{"nombre": "Definición Formal", "detalle": "Concepto central."}, {"nombre": "Marco Histórico", "detalle": "Evolución y descubrimiento."}, {"nombre": "Principios Básicos", "detalle": "Reglas y leyes esenciales."}]},
                        {"titulo": "2. Estructura & Componentes", "descripcion": "Arquitectura y partes constitutivas.", "sub_conceptos": [{"nombre": "Elemento Núcleo", "detalle": "Módulo principal."}, {"nombre": "Subsistemas", "detalle": "Componentes de apoyo."}, {"nombre": "Interrelaciones", "detalle": "Flujos de comunicación."}]},
                        {"titulo": "3. Metodología & Fases", "descripcion": "Etapas y procedimientos de aplicación.", "sub_conceptos": [{"nombre": "Fase Inicial", "detalle": "Preparación y análisis."}, {"nombre": "Desarrollo", "detalle": "Ejecución técnica."}, {"nombre": "Evaluación", "detalle": "Control y validación."}]},
                        {"titulo": "4. Tipos & Clasificación", "descripcion": "Taxonomía y categorías principales.", "sub_conceptos": [{"nombre": "Variante Principal", "detalle": "Características estándar."}, {"nombre": "Variante Avanzada", "detalle": "Propiedades extendidas."}, {"nombre": "Modelos Híbridos", "detalle": "Combinaciones prácticas."}]},
                        {"titulo": "5. Aplicaciones Prácticas", "descripcion": "Casos de uso e impacto real.", "sub_conceptos": [{"nombre": "Ámbito Educativo", "detalle": "Investigación y pedagogía."}, {"nombre": "Sector Profesional", "detalle": "Implementación práctica."}, {"nombre": "Casos Demostrados", "detalle": "Ejemplos reales de éxito."}]},
                        {"titulo": "6. Retos & Futuro", "descripcion": "Perspectivas y áreas de innovación.", "sub_conceptos": [{"nombre": "Desafíos Actuales", "detalle": "Limitaciones y soluciones."}, {"nombre": "Nuevas Tendencias", "detalle": "Evolución tecnológica."}, {"nombre": "Impacto Global", "detalle": "Proyección a largo plazo."}]}
                    ]
                }

            self.after(0, lambda: self._mostrar_mapa_generado(datos))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error al Generar", f"No se pudo generar el mapa mental: {e}"))
        finally:
            self.after(0, lambda: [
                self.btn_generar.configure(state="normal"),
                self.lbl_status.configure(text="Listo", text_color=COLOR_SUCCESS)
            ])

    def _extraer_json(self, texto):
        if not texto:
            return None
        texto = re.sub(r'<think>.*?</think>', '', texto, flags=re.DOTALL)
        texto = re.sub(r'```(?:think|thought).*?```', '', texto, flags=re.DOTALL)

        match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", texto, re.DOTALL)
        if match:
            candidato = match.group(1)
            try:
                return json.loads(candidato)
            except Exception:
                try:
                    c_limpio = re.sub(r',\s*([\]}])', r'\1', candidato)
                    return json.loads(c_limpio)
                except Exception:
                    pass

        start = texto.find("{")
        end = texto.rfind("}")
        if start != -1 and end != -1 and end > start:
            candidato = texto[start:end+1]
            try:
                return json.loads(candidato)
            except Exception:
                try:
                    c_limpio = re.sub(r',\s*([\]}])', r'\1', candidato)
                    return json.loads(c_limpio)
                except Exception:
                    pass
        return None

    def _mostrar_mapa_generado(self, datos):
        self.datos_mapa = datos

        self.txt_estructura.delete("1.0", "end")
        self.txt_estructura.insert("end", json.dumps(datos, ensure_ascii=False, indent=2))

        self._dibujar_mapa_visual(datos)

        self.btn_exportar_word.configure(state="normal")
        self.btn_exportar_img.configure(state="normal")
        ramas_count = len(datos.get("ramas", []))
        total_subs = sum(len(r.get("sub_conceptos", [])) for r in datos.get("ramas", []))
        self.lbl_info_ramas.configure(
            text=f"🧠 {ramas_count} Ramas • {total_subs} Subconceptos",
            text_color=COLOR_ACCENT_CYAN
        )

    def redibujar_desde_editor(self):
        contenido = self.txt_estructura.get("1.0", "end-1c").strip()
        if not contenido:
            return
        try:
            datos = json.loads(contenido)
            self.datos_mapa = datos
            self._dibujar_mapa_visual(datos)
            self.lbl_status.configure(text="Mapa visual actualizado", text_color=COLOR_SUCCESS)
        except Exception as e:
            messagebox.showerror("Error JSON", f"El formato JSON no es válido:\n{e}")

    def _dibujar_mapa_visual(self, datos):
        self.ax.clear()
        self.ax.set_facecolor("#060a14")

        tema_central = datos.get("tema_central", "Tema Central")
        ramas = datos.get("ramas", [])
        num_ramas = len(ramas)

        if num_ramas == 0:
            self._inicializar_canvas_vacio()
            return

        paleta_colores = [
            "#06b6d4", "#38bdf8", "#6366f1", "#10b981",
            "#f59e0b", "#ec4899", "#8b5cf6", "#14b8a6", "#f43f5e"
        ]

        tema_fmt = "\n".join([tema_central[i:i+18] for i in range(0, len(tema_central), 18)])
        self.ax.text(0, 0, f"🌟\n{tema_fmt}",
                     color="#ffffff", fontsize=11, fontweight="bold", ha="center", va="center",
                     bbox=dict(boxstyle="round,pad=0.75", facecolor="#1e3a8a", edgecolor="#38bdf8", lw=2.5, alpha=0.96))

        for i, rama in enumerate(ramas):
            color = paleta_colores[i % len(paleta_colores)]
            angulo = (2 * np.pi * i / num_ramas) + (np.pi / (num_ramas * 2))

            r_rama = 4.1 if (i % 2 == 0) else 4.7
            x_rama = r_rama * np.cos(angulo)
            y_rama = r_rama * np.sin(angulo)

            self.ax.annotate("", xy=(x_rama, y_rama), xytext=(0, 0),
                             arrowprops=dict(arrowstyle="-", color=color, lw=2.4, alpha=0.85,
                                             connectionstyle="arc3,rad=0.08"))

            titulo_rama = rama.get("titulo", f"Rama {i+1}")
            titulo_fmt = "\n".join([titulo_rama[k:k+14] for k in range(0, len(titulo_rama), 14)])

            self.ax.text(x_rama, y_rama, titulo_fmt,
                         color="#f8fafc", fontsize=9.5, fontweight="bold", ha="center", va="center",
                         bbox=dict(boxstyle="round,pad=0.5", facecolor="#0c162c", edgecolor=color, lw=2.0, alpha=0.96))

            subs = rama.get("sub_conceptos", [])
            num_subs = len(subs)

            for j, sub in enumerate(subs):
                offset_ang = (j - (num_subs - 1) / 2) * (0.30 if num_subs > 1 else 0)
                sub_ang = angulo + offset_ang

                r_sub = 7.1 if (j % 2 == 0) else 7.9
                x_sub = r_sub * np.cos(sub_ang)
                y_sub = r_sub * np.sin(sub_ang)

                self.ax.annotate("", xy=(x_sub, y_sub), xytext=(x_rama, y_rama),
                                 arrowprops=dict(arrowstyle="-", color=color, lw=1.3, ls="--", alpha=0.65,
                                                 connectionstyle="arc3,rad=-0.06"))

                nombre_sub = sub.get("nombre", f"Punto {j+1}")
                nombre_fmt = "\n".join([nombre_sub[k:k+15] for k in range(0, len(nombre_sub), 15)])

                self.ax.text(x_sub, y_sub, nombre_fmt,
                             color="#e2e8f0", fontsize=8, ha="center", va="center",
                             bbox=dict(boxstyle="round,pad=0.35", facecolor="#060d1d", edgecolor=color, lw=1.1, alpha=0.92))

        self.ax.set_xlim(-9.8, 9.8)
        self.ax.set_ylim(-9.8, 9.8)
        self.ax.set_aspect("equal")
        self.ax.axis("off")
        self.canvas_grafico.draw()

    def exportar_imagen(self):
        if not self.datos_mapa or not self.fig:
            messagebox.showwarning("Sin mapa", "Primero genera o redibuja un mapa mental.")
            return

        tema = self.datos_mapa.get("tema_central", "Mapa_Mental").replace(" ", "_")
        path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("Imagen PNG", "*.png"), ("Todos los archivos", "*.*")],
            initialfile=f"MapaMental_{tema}_{datetime.now().strftime('%Y%m%d')}.png"
        )
        if not path:
            return

        try:
            self.fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="#070c18", edgecolor="none")
            messagebox.showinfo("Imagen Guardada", f"Mapa mental exportado con éxito en alta resolución (300 DPI):\n{path}")
        except Exception as e:
            messagebox.showerror("Error al Guardar", f"No se pudo guardar la imagen:\n{e}")

    def exportar_word(self):
        if not self.datos_mapa:
            messagebox.showwarning("Sin mapa", "Primero genera o redibuja un mapa mental.")
            return

        tema = self.datos_mapa.get("tema_central", "Mapa Mental")
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documento Word", "*.docx"), ("Todos los archivos", "*.*")],
            initialfile=f"MapaMental_{tema.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        )
        if not path:
            return

        try:
            doc = Document()
            doc.add_heading(f"Mapa Mental: {tema}", 0)
            doc.add_paragraph(f"Nivel Académico: {self.combo_nivel.get()}  |  Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            
            desc_gen = self.datos_mapa.get("descripcion_general", "")
            if desc_gen:
                p_res = doc.add_paragraph()
                p_res.add_run("Resumen Conceptual: ").bold = True
                p_res.add_run(desc_gen)

            temp_img = os.path.expanduser("~/.temp_mapa_mental_export.png")
            self.fig.savefig(temp_img, dpi=300, bbox_inches="tight", facecolor="#070c18", edgecolor="none")

            doc.add_paragraph("")
            doc.add_heading("Estructura Gráfica del Mapa Mental", level=1)
            doc.add_picture(temp_img, width=Inches(6.2))
            doc.add_paragraph("")

            if os.path.exists(temp_img):
                os.remove(temp_img)

            doc.add_heading("Desglose Detallado de Ramas y Conceptos", level=1)

            for rama in self.datos_mapa.get("ramas", []):
                doc.add_heading(f"📌 {rama.get('titulo', 'Rama')}", level=2)
                if rama.get("descripcion"):
                    doc.add_paragraph(rama.get("descripcion"))

                subs = rama.get("sub_conceptos", [])
                if subs:
                    for sub in subs:
                        p_sub = doc.add_paragraph(style="List Bullet")
                        r_bold = p_sub.add_run(f"{sub.get('nombre', 'Concepto')}: ")
                        r_bold.bold = True
                        p_sub.add_run(sub.get("detalle", ""))

            doc.save(path)
            messagebox.showinfo("Exportado a Word", f"Documento Word creado con éxito con gráfico HD e información detallada:\n{path}")
        except Exception as e:
            messagebox.showerror("Error al Exportar", f"No se pudo generar el documento Word:\n{e}")
