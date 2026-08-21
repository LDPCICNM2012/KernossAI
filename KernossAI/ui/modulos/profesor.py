"""
KernossAI - Módulo: Herramientas del Docente (Profesor)
Generador de ejercicios pedagógicos estructurados y corrector automático con IA.
"""

import threading
from datetime import datetime
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document

from KernossAI.core.theme import (
    COLOR_BG_DARK,
    COLOR_BG_CARD,
    COLOR_BG_CARD_LIGHT,
    COLOR_BG_SURFACE,
    COLOR_BORDER,
    COLOR_ACCENT_PRIMARY,
    COLOR_ACCENT_HOVER,
    COLOR_ACCENT_CYAN,
    COLOR_ACCENT_SKY,
    COLOR_ACCENT_PURPLE,
    COLOR_ACCENT_PURPLE_HOVER,
    COLOR_TEXT_MAIN,
    COLOR_TEXT_MUTED,
    COLOR_SUCCESS,
    COLOR_SUCCESS_HOVER,
    COLOR_DANGER,
    COLOR_DANGER_HOVER,
    COLOR_WARNING,
)
from KernossAI.core.i18n import t
from KernossAI.core.auth import llamar_groq, llamar_gemini
from KernossAI.core.tts import tts_engine

INSTRUCCIONES_EJERCICIO = """Eres un profesor experto creando material educativo.
Crea ejercicios bien estructurados con:
1. Título claro y objetivo de aprendizaje.
2. Enunciado detallado.
3. Apartados o preguntas numerados.
4. Nivel de dificultad indicado.
5. Tiempo estimado.
NO incluyas las soluciones a menos que se te pida explícitamente.
El formato debe ser limpio y listo para imprimir o enviar a alumnos."""

INSTRUCCIONES_SOLUCIONES = """Eres un profesor. Genera el solucionario completo y detallado
del ejercicio que te proporcionan. Explica cada paso y por qué es correcto."""

INSTRUCCIONES_CORRECTOR = """Eres un profesor corrector experto y riguroso.
Tu tarea es evaluar las respuestas de un alumno comparándolas con el ejercicio/criterios dados.
Para cada pregunta debes:
1. Indicar si está CORRECTA, PARCIALMENTE CORRECTA o INCORRECTA.
2. Puntuación obtenida sobre la puntuación máxima de esa pregunta.
3. Comentario breve explicando el acierto o el error.
4. Sugerencia de mejora si aplica.

Al final incluye:
- NOTA TOTAL (sobre 10).
- Resumen general de puntos fuertes y débiles del alumno.
- Recomendaciones de estudio personalizadas.

Sé justo, constructivo y motivador en el tono."""


class ModuloCreadorEjercicios(ctk.CTkFrame):
    """Módulo generador de ejercicios didácticos con IA para docentes."""
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.modelo_actual = "groq"
        self.ejercicio_actual = ""
        self.solucionario_actual = ""
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        sidebar = ctk.CTkFrame(self, width=290, corner_radius=0,
                               fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        sidebar.grid(row=0, column=0, sticky="nsew")
        sidebar.grid_propagate(False)
        ctk.CTkLabel(sidebar, text=t("cread_titulo"),
                     font=("Segoe UI", 16, "bold"), text_color=COLOR_ACCENT_PURPLE).pack(pady=(20, 5), padx=10)
        ctk.CTkLabel(sidebar, text="Teacher Tools",
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(pady=(0, 15))

        ctk.CTkLabel(sidebar, text="AI Engine:", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        fm = ctk.CTkFrame(sidebar, fg_color="transparent")
        fm.pack(fill="x", padx=15, pady=5)
        self.btn_groq = ctk.CTkButton(fm, text="⚡ Groq", height=38,
                                       fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                       font=("Segoe UI", 10, "bold"),
                                       command=lambda: self._set_modelo("groq"))
        self.btn_groq.pack(side="left", expand=True, fill="x", padx=(0, 3))
        self.btn_gemini = ctk.CTkButton(fm, text="🧠 Gemini", height=38,
                                         fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                         font=("Segoe UI", 10, "bold"),
                                         command=lambda: self._set_modelo("gemini"))
        self.btn_gemini.pack(side="left", expand=True, fill="x", padx=(3, 0))
        self.lbl_limite = ctk.CTkLabel(sidebar, text="Groq / Gemini",
                                        font=("Segoe UI", 9), text_color=COLOR_TEXT_MUTED)
        self.lbl_limite.pack(pady=(2, 10))
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=5)

        ctk.CTkLabel(sidebar, text="Topic / Subject", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15, pady=(5, 2))
        self.entry_tema = ctk.CTkEntry(sidebar, placeholder_text="e.g. Physics, History...",
                                       fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_tema.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Level", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.combo_nivel = ctk.CTkComboBox(sidebar, values=[
            "Primary", "Middle School", "High School", "University"],
            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.combo_nivel.set("High School")
        self.combo_nivel.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Type", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.combo_tipo = ctk.CTkComboBox(sidebar, values=[
            "Step-by-step problems", "Theoretical questions", "Multiple Choice", "Practical"],
            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.combo_tipo.set("Step-by-step problems")
        self.combo_tipo.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Questions", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.spin_preguntas = ctk.CTkSlider(sidebar, from_=3, to=20, number_of_steps=17, progress_color=COLOR_ACCENT_PURPLE)
        self.spin_preguntas.set(8)
        self.spin_preguntas.pack(fill="x", padx=15, pady=(0, 2))
        self.lbl_n_preguntas = ctk.CTkLabel(sidebar, text="8 questions", font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED)
        self.lbl_n_preguntas.pack()
        self.spin_preguntas.configure(command=lambda v: self.lbl_n_preguntas.configure(text=f"{int(v)} questions"))

        self.check_soluciones = ctk.CTkCheckBox(sidebar, text="Include solutions",
                                                fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER)
        self.check_soluciones.pack(padx=15, pady=8, anchor="w")
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=5)

        self.btn_generar = ctk.CTkButton(sidebar, text=t("cread_btn_generar"), height=42,
                                          font=("Segoe UI", 13, "bold"),
                                          fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                          command=self.generar_ejercicio)
        self.btn_generar.pack(fill="x", padx=15, pady=5)

        self.btn_regenerar = ctk.CTkButton(sidebar, text="🔄 Regenerar", height=36,
                                            fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_HOVER,
                                            command=self.generar_ejercicio, state="disabled")
        self.btn_regenerar.pack(fill="x", padx=15, pady=2)

        self.btn_editar = ctk.CTkButton(sidebar, text="✏️ Editar manualmente", height=36,
                                         fg_color="transparent", border_width=1, border_color=COLOR_BORDER,
                                         command=self.activar_edicion, state="disabled")
        self.btn_editar.pack(fill="x", padx=15, pady=2)

        self.btn_solucionario = ctk.CTkButton(sidebar, text="💡 Generar Solucionario", height=36,
                                               fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                                               command=self.generar_solucionario, state="disabled")
        self.btn_solucionario.pack(fill="x", padx=15, pady=2)

        self.btn_tts_ejercicio = ctk.CTkButton(sidebar, text="🔊 Escuchar Ejercicio", height=36,
                                              fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                              hover_color=COLOR_ACCENT_HOVER,
                                              command=self._toggle_tts, state="disabled")
        self.btn_tts_ejercicio.pack(fill="x", padx=15, pady=2)

        self.btn_exportar = ctk.CTkButton(sidebar, text="📄 Exportar con solución", height=36,
                                           fg_color=COLOR_BG_CARD_LIGHT, hover_color=COLOR_BG_SURFACE,
                                           command=lambda: self.exportar_word(True), state="disabled")
        self.btn_exportar.pack(fill="x", padx=15, pady=2)

        self.status = ctk.CTkLabel(sidebar, text="Listo", text_color=COLOR_TEXT_MUTED, font=("Segoe UI", 10))
        self.status.pack(side="bottom", pady=8)

        panel = ctk.CTkFrame(self, fg_color="transparent")
        panel.grid(row=0, column=1, sticky="nsew", padx=15, pady=15)
        panel.grid_rowconfigure(1, weight=1)
        panel.grid_columnconfigure(0, weight=1)

        cab = ctk.CTkFrame(panel, fg_color="transparent")
        cab.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        ctk.CTkLabel(cab, text="Ejercicio Pedagógico Generado", font=("Segoe UI", 18, "bold"), text_color=COLOR_TEXT_MAIN).pack(side="left")
        self.lbl_modo_edicion = ctk.CTkLabel(cab, text="", font=("Segoe UI", 11), text_color=COLOR_WARNING)
        self.lbl_modo_edicion.pack(side="right")

        self.txt_ejercicio = ctk.CTkTextbox(panel, font=("Consolas", 13), fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_ejercicio.grid(row=1, column=0, sticky="nsew")
        self.txt_ejercicio.insert("end", "Configura los parámetros en el panel izquierdo y pulsa 'Generar Ejercicio'.")
        self.txt_ejercicio.configure(state="disabled")

        self.entry_cambio = ctk.CTkEntry(panel, placeholder_text="Ajustes con IA: Ej: 'Hazlo más difícil' / 'Añade 2 preguntas teóricas'", height=42,
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_cambio.grid(row=2, column=0, sticky="ew", pady=(10, 0))
        self.entry_cambio.bind("<Return>", lambda e: self.aplicar_cambio_ia())

        btn_row = ctk.CTkFrame(panel, fg_color="transparent")
        btn_row.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        ctk.CTkButton(btn_row, text="⚡ Aplicar cambio con IA", fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                      height=38, font=("Segoe UI", 12, "bold"),
                      command=self.aplicar_cambio_ia).pack(side="left", padx=(0, 8))

    def _set_modelo(self, m):
        self.modelo_actual = m
        if m == "groq":
            self.btn_groq.configure(fg_color=COLOR_ACCENT_PRIMARY)
            self.btn_gemini.configure(fg_color=COLOR_BG_SURFACE)
            self.lbl_limite.configure(text="Groq: ~1.000 msgs/día")
        else:
            self.btn_groq.configure(fg_color=COLOR_BG_SURFACE)
            self.btn_gemini.configure(fg_color=COLOR_ACCENT_PURPLE)
            self.lbl_limite.configure(text="Gemini: 15 msgs/minuto")

    def _llamar_ia(self, prompt):
        full_prompt = f"{INSTRUCCIONES_EJERCICIO}\n\n{prompt}"
        if self.modelo_actual == "groq":
            return llamar_groq(full_prompt)
        return llamar_gemini(full_prompt)

    def _construir_prompt(self):
        tema  = self.entry_tema.get().strip() or "tema libre"
        nivel = self.combo_nivel.get()
        tipo  = self.combo_tipo.get()
        n     = int(self.spin_preguntas.get())
        inc   = self.check_soluciones.get()
        return (f"Crea un ejercicio de {tipo} sobre '{tema}' para alumnos de {nivel}. "
                f"Debe tener exactamente {n} preguntas/apartados. "
                f"{'Incluye el solucionario completo al final.' if inc else 'NO incluyas las soluciones.'} "
                f"Formato limpio y listo para entregar a los alumnos.")

    def generar_ejercicio(self):
        if not self.entry_tema.get().strip():
            messagebox.showwarning("Tema vacío", "Introduce el tema del ejercicio.")
            return
        self.btn_generar.configure(state="disabled")
        self.status.configure(text="Generando...", text_color="orange")
        threading.Thread(target=self._thread_generar, daemon=True).start()

    def _thread_generar(self):
        try:
            resultado = self._llamar_ia(self._construir_prompt())
            self.ejercicio_actual = resultado
            self.after(0, self._mostrar_ejercicio, resultado)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    def _mostrar_ejercicio(self, texto):
        self.txt_ejercicio.configure(state="normal")
        self.txt_ejercicio.delete("1.0", "end")
        self.txt_ejercicio.insert("end", texto)
        self.txt_ejercicio.configure(state="disabled")
        for btn in [self.btn_regenerar, self.btn_editar, self.btn_solucionario,
                    self.btn_tts_ejercicio, self.btn_exportar]:
            btn.configure(state="normal")

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_ejercicio.configure(text="🔊 Escuchar Ejercicio", fg_color=COLOR_BG_CARD_LIGHT)
        else:
            texto = self.txt_ejercicio.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin ejercicio", "Primero genera un ejercicio para escucharlo.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_ejercicio.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_ejercicio.configure(text="🔊 Escuchar Ejercicio", fg_color=COLOR_BG_CARD_LIGHT)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def activar_edicion(self):
        self.txt_ejercicio.configure(state="normal")
        self.lbl_modo_edicion.configure(text="✏️ Modo edición manual activo")
        self.btn_editar.configure(text="Bloquear edición", command=self.desactivar_edicion)

    def desactivar_edicion(self):
        self.ejercicio_actual = self.txt_ejercicio.get("1.0", "end-1c")
        self.txt_ejercicio.configure(state="disabled")
        self.lbl_modo_edicion.configure(text="")
        self.btn_editar.configure(text="Editar manualmente", command=self.activar_edicion)

    def aplicar_cambio_ia(self):
        cambio = self.entry_cambio.get().strip()
        if not cambio:
            messagebox.showwarning("Atención", "Escribe qué quieres cambiar.")
            return
        if not self.ejercicio_actual:
            messagebox.showwarning("Atención", "Primero genera un ejercicio.")
            return
        self.status.configure(text="Aplicando cambio...", text_color="orange")
        prompt = (f"Este es el ejercicio actual:\n\n{self.ejercicio_actual}\n\n"
                  f"Aplica el siguiente cambio y devuelve el ejercicio completo actualizado: {cambio}")
        self.entry_cambio.delete(0, "end")
        threading.Thread(target=lambda: self._thread_cambio(prompt), daemon=True).start()

    def _thread_cambio(self, prompt):
        try:
            resultado = self._llamar_ia(prompt)
            self.ejercicio_actual = resultado
            self.after(0, self._mostrar_ejercicio, resultado)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    def generar_solucionario(self):
        if not self.ejercicio_actual:
            messagebox.showwarning("Atención", "Primero genera un ejercicio.")
            return
        self.status.configure(text="Generando solucionario...", text_color="orange")
        prompt = f"{INSTRUCCIONES_SOLUCIONES}\n\nGenera el solucionario completo de este ejercicio:\n\n{self.ejercicio_actual}"
        def _thread():
            try:
                sol = llamar_groq(prompt) if self.modelo_actual == "groq" else llamar_gemini(prompt)
                self.solucionario_actual = sol
                def _mostrar():
                    ven = ctk.CTkToplevel(self)
                    ven.title("Solucionario – KernossAI")
                    ven.geometry("800x600")
                    ven.configure(fg_color=COLOR_BG_DARK)
                    ctk.CTkLabel(ven, text="💡 Solucionario Completo",
                                 font=("Segoe UI", 18, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=15)
                    txt = ctk.CTkTextbox(ven, font=("Consolas", 12), fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
                    txt.pack(fill="both", expand=True, padx=15, pady=(0, 10))
                    txt.insert("end", sol)
                    txt.configure(state="disabled")
                    ctk.CTkButton(ven, text="📄 Exportar Solucionario a Word",
                                  fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                                  command=lambda: self._exportar_solucionario(sol)).pack(fill="x", padx=15, pady=(0, 15))
                self.after(0, _mostrar)
            except Exception as e:
                self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            finally:
                self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))
        threading.Thread(target=_thread, daemon=True).start()

    def exportar_word(self, con_solucion=True):
        texto = self.txt_ejercicio.get("1.0", "end-1c")
        if not texto.strip():
            messagebox.showwarning("Vacío", "No hay ejercicio que exportar.")
            return
        tema   = self.entry_tema.get().strip() or "Ejercicio"
        sufijo = "ConSolucion" if con_solucion else "SinSolucion"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Ejercicio_{tema}_{sufijo}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return
        doc = Document()
        doc.add_heading(f"Ejercicio — {tema} (KernossAI)", 0)
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()} | Tipo: {self.combo_tipo.get()}")
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")
        doc.add_paragraph(texto)
        if con_solucion and self.solucionario_actual:
            doc.add_page_break()
            doc.add_heading("SOLUCIONARIO", level=1)
            doc.add_paragraph(self.solucionario_actual)
        doc.save(path)
        messagebox.showinfo("Exportado", f"Guardado en:\n{path}")

    def _exportar_solucionario(self, sol):
        tema = self.entry_tema.get().strip() or "Ejercicio"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Solucionario_{tema}_{datetime.now().strftime('%Y%m%d')}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Solucionario — {tema} (KernossAI)", 0)
            doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
            doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph("")
            doc.add_paragraph(sol)
            doc.save(path)
            messagebox.showinfo("Exportado", f"Solucionario guardado en:\n{path}")


class ModuloCorrectorExamenes(ctk.CTkFrame):
    """Módulo corrector automatizado de exámenes y respuestas de alumnos."""
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.modelo_actual = "groq"
        self.correccion_actual = ""
        self.alumnos = {}
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(2, weight=1)
        self.grid_rowconfigure(0, weight=1)

        sidebar = ctk.CTkFrame(self, width=270, corner_radius=0,
                               fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        sidebar.grid(row=0, column=0, sticky="nsew")
        sidebar.grid_propagate(False)
        ctk.CTkLabel(sidebar, text="📋 Corrector de\nExámenes",
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(20, 5))
        ctk.CTkLabel(sidebar, text="Modo Docente",
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED).pack(pady=(0, 15))

        ctk.CTkLabel(sidebar, text="Motor de IA:", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        fm = ctk.CTkFrame(sidebar, fg_color="transparent")
        fm.pack(fill="x", padx=15, pady=5)
        self.btn_groq = ctk.CTkButton(fm, text="⚡ Groq", height=42,
                                       fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                       font=("Segoe UI", 10, "bold"),
                                       command=lambda: self._set_modelo("groq"))
        self.btn_groq.pack(side="left", expand=True, fill="x", padx=(0, 3))
        self.btn_gemini = ctk.CTkButton(fm, text="🧠 Gemini", height=42,
                                         fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                         font=("Segoe UI", 10, "bold"),
                                         command=lambda: self._set_modelo("gemini"))
        self.btn_gemini.pack(side="left", expand=True, fill="x", padx=(3, 0))
        self.lbl_limite = ctk.CTkLabel(sidebar, text="Groq: ~1.000 msgs/día",
                                        font=("Segoe UI", 9), text_color=COLOR_TEXT_MUTED)
        self.lbl_limite.pack(pady=(2, 8))
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=5)

        ctk.CTkLabel(sidebar, text="Materia / Examen", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15, pady=(5, 2))
        self.entry_materia = ctk.CTkEntry(sidebar, placeholder_text="Ej: Historia — Tema 4",
                                          fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_materia.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Nivel educativo", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.combo_nivel = ctk.CTkComboBox(sidebar, values=[
            "Primaria", "1º ESO", "2º ESO", "3º ESO", "4º ESO",
            "1º Bachillerato", "2º Bachillerato", "Universidad"],
            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.combo_nivel.set("2º ESO")
        self.combo_nivel.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Nombre del alumno", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.entry_alumno = ctk.CTkEntry(sidebar, placeholder_text="Nombre y apellidos",
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_alumno.pack(fill="x", padx=15, pady=(0, 8))

        ctk.CTkLabel(sidebar, text="Criterios de puntuación", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w", padx=15)
        self.entry_criterios = ctk.CTkEntry(sidebar, placeholder_text="Ej: P1=2pts, P2=3pts, P3=5pts",
                                            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_criterios.pack(fill="x", padx=15, pady=(0, 8))
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=5)

        self.btn_corregir = ctk.CTkButton(sidebar, text=t("corr_btn_corregir"),
                                           height=44, font=("Segoe UI", 13, "bold"),
                                           fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                                           command=self.corregir_examen)
        self.btn_corregir.pack(fill="x", padx=15, pady=5)

        self.btn_guardar_alumno = ctk.CTkButton(sidebar, text=t("apunt_btn_guardar"),
                                                 height=36, fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_HOVER,
                                                 command=self.guardar_alumno, state="disabled")
        self.btn_guardar_alumno.pack(fill="x", padx=15, pady=2)

        self.btn_tts_corrector = ctk.CTkButton(sidebar, text=t("btn_escuchar"), height=36,
                                              fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                              hover_color=COLOR_ACCENT_HOVER,
                                              command=self._toggle_tts, state="disabled")
        self.btn_tts_corrector.pack(fill="x", padx=15, pady=2)

        self.btn_exportar_uno = ctk.CTkButton(sidebar, text=t("btn_word"),
                                               height=36, fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                               command=self.exportar_correccion, state="disabled")
        self.btn_exportar_uno.pack(fill="x", padx=15, pady=2)

        self.btn_exportar_clase = ctk.CTkButton(sidebar, text=t("calc_btn_exportar"),
                                                 height=36, fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                                 command=self.exportar_informe_clase, state="disabled")
        self.btn_exportar_clase.pack(fill="x", padx=15, pady=2)

        ctk.CTkButton(sidebar, text=t("btn_nuevo_chat"), height=32,
                      fg_color="transparent", border_width=1, border_color=COLOR_BORDER,
                      command=self.limpiar_todo).pack(fill="x", padx=15, pady=2)

        ctk.CTkLabel(sidebar, text="Alumnos corregidos",
                     font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED).pack(pady=(8, 2))
        self.frame_alumnos = ctk.CTkScrollableFrame(sidebar, fg_color="transparent", height=100)
        self.frame_alumnos.pack(fill="x", padx=10, pady=2)
        self.status = ctk.CTkLabel(sidebar, text="Listo", text_color=COLOR_TEXT_MUTED, font=("Segoe UI", 10))
        self.status.pack(side="bottom", pady=6)

        panel_izq = ctk.CTkFrame(self, fg_color="transparent")
        panel_izq.grid(row=0, column=1, sticky="nsew", padx=(15, 7), pady=15)
        panel_izq.grid_rowconfigure(1, weight=1)
        panel_izq.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(panel_izq, text="📝 Enunciado y Criterios del examen",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.txt_enunciado = ctk.CTkTextbox(panel_izq, font=("Consolas", 12), fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_enunciado.grid(row=1, column=0, sticky="nsew")
        self.txt_enunciado.insert("end",
            "Pega aquí el enunciado del examen y/o los criterios de corrección...\n\n"
            "Ejemplo:\n"
            "Pregunta 1 (2 pts): ¿Qué es la fotosíntesis? Explícala.\n"
            "Pregunta 2 (3 pts): Nombra 3 diferencias entre células animales y vegetales.")

        panel_der = ctk.CTkFrame(self, fg_color="transparent")
        panel_der.grid(row=0, column=2, sticky="nsew", padx=(7, 15), pady=15)
        panel_der.grid_rowconfigure(1, weight=2)
        panel_der.grid_rowconfigure(3, weight=3)
        panel_der.grid_columnconfigure(0, weight=1)

        ctk.CTkLabel(panel_der, text="✍️ Respuestas del alumno",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_TEXT_MAIN).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.txt_respuestas = ctk.CTkTextbox(panel_der, font=("Consolas", 12), fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_respuestas.grid(row=1, column=0, sticky="nsew", pady=(0, 10))
        self.txt_respuestas.insert("end", "Pega o escribe aquí las respuestas del alumno...")

        sep_frame = ctk.CTkFrame(panel_der, fg_color="transparent")
        sep_frame.grid(row=2, column=0, sticky="ew", pady=4)
        ctk.CTkLabel(sep_frame, text="📋 Corrección Inteligente",
                     font=("Segoe UI", 15, "bold"), text_color=COLOR_ACCENT_CYAN).pack(side="left")
        self.lbl_nota = ctk.CTkLabel(sep_frame, text="", font=("Segoe UI", 16, "bold"), text_color=COLOR_SUCCESS)
        self.lbl_nota.pack(side="right")

        self.txt_correccion = ctk.CTkTextbox(panel_der, font=("Consolas", 12), state="disabled",
                                            fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_correccion.grid(row=3, column=0, sticky="nsew")

    def _set_modelo(self, m):
        self.modelo_actual = m
        if m == "groq":
            self.btn_groq.configure(fg_color=COLOR_ACCENT_PRIMARY)
            self.btn_gemini.configure(fg_color=COLOR_BG_SURFACE)
            self.lbl_limite.configure(text="Groq: ~1.000 msgs/día")
        else:
            self.btn_groq.configure(fg_color=COLOR_BG_SURFACE)
            self.btn_gemini.configure(fg_color=COLOR_ACCENT_PURPLE)
            self.lbl_limite.configure(text="Gemini: 15 msgs/minuto")

    def _llamar_ia(self, prompt):
        full_prompt = f"{INSTRUCCIONES_CORRECTOR}\n\n{prompt}"
        if self.modelo_actual == "groq":
            return llamar_groq(full_prompt)
        return llamar_gemini(full_prompt)

    def corregir_examen(self):
        enunciado = self.txt_enunciado.get("1.0", "end-1c").strip()
        respuestas = self.txt_respuestas.get("1.0", "end-1c").strip()
        alumno = self.entry_alumno.get().strip() or "Alumno sin nombre"
        materia = self.entry_materia.get().strip() or "Examen"
        criterios = self.entry_criterios.get().strip()
        if not enunciado or "Pega aquí" in enunciado:
            messagebox.showwarning("Falta enunciado", "Introduce el enunciado del examen.")
            return
        if not respuestas or "Pega o escribe" in respuestas:
            messagebox.showwarning("Faltan respuestas", "Introduce las respuestas del alumno.")
            return
        self.btn_corregir.configure(state="disabled")
        self.status.configure(text="Corrigiendo...", text_color="orange")
        self.lbl_nota.configure(text="")
        prompt = (f"MATERIA: {materia}\nNIVEL: {self.combo_nivel.get()}\nALUMNO: {alumno}\n"
                  + (f"CRITERIOS DE PUNTUACIÓN: {criterios}\n" if criterios else "")
                  + f"\nENUNCIADO DEL EXAMEN:\n{enunciado}\n\nRESPUESTAS DEL ALUMNO:\n{respuestas}")
        threading.Thread(target=self._thread_corregir, args=(prompt, alumno, respuestas), daemon=True).start()

    def _thread_corregir(self, prompt, alumno, respuestas):
        try:
            resultado = self._llamar_ia(prompt)
            self.correccion_actual = resultado
            nota_str = ""
            for linea in resultado.splitlines():
                if "NOTA TOTAL" in linea.upper() or "NOTA FINAL" in linea.upper():
                    partes = linea.split(":")
                    if len(partes) > 1:
                        nota_str = partes[-1].strip().split()[0]
                    break
            def _mostrar():
                self.txt_correccion.configure(state="normal")
                self.txt_correccion.delete("1.0", "end")
                self.txt_correccion.insert("end", resultado)
                self.txt_correccion.configure(state="disabled")
                if nota_str:
                    self.lbl_nota.configure(text=f"Nota: {nota_str}")
                for btn in [self.btn_guardar_alumno, self.btn_exportar_uno, self.btn_tts_corrector]:
                    btn.configure(state="normal")
            self.after(0, _mostrar)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
        finally:
            self.after(0, lambda: self.btn_corregir.configure(state="normal"))
            self.after(0, lambda: self.status.configure(text="Listo", text_color="gray"))

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_corrector.configure(text="🔊 Escuchar Corrección", fg_color=COLOR_BG_CARD_LIGHT)
        else:
            texto = self.txt_correccion.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin corrección", "Primero realiza una corrección para escucharla en voz alta.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_corrector.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_corrector.configure(text="🔊 Escuchar Corrección", fg_color=COLOR_BG_CARD_LIGHT)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def guardar_alumno(self):
        alumno = self.entry_alumno.get().strip() or "Alumno sin nombre"
        if not self.correccion_actual:
            messagebox.showwarning("Sin corrección", "Primero corrige un examen.")
            return
        self.alumnos[alumno] = {
            "respuestas": self.txt_respuestas.get("1.0", "end-1c"),
            "correccion": self.correccion_actual,
            "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M")}
        self._actualizar_lista_alumnos()
        self.btn_exportar_clase.configure(state="normal")
        messagebox.showinfo("Guardado", f"Resultado de '{alumno}' guardado.\nTotal alumnos: {len(self.alumnos)}")

    def _actualizar_lista_alumnos(self):
        for w in self.frame_alumnos.winfo_children():
            w.destroy()
        for nombre in self.alumnos:
            fila = ctk.CTkFrame(self.frame_alumnos, fg_color="transparent")
            fila.pack(fill="x", pady=2)
            ctk.CTkButton(fila, text=nombre, fg_color="transparent",
                          text_color=COLOR_TEXT_MAIN, anchor="w", hover_color=COLOR_BG_SURFACE, height=28,
                          command=lambda n=nombre: self._cargar_alumno(n)).pack(side="left", fill="x", expand=True)
            ctk.CTkButton(fila, text="❌", width=28, height=28,
                          fg_color=COLOR_DANGER, hover_color=COLOR_DANGER_HOVER,
                          command=lambda n=nombre: self._borrar_alumno(n)).pack(side="right")

    def _cargar_alumno(self, nombre):
        datos = self.alumnos[nombre]
        self.entry_alumno.delete(0, "end")
        self.entry_alumno.insert(0, nombre)
        self.txt_respuestas.delete("1.0", "end")
        self.txt_respuestas.insert("end", datos["respuestas"])
        self.txt_correccion.configure(state="normal")
        self.txt_correccion.delete("1.0", "end")
        self.txt_correccion.insert("end", datos["correccion"])
        self.txt_correccion.configure(state="disabled")
        self.correccion_actual = datos["correccion"]
        for btn in [self.btn_guardar_alumno, self.btn_exportar_uno]:
            btn.configure(state="normal")

    def _borrar_alumno(self, nombre):
        if messagebox.askyesno("Confirmar", f"¿Borrar resultado de '{nombre}'?"):
            del self.alumnos[nombre]
            self._actualizar_lista_alumnos()
            if not self.alumnos:
                self.btn_exportar_clase.configure(state="disabled")

    def limpiar_todo(self):
        self.txt_respuestas.delete("1.0", "end")
        self.txt_correccion.configure(state="normal")
        self.txt_correccion.delete("1.0", "end")
        self.txt_correccion.configure(state="disabled")
        self.entry_alumno.delete(0, "end")
        self.correccion_actual = ""
        self.lbl_nota.configure(text="")
        for btn in [self.btn_guardar_alumno, self.btn_exportar_uno]:
            btn.configure(state="disabled")

    def exportar_correccion(self):
        if not self.correccion_actual:
            messagebox.showwarning("Vacío", "No hay corrección que exportar.")
            return
        alumno  = self.entry_alumno.get().strip() or "Alumno"
        materia = self.entry_materia.get().strip() or "Examen"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"Correccion_{alumno}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return
        doc = Document()
        doc.add_heading(f"Corrección — {materia} (KernossAI)", 0)
        doc.add_paragraph(f"Alumno: {alumno}")
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
        doc.add_paragraph(f"Fecha corrección: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")
        doc.add_heading("Respuestas del alumno", level=1)
        doc.add_paragraph(self.txt_respuestas.get("1.0", "end-1c"))
        doc.add_page_break()
        doc.add_heading("Corrección y calificación", level=1)
        doc.add_paragraph(self.correccion_actual)
        doc.save(path)
        messagebox.showinfo("Exportado", f"Corrección guardada en:\n{path}")

    def exportar_informe_clase(self):
        if not self.alumnos:
            messagebox.showwarning("Sin datos", "No hay alumnos guardados todavía.")
            return
        materia = self.entry_materia.get().strip() or "Examen"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=f"InformeClase_{materia}_{datetime.now().strftime('%Y%m%d')}.docx")
        if not path:
            return
        doc = Document()
        doc.add_heading(f"Informe de Clase — {materia} (KernossAI)", 0)
        doc.add_paragraph(f"Nivel: {self.combo_nivel.get()}")
        doc.add_paragraph(f"Total alumnos corregidos: {len(self.alumnos)}")
        doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph("")
        for nombre, datos in self.alumnos.items():
            doc.add_heading(nombre, level=1)
            doc.add_paragraph(f"Corregido: {datos['timestamp']}")
            doc.add_paragraph("")
            doc.add_heading("Corrección", level=2)
            doc.add_paragraph(datos["correccion"])
            doc.add_page_break()
        doc.save(path)
        messagebox.showinfo("Exportado", f"Informe de clase guardado en:\n{path}")


class ModuloHerramientasProfesor(ctk.CTkFrame):
    """Contenedor de pestañas docentes que integra el generador de ejercicios y el corrector."""
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.tabview = ctk.CTkTabview(
            self,
            fg_color="transparent",
            segmented_button_fg_color=COLOR_BG_CARD,
            segmented_button_selected_color=COLOR_ACCENT_PURPLE,
            segmented_button_selected_hover_color=COLOR_ACCENT_PURPLE_HOVER,
            segmented_button_unselected_hover_color=COLOR_BG_SURFACE
        )
        self.tabview.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)

        tab_ejercicios = self.tabview.add(f"📝 {t('mod_creador')}")
        tab_corrector = self.tabview.add(f"📋 {t('mod_corrector')}")

        tab_ejercicios.grid_columnconfigure(0, weight=1)
        tab_ejercicios.grid_rowconfigure(0, weight=1)
        self.creador = ModuloCreadorEjercicios(tab_ejercicios)
        self.creador.grid(row=0, column=0, sticky="nsew")

        tab_corrector.grid_columnconfigure(0, weight=1)
        tab_corrector.grid_rowconfigure(0, weight=1)
        self.corrector = ModuloCorrectorExamenes(tab_corrector)
        self.corrector.grid(row=0, column=0, sticky="nsew")
