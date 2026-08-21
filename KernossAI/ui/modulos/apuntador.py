"""
KernossAI - Módulo: Apuntador de Notas
Gestor local de notas, apuntes de clase, lectura con voz TTS y exportación a Word.
"""

import os
import json
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document

from KernossAI.core.theme import (
    COLOR_BG_CARD,
    COLOR_BG_CARD_LIGHT,
    COLOR_BG_SURFACE,
    COLOR_BORDER,
    COLOR_ACCENT_PRIMARY,
    COLOR_ACCENT_HOVER,
    COLOR_ACCENT_CYAN,
    COLOR_ACCENT_SKY,
    COLOR_ACCENT_PURPLE,
    COLOR_ACCENT_PURPLE_HOVER,
    COLOR_TEXT_MAIN,
    COLOR_SUCCESS,
    COLOR_SUCCESS_HOVER,
    COLOR_DANGER,
    COLOR_DANGER_HOVER,
)
from KernossAI.core.i18n import t
from KernossAI.core.tts import tts_engine


class ModuloApuntador(ctk.CTkFrame):
    """Módulo de creación, gestión y exportación de notas académicas con TTS."""
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.archivo_actual = None
        self.notas_guardadas = {}
        self._cargar_notas()
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.frame_lateral = ctk.CTkFrame(self, width=270, corner_radius=0,
                                          fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        self.frame_lateral.grid(row=0, column=0, sticky="nsew")
        ctk.CTkLabel(self.frame_lateral, text=t("apunt_titulo"),
                     font=("Segoe UI", 16, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=20, padx=10)

        self.btn_nueva_nota = ctk.CTkButton(self.frame_lateral, text=t("apunt_btn_nueva"),
                                            fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                            height=38, font=("Segoe UI", 12, "bold"),
                                            command=self.nueva_nota)
        self.btn_nueva_nota.pack(fill="x", padx=15, pady=(0, 10))

        self.lista_notas_frame = ctk.CTkScrollableFrame(self.frame_lateral, fg_color="transparent")
        self.lista_notas_frame.pack(fill="both", expand=True, padx=10, pady=5)
        self.actualizar_listbox()

        self.frame_editor = ctk.CTkFrame(self, fg_color="transparent")
        self.frame_editor.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.frame_editor.grid_rowconfigure(1, weight=1)
        self.frame_editor.grid_columnconfigure(0, weight=1)

        frame_top = ctk.CTkFrame(self.frame_editor, fg_color="transparent")
        frame_top.grid(row=0, column=0, sticky="ew", pady=(0, 12))
        self.label_nota_abierta = ctk.CTkLabel(frame_top, text=t("apunt_lbl_sin_nota"),
                                               font=("Segoe UI", 18, "bold"), text_color=COLOR_TEXT_MAIN)
        self.label_nota_abierta.pack(side="left")

        btn_bar = ctk.CTkFrame(frame_top, fg_color="transparent")
        btn_bar.pack(side="right")

        self.btn_tts_nota = ctk.CTkButton(btn_bar, text=t("apunt_btn_leer"), width=85, height=36,
                                          fg_color=COLOR_BG_SURFACE, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                          hover_color=COLOR_ACCENT_HOVER,
                                          command=self._toggle_tts)
        self.btn_tts_nota.pack(side="left", padx=4)

        ctk.CTkButton(btn_bar, text=t("apunt_btn_guardar"), width=95, height=36,
                      fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                      command=self.guardar_nota).pack(side="left", padx=4)
        ctk.CTkButton(btn_bar, text=t("apunt_btn_word"), width=95, height=36,
                      fg_color=COLOR_ACCENT_PURPLE, hover_color=COLOR_ACCENT_PURPLE_HOVER,
                      command=self.exportar_nota_word).pack(side="left", padx=4)
        ctk.CTkButton(btn_bar, text=t("apunt_btn_borrar"), width=95, height=36,
                      fg_color=COLOR_DANGER, hover_color=COLOR_DANGER_HOVER,
                      command=self.eliminar_nota).pack(side="left", padx=4)

        self.editor_texto = ctk.CTkTextbox(self.frame_editor, font=("Segoe UI", 14),
                                           fg_color=COLOR_BG_CARD_LIGHT, border_width=1,
                                           border_color=COLOR_BORDER, wrap="word")
        self.editor_texto.grid(row=1, column=0, sticky="nsew")

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_nota.configure(text=t("apunt_btn_leer"), fg_color=COLOR_BG_SURFACE)
        else:
            texto = self.editor_texto.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin texto", "No hay texto escrito en la nota para leer.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_nota.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_nota.configure(text=t("apunt_btn_leer"), fg_color=COLOR_BG_SURFACE)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def nueva_nota(self):
        d = ctk.CTkInputDialog(text="Nombre de la nueva nota:", title="Nueva Nota")
        nombre = d.get_input()
        if nombre and nombre.strip():
            nombre = nombre.strip()
            self.archivo_actual = nombre
            self.notas_guardadas[nombre] = ""
            self._guardar_notas()
            self.actualizar_listbox()
            self.abrir_nota(nombre)

    def abrir_nota(self, nombre):
        self.archivo_actual = nombre
        self.label_nota_abierta.configure(text=nombre)
        self.editor_texto.delete("1.0", "end")
        self.editor_texto.insert("1.0", self.notas_guardadas.get(nombre, ""))

    def actualizar_listbox(self):
        for w in self.lista_notas_frame.winfo_children():
            w.destroy()
        for nombre in self.notas_guardadas:
            ctk.CTkButton(self.lista_notas_frame, text=f"• {nombre}", anchor="w",
                          fg_color="transparent", text_color=COLOR_TEXT_MAIN,
                          hover_color=COLOR_BG_SURFACE, height=32,
                          command=lambda n=nombre: self.abrir_nota(n)).pack(fill="x", pady=2)

    def guardar_nota(self):
        if self.archivo_actual:
            self.notas_guardadas[self.archivo_actual] = self.editor_texto.get("1.0", "end-1c")
            self._guardar_notas()
            messagebox.showinfo("Guardado", "Nota guardada correctamente.")

    def eliminar_nota(self):
        if self.archivo_actual and messagebox.askyesno("Confirmar", f"¿Eliminar '{self.archivo_actual}'?"):
            del self.notas_guardadas[self.archivo_actual]
            self._guardar_notas()
            self.archivo_actual = None
            self.editor_texto.delete("1.0", "end")
            self.label_nota_abierta.configure(text=t("apunt_lbl_sin_nota"))
            self.actualizar_listbox()

    def exportar_nota_word(self):
        if not self.archivo_actual:
            return
        doc = Document()
        doc.add_heading(self.archivo_actual, 0)
        doc.add_paragraph(self.editor_texto.get("1.0", "end-1c"))
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"{self.archivo_actual}.docx")
        if path:
            doc.save(path)
            messagebox.showinfo("Éxito", "Exportado correctamente.")

    def _guardar_notas(self):
        ruta = os.path.expanduser("~/.apuntador_notas.json")
        with open(ruta, 'w', encoding='utf-8') as f:
            json.dump(self.notas_guardadas, f, ensure_ascii=False, indent=2)

    def _cargar_notas(self):
        ruta = os.path.expanduser("~/.apuntador_notas.json")
        if os.path.exists(ruta):
            try:
                with open(ruta, 'r', encoding='utf-8') as f:
                    self.notas_guardadas = json.load(f)
            except Exception:
                self.notas_guardadas = {}
