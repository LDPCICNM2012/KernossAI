"""
KernossAI - Módulo: Ayudante de Resolución de Problemas y Chat IA
Asistente interactivo multitarea con soporte Groq / Gemini, exportación Word e historial.
"""

import os
import json
import threading
from datetime import datetime
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document

from KernossAI.core.theme import (
    COLOR_BG_CARD,
    COLOR_BG_CARD_LIGHT,
    COLOR_BG_SURFACE,
    COLOR_BORDER,
    COLOR_ACCENT_PRIMARY,
    COLOR_ACCENT_HOVER,
    COLOR_ACCENT_CYAN,
    COLOR_ACCENT_SKY,
    COLOR_ACCENT_PURPLE,
    COLOR_ACCENT_PURPLE_HOVER,
    COLOR_TEXT_MAIN,
    COLOR_TEXT_MUTED,
    COLOR_SUCCESS,
    COLOR_DANGER,
    COLOR_DANGER_HOVER,
    construir_prompt,
)
from KernossAI.core.i18n import t
from KernossAI.core.auth import llamar_groq, llamar_gemini
from KernossAI.core.tts import tts_engine


class ModuloAyudador(ctk.CTkFrame):
    """Módulo tutor y asistente de resolución académica con selector de modelo IA."""
    def __init__(self, master, sesion=None):
        super().__init__(master, fg_color="transparent")
        self.sesion = sesion or {}
        self.modelo_actual = "groq"
        self.historial_conversacion = []
        self.instrucciones_groq = (
            "Eres un asistente BÁSICO y RÁPIDO (Groq). "
            "Tu objetivo es ayudar y resolver de la mejor manera posible lo que te pida el usuario. "
            "Sé directo pero si hace falta explicar cualquier cosa hazlo."
        )
        self.instrucciones_gemini = (
            "Eres un asistente AVANZADO y PROFUNDO (Gemini). "
            "Tu función es resolver de la forma más inteligente cualquier cosa que te pidan. "
            "Siempre explica el desarrollo y el porqué de las cosas."
        )
        suffix = self.sesion.get("email", "default").replace("@", "_").replace(".", "_")
        self.ruta_historial = os.path.expanduser(f"~/.historial_solver_{suffix}.json")
        self.chat_actual_id = None
        self.todo_el_historial = self._cargar_historial()
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        sidebar = ctk.CTkFrame(self, width=280, corner_radius=0,
                               fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        sidebar.grid(row=0, column=0, sticky="nsew")
        ctk.CTkLabel(sidebar, text=t("ayud_titulo"), font=("Segoe UI", 16, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(20, 5), padx=10)
        self.entry_nombre = ctk.CTkEntry(sidebar, placeholder_text=t("placeholder_nombre"),
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_nombre.pack(fill="x", padx=15, pady=8)
        if self.sesion.get("nombre"):
            self.entry_nombre.insert(0, self.sesion["nombre"])

        ctk.CTkLabel(sidebar, text="⚙️ System Prompt",
                     font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(10, 2))
        self.txt_instrucciones = ctk.CTkTextbox(sidebar, height=110, wrap="word", font=("Segoe UI", 11),
                                                fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.txt_instrucciones.pack(fill="x", padx=15, pady=5)
        self.txt_instrucciones.insert("1.0", self.instrucciones_groq)

        ctk.CTkLabel(sidebar, text="AI Engine", font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(10, 2))
        frame_modelo = ctk.CTkFrame(sidebar, fg_color=COLOR_BG_CARD_LIGHT, corner_radius=10, border_width=1, border_color=COLOR_BORDER)
        frame_modelo.pack(fill="x", padx=15, pady=5)
        self.btn_basico = ctk.CTkButton(frame_modelo, text="⚡ Groq", height=36,
                                         fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                         command=lambda: self._cambiar_modelo("groq"))
        self.btn_basico.pack(fill="x", padx=8, pady=(8, 4))
        self.btn_avanzado = ctk.CTkButton(frame_modelo, text="🧠 Gemini", height=36,
                                           fg_color="transparent", border_width=1, border_color=COLOR_ACCENT_PURPLE,
                                           hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                           command=lambda: self._cambiar_modelo("gemini"))
        self.btn_avanzado.pack(fill="x", padx=8, pady=(4, 8))
        self.lbl_limite = ctk.CTkLabel(sidebar, text="Groq / Gemini",
                                        font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED)
        self.lbl_limite.pack(pady=(0, 5))
        ctk.CTkFrame(sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=10, pady=8)

        ctk.CTkButton(sidebar, text=t("btn_word"), fg_color=COLOR_BG_SURFACE, hover_color=COLOR_ACCENT_PRIMARY,
                      border_width=1, border_color=COLOR_BORDER,
                      command=self._exportar_word).pack(fill="x", padx=15, pady=4)
        ctk.CTkButton(sidebar, text=t("btn_nuevo_chat"), fg_color="transparent", border_width=1,
                      border_color=COLOR_BORDER, hover_color=COLOR_BG_SURFACE,
                      command=self._limpiar_chat).pack(fill="x", padx=15, pady=4)

        ctk.CTkLabel(sidebar, text=t("hdr_historial"), font=("Segoe UI", 10, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(12, 3))
        self.frame_historial = ctk.CTkScrollableFrame(sidebar, fg_color="transparent")
        self.frame_historial.pack(fill="both", expand=True, padx=10, pady=5)
        self.status_label = ctk.CTkLabel(sidebar, text="🟢 Ready",
                                          text_color=COLOR_SUCCESS, font=("Segoe UI", 11))
        self.status_label.pack(side="bottom", pady=(0, 15))
        self._actualizar_historial_ui()

        chat_frame = ctk.CTkFrame(self, fg_color="transparent")
        chat_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        chat_frame.grid_rowconfigure(1, weight=1)
        chat_frame.grid_columnconfigure(0, weight=1)

        self.lbl_banner = ctk.CTkLabel(chat_frame,
                                        text="AI Active Assistant",
                                        font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED)
        self.lbl_banner.grid(row=0, column=0, sticky="w", pady=(0, 6))

        self.txt_chat = ctk.CTkTextbox(chat_frame, font=("Segoe UI", 14), state="disabled", wrap="word",
                                       fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.txt_chat.grid(row=1, column=0, sticky="nsew", pady=(0, 12))

        input_frame = ctk.CTkFrame(chat_frame, fg_color="transparent")
        input_frame.grid(row=2, column=0, sticky="ew")
        input_frame.grid_columnconfigure(0, weight=1)

        self.entry_pregunta = ctk.CTkEntry(input_frame, placeholder_text=t("ayud_lbl_enunciado"), height=44,
                                           fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_pregunta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_pregunta.bind("<Return>", lambda e: self._enviar())

        ctk.CTkButton(input_frame, text=t("ayud_btn_resolver"), width=140, height=44,
                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                      font=("Segoe UI", 12, "bold"),
                      command=self._enviar).grid(row=0, column=1, padx=(0, 6))

        self.btn_tts_ayudador = ctk.CTkButton(input_frame, text=t("btn_escuchar"), width=110, height=44,
                                              font=("Segoe UI", 12, "bold"),
                                              fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                              hover_color=COLOR_ACCENT_HOVER,
                                              command=self._toggle_tts)
        self.btn_tts_ayudador.grid(row=0, column=2)

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_ayudador.configure(text=t("btn_escuchar"), fg_color=COLOR_BG_CARD)
        else:
            texto = getattr(self, "ultima_respuesta_ayudador", "")
            if not texto:
                texto = self.txt_chat.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin respuesta", "No hay respuesta para escuchar en voz alta.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_ayudador.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_ayudador.configure(text=t("btn_escuchar"), fg_color=COLOR_BG_CARD)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def _cambiar_modelo(self, modelo):
        texto_actual = self.txt_instrucciones.get("1.0", "end-1c").strip()
        if self.modelo_actual == "groq":
            self.instrucciones_groq = texto_actual
        else:
            self.instrucciones_gemini = texto_actual
        self.modelo_actual = modelo
        self.txt_instrucciones.delete("1.0", "end")
        if modelo == "groq":
            self.btn_basico.configure(fg_color=COLOR_ACCENT_PRIMARY)
            self.btn_avanzado.configure(fg_color="transparent")
            self.lbl_limite.configure(text="ℹ️ Groq: ~1.000 msgs/día")
            self.lbl_banner.configure(text="Modelo activo: Groq (Básico)  •  Límite aprox: 1.000 mensajes / día")
            self.status_label.configure(text="🟢 Groq Listo", text_color=COLOR_SUCCESS)
            self.txt_instrucciones.insert("1.0", self.instrucciones_groq)
        else:
            self.btn_avanzado.configure(fg_color=COLOR_ACCENT_PURPLE)
            self.btn_basico.configure(fg_color="transparent")
            self.lbl_limite.configure(text="ℹ️ Gemini: 5 msgs/minuto")
            self.lbl_banner.configure(text="Modelo activo: Gemini  •  Límite: 5 mensajes / minuto")
            self.status_label.configure(text="🟣 Gemini Listo", text_color="#a5b4fc")
            self.txt_instrucciones.insert("1.0", self.instrucciones_gemini)

    def _agregar_texto(self, emisor, texto):
        self.txt_chat.configure(state="normal")
        if emisor == "IA":
            self.txt_chat.insert("end", f"\n 🤖 IA ({self.modelo_actual.upper()}):\n{texto}\n")
        else:
            self.txt_chat.insert("end", f"\n\n 👤 {emisor}:\n{texto}\n")
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")

    def _enviar(self):
        pregunta = self.entry_pregunta.get().strip()
        nombre   = self.entry_nombre.get().strip() or "Tú"
        if not pregunta:
            return
        self._agregar_texto(nombre, pregunta)
        self.entry_pregunta.delete(0, "end")
        self.status_label.configure(text="🟡 Pensando...", text_color="orange")
        self.historial_conversacion.append({"role": "user", "content": pregunta})
        threading.Thread(target=self._proceso_ia, daemon=True).start()

    def _proceso_ia(self):
        try:
            self.after(0, lambda: self._stream_update(f"\n 🤖 IA ({self.modelo_actual.upper()}):\n"))
            instrucciones = self.txt_instrucciones.get("1.0", "end-1c").strip()
            prompt = construir_prompt(instrucciones, self.historial_conversacion)
            if self.modelo_actual == "groq":
                respuesta_completa = llamar_groq(prompt)
            else:
                respuesta_completa = llamar_gemini(prompt)
            self.ultima_respuesta_ayudador = respuesta_completa
            self.historial_conversacion.append({"role": "assistant", "content": respuesta_completa})
            self.after(0, self._stream_update, respuesta_completa)
            self._guardar_historial()
            color = COLOR_SUCCESS if self.modelo_actual == "groq" else "#a5b4fc"
            texto = "🟢 Groq Listo" if self.modelo_actual == "groq" else "🟣 Gemini Listo"
            self.after(0, lambda: self.status_label.configure(text=texto, text_color=color))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            self.after(0, lambda: self.status_label.configure(text="🔴 Error", text_color=COLOR_DANGER))
        finally:
            self.after(0, lambda: self.txt_chat.configure(state="disabled"))

    def _stream_update(self, contenido):
        self.txt_chat.configure(state="normal")
        self.txt_chat.insert("end", contenido)
        self.txt_chat.see("end")
        self.txt_chat.configure(state="disabled")

    def _limpiar_chat(self):
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        self.txt_chat.configure(state="disabled")
        self.historial_conversacion = []
        self.chat_actual_id = None

    def _exportar_word(self):
        if not self.historial_conversacion:
            messagebox.showwarning("Vacío", "No hay nada que exportar.")
            return
        nombre = self.entry_nombre.get().strip() or "Usuario"
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             initialfile=f"Solucion_{nombre}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Informe de Solución – {nombre} (KernossAI)", 0)
            for msg in self.historial_conversacion:
                rol = "Tú" if msg["role"] == "user" else "IA"
                p = doc.add_paragraph()
                p.add_run(f"{rol}: ").bold = True
                p.add_run(msg["content"])
            doc.save(path)
            messagebox.showinfo("Éxito", f"Guardado en:\n{path}")

    def _cargar_historial(self):
        if os.path.exists(self.ruta_historial):
            try:
                with open(self.ruta_historial, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def _guardar_historial(self):
        if not self.historial_conversacion:
            return
        if self.chat_actual_id is None:
            primer = next((m["content"] for m in self.historial_conversacion if m["role"] == "user"), "Chat")
            resumen = primer[:25] + "..." if len(primer) > 25 else primer
            self.chat_actual_id = f"[{datetime.now().strftime('%H:%M')}] {resumen}"
        self.todo_el_historial[self.chat_actual_id] = self.historial_conversacion
        try:
            with open(self.ruta_historial, "w", encoding="utf-8") as f:
                json.dump(self.todo_el_historial, f, ensure_ascii=False, indent=2)
            self.after(0, self._actualizar_historial_ui)
        except Exception as e:
            print(f"Error guardando historial: {e}")

    def _actualizar_historial_ui(self):
        for w in self.frame_historial.winfo_children():
            w.destroy()
        for chat_id in reversed(list(self.todo_el_historial.keys())):
            fila = ctk.CTkFrame(self.frame_historial, fg_color="transparent")
            fila.pack(fill="x", pady=3)
            ctk.CTkButton(fila, text=chat_id, fg_color="transparent",
                          text_color=COLOR_TEXT_MAIN, anchor="w", hover_color=COLOR_BG_SURFACE, height=32,
                          command=lambda cid=chat_id: self._cargar_chat(cid)).pack(side="left", fill="x", expand=True, padx=(0, 4))
            ctk.CTkButton(fila, text="❌", width=32, height=32,
                          fg_color=COLOR_DANGER, hover_color=COLOR_DANGER_HOVER,
                          command=lambda cid=chat_id: self._borrar_chat(cid)).pack(side="right")

    def _cargar_chat(self, chat_id):
        self.chat_actual_id = chat_id
        self.historial_conversacion = self.todo_el_historial[chat_id]
        self.txt_chat.configure(state="normal")
        self.txt_chat.delete("1.0", "end")
        nombre = self.entry_nombre.get().strip() or "Tú"
        for msg in self.historial_conversacion:
            if msg["role"] == "user":
                self.txt_chat.insert("end", f"\n\n 👤 {nombre}:\n{msg['content']}\n")
            else:
                self.txt_chat.insert("end", f"\n 🤖 IA:\n{msg['content']}\n")
        self.txt_chat.configure(state="disabled")
        self.txt_chat.see("end")

    def _borrar_chat(self, chat_id):
        if messagebox.askyesno("Confirmar", f"¿Borrar este chat?\n'{chat_id}'"):
            del self.todo_el_historial[chat_id]
            try:
                with open(self.ruta_historial, "w", encoding="utf-8") as f:
                    json.dump(self.todo_el_historial, f, ensure_ascii=False, indent=2)
            except Exception:
                pass
            if self.chat_actual_id == chat_id:
                self._limpiar_chat()
            self._actualizar_historial_ui()
