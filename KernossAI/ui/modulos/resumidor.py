"""
KernossAI - Módulo: Resumidor Inteligente de Textos & PDFs
Síntesis académica avanzada con IA, control de temas, progreso visual y exportación.
"""

import threading
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document

from KernossAI.core.theme import (
    COLOR_BG_CARD,
    COLOR_BG_CARD_LIGHT,
    COLOR_BORDER,
    COLOR_ACCENT_PRIMARY,
    COLOR_ACCENT_HOVER,
    COLOR_ACCENT_CYAN,
    COLOR_ACCENT_SKY,
    COLOR_TEXT_MAIN,
    COLOR_SUCCESS,
    COLOR_SUCCESS_HOVER,
    COLOR_DANGER,
)
from KernossAI.core.i18n import t
from KernossAI.core.auth import llamar_groq
from KernossAI.core.tts import tts_engine


class ModuloResumidor(ctk.CTkFrame):
    """Módulo de resumen y síntesis conceptual con LLM y lectura de voz."""
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.modelo = "groq"
        self.instrucciones = (
            "Eres un experto en el tema proporcionado. Tu conocimiento se basa estrictamente en hechos reales. "
            "REGLA DE SEGURIDAD ABSOLUTA: Solo puedes responder a temas que pertenezcan al ámbito educativo, "
            "académico, histórico o laboral. Si el usuario te pide algo fuera de estos ámbitos, DEBES responder "
            "ÚNICAMENTE con la frase: 'ERROR: La petición no pertenece al ámbito educativo o laboral.' "
            "Si la petición es válida, redacta un texto muy extenso, preciso y con párrafos bien estructurados "
            "explicando el contexto, las causas y las consecuencias. No inventes datos bajo ninguna circunstancia."
        )
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        header = ctk.CTkFrame(self, fg_color="transparent")
        header.grid(row=0, column=0, padx=20, pady=20, sticky="ew")
        ctk.CTkLabel(header, text=t("resum_titulo"),
                     font=("Segoe UI", 28, "bold"), text_color=COLOR_ACCENT_SKY).pack(side="left")
        self.entry_nombre = ctk.CTkEntry(header, placeholder_text=t("placeholder_nombre"), width=200,
                                         fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_nombre.pack(side="right", padx=10)

        main = ctk.CTkFrame(self, fg_color="transparent")
        main.grid(row=1, column=0, padx=20, sticky="nsew")
        main.grid_columnconfigure((0, 1), weight=1)
        main.grid_rowconfigure(0, weight=1)

        input_f = ctk.CTkFrame(main, fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER, corner_radius=14)
        input_f.grid(row=0, column=0, padx=(0, 10), sticky="nsew")
        ctk.CTkLabel(input_f, text=t("resum_lbl_entrada"), font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=10)
        self.txt_input = ctk.CTkTextbox(input_f, font=("Segoe UI", 13), fg_color=COLOR_BG_CARD_LIGHT, wrap="word")
        self.txt_input.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        output_f = ctk.CTkFrame(main, fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER, corner_radius=14)
        output_f.grid(row=0, column=1, padx=(10, 0), sticky="nsew")
        ctk.CTkLabel(output_f, text=t("resum_lbl_salida"), font=("Segoe UI", 13, "bold"), text_color=COLOR_ACCENT_CYAN).pack(pady=10)
        self.txt_output = ctk.CTkTextbox(output_f, font=("Segoe UI", 13), fg_color=COLOR_BG_CARD_LIGHT, wrap="word")
        self.txt_output.pack(fill="both", expand=True, padx=15, pady=(0, 15))

        footer = ctk.CTkFrame(self, fg_color="transparent")
        footer.grid(row=2, column=0, padx=20, pady=20, sticky="ew")
        self.progress_bar = ctk.CTkProgressBar(footer, progress_color=COLOR_ACCENT_CYAN)
        self.progress_bar.pack(fill="x", pady=(0, 15))
        self.progress_bar.set(0)

        self.btn_procesar = ctk.CTkButton(footer, text=t("resum_btn_resumir"),
                                          height=45, font=("Segoe UI", 14, "bold"),
                                          fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                          command=self.iniciar_proceso)
        self.btn_procesar.pack(side="left", fill="x", expand=True, padx=(0, 10))

        self.btn_tts_resumen = ctk.CTkButton(footer, text=t("btn_escuchar"), height=45, width=140,
                                            font=("Segoe UI", 12, "bold"),
                                            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                            hover_color=COLOR_ACCENT_HOVER,
                                            command=self._toggle_tts)
        self.btn_tts_resumen.pack(side="left", padx=(0, 10))

        ctk.CTkButton(footer, text=t("btn_word"), height=45, width=180,
                      fg_color=COLOR_SUCCESS, hover_color=COLOR_SUCCESS_HOVER,
                      font=("Segoe UI", 13, "bold"),
                      command=self.exportar_word).pack(side="right")

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_resumen.configure(text=t("btn_escuchar"), fg_color=COLOR_BG_CARD)
        else:
            texto = self.txt_output.get("1.0", "end-1c").strip()
            if not texto or "ERROR:" in texto:
                messagebox.showinfo("Sin resumen", "Primero genera un resumen para escucharlo en voz alta.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_resumen.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_resumen.configure(text=t("btn_escuchar"), fg_color=COLOR_BG_CARD)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def iniciar_proceso(self):
        texto = self.txt_input.get("1.0", "end-1c").strip()
        if not texto:
            messagebox.showwarning("Atención", "Por favor, introduce el texto que deseas resumir.")
            return
        self.txt_output.delete("1.0", "end")
        self.btn_procesar.configure(state="disabled")
        self.progress_bar.configure(mode="indeterminate")
        self.progress_bar.start()
        threading.Thread(target=self._ejecutar_ia, args=(texto,), daemon=True).start()

    def _ejecutar_ia(self, texto):
        try:
            resultado = llamar_groq(
                f"{self.instrucciones}\n\nDesarrolla o resume de manera extensa y rigurosa: {texto}"
            )
            self.after(0, self._escribir_output, resultado)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error Cloud", f"Fallo al conectar con IA: {e}"))
        finally:
            self.after(0, self._finalizar)

    def _escribir_output(self, char):
        self.txt_output.insert("end", char)
        self.txt_output.see("end")

    def _finalizar(self):
        self.progress_bar.stop()
        self.progress_bar.set(1)
        self.btn_procesar.configure(state="normal")

    def exportar_word(self):
        contenido = self.txt_output.get("1.0", "end-1c").strip()
        if not contenido or "ERROR:" in contenido:
            messagebox.showwarning("No se puede guardar", "No hay un resumen válido para exportar.")
            return
        nombre = self.entry_nombre.get().strip() or "Usuario"
        ruta = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"Resumen_{nombre}.docx")
        if ruta:
            try:
                doc = Document()
                doc.add_heading(f'Resumen Académico de {nombre}', 0)
                doc.add_paragraph(contenido)
                doc.save(ruta)
                messagebox.showinfo("Éxito", f"Archivo guardado en:\n{ruta}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar: {e}")
