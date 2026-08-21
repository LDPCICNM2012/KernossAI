"""
KernossAI - Dashboard Principal de Estudios
Contenedor principal con chat interactivo, historial cloud, conmutador de módulos y actualizaciones.
"""

import os
import sys
import json
import tempfile
import threading
import subprocess
import webbrowser
from datetime import datetime
import requests
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document

from KernossAI.core.theme import (
    VERSION_APP,
    COLOR_BG_DARK,
    COLOR_BG_SIDEBAR,
    COLOR_BG_CARD,
    COLOR_BG_CARD_LIGHT,
    COLOR_BG_SURFACE,
    COLOR_BORDER,
    COLOR_ACCENT_PRIMARY,
    COLOR_ACCENT_HOVER,
    COLOR_ACCENT_CYAN,
    COLOR_ACCENT_SKY,
    COLOR_ACCENT_PURPLE,
    COLOR_ACCENT_PURPLE_HOVER,
    COLOR_TEXT_MAIN,
    COLOR_TEXT_MUTED,
    COLOR_TEXT_DIM,
    COLOR_SUCCESS,
    COLOR_DANGER,
    COLOR_DANGER_HOVER,
    COLOR_WARNING,
    aplicar_icono,
    construir_prompt,
    es_version_superior,
)
from KernossAI.core.config import (
    obtener_fecha_instalacion,
)
from KernossAI.core.i18n import t
from KernossAI.core.tts import tts_engine
from KernossAI.core.auth import (
    borrar_token,
    verificar_estado_baneo,
    obtener_estado_casa,
    fijar_red_hogar_actual,
    obtener_chats_cloud,
    guardar_chat_cloud,
    borrar_chat_cloud,
    llamar_groq,
    llamar_gemini,
)
from KernossAI.ui.login.home_alert import ModalAlertaCasa
from KernossAI.ui.modals import (
    VentanaAjustes,
    VentanaNovedadesIA,
    VentanaSoporteE2EE,
    VentanaBandejaSoporte,
    VentanaAdminModeracion,
    VentanaTutoriaAlumnoProfesor,
)
from KernossAI.ui.modulos.calculador import ModuloCalculador
from KernossAI.ui.modulos.resumidor import ModuloResumidor
from KernossAI.ui.modulos.mapa_mental import ModuloMapaMental
from KernossAI.ui.modulos.examenes import ModuloExamen
from KernossAI.ui.modulos.apuntador import ModuloApuntador
from KernossAI.ui.modulos.ayudante import ModuloAyudador
from KernossAI.ui.modulos.agenda import ModuloCalendario
from KernossAI.ui.modulos.profesor import ModuloCreadorEjercicios, ModuloCorrectorExamenes


class DashboardEstudios(ctk.CTk):
    """Ventana principal de la aplicación educativa KernossAI."""
    def __init__(self, sesion):
        super().__init__()
        self.sesion = sesion
        self.rol = sesion.get("rol", "Alumno")
        self.nombre = sesion.get("nombre", "Usuario")
        self.email = sesion.get("email", "")

        self.title(f"KernossAI – {self.rol}: {self.nombre}")
        self.geometry("1420x840")
        self.minsize(1050, 680)
        self.configure(fg_color=COLOR_BG_DARK)
        aplicar_icono(self)

        self._modulos: dict[str, ctk.CTkFrame] = {}
        self._modulo_activo = None

        self.modelo_chat_home = "groq"
        self.historial_chat_home = []
        suffix = self.email.replace("@", "_").replace(".", "_") if self.email else "default"
        self.ruta_historial_home = os.path.expanduser(f"~/.historial_home_{suffix}.json")
        self.todo_el_historial_home = self._cargar_historial_home()
        self.chat_home_id_actual = None

        self._botones_modulos = {}
        self._datos_actualizacion = None
        self.frame_banner_update = None
        self._build_ui()
        self._mostrar_home_chat()

        threading.Thread(target=self._sincronizar_chats_cloud, daemon=True).start()
        threading.Thread(target=self._comprobar_actualizaciones, daemon=True).start()

        self._vigilante_activo = True
        self._iniciar_vigilante_baneos()

        self.after(600, self._comprobar_changelog_post_actualizacion)
        self.after(900, self._comprobar_politica_hogar_inicio)

    def _build_ui(self):
        # Sidebar
        self.sidebar = ctk.CTkFrame(self, width=285, corner_radius=0,
                                   fg_color=COLOR_BG_SIDEBAR,
                                   border_width=1, border_color=COLOR_BORDER)
        self.sidebar.pack(side="left", fill="y")
        self.sidebar.pack_propagate(False)

        frame_brand = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        frame_brand.pack(fill="x", padx=18, pady=(24, 10))

        ctk.CTkLabel(frame_brand, text=t("app_nombre"),
                     font=("Segoe UI", 24, "bold"), text_color=COLOR_ACCENT_SKY).pack(anchor="w")
        ctk.CTkLabel(frame_brand, text=t("app_tagline"),
                     font=("Segoe UI", 11), text_color=COLOR_TEXT_DIM).pack(anchor="w")

        frame_user = ctk.CTkFrame(self.sidebar, fg_color=COLOR_BG_CARD, corner_radius=12,
                                  border_width=1, border_color=COLOR_BORDER)
        frame_user.pack(fill="x", padx=15, pady=(10, 10))
        icono = "🎓" if self.rol == "Alumno" else "👨‍🏫"
        self.lbl_perfil_nombre = ctk.CTkLabel(frame_user, text=f"{icono} {self.nombre}",
                                              font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MAIN, anchor="w")
        self.lbl_perfil_nombre.pack(fill="x", padx=12, pady=(10, 2))

        self.lbl_perfil_email = ctk.CTkLabel(frame_user, text=self.email,
                                             font=("Segoe UI", 10), text_color=COLOR_TEXT_MUTED, anchor="w")
        self.lbl_perfil_email.pack(fill="x", padx=12)

        rol_texto = t("lbl_rol_alumno") if self.rol == "Alumno" else t("lbl_rol_profesor")
        badge_color = COLOR_ACCENT_PRIMARY if self.rol == "Alumno" else COLOR_ACCENT_PURPLE
        self.lbl_perfil_rol = ctk.CTkLabel(frame_user, text=f"  {rol_texto}  ",
                                           font=("Segoe UI", 10, "bold"), fg_color=badge_color,
                                           corner_radius=8, text_color="white")
        self.lbl_perfil_rol.pack(anchor="w", padx=12, pady=(6, 10))

        self.btn_home = ctk.CTkButton(
            self.sidebar, text=t("btn_home"),
            font=("Segoe UI", 13, "bold"), height=42, anchor="w",
            fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
            command=self._mostrar_home_chat
        )
        self.btn_home.pack(fill="x", padx=15, pady=(4, 6))

        frame_historial_header = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        frame_historial_header.pack(fill="x", padx=16, pady=(8, 2))
        ctk.CTkLabel(frame_historial_header, text=t("hdr_historial"),
                     font=("Segoe UI", 10, "bold"), text_color=COLOR_ACCENT_CYAN).pack(side="left")

        ctk.CTkButton(frame_historial_header, text=t("btn_nuevo_chat"), width=55, height=22,
                      font=("Segoe UI", 10, "bold"), fg_color=COLOR_BG_SURFACE,
                      hover_color=COLOR_ACCENT_HOVER,
                      command=self._nuevo_chat_home).pack(side="right")

        self.scroll_historial_home = ctk.CTkScrollableFrame(self.sidebar, height=95, fg_color=COLOR_BG_CARD,
                                                           border_width=1, border_color=COLOR_BORDER,
                                                           corner_radius=8)
        self.scroll_historial_home.pack(fill="x", padx=15, pady=(2, 8))
        self._actualizar_historial_home_ui()

        ctk.CTkFrame(self.sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=4)

        ctk.CTkLabel(self.sidebar, text=t("hdr_modulos_estudio"),
                     font=("Segoe UI", 10, "bold"), text_color=COLOR_TEXT_DIM).pack(anchor="w", padx=20, pady=(6, 2))

        self._btn(t("mod_mapas"), "mapa_mental")
        self._btn(t("mod_calculador"), "calculador")
        self._btn(t("mod_apuntador"), "apuntador")
        self._btn(t("mod_resumidor"), "resumidor")
        self._btn(t("mod_examenes"), "examen")
        self._btn(t("mod_ayudante"), "ayudador")
        self._btn(t("mod_agenda"), "calendario")

        if self.rol == "Profesor":
            ctk.CTkFrame(self.sidebar, height=1, fg_color=COLOR_BORDER).pack(fill="x", padx=15, pady=6)
            ctk.CTkLabel(self.sidebar, text=t("hdr_herramientas_docente"),
                         font=("Segoe UI", 10, "bold"), text_color=COLOR_ACCENT_PURPLE).pack(anchor="w", padx=20, pady=(2, 2))
            self._btn(t("mod_ejercicios"), "creador", color=COLOR_BG_SURFACE)
            self._btn(t("mod_corrector"), "corrector", color=COLOR_BG_SURFACE)

        ctk.CTkButton(self.sidebar, text=t("btn_ajustes"), height=36,
                      fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER,
                      text_color=COLOR_TEXT_MAIN, hover_color=COLOR_ACCENT_HOVER,
                      command=self._abrir_ajustes).pack(fill="x", padx=15, pady=(4, 6), side="bottom")

        ctk.CTkButton(self.sidebar, text=t("btn_cerrar_sesion"), height=36,
                      fg_color="transparent", border_width=1, border_color=COLOR_BORDER,
                      text_color=COLOR_TEXT_MUTED, hover_color=COLOR_DANGER_HOVER,
                      command=self._cerrar_sesion).pack(fill="x", padx=15, pady=(4, 16), side="bottom")

        # Contenedor derecho
        self.contenedor = ctk.CTkFrame(self, corner_radius=0, fg_color=COLOR_BG_DARK)
        self.contenedor.pack(side="right", fill="both", expand=True)
        self.contenedor.grid_rowconfigure(1, weight=1)
        self.contenedor.grid_columnconfigure(0, weight=1)

        self.header_top = ctk.CTkFrame(self.contenedor, fg_color=COLOR_BG_DARK, height=44)
        self.header_top.grid(row=0, column=0, sticky="ew", padx=25, pady=(12, 4))
        self.header_top.grid_columnconfigure(0, weight=1)

        self.lbl_seccion_actual = ctk.CTkLabel(
            self.header_top, text=t("btn_home"),
            font=("Segoe UI", 13, "bold"), text_color=COLOR_TEXT_MUTED
        )
        self.lbl_seccion_actual.grid(row=0, column=0, sticky="w")

        frame_top_derecha = ctk.CTkFrame(self.header_top, fg_color="transparent")
        frame_top_derecha.grid(row=0, column=1, sticky="e")

        es_admin = (self.email.lower() in ["kernossai@support.com", "admin@kernosai.com", "soporte@kernosai.com"]) or self.sesion.get("is_premium", False)
        if es_admin:
            self.btn_admin_top = ctk.CTkButton(
                frame_top_derecha, text="👑 Moderación & Bans",
                font=("Segoe UI", 11, "bold"), height=32,
                fg_color="#3b0764", border_width=1, border_color="#c084fc",
                text_color="#f3e8ff", hover_color="#6b21a8",
                command=self._abrir_admin_moderacion
            )
            self.btn_admin_top.pack(side="right", padx=(6, 0))

        texto_tutoria = t("btn_tutoria_alumno") if self.rol == "Alumno" else t("btn_tutoria_profesor")
        self.btn_tutoria_top = ctk.CTkButton(
            frame_top_derecha, text=texto_tutoria,
            font=("Segoe UI", 11, "bold"), height=32,
            fg_color="#1e1b4b" if self.rol == "Alumno" else "#312e81",
            border_width=1, border_color="#818cf8",
            text_color="#e0e7ff", hover_color="#4338ca",
            command=self._abrir_tutoria_profesor
        )
        self.btn_tutoria_top.pack(side="right", padx=(6, 0))

        self.btn_soporte_top = ctk.CTkButton(
            frame_top_derecha, text="🛡️ Soporte Oficial",
            font=("Segoe UI", 11, "bold"), height=32,
            fg_color="#064e3b", border_width=1, border_color="#34d399",
            text_color="#a7f3d0", hover_color="#059669",
            command=self._abrir_soporte_e2ee
        )
        self.btn_soporte_top.pack(side="right", padx=(6, 0))

        self.btn_novedades_top = ctk.CTkButton(
            frame_top_derecha, text=t("btn_novedades"),
            font=("Segoe UI", 11, "bold"), height=32,
            fg_color="#0c234a", border_width=1, border_color="#38bdf8",
            text_color="#38bdf8", hover_color="#0284c7",
            command=self._abrir_modal_novedades_ia
        )
        self.btn_novedades_top.pack(side="right", padx=(6, 0))

        self.frame_contenido = ctk.CTkFrame(self.contenedor, fg_color="transparent")
        self.frame_contenido.grid(row=1, column=0, sticky="nsew")
        self.frame_contenido.grid_rowconfigure(0, weight=1)
        self.frame_contenido.grid_columnconfigure(0, weight=1)

        self._crear_vista_home_chat()

    def _abrir_ajustes(self):
        VentanaAjustes(self)

    def _abrir_modal_novedades_ia(self):
        VentanaNovedadesIA(self)

    def _abrir_soporte_e2ee(self):
        es_admin = (self.email.lower() in ["kernossai@support.com", "admin@kernosai.com", "soporte@kernosai.com"]) or self.sesion.get("is_premium", False)
        if es_admin:
            VentanaBandejaSoporte(self, self.sesion)
        else:
            VentanaSoporteE2EE(self, self.sesion)

    def _abrir_tutoria_profesor(self):
        VentanaTutoriaAlumnoProfesor(self, self.sesion)

    def _al_cambiar_cuenta(self, nueva_sesion):
        self.sesion = nueva_sesion
        self.email = nueva_sesion.get("email", "").lower()
        self.nombre = nueva_sesion.get("nombre", "Usuario")
        self.rol = nueva_sesion.get("rol", "Alumno")
        self.is_premium = nueva_sesion.get("is_premium", False)

        icono = "🎓" if self.rol == "Alumno" else "👨‍🏫"
        if hasattr(self, "lbl_perfil_nombre"):
            self.lbl_perfil_nombre.configure(text=f"{icono} {self.nombre}")
        if hasattr(self, "lbl_perfil_email"):
            self.lbl_perfil_email.configure(text=self.email)
        if hasattr(self, "lbl_perfil_rol"):
            rol_texto = t("lbl_rol_alumno") if self.rol == "Alumno" else t("lbl_rol_profesor")
            badge_color = COLOR_ACCENT_PRIMARY if self.rol == "Alumno" else COLOR_ACCENT_PURPLE
            self.lbl_perfil_rol.configure(text=f"  {rol_texto}  ", fg_color=badge_color)

        if hasattr(self, "btn_tutoria_top"):
            txt_t = t("btn_tutoria_alumno") if self.rol == "Alumno" else t("btn_tutoria_profesor")
            fg_t = "#1e1b4b" if self.rol == "Alumno" else "#312e81"
            self.btn_tutoria_top.configure(text=txt_t, fg_color=fg_t)

        es_admin = (self.email in ["kernossai@support.com", "admin@kernosai.com", "soporte@kernosai.com"]) or self.is_premium
        if hasattr(self, "btn_admin_top"):
            if es_admin:
                self.btn_admin_top.pack(side="right", padx=(6, 0))
            else:
                self.btn_admin_top.pack_forget()

        self._mostrar_home_chat()

    def _recargar_idioma_ui(self):
        if self._modulo_activo:
            try:
                self._modulo_activo.grid_forget()
            except Exception:
                pass
            self._modulo_activo = None
        for m in list(self._modulos.values()):
            try:
                m.destroy()
            except Exception:
                pass
        self._modulos.clear()
        self._botones_modulos.clear()

        for w in self.winfo_children():
            try:
                w.destroy()
            except Exception:
                pass

        self._build_ui()
        self._mostrar_home_chat()

    def _iniciar_vigilante_baneos(self):
        def _loop():
            while self._vigilante_activo:
                threading.Event().wait(3.5)
                if not self._vigilante_activo:
                    break
                try:
                    ok_clean, motivo_ban, tipo_ban = verificar_estado_baneo(self.email)
                    if not ok_clean:
                        self._vigilante_activo = False
                        self.after(0, lambda t=tipo_ban, m=motivo_ban: self._mostrar_pantalla_baneado(t, m))
                        break
                except Exception:
                    pass
        
        threading.Thread(target=_loop, daemon=True).start()

    def _mostrar_pantalla_baneado(self, tipo_ban: str, motivo_ban: str):
        borrar_token()
        for w in self.sidebar.winfo_children():
            try:
                w.configure(state="disabled")
            except Exception:
                pass

        for w in self.contenedor.winfo_children():
            w.destroy()

        frame_lockout = ctk.CTkFrame(self.contenedor, fg_color="#3b0707", corner_radius=0)
        frame_lockout.pack(fill="both", expand=True)

        card_bloqueo = ctk.CTkFrame(
            frame_lockout, fg_color="#180404", corner_radius=20,
            border_width=2, border_color="#ef4444"
        )
        card_bloqueo.place(relx=0.5, rely=0.5, anchor="center")

        ctk.CTkLabel(card_bloqueo, text="⛔", font=("Segoe UI", 56)).pack(padx=50, pady=(26, 4))
        ctk.CTkLabel(
            card_bloqueo, text="ACCESO BLOQUEADO — CUENTA BANEADA",
            font=("Segoe UI", 18, "bold"), text_color="#ef4444"
        ).pack(padx=50, pady=(0, 4))

        ctk.CTkLabel(
            card_bloqueo, text="Has sido sancionado por la administración de KernossAI.\nTu sesión ha sido revocada y tu acceso a la IA queda suspendido.",
            font=("Segoe UI", 11), text_color="#fca5a5", justify="center"
        ).pack(padx=50, pady=(0, 16))

        box_detalles = ctk.CTkFrame(card_bloqueo, fg_color="#260505", corner_radius=10, border_width=1, border_color="#7f1d1d")
        box_detalles.pack(fill="x", padx=30, pady=(0, 18))

        tipo_str = tipo_ban or "Baneo de Cuenta Permanente"
        motivo_str = motivo_ban or "Infracción de las normas de la comunidad"

        ctk.CTkLabel(box_detalles, text=f"🔒 Tipo de Sanción: {tipo_str}", font=("Segoe UI", 11, "bold"), text_color="#fca5a5").pack(anchor="w", padx=16, pady=(12, 3))
        ctk.CTkLabel(box_detalles, text=f"📋 Motivo: {motivo_str}", font=("Segoe UI", 11), text_color="#fee2e2", wraplength=440, justify="left").pack(anchor="w", padx=16, pady=(0, 3))
        ctk.CTkLabel(box_detalles, text="⚡ Estado: Sesión Cancelada • Inteligencia Artificial Bloqueada.", font=("Segoe UI", 10, "bold"), text_color="#f87171").pack(anchor="w", padx=16, pady=(0, 12))

        btn_cerrar = ctk.CTkButton(
            card_bloqueo, text="🚪 Cerrar Sesión y Salir", height=42,
            font=("Segoe UI", 12, "bold"), fg_color="#dc2626", hover_color="#b91c1c",
            command=self._ejecutar_cierre_por_ban
        )
        btn_cerrar.pack(fill="x", padx=30, pady=(0, 26))

    def _ejecutar_cierre_por_ban(self):
        borrar_token()
        self.quit()
        self.destroy()
        os._exit(0)

    def _abrir_admin_moderacion(self):
        VentanaAdminModeracion(self, self.sesion)

    def _comprobar_politica_hogar_inicio(self):
        def _thread():
            _, dias_transcurridos, dias_restantes_gracia = obtener_fecha_instalacion()
            codigo, texto, detalles = obtener_estado_casa(self.email)

            def _evaluar():
                if not self.winfo_exists():
                    return
                if codigo == "desconocido":
                    if dias_restantes_gracia > 0:
                        ModalAlertaCasa(self, tipo="sin_casa", dias_restantes=dias_restantes_gracia, on_casa_establecida=self._al_establecer_casa)
                    else:
                        self._bloquear_dashboard_por_casa("sin_casa")
                elif codigo == "no":
                    if dias_restantes_gracia > 0:
                        ModalAlertaCasa(self, tipo="fuera_de_casa", dias_restantes=dias_restantes_gracia, on_casa_establecida=self._al_establecer_casa)
                    else:
                        self._bloquear_dashboard_por_casa("fuera_de_casa")

            self.after(0, _evaluar)

        threading.Thread(target=_thread, daemon=True).start()

    def _al_establecer_casa(self):
        self._desbloquear_dashboard()

    def _bloquear_dashboard_por_casa(self, motivo: str):
        for mid, btn in self._botones_modulos.items():
            btn.configure(state="disabled")

        for child in self.frame_contenido.winfo_children():
            child.grid_forget()

        if hasattr(self, "frame_bloqueo_casa") and self.frame_bloqueo_casa:
            self.frame_bloqueo_casa.destroy()

        self.frame_bloqueo_casa = ctk.CTkFrame(self.frame_contenido, fg_color=COLOR_BG_DARK)
        self.frame_bloqueo_casa.grid(row=0, column=0, sticky="nsew")
        self.frame_bloqueo_casa.grid_rowconfigure(0, weight=1)
        self.frame_bloqueo_casa.grid_columnconfigure(0, weight=1)

        card_bloqueo = ctk.CTkFrame(self.frame_bloqueo_casa, fg_color=COLOR_BG_CARD, corner_radius=16,
                                    border_width=2, border_color=COLOR_DANGER)
        card_bloqueo.grid(row=0, column=0, padx=40, pady=40)

        ctk.CTkLabel(card_bloqueo, text="🚫", font=("Segoe UI", 48)).pack(pady=(24, 6))

        if motivo == "sin_casa":
            titulo_b = "Servicios Bloqueados: No hay Casa Establecida"
            desc_b = (
                "Han transcurrido los 15 días de gracia desde la instalación de KernossAI.\n\n"
                "Para desbloquear todos los servicios y módulos, debes registrar tu conexión\n"
                "actual como tu Hogar Principal de Estudio."
            )
        else:
            titulo_b = "Servicios Bloqueados: Fuera de tu Hogar Principal"
            desc_b = (
                "No estás conectado en tu red de Hogar Principal de Estudio registrada.\n\n"
                "Para desbloquear el acceso puedes:\n"
                "• Establecer esta red como tu nuevo Hogar Principal.\n"
                "• O activar el Pase de Hogar Temporal por 7 días (disponible 1 vez al mes desde Ajustes)."
            )

        ctk.CTkLabel(card_bloqueo, text=titulo_b, font=("Segoe UI", 16, "bold"), text_color=COLOR_DANGER).pack(padx=24, pady=(0, 8))
        ctk.CTkLabel(card_bloqueo, text=desc_b, font=("Segoe UI", 11), text_color=COLOR_TEXT_MUTED, justify="center").pack(padx=28, pady=(0, 20))

        btn_desbloquear = ctk.CTkButton(
            card_bloqueo, text="🏡 Establecer Esta Red como mi Casa y Desbloquear", height=42,
            font=("Segoe UI", 12, "bold"), fg_color=COLOR_SUCCESS, hover_color="#15803d",
            command=self._desbloquear_fijando_casa
        )
        btn_desbloquear.pack(fill="x", padx=28, pady=(0, 8))

        btn_ir_ajustes = ctk.CTkButton(
            card_bloqueo, text="⚙️ Abrir Ajustes de Red & Pases", height=36,
            font=("Segoe UI", 11, "bold"), fg_color=COLOR_BG_SURFACE, border_width=1, border_color=COLOR_BORDER,
            hover_color=COLOR_BORDER, command=self._abrir_ajustes
        )
        btn_ir_ajustes.pack(fill="x", padx=28, pady=(0, 24))

    def _desbloquear_fijando_casa(self):
        def _thread():
            ok, msg = fijar_red_hogar_actual(self.email)
            if ok:
                self.after(0, lambda: messagebox.showinfo("Hogar Actualizado", msg))
                self.after(0, self._desbloquear_dashboard)
            else:
                self.after(0, lambda: messagebox.showerror("Error", msg))
        threading.Thread(target=_thread, daemon=True).start()

    def _desbloquear_dashboard(self):
        for mid, btn in self._botones_modulos.items():
            btn.configure(state="normal")

        if hasattr(self, "frame_bloqueo_casa") and self.frame_bloqueo_casa:
            self.frame_bloqueo_casa.destroy()
            self.frame_bloqueo_casa = None

        self._mostrar_home_chat()

    def _btn(self, texto, modulo_id, color=COLOR_BG_SURFACE):
        btn = ctk.CTkButton(
            self.sidebar, text=texto,
            font=("Segoe UI", 12, "bold"), height=36, anchor="w",
            fg_color=color, hover_color=COLOR_ACCENT_HOVER,
            command=lambda mid=modulo_id: self._abrir_modulo(mid)
        )
        btn.pack(fill="x", padx=15, pady=2)
        self._botones_modulos[modulo_id] = btn

    def _crear_vista_home_chat(self):
        self.frame_home = ctk.CTkFrame(self.frame_contenido, fg_color="transparent")
        self.frame_home.grid_rowconfigure(2, weight=1)
        self.frame_home.grid_columnconfigure(0, weight=1)

        header_home = ctk.CTkFrame(self.frame_home, fg_color="transparent")
        header_home.grid(row=0, column=0, sticky="ew", padx=25, pady=(10, 10))

        rol_nombre = t("lbl_rol_alumno") if self.rol == "Alumno" else t("lbl_rol_profesor")
        saludo = t("home_bienvenida", rol=rol_nombre, nombre=self.nombre)
        ctk.CTkLabel(header_home, text=saludo, font=("Segoe UI", 26, "bold"), text_color=COLOR_TEXT_MAIN).pack(anchor="w")

        sub_header = ctk.CTkFrame(header_home, fg_color="transparent")
        sub_header.pack(fill="x", pady=(2, 0))

        self.lbl_estado_ia_home = ctk.CTkLabel(sub_header, text=t("home_estado_ia"),
                                               font=("Segoe UI", 12), text_color=COLOR_SUCCESS)
        self.lbl_estado_ia_home.pack(side="left")

        frame_selector_ia = ctk.CTkFrame(sub_header, fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER, corner_radius=8)
        frame_selector_ia.pack(side="right")

        self.btn_home_groq = ctk.CTkButton(frame_selector_ia, text=t("opt_groq"), height=28, width=115,
                                           font=("Segoe UI", 11, "bold"),
                                           fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                           command=lambda: self._set_home_modelo("groq"))
        self.btn_home_groq.pack(side="left", padx=2, pady=2)

        self.btn_home_gemini = ctk.CTkButton(frame_selector_ia, text=t("opt_gemini"), height=28, width=125,
                                             font=("Segoe UI", 11, "bold"),
                                             fg_color="transparent", hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                             command=lambda: self._set_home_modelo("gemini"))
        self.btn_home_gemini.pack(side="left", padx=2, pady=2)

        bar_modulos = ctk.CTkFrame(self.frame_home, fg_color=COLOR_BG_CARD, corner_radius=12,
                                   border_width=1, border_color=COLOR_BORDER)
        bar_modulos.grid(row=1, column=0, sticky="ew", padx=25, pady=(0, 12))

        ctk.CTkLabel(bar_modulos, text="⚡", font=("Segoe UI", 11, "bold"),
                     text_color=COLOR_TEXT_MUTED).pack(side="left", padx=(14, 4), pady=8)

        modulos_rapidos = [
            (t("mod_mapas"), "mapa_mental"),
            (t("mod_calculador"), "calculador"),
            (t("mod_apuntador"), "apuntador"),
            (t("mod_resumidor"), "resumidor"),
            (t("mod_examenes"), "examen"),
            (t("mod_ayudante"), "ayudador"),
            (t("mod_agenda"), "calendario")
        ]
        if self.rol == "Profesor":
            modulos_rapidos.extend([(t("mod_ejercicios"), "creador"), (t("mod_corrector"), "corrector")])

        for label, mid in modulos_rapidos:
            ctk.CTkButton(bar_modulos, text=label, height=28, font=("Segoe UI", 11, "bold"),
                          fg_color=COLOR_BG_CARD_LIGHT, hover_color=COLOR_ACCENT_HOVER,
                          border_width=1, border_color=COLOR_BORDER,
                          command=lambda m=mid: self._abrir_modulo(m)).pack(side="left", padx=3, pady=6)

        self.txt_home_chat = ctk.CTkTextbox(self.frame_home, font=("Segoe UI", 14), state="disabled", wrap="word",
                                            fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER,
                                            corner_radius=14)
        self.txt_home_chat.grid(row=2, column=0, sticky="nsew", padx=25, pady=(0, 12))

        self._inicializar_chat_bienvenida_texto()

        input_container = ctk.CTkFrame(self.frame_home, fg_color="transparent")
        input_container.grid(row=3, column=0, sticky="ew", padx=25, pady=(0, 20))
        input_container.grid_columnconfigure(0, weight=1)

        self.entry_home_pregunta = ctk.CTkEntry(input_container, placeholder_text=t("home_placeholder_input"),
                                                height=46, font=("Segoe UI", 13),
                                                fg_color=COLOR_BG_CARD, border_color=COLOR_BORDER)
        self.entry_home_pregunta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_home_pregunta.bind("<Return>", lambda e: self._enviar_chat_home())

        btn_enviar_home = ctk.CTkButton(input_container, text=t("btn_consultar_ia"), width=130, height=46,
                                       font=("Segoe UI", 13, "bold"),
                                       fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                       command=self._enviar_chat_home)
        btn_enviar_home.grid(row=0, column=1, padx=(0, 6))

        self.btn_tts_home = ctk.CTkButton(input_container, text=t("btn_escuchar"), width=105, height=46,
                                          font=("Segoe UI", 12, "bold"),
                                          fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_ACCENT_CYAN,
                                          hover_color=COLOR_ACCENT_HOVER,
                                          command=self._toggle_tts_home)
        self.btn_tts_home.grid(row=0, column=2, padx=(0, 6))

        btn_word_home = ctk.CTkButton(input_container, text=t("btn_word"), width=75, height=46,
                                      font=("Segoe UI", 12, "bold"),
                                      fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER,
                                      hover_color=COLOR_ACCENT_PURPLE_HOVER,
                                      command=self._exportar_home_word)
        btn_word_home.grid(row=0, column=3)

    def _inicializar_chat_bienvenida_texto(self):
        self.txt_home_chat.configure(state="normal")
        self.txt_home_chat.delete("1.0", "end")
        texto_intro = (
            f"👋 ¡Hola, {self.nombre}! Soy tu asistente académico inteligente en KernossAI.\n\n"
            f"• Pregúntame dudas de cualquier asignatura (Matemáticas, Historia, Ciencias, Idiomas...).\n"
            f"• Pídeme resúmenes, esquemas explicativos o resolución paso a paso de problemas.\n"
            f"• O cambia a cualquiera de los módulos dedicados en el menú lateral o en la barra superior.\n"
            f"──────────────────────────────────────────────────────────────────────────\n"
        )
        self.txt_home_chat.insert("end", texto_intro)
        self.txt_home_chat.configure(state="disabled")

    def _set_home_modelo(self, m):
        self.modelo_chat_home = m
        if m == "groq":
            self.btn_home_groq.configure(fg_color=COLOR_ACCENT_PRIMARY)
            self.btn_home_gemini.configure(fg_color="transparent")
            self.lbl_estado_ia_home.configure(text="🟢 Asistente IA Activo (Groq - LLaMA 3.3 70B Ultrarrápido)", text_color=COLOR_SUCCESS)
        else:
            self.btn_home_groq.configure(fg_color="transparent")
            self.btn_home_gemini.configure(fg_color=COLOR_ACCENT_PURPLE)
            self.lbl_estado_ia_home.configure(text="🟣 Asistente IA Activo (Google Gemini - Análisis Avanzado)", text_color="#a5b4fc")

    def _enviar_chat_home(self):
        pregunta = self.entry_home_pregunta.get().strip()
        if not pregunta:
            return

        self.txt_home_chat.configure(state="normal")
        self.txt_home_chat.insert("end", f"\n\n👤 {self.nombre}:\n{pregunta}\n")
        self.txt_home_chat.see("end")
        self.txt_home_chat.configure(state="disabled")

        self.entry_home_pregunta.delete(0, "end")
        self.lbl_estado_ia_home.configure(text="🟡 Pensando y analizando respuesta...", text_color=COLOR_WARNING)
        self.historial_chat_home.append({"role": "user", "content": pregunta})

        threading.Thread(target=self._thread_proceso_home, daemon=True).start()

    def _thread_proceso_home(self):
        try:
            self.after(0, lambda: self._stream_home_texto(f"\n🤖 KernossAI ({self.modelo_chat_home.upper()}):\n"))
            instrucciones = (
                "Eres el tutor y asistente de estudio principal de la plataforma KernossAI. "
                "Responde de forma clara, didáctica, amigable y rigurosa a las consultas académicas del usuario."
            )
            prompt = construir_prompt(instrucciones, self.historial_chat_home)
            if self.modelo_chat_home == "groq":
                respuesta = llamar_groq(prompt)
            else:
                respuesta = llamar_gemini(prompt)

            self.ultima_respuesta_ia_home = respuesta
            self.historial_chat_home.append({"role": "assistant", "content": respuesta})
            self.after(0, self._stream_home_texto, respuesta + "\n")
            self._guardar_historial_home()

            color = COLOR_SUCCESS if self.modelo_chat_home == "groq" else "#a5b4fc"
            self.after(0, lambda: self.lbl_estado_ia_home.configure(
                text=f"🟢 Listo • {self.modelo_chat_home.upper()} respondió correctamente", text_color=color))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error IA", str(e)))
            self.after(0, lambda: self.lbl_estado_ia_home.configure(text="🔴 Error al procesar respuesta", text_color=COLOR_DANGER))
        finally:
            self.after(0, lambda: self.txt_home_chat.configure(state="disabled"))

    def _toggle_tts_home(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_home.configure(text="🔊 Escuchar", fg_color=COLOR_BG_CARD)
        else:
            texto = getattr(self, "ultima_respuesta_ia_home", "")
            if not texto:
                texto = self.txt_home_chat.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("Sin texto", "No hay respuesta para reproducir en voz alta.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_home.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_home.configure(text="🔊 Escuchar", fg_color=COLOR_BG_CARD)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def _stream_home_texto(self, texto):
        self.txt_home_chat.configure(state="normal")
        self.txt_home_chat.insert("end", texto)
        self.txt_home_chat.see("end")
        self.txt_home_chat.configure(state="disabled")

    def _nuevo_chat_home(self):
        self._mostrar_home_chat()
        self.historial_chat_home = []
        self.chat_home_id_actual = None
        self._inicializar_chat_bienvenida_texto()

    def _exportar_home_word(self):
        if not self.historial_chat_home:
            messagebox.showwarning("Vacío", "No hay mensajes en el chat para exportar.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"Conversacion_KernossAI_{datetime.now().strftime('%Y%m%d_%H%M')}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Sesión de Estudio – {self.nombre} (KernossAI)", 0)
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph("")
            for msg in self.historial_chat_home:
                rol = self.nombre if msg["role"] == "user" else "KernossAI Tutor"
                p = doc.add_paragraph()
                p.add_run(f"{rol}: ").bold = True
                p.add_run(msg["content"])
            doc.save(path)
            messagebox.showinfo("Éxito", f"Conversación guardada en:\n{path}")

    def _cargar_historial_home(self):
        if os.path.exists(self.ruta_historial_home):
            try:
                with open(self.ruta_historial_home, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                return {}
        return {}

    def _sincronizar_chats_cloud(self):
        try:
            chats_remotos = obtener_chats_cloud()
            if chats_remotos and isinstance(chats_remotos, dict):
                cambios = False
                for chat_id, data in chats_remotos.items():
                    mensajes = data.get("mensajes", [])
                    if chat_id not in self.todo_el_historial_home and mensajes:
                        self.todo_el_historial_home[chat_id] = mensajes
                        cambios = True
                if cambios:
                    with open(self.ruta_historial_home, "w", encoding="utf-8") as f:
                        json.dump(self.todo_el_historial_home, f, ensure_ascii=False, indent=2)
                    self.after(0, self._actualizar_historial_home_ui)
        except Exception as e:
            print(f"[Sync Cloud] {e}")

    def _guardar_historial_home(self):
        if not self.historial_chat_home:
            return
        if self.chat_home_id_actual is None:
            primer = next((m["content"] for m in self.historial_chat_home if m["role"] == "user"), "Consulta")
            resumen = primer[:20] + "..." if len(primer) > 20 else primer
            self.chat_home_id_actual = f"[{datetime.now().strftime('%H:%M')}] {resumen}"
        self.todo_el_historial_home[self.chat_home_id_actual] = self.historial_chat_home
        
        try:
            with open(self.ruta_historial_home, "w", encoding="utf-8") as f:
                json.dump(self.todo_el_historial_home, f, ensure_ascii=False, indent=2)
            self.after(0, self._actualizar_historial_home_ui)
        except Exception as e:
            print(f"Error guardando historial local: {e}")

        chat_id_actual = self.chat_home_id_actual
        mensajes_actuales = list(self.historial_chat_home)
        threading.Thread(
            target=lambda: guardar_chat_cloud(chat_id_actual, chat_id_actual, mensajes_actuales),
            daemon=True
        ).start()

    def _actualizar_historial_home_ui(self):
        for w in self.scroll_historial_home.winfo_children():
            w.destroy()
        for chat_id in reversed(list(self.todo_el_historial_home.keys())):
            fila = ctk.CTkFrame(self.scroll_historial_home, fg_color="transparent")
            fila.pack(fill="x", pady=2)
            ctk.CTkButton(fila, text=chat_id, fg_color="transparent",
                          text_color=COLOR_TEXT_MAIN, anchor="w", hover_color=COLOR_BG_SURFACE, height=26,
                          font=("Segoe UI", 10),
                          command=lambda cid=chat_id: self._cargar_chat_home(cid)).pack(side="left", fill="x", expand=True)
            ctk.CTkButton(fila, text="✕", width=22, height=22,
                          fg_color="transparent", hover_color=COLOR_DANGER_HOVER,
                          text_color=COLOR_TEXT_MUTED, font=("Segoe UI", 9, "bold"),
                          command=lambda cid=chat_id: self._borrar_chat_home(cid)).pack(side="right")

    def _cargar_chat_home(self, chat_id):
        self._mostrar_home_chat()
        self.chat_home_id_actual = chat_id
        self.historial_chat_home = self.todo_el_historial_home.get(chat_id, [])
        self.txt_home_chat.configure(state="normal")
        self.txt_home_chat.delete("1.0", "end")
        self.txt_home_chat.insert("end", f"💬 Cargado: {chat_id}\n{'─'*55}\n")
        for msg in self.historial_chat_home:
            if msg["role"] == "user":
                self.txt_home_chat.insert("end", f"\n👤 {self.nombre}:\n{msg['content']}\n")
            else:
                self.txt_home_chat.insert("end", f"\n🤖 KernossAI:\n{msg['content']}\n")
        self.txt_home_chat.configure(state="disabled")
        self.txt_home_chat.see("end")

    def _borrar_chat_home(self, chat_id):
        if messagebox.askyesno("Confirmar", f"¿Eliminar este chat?\n'{chat_id}'"):
            if chat_id in self.todo_el_historial_home:
                del self.todo_el_historial_home[chat_id]
                try:
                    with open(self.ruta_historial_home, "w", encoding="utf-8") as f:
                        json.dump(self.todo_el_historial_home, f, ensure_ascii=False, indent=2)
                except Exception:
                    pass
                if self.chat_home_id_actual == chat_id:
                    self._nuevo_chat_home()
                self._actualizar_historial_home_ui()

                threading.Thread(
                    target=lambda cid=chat_id: borrar_chat_cloud(cid),
                    daemon=True
                ).start()

    def _mostrar_home_chat(self):
        if self._modulo_activo:
            self._modulo_activo.grid_forget()
            self._modulo_activo = None

        self.btn_home.configure(fg_color=COLOR_ACCENT_PRIMARY)
        for btn in self._botones_modulos.values():
            btn.configure(fg_color=COLOR_BG_SURFACE)

        if hasattr(self, "lbl_seccion_actual"):
            self.lbl_seccion_actual.configure(text="🏠 Inicio / Chat Asistente IA")

        self.frame_home.grid(row=0, column=0, sticky="nsew")

    def _abrir_modulo(self, modulo_id):
        self.frame_home.grid_forget()

        if self._modulo_activo:
            self._modulo_activo.grid_forget()

        self.btn_home.configure(fg_color=COLOR_BG_SURFACE)
        for mid, btn in self._botones_modulos.items():
            if mid == modulo_id:
                btn.configure(fg_color=COLOR_ACCENT_PRIMARY)
            else:
                btn.configure(fg_color=COLOR_BG_SURFACE)

        nombres_secciones = {
            "mapa_mental": "🧠 Generador de Mapas Mentales con IA",
            "calculador": "📊 Calculador de Medias y Ponderaciones",
            "apuntador": "📝 Apuntador de Notas de Clase",
            "resumidor": "🔍 Resumidor Inteligente de Textos",
            "examen": "🎯 Generador y Práctica de Exámenes",
            "ayudador": "🤖 Tutor y Ayudante de Dudas Académicas",
            "calendario": "📅 Agenda y Planificador de Estudios",
            "creador": "✏️ Creador de Ejercicios Didácticos",
            "corrector": "📋 Corrector Inteligente de Exámenes"
        }
        if hasattr(self, "lbl_seccion_actual"):
            self.lbl_seccion_actual.configure(text=nombres_secciones.get(modulo_id, "📂 Módulo de Estudio"))

        if modulo_id not in self._modulos:
            if modulo_id == "mapa_mental":
                self._modulos[modulo_id] = ModuloMapaMental(self.frame_contenido)
            elif modulo_id == "calculador":
                self._modulos[modulo_id] = ModuloCalculador(self.frame_contenido)
            elif modulo_id == "apuntador":
                self._modulos[modulo_id] = ModuloApuntador(self.frame_contenido)
            elif modulo_id == "resumidor":
                self._modulos[modulo_id] = ModuloResumidor(self.frame_contenido)
            elif modulo_id == "examen":
                self._modulos[modulo_id] = ModuloExamen(self.frame_contenido)
            elif modulo_id == "ayudador":
                self._modulos[modulo_id] = ModuloAyudador(self.frame_contenido, sesion=self.sesion)
            elif modulo_id == "calendario":
                self._modulos[modulo_id] = ModuloCalendario(self.frame_contenido)
            elif modulo_id == "creador":
                self._modulos[modulo_id] = ModuloCreadorEjercicios(self.frame_contenido)
            elif modulo_id == "corrector":
                self._modulos[modulo_id] = ModuloCorrectorExamenes(self.frame_contenido)

        modulo = self._modulos[modulo_id]
        modulo.grid(row=0, column=0, sticky="nsew")
        self._modulo_activo = modulo

    def _cerrar_sesion(self):
        if messagebox.askyesno("Cerrar Sesión", "¿Seguro que quieres cerrar sesión?"):
            borrar_token()
            self.quit()
            self.destroy()
            os._exit(0)

    def _comprobar_actualizaciones(self):
        try:
            url = "https://api.github.com/repos/LDPCICNM2012/KernossAI/releases/latest"
            headers = {"User-Agent": "KernossAI-Desktop-App"}
            resp = requests.get(url, headers=headers, timeout=6)
            if resp.status_code == 200:
                data = resp.json()
                tag_remoto = data.get("tag_name", "")
                if tag_remoto and es_version_superior(tag_remoto, VERSION_APP):
                    assets = data.get("assets", [])
                    download_url = data.get("html_url", "https://github.com/LDPCICNM2012/KernossAI/releases/latest")
                    
                    if sys.platform == "win32":
                        for a in assets:
                            if a.get("name", "").lower().endswith(".exe"):
                                download_url = a.get("browser_download_url", download_url)
                                break
                    elif sys.platform == "darwin":
                        for a in assets:
                            if a.get("name", "").lower().endswith((".zip", ".dmg")):
                                download_url = a.get("browser_download_url", download_url)
                                break

                    info = {
                        "tag": tag_remoto,
                        "url": download_url,
                        "html_url": data.get("html_url", ""),
                        "notas": data.get("body", "Mejoras de rendimiento, corrección de errores y nuevas funciones.")
                    }
                    self._datos_actualizacion = info
                    self.after(500, lambda: self._mostrar_notificacion_actualizacion(info))
        except Exception:
            pass

    def _mostrar_notificacion_actualizacion(self, info):
        tag = info["tag"]
        dl_url = info["url"]

        if hasattr(self, "sidebar") and not hasattr(self, "btn_update_sidebar"):
            self.btn_update_sidebar = ctk.CTkButton(
                self.sidebar,
                text=f"🚀 Actualizar a {tag}",
                font=("Segoe UI", 11, "bold"),
                height=32,
                fg_color="#0284c7",
                hover_color="#0ea5e9",
                command=lambda: self._abrir_modal_actualizacion(info)
            )
            self.btn_update_sidebar.pack(fill="x", padx=15, pady=(2, 4), after=self.btn_home)

        if hasattr(self, "frame_home") and self.frame_banner_update is None:
            self.frame_banner_update = ctk.CTkFrame(
                self.frame_home,
                fg_color="#0c234a",
                border_width=1,
                border_color="#38bdf8",
                corner_radius=10
            )
            self.frame_banner_update.grid(row=0, column=0, sticky="ew", padx=25, pady=(15, 0))

            frame_txt = ctk.CTkFrame(self.frame_banner_update, fg_color="transparent")
            frame_txt.pack(side="left", padx=15, pady=10)

            ctk.CTkLabel(
                frame_txt,
                text=f"🎉 ¡Nueva actualización disponible ({tag})!",
                font=("Segoe UI", 13, "bold"),
                text_color="#38bdf8"
            ).pack(anchor="w")

            ctk.CTkLabel(
                frame_txt,
                text=f"Tu versión actual es v{VERSION_APP}. Haz clic en 'Descargar' para actualizar tu app.",
                font=("Segoe UI", 11),
                text_color=COLOR_TEXT_MUTED
            ).pack(anchor="w")

            frame_btns = ctk.CTkFrame(self.frame_banner_update, fg_color="transparent")
            frame_btns.pack(side="right", padx=15, pady=10)

            ctk.CTkButton(
                frame_btns,
                text="⚡ Descargar",
                font=("Segoe UI", 12, "bold"),
                fg_color=COLOR_ACCENT_PRIMARY,
                hover_color=COLOR_ACCENT_HOVER,
                height=30,
                command=lambda: webbrowser.open(dl_url)
            ).pack(side="left", padx=4)

            ctk.CTkButton(
                frame_btns,
                text="ℹ️ Novedades",
                font=("Segoe UI", 11),
                fg_color=COLOR_BG_SURFACE,
                hover_color=COLOR_BORDER,
                height=30,
                command=lambda: self._abrir_modal_actualizacion(info)
            ).pack(side="left", padx=4)

            ctk.CTkButton(
                frame_btns,
                text="✕",
                width=26,
                height=30,
                fg_color="transparent",
                hover_color=COLOR_DANGER_HOVER,
                command=self._cerrar_banner_update
            ).pack(side="left", padx=2)

    def _cerrar_banner_update(self):
        if self.frame_banner_update:
            self.frame_banner_update.destroy()
            self.frame_banner_update = None

    def _descargar_e_instalar_ota(self, info, progress_bar, lbl_status, btn_ota, btn_web, modal):
        url_descarga = info.get("url", "")
        if sys.platform != "win32" or not url_descarga.lower().endswith(".exe"):
            webbrowser.open(info.get("html_url", url_descarga))
            modal.destroy()
            return

        btn_ota.configure(state="disabled", text=t("ota_descargando"))
        btn_web.configure(state="disabled")
        progress_bar.pack(fill="x", padx=30, pady=(6, 4))
        progress_bar.set(0.0)
        lbl_status.pack(anchor="w", padx=30, pady=(0, 10))
        lbl_status.configure(text=t("ota_descargando"), text_color=COLOR_ACCENT_SKY)

        def _worker():
            try:
                tag_limpio = info.get("tag", "latest").replace(" ", "_")
                ruta_temp = os.path.join(tempfile.gettempdir(), f"KernossAI_Setup_OTA_{tag_limpio}.exe")

                resp = requests.get(url_descarga, stream=True, timeout=90)
                total_bytes = int(resp.headers.get("content-length", 0))
                descargados = 0

                with open(ruta_temp, "wb") as f:
                    for chunk in resp.iter_content(chunk_size=65536):
                        if chunk:
                            f.write(chunk)
                            descargados += len(chunk)
                            if total_bytes > 0:
                                p = descargados / total_bytes
                                mb_act = descargados / (1024 * 1024)
                                mb_tot = total_bytes / (1024 * 1024)
                                pct = int(p * 100)
                                self.after(0, lambda p_val=p, txt=f"⬇️ {mb_act:.1f} MB / {mb_tot:.1f} MB ({pct}%)...": (
                                    progress_bar.set(p_val),
                                    lbl_status.configure(text=txt)
                                ))

                self.after(0, lambda: (
                    progress_bar.set(1.0),
                    lbl_status.configure(text=t("ota_instalando"), text_color=COLOR_SUCCESS)
                ))

                cmd = f'"{ruta_temp}" /VERYSILENT /SUPPRESSMSGBOXES /FORCECLOSEAPPLICATIONS'
                subprocess.Popen(cmd, shell=True)

                def _salir():
                    try:
                        self.quit()
                        self.destroy()
                    except Exception:
                        pass
                    os._exit(0)

                self.after(1200, _salir)

            except Exception as e:
                err_msg = str(e)
                self.after(0, lambda: (
                    lbl_status.configure(text=t("ota_error", err=err_msg), text_color=COLOR_DANGER),
                    btn_ota.configure(state="normal", text="🔄 Reintentar OTA"),
                    btn_web.configure(state="normal")
                ))

        threading.Thread(target=_worker, daemon=True).start()

    def _abrir_modal_actualizacion(self, info):
        modal = ctk.CTkToplevel(self)
        modal.title(f"Actualización de KernossAI ({info['tag']})")
        modal.geometry("560x490")
        modal.minsize(500, 420)
        modal.configure(fg_color=COLOR_BG_DARK)
        modal.transient(self)
        modal.grab_set()

        ctk.CTkLabel(modal, text="🚀 Nueva Versión Disponible",
                     font=("Segoe UI", 18, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(18, 4))
        ctk.CTkLabel(modal, text=f"Versión instalada: v{VERSION_APP}   ➜   Última versión: {info['tag']}",
                     font=("Segoe UI", 12), text_color=COLOR_TEXT_MAIN).pack(pady=(0, 10))

        ctk.CTkLabel(modal, text="Novedades de esta versión:",
                     font=("Segoe UI", 11, "bold"), text_color=COLOR_TEXT_MUTED).pack(anchor="w", padx=30, pady=(4, 2))

        txt_notas = ctk.CTkTextbox(modal, height=140, fg_color=COLOR_BG_CARD,
                                  border_width=1, border_color=COLOR_BORDER,
                                  text_color=COLOR_TEXT_MAIN, font=("Segoe UI", 11))
        txt_notas.pack(fill="both", expand=True, padx=30, pady=(0, 8))
        txt_notas.insert("end", info.get("notas", "Mejoras de rendimiento, estabilidad y nuevas funciones."))
        txt_notas.configure(state="disabled")

        progress_bar = ctk.CTkProgressBar(modal, height=12, progress_color=COLOR_ACCENT_CYAN)
        lbl_status = ctk.CTkLabel(modal, text="", font=("Segoe UI", 10, "bold"), text_color=COLOR_TEXT_MUTED)

        frame_modal_btns = ctk.CTkFrame(modal, fg_color="transparent")
        frame_modal_btns.pack(fill="x", padx=30, pady=(8, 20))

        btn_ota = ctk.CTkButton(
            frame_modal_btns,
            text=t("ota_btn_actualizar"),
            font=("Segoe UI", 12, "bold"),
            fg_color=COLOR_ACCENT_PRIMARY,
            hover_color=COLOR_ACCENT_HOVER,
            height=38,
            command=lambda: self._descargar_e_instalar_ota(info, progress_bar, lbl_status, btn_ota, btn_web, modal)
        )
        btn_ota.pack(side="left", fill="x", expand=True, padx=(0, 6))

        btn_web = ctk.CTkButton(
            frame_modal_btns,
            text=t("ota_btn_web"),
            font=("Segoe UI", 11),
            fg_color=COLOR_BG_CARD,
            border_width=1,
            border_color=COLOR_BORDER,
            hover_color=COLOR_BG_SURFACE,
            height=38,
            command=lambda: [webbrowser.open(info.get("html_url", info["url"])), modal.destroy()]
        )
        btn_web.pack(side="left", padx=(0, 6))

        ctk.CTkButton(
            frame_modal_btns,
            text="✕",
            font=("Segoe UI", 12),
            fg_color=COLOR_BG_SURFACE,
            hover_color=COLOR_DANGER_HOVER,
            height=38,
            width=40,
            command=modal.destroy
        ).pack(side="right")

    def _comprobar_changelog_post_actualizacion(self):
        ruta_version_vista = os.path.expanduser("~/.kernoss_version_seen.json")
        version_vista = ""
        try:
            if os.path.exists(ruta_version_vista):
                with open(ruta_version_vista, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    version_vista = data.get("ultima_version_vista", "")
        except Exception:
            pass

        if version_vista != VERSION_APP:
            try:
                with open(ruta_version_vista, "w", encoding="utf-8") as f:
                    json.dump({"ultima_version_vista": VERSION_APP}, f)
            except Exception:
                pass

            self._mostrar_modal_bienvenida_changelog()

    def _mostrar_modal_bienvenida_changelog(self):
        modal = ctk.CTkToplevel(self)
        modal.title(t("modal_novedades_titulo", version=VERSION_APP))
        modal.geometry("620x540")
        modal.minsize(540, 460)
        modal.configure(fg_color=COLOR_BG_DARK)
        modal.transient(self)
        modal.grab_set()

        ctk.CTkLabel(modal, text=t("modal_novedades_titulo", version=VERSION_APP),
                     font=("Segoe UI", 20, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=(22, 4))
        ctk.CTkLabel(modal, text=t("modal_novedades_subtitulo"),
                     font=("Segoe UI", 12), text_color=COLOR_TEXT_MUTED).pack(pady=(0, 14))

        frame_box = ctk.CTkScrollableFrame(modal, fg_color=COLOR_BG_CARD,
                                           border_width=1, border_color=COLOR_BORDER,
                                           corner_radius=12)
        frame_box.pack(fill="both", expand=True, padx=25, pady=(0, 16))

        novedades = [
            (t("nov_item_soporte_tit"), t("nov_item_soporte_desc")),
            (t("nov_item_multi_tit"), t("nov_item_multi_desc")),
            (t("nov_item_rol_tit"), t("nov_item_rol_desc")),
            (t("nov_item_mod_tit"), t("nov_item_mod_desc")),
            (t("nov_item_cloud_tit"), t("nov_item_cloud_desc")),
            (t("nov_item_hwid_tit"), t("nov_item_hwid_desc")),
        ]

        for titulo, desc in novedades:
            item = ctk.CTkFrame(frame_box, fg_color=COLOR_BG_CARD_LIGHT, corner_radius=8,
                                border_width=1, border_color=COLOR_BORDER)
            item.pack(fill="x", pady=5, padx=2)
            ctk.CTkLabel(item, text=titulo, font=("Segoe UI", 12, "bold"),
                         text_color=COLOR_ACCENT_CYAN).pack(anchor="w", padx=12, pady=(8, 2))
            ctk.CTkLabel(item, text=desc, font=("Segoe UI", 11),
                         text_color=COLOR_TEXT_MAIN, wraplength=500, justify="left").pack(anchor="w", padx=12, pady=(0, 8))

        ctk.CTkButton(modal, text=t("modal_novedades_btn_empezar"),
                      font=("Segoe UI", 13, "bold"), height=42,
                      fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                      command=modal.destroy).pack(fill="x", padx=25, pady=(0, 20))

    def _al_cerrar(self):
        self._vigilante_activo = False
        self.quit()
        self.destroy()
        os._exit(0)
