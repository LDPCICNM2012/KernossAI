"""
KernossAI - Módulo: Generador y Corrector de Exámenes
Creación de evaluaciones con preguntas tipo test y desarrollo, autocorrección y exportación.
"""

import threading
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx import Document

from KernossAI.core.theme import (
    COLOR_BG_CARD,
    COLOR_BG_CARD_LIGHT,
    COLOR_BG_SURFACE,
    COLOR_BORDER,
    COLOR_ACCENT_PRIMARY,
    COLOR_ACCENT_HOVER,
    COLOR_ACCENT_CYAN,
    COLOR_ACCENT_CYAN_HOVER,
    COLOR_ACCENT_SKY,
    COLOR_TEXT_MAIN,
    COLOR_SUCCESS,
    COLOR_DANGER,
    construir_prompt,
)
from KernossAI.core.i18n import t
from KernossAI.core.auth import llamar_groq
from KernossAI.core.tts import tts_engine


class ModuloExamen(ctk.CTkFrame):
    """Módulo de evaluación y simulacros de examen con retroalimentación inmediata."""
    def __init__(self, master):
        super().__init__(master, fg_color="transparent")
        self.examen_en_memoria = ""
        self.historial_conversacion = []
        self.instrucciones_base = (
            "Eres un evaluador académico profesional. El examen debe tener:\n"
            "1. Un título relevante.\n"
            "2. Preguntas de opción múltiple (A-E) o completar huecos.\n"
            "3. La mitad de preguntas de desarrollo.\n"
            "REGLA CRÍTICA: No des las respuestas hasta que el usuario responda. "
            "No pongas las respuestas correctas en el examen."
        )
        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        sidebar = ctk.CTkFrame(self, width=280, corner_radius=0,
                               fg_color=COLOR_BG_CARD, border_width=1, border_color=COLOR_BORDER)
        sidebar.grid(row=0, column=0, sticky="nsew")
        ctk.CTkLabel(sidebar, text=t("exam_titulo"), font=("Segoe UI", 16, "bold"), text_color=COLOR_ACCENT_SKY).pack(pady=20, padx=10)
        self.entry_nombre = ctk.CTkEntry(sidebar, placeholder_text=t("placeholder_nombre"),
                                         fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_nombre.pack(fill="x", padx=20, pady=8)

        ctk.CTkLabel(sidebar, text="🟢 AI Ready", text_color=COLOR_SUCCESS,
                     font=("Segoe UI", 12, "bold")).pack(fill="x", padx=20, pady=5)
        ctk.CTkLabel(sidebar, text=t("exam_lbl_tema"), font=("Segoe UI", 12, "bold"), text_color=COLOR_TEXT_MAIN).pack(pady=(15, 0), anchor="w", padx=20)
        self.txt_tema = ctk.CTkTextbox(sidebar, height=130, fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.txt_tema.pack(fill="x", padx=20, pady=8)

        self.btn_generar = ctk.CTkButton(sidebar, text=t("exam_btn_generar"),
                                         fg_color=COLOR_ACCENT_PRIMARY, hover_color=COLOR_ACCENT_HOVER,
                                         height=40, font=("Segoe UI", 12, "bold"),
                                         command=self.iniciar_generacion)
        self.btn_generar.pack(fill="x", padx=20, pady=8)

        self.btn_tts_examen = ctk.CTkButton(sidebar, text=t("btn_escuchar"), fg_color=COLOR_BG_CARD_LIGHT,
                                            border_width=1, border_color=COLOR_ACCENT_CYAN,
                                            hover_color=COLOR_ACCENT_HOVER,
                                            height=36, font=("Segoe UI", 11, "bold"),
                                            command=self._toggle_tts)
        self.btn_tts_examen.pack(fill="x", padx=20, pady=4)

        ctk.CTkButton(sidebar, text=t("btn_word"), fg_color=COLOR_BG_SURFACE, border_width=1,
                      border_color=COLOR_BORDER, hover_color=COLOR_ACCENT_HOVER,
                      height=36, command=self.exportar_word).pack(fill="x", padx=20, pady=4)

        self.status_label = ctk.CTkLabel(sidebar, text="🟢 Ready",
                                         text_color=COLOR_SUCCESS, font=("Segoe UI", 11))
        self.status_label.pack(side="bottom", pady=15)

        main_f = ctk.CTkFrame(self, fg_color="transparent")
        main_f.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        main_f.grid_rowconfigure(0, weight=1)
        main_f.grid_columnconfigure(0, weight=1)

        self.output_text = ctk.CTkTextbox(main_f, font=("Consolas", 13),
                                          fg_color=COLOR_BG_CARD_LIGHT, border_width=1, border_color=COLOR_BORDER)
        self.output_text.grid(row=0, column=0, sticky="nsew", pady=(0, 10))

        input_f = ctk.CTkFrame(main_f, fg_color="transparent")
        input_f.grid(row=1, column=0, sticky="ew")
        input_f.grid_columnconfigure(0, weight=1)
        self.entry_respuesta = ctk.CTkEntry(input_f, placeholder_text="Escribe aquí tus respuestas...", height=42,
                                            fg_color=COLOR_BG_CARD_LIGHT, border_color=COLOR_BORDER)
        self.entry_respuesta.grid(row=0, column=0, sticky="ew", padx=(0, 10))
        self.entry_respuesta.bind("<Return>", lambda e: self.enviar_respuesta())
        ctk.CTkButton(input_f, text="Enviar Respuestas", width=140, height=42,
                      fg_color=COLOR_ACCENT_CYAN, hover_color=COLOR_ACCENT_CYAN_HOVER,
                      font=("Segoe UI", 12, "bold"),
                      command=self.enviar_respuesta).grid(row=0, column=1)

    def _toggle_tts(self):
        if tts_engine.esta_reproduciendo():
            tts_engine.detener()
            self.btn_tts_examen.configure(text=t("btn_escuchar"), fg_color=COLOR_BG_CARD_LIGHT)
        else:
            texto = self.output_text.get("1.0", "end-1c").strip()
            if not texto:
                messagebox.showinfo("KernossAI", "Primero genera un examen para escucharlo.")
                return

            def _cb(rep):
                if rep:
                    self.btn_tts_examen.configure(text="⏹️ Detener", fg_color=COLOR_DANGER)
                else:
                    self.btn_tts_examen.configure(text=t("btn_escuchar"), fg_color=COLOR_BG_CARD_LIGHT)

            tts_engine.hablar(texto, callback_estado=lambda r: self.after(0, lambda: _cb(r)))

    def iniciar_generacion(self):
        tema = self.txt_tema.get("1.0", "end-1c").strip()
        if not tema:
            messagebox.showwarning("Atención", "Escribe un tema para el examen.")
            return
        self.output_text.delete("1.0", "end")
        self.output_text.insert("end", f"Generando examen sobre: {tema}...\n\n")
        self.btn_generar.configure(state="disabled")
        self.status_label.configure(text="🟡 Elaborando Examen...", text_color="orange")
        threading.Thread(target=self._proceso_groq, args=(tema,), daemon=True).start()

    def _proceso_groq(self, tema):
        try:
            prompt = f"{self.instrucciones_base}\n\nHazme un examen sobre: {tema}"
            full = llamar_groq(prompt)
            self.examen_en_memoria = full
            self.historial_conversacion = [
                {'role': 'system', 'content': self.instrucciones_base},
                {'role': 'assistant', 'content': full}
            ]
            self.after(0, self._update_output, full)
            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status_label.configure(text="🟢 Examen Listo", text_color=COLOR_SUCCESS))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error Cloud", str(e)))
            self.after(0, lambda: self.btn_generar.configure(state="normal"))
            self.after(0, lambda: self.status_label.configure(text="🔴 Error de red", text_color=COLOR_DANGER))

    def enviar_respuesta(self):
        msg = self.entry_respuesta.get().strip()
        if not msg:
            return
        self.output_text.insert("end", f"\n\n👤 TÚ: {msg}\n\n🤖 CORRECCIÓN IA: ")
        self.entry_respuesta.delete(0, "end")
        self.status_label.configure(text="🟡 Corrigiendo...", text_color="orange")
        self.historial_conversacion.append({'role': 'user', 'content': msg})
        threading.Thread(target=self._proceso_respuesta, daemon=True).start()

    def _proceso_respuesta(self):
        try:
            full = llamar_groq(construir_prompt(self.instrucciones_base, self.historial_conversacion))
            self.historial_conversacion.append({'role': 'assistant', 'content': full})
            self.after(0, self._update_output, full)
            self.after(0, lambda: self.status_label.configure(text="🟢 Corrección Completada", text_color=COLOR_SUCCESS))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Error", str(e)))

    def _update_output(self, texto):
        self.output_text.insert("end", texto)
        self.output_text.see("end")

    def exportar_word(self):
        texto = self.output_text.get("1.0", "end-1c")
        if not texto.strip():
            messagebox.showwarning("Vacío", "No hay contenido para exportar.")
            return
        nombre = self.entry_nombre.get().strip() or "Usuario"
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"Examen_{nombre}.docx")
        if path:
            doc = Document()
            doc.add_heading(f"Examen de {nombre} – KernossAI", 0)
            doc.add_paragraph(texto)
            doc.save(path)
            messagebox.showinfo("Éxito", f"Archivo guardado en:\n{path}")
